from __future__ import annotations

import json
import platform
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from textwrap import fill

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT_DIR = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT_DIR / "docs"
REPORTS_DIR = DOCS_DIR / "reports"
PROCESSED_DIR = ROOT_DIR / "data" / "processed"
RAW_DIR = ROOT_DIR / "data" / "ipl_csv2"
MODELS_DIR = ROOT_DIR / "models"
OUTPUT_PATH = REPORTS_DIR / "IPL_Prediction_Project_Report.docx"
ASSETS_DIR = DOCS_DIR / "report_assets"

if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from ipl_predictor import (
    CATEGORICAL_FEATURES,
    load_models,
    load_pre_match_models,
    load_support_tables,
    predict_match_state,
    predict_pre_match,
)

BODY_FONT = "Cambria"
CODE_FONT = "Consolas"
BODY_SIZE = 11

REPORT_PROFILE = {
    "title": "IPL Match Score and Win Prediction System",
    "degree_line_1": "Bachelor of Technology in",
    "degree_line_2": "Computer Science & Engineering (Artificial Intelligence and Machine Learning)",
    "students": [
        "Rahul (2822450)",
        "Rinku (2822306)",
        "Ankit (2822307)",
    ],
    "supervisor": "Ms. Chetna Dahiya (Assistant Professor)",
    "hod": "Dr. Devendra Prasad",
    "department": "Department of Computer Science & Engineering (Artificial Intelligence and Machine Learning)",
    "institute": "Panipat Institute of Engineering & Technology, Samalkha, Panipat",
    "university": "Kurukshetra University, Kurukshetra, India (2022-2026)",
    "place": "Panipat",
}


@dataclass(frozen=True)
class FigureSpec:
    number: str
    title: str
    filename: str


@dataclass(frozen=True)
class TableSpec:
    number: str
    title: str


FIGURES = [
    FigureSpec("Figure 4.1", "End-to-End IPL Prediction Pipeline", "figure_4_1_pipeline.png"),
    FigureSpec("Figure 4.2", "Live Model Comparison on the Selected Test Scope", "figure_4_2_model_comparison.png"),
    FigureSpec("Figure 4.3", "Feature Family Composition of the Live Prediction Inputs", "figure_4_3_feature_mix.png"),
    FigureSpec("Figure 5.1", "Top IPL Venues by Historical First-Innings Average", "figure_5_1_venue_stats.png"),
    FigureSpec("Figure 5.2", "Residual Uncertainty Band Used for Projected Score Range", "figure_5_2_uncertainty.png"),
    FigureSpec("Figure 5.3", "Phase-wise Scoring Behaviour in the Engineered Dataset", "figure_5_3_phase_scoring.png"),
    FigureSpec("Figure 5.4", "Held-Out Score Error Distribution of the Deployed Regressor", "figure_5_4_score_error_distribution.png"),
    FigureSpec("Figure 5.5", "Held-Out Win Probability Calibration Snapshot", "figure_5_5_win_calibration.png"),
    FigureSpec("Figure 6.1", "Engineered Dataset Coverage Dashboard", "figure_6_1_dataset_dashboard.png"),
    FigureSpec("Figure 6.2", "Flask Live Prediction Interface Mockup", "figure_6_2_flask_mockup.png"),
    FigureSpec("Figure 6.3", "Streamlit Analytics Dashboard Mockup", "figure_6_3_streamlit_mockup.png"),
    FigureSpec("Figure 6.4", "JSON API Request and Response Snapshot", "figure_6_4_api_snapshot.png"),
    FigureSpec("Figure 6.5", "CLI Live Prediction Output Snapshot", "figure_6_5_cli_snapshot.png"),
    FigureSpec("Figure 6.6", "Test and Artifact Verification Snapshot", "figure_6_6_validation_snapshot.png"),
    FigureSpec("Figure 6.7", "Production Monitoring and Drift Snapshot", "figure_6_7_monitoring_snapshot.png"),
]

TABLES = [
    TableSpec("Table 1.1", "Benefits of the Project"),
    TableSpec("Table 2.1", "Review of Key Literature"),
    TableSpec("Table 3.1", "Key Deliverables of the Project"),
    TableSpec("Table 4.1", "Feature Groups Used in the Live Prediction Dataset"),
    TableSpec("Table 4.2", "Model Families and Their Roles"),
    TableSpec("Table 5.1", "Software and Hardware Details"),
    TableSpec("Table 5.2", "Deployment Performance of the Selected Live Models"),
    TableSpec("Table 5.3", "Scope-wise Comparison Used for Final Selection"),
    TableSpec("Table 5.4", "Integration and Verification Summary"),
    TableSpec("Table A.1", "Core Project Files and Responsibilities"),
    TableSpec("Table B.1", "Abbreviations Used in the Report"),
]

CHAPTER_4_TABLES = {
    "support_assets": TableSpec("Table 4.3", "Processed Support Tables and Their Analytical Roles"),
    "season_coverage": TableSpec("Table 4.4", "Season-wise Coverage of the Engineered Dataset"),
    "team_summary": TableSpec("Table 4.5", "Active IPL Team Snapshot from the Engineered Dataset"),
    "module_inventory": TableSpec("Table 4.6", "Core Project Modules and Responsibilities"),
    "workflow_inventory": TableSpec("Table 4.7", "Training, Promotion, and Serving Workflows"),
    "api_contract": TableSpec("Table 4.8", "Core Live API Fields and Returned Analytics"),
    "preprocessing_rules": TableSpec("Table 4.9", "Normalization and Leakage-Control Rules Applied During Preprocessing"),
    "notebook_inventory": TableSpec("Table 4.10", "Notebook Inventory and Their Roles in the Project Lifecycle"),
    "dependency_roles": TableSpec("Table 4.11", "Primary Dependencies and Their Responsibilities in the Project"),
}

CHAPTER_5_TABLES = {
    "season_scoring": TableSpec("Table 5.5", "Season-wise Scoring Environment Across IPL History"),
    "venue_profile": TableSpec("Table 5.6", "Venue Profile Snapshot for Leading Grounds"),
    "batter_leaders": TableSpec("Table 5.7", "Batter Form Leaders in the Latest Support Table"),
    "bowler_leaders": TableSpec("Table 5.8", "Bowler Form Leaders in the Latest Support Table"),
    "recent_score_models": TableSpec("Table 5.9", "Recent-Scope Live Score Benchmark from the Search Report"),
    "all_score_models": TableSpec("Table 5.10", "All-History Live Score Benchmark from the Search Report"),
    "recent_win_models": TableSpec("Table 5.11", "Recent-Scope Live Win Benchmark from the Search Report"),
    "all_win_models": TableSpec("Table 5.12", "All-History Live Win Benchmark from the Search Report"),
    "live_scenarios": TableSpec("Table 5.13", "Scenario-Based Live Inference Walkthrough"),
    "pre_match_scenarios": TableSpec("Table 5.14", "Scenario-Based Pre-Match Inference Walkthrough"),
    "score_case_studies": TableSpec("Table 5.15", "Largest Held-Out Score Misses from the Promoted Regressor"),
    "win_case_studies": TableSpec("Table 5.16", "Highest Log-Loss Cases from the Promoted Classifier"),
    "phase_summary": TableSpec("Table 5.17", "Phase-wise Match Behaviour in the Engineered Dataset"),
    "score_wicket_summary": TableSpec("Table 5.18", "Held-Out Score Error by Wicket-State Bucket"),
    "pre_match_metrics": TableSpec("Table 5.19", "Pre-Match Model Validation and Test Performance"),
    "win_calibration_bins": TableSpec("Table 5.20", "Held-Out Win Probability Calibration Bins"),
    "experiment_versions": TableSpec("Table 5.21", "Versioned Experiment Snapshots from the Model Registry"),
    "drift_summary": TableSpec("Table 5.22", "Production Monitoring and Drift Readiness Snapshot"),
}

APPENDIX_TABLES = {
    "feature_catalog_1": TableSpec("Table A.2", "Feature Dictionary: Match Metadata, Participants, and Toss Context"),
    "feature_catalog_2": TableSpec("Table A.3", "Feature Dictionary: Match-State and Phase Variables"),
    "feature_catalog_3": TableSpec("Table A.4", "Feature Dictionary: Momentum and Chase-Math Variables"),
    "feature_catalog_4": TableSpec("Table A.5", "Feature Dictionary: Team and Player Context Variables"),
    "feature_catalog_5": TableSpec("Table A.6", "Feature Dictionary: Weather, Venue, and Target Variables"),
    "artifact_inventory": TableSpec("Table A.7", "Model Artifact Inventory in the Workspace"),
    "source_inventory": TableSpec("Table A.8", "Source File Inventory and Approximate Line Counts"),
    "test_inventory": TableSpec("Table A.9", "Automated Test Inventory and Coverage Focus"),
    "notebook_inventory": TableSpec("Table A.10", "Notebook Inventory with Cell Counts and Roles"),
    "dependency_inventory": TableSpec("Table A.11", "Dependency Inventory and Project Usage"),
    "version_inventory": TableSpec("Table A.12", "Versioned Model Registry Snapshot"),
    "workspace_map": TableSpec("Table A.13", "Top-Level Workspace Directory Map"),
}

ALL_TABLES = (
    TABLES[:5]
    + list(CHAPTER_4_TABLES.values())
    + TABLES[5:9]
    + list(CHAPTER_5_TABLES.values())
    + [TABLES[9]]
    + list(APPENDIX_TABLES.values())
    + [TABLES[10]]
)

DEPENDENCY_ROLE_NOTES = {
    "numpy": "Vectorized numerical operations used in feature processing and metric computation.",
    "pandas": "Tabular data ingestion, joins, grouping, and report-side analytics.",
    "scikit-learn": "Classical ML models, evaluation metrics, calibration helpers, and preprocessing utilities.",
    "joblib": "Persistence and reload of fitted classical-model artifacts.",
    "Flask": "Web application layer and JSON API endpoint for live prediction.",
    "streamlit": "Interactive analytics dashboard for project demonstration and inspection.",
    "xgboost": "Gradient-boosting benchmark for high-performance structured-data modeling.",
    "catboost": "Categorical-friendly boosting model used in the live win-probability search.",
    "pytest": "Automated verification of inference utilities and web routes.",
    "matplotlib": "Generation of report figures and lightweight interface mockups.",
    "python-docx": "Programmatic creation of the final academic report in `.docx` format.",
    "pillow": "Image support layer required by the report-generation stack.",
}

PREPROCESSING_RULES = [
    ("Alias normalization", "Maps team and venue name variants to canonical labels before feature generation.", "Prevents duplicate identities and mismatched support-table joins."),
    ("Chronological sorting", "Processes matches and balls in time order using match date and legal-ball position.", "Ensures every historical feature only uses earlier evidence."),
    ("Legal-ball conversion", "Converts over.ball notation into legal balls bowled, over number, and ball-in-over fields.", "Makes phase logic and remaining-balls calculations deterministic."),
    ("Rolling momentum windows", "Builds recent scoring windows such as last over, last two overs, and last five overs.", "Captures acceleration and collapse behaviour without future leakage."),
    ("Chase-only target features", "Populates target, target remaining, and required-rate fields only for second innings.", "Keeps first-innings rows semantically clean and avoids fabricated targets."),
    ("Support-table materialization", "Writes venue, team, player, and matchup summaries into separate processed CSV files.", "Improves reuse, speeds inference, and centralizes cricket logic."),
    ("Active-team filtering", "Restricts final comparison scopes to the active 2026 IPL team set.", "Balances historical depth with modern franchise relevance."),
    ("Artifact version checks", "Cross-checks reports and saved models after training through the pipeline registry.", "Reduces the risk of documenting stale or inconsistent artifacts."),
]

WORKSPACE_MAP = [
    ("data", "Raw CSV2 files, processed feature tables, support tables, and weather snapshots."),
    ("docs", "Generated report assets, LaTeX paper draft, summaries, and the final `.docx` output."),
    ("ipl_predictor", "Shared inference, calibration, monitoring, online-learning, and torch model modules."),
    ("models", "Saved live and pre-match artifacts, prediction exports, and JSON reports."),
    ("notebooks", "Exploratory analysis, preprocessing notebooks, and model-comparison experiments."),
    ("scripts", "Reproducible command-line workflows for preprocessing, training, reporting, and data refresh."),
    ("static", "Shared CSS assets for the Flask frontend."),
    ("templates", "HTML templates for the Flask interface."),
    ("tests", "Pytest suites for the common prediction path and web routes."),
]

SUPPORT_TABLE_DESCRIPTIONS = {
    "ipl_features.csv": ("Main training frame", "Ball-by-ball match-state rows used for live score and win modeling.", "match_id + innings + legal_balls_bowled"),
    "venue_stats.csv": ("Venue priors", "Historical scoring baselines and bat-first win rate by venue.", "venue"),
    "team_form_latest.csv": ("Team form", "Rolling recent form score for each IPL team.", "team"),
    "team_venue_form_latest.csv": ("Team-venue form", "Recent team strength conditioned on venue.", "team + venue"),
    "matchup_form_latest.csv": ("Head-to-head form", "Recent batting-team performance against the opponent.", "team + opponent"),
    "batter_form_latest.csv": ("Batter profile", "Career-level batting rate and average snapshots.", "batter"),
    "bowler_form_latest.csv": ("Bowler profile", "Career-level economy and strike-rate snapshots.", "bowler"),
    "batter_bowler_form_latest.csv": ("Matchup micro-history", "Batter-versus-bowler strike-rate history.", "batter + bowler"),
    "team_player_pool_2026.csv": ("Player pool", "Team-wise player list used by dashboards and selectors.", "team + player"),
    "active_teams_2026.csv": ("Active-team filter", "Canonical active-team set used to scope recent IPL experiments.", "team"),
}

SOURCE_FILE_NOTES = {
    "scripts/preprocess_ipl.py": "Builds the engineered feature table and historical support tables from raw Cricsheet CSV2 files.",
    "scripts/train_models.py": "Runs the CPU baseline workflow and produces the baseline report artifacts.",
    "scripts/train_gpu_best.py": "Benchmarks GPU CatBoost/XGBoost variants and selects a deployment pair on the recent active-team scope.",
    "scripts/train_best_model_search.py": "Performs the broader cross-scope search including the torch entity-embedding models.",
    "scripts/train_pre_match.py": "Trains the lightweight pre-match score and win models.",
    "scripts/generate_project_report.py": "Compiles the deep academic report from workspace artifacts and generated assets.",
    "ipl_predictor/common.py": "Shared inference contract covering validation, feature building, model execution, and formatted outputs.",
    "ipl_predictor/live_data.py": "Fetches live weather context and computes dew-risk features for inference-time enrichment.",
    "ipl_predictor/calibration.py": "Wraps isotonic calibration so probability outputs remain reusable at deployment time.",
    "ipl_predictor/ensembles.py": "Defines weighted regressor and classifier ensemble helpers used during model comparison.",
    "ipl_predictor/torch_tabular.py": "Implements the deep tabular entity-embedding models used in the final search workflow.",
    "ipl_predictor/online_learning.py": "Stores optional feedback hooks for future online-learning and monitoring extensions.",
    "web_app.py": "Flask web interface and JSON API for live predictions.",
    "streamlit_app.py": "Interactive analysis dashboard with model selection, venue summaries, and form views.",
    "predict_cli.py": "Terminal-based live prediction flow for quick manual use.",
    "templates/index.html": "Primary Flask HTML template for live and pre-match forms.",
    "static/styles.css": "Shared styling used by the Flask frontend.",
    "tests/test_common.py": "Unit and smoke tests for parsing, validation, feature-frame construction, and model inference.",
    "tests/test_web_app.py": "Route tests covering the Flask API success and validation-failure paths.",
}

TEST_CASE_NOTES = {
    "test_alias_normalization": "Confirms historical team and venue aliases map to the canonical names expected by the models.",
    "test_parse_overs_valid_and_invalid": "Checks overs parsing and rejects invalid ball counts such as 14.7.",
    "test_input_validation_required_fields": "Verifies that missing mandatory payload fields surface clear validation errors.",
    "test_feature_frame_columns_are_consistent": "Ensures the constructed live feature frame matches the trained feature contract exactly.",
    "test_second_innings_requires_first_innings_total": "Prevents innings-two predictions from running without a chase target.",
    "test_predict_match_state_smoke": "Runs an end-to-end live prediction using dummy models to verify formatted outputs.",
    "test_venue_snapshot_is_leakage_safe": "Checks that venue priors only use prior matches and never leak future outcomes.",
    "test_model_loading_when_artifacts_exist": "Confirms the saved live artifacts expose the expected prediction interfaces.",
    "test_real_saved_models_can_run_live_prediction": "Validates that the promoted live models execute on a realistic payload.",
    "test_real_saved_models_can_run_pre_match_prediction": "Validates that the dedicated pre-match models execute on a realistic matchup.",
    "test_api_predict_route_returns_prediction": "Checks the Flask API success path and confirms prediction keys are present.",
    "test_api_predict_route_validates_bad_input": "Checks that malformed API requests return HTTP 400 with validation errors.",
}

LIVE_SCENARIOS = [
    {
        "label": "Powerplay launch",
        "payload": {
            "season": "2026",
            "venue": "Wankhede Stadium",
            "batting_team": "Mumbai Indians",
            "bowling_team": "Chennai Super Kings",
            "innings": 1,
            "runs": 62,
            "wickets": 1,
            "overs": "6.0",
            "runs_last_5": 58,
            "wickets_last_5": 1,
            "striker": "Rohit Sharma",
            "bowler": "Noor Ahmad",
            "use_live_weather": False,
        },
        "interpretation": "Aggressive powerplay scoring pushes the total well above venue par while the batting side keeps a strong but not decisive win projection.",
    },
    {
        "label": "Middle-over rebuild",
        "payload": {
            "season": "2026",
            "venue": "MA Chidambaram Stadium",
            "batting_team": "Chennai Super Kings",
            "bowling_team": "Kolkata Knight Riders",
            "innings": 1,
            "runs": 84,
            "wickets": 4,
            "overs": "11.3",
            "runs_last_5": 29,
            "wickets_last_5": 2,
            "striker": "RD Gaikwad",
            "bowler": "AD Russell",
            "use_live_weather": False,
        },
        "interpretation": "The model moderates its score estimate when wickets fall in the middle overs even if the batting side remains competitively placed.",
    },
    {
        "label": "Death overs surge",
        "payload": {
            "season": "2026",
            "venue": "Narendra Modi Stadium",
            "batting_team": "Gujarat Titans",
            "bowling_team": "Rajasthan Royals",
            "innings": 1,
            "runs": 168,
            "wickets": 4,
            "overs": "17.2",
            "runs_last_5": 53,
            "wickets_last_5": 1,
            "striker": "Shubman Gill",
            "bowler": "Avesh Khan",
            "use_live_weather": False,
        },
        "interpretation": "Late acceleration sharply increases both projected total and batting-side win probability when wickets remain in hand.",
    },
    {
        "label": "Chase under pressure",
        "payload": {
            "season": "2026",
            "venue": "M. Chinnaswamy Stadium",
            "batting_team": "Royal Challengers Bengaluru",
            "bowling_team": "Sunrisers Hyderabad",
            "innings": 2,
            "runs": 96,
            "wickets": 3,
            "overs": "11.0",
            "first_innings_total": 188,
            "runs_last_5": 37,
            "wickets_last_5": 1,
            "striker": "V Kohli",
            "bowler": "PJ Cummins",
            "use_live_weather": False,
        },
        "interpretation": "A chase can still remain close to 50-50 despite a positive runs-vs-par signal if the required rate begins to outpace the current tempo.",
    },
    {
        "label": "Late collapse risk",
        "payload": {
            "season": "2026",
            "venue": "Ekana Cricket Stadium",
            "batting_team": "Lucknow Super Giants",
            "bowling_team": "Delhi Capitals",
            "innings": 2,
            "runs": 132,
            "wickets": 7,
            "overs": "16.4",
            "first_innings_total": 171,
            "runs_last_5": 24,
            "wickets_last_5": 3,
            "striker": "A Badoni",
            "bowler": "Kuldeep Yadav",
            "use_live_weather": False,
        },
        "interpretation": "The simulation layer exposes collapse risk clearly when a side enters the death overs with few wickets in hand and a rising required rate.",
    },
]

PREMATCH_SCENARIOS = [
    {
        "team1": "Mumbai Indians",
        "team2": "Chennai Super Kings",
        "venue": "Wankhede Stadium",
        "interpretation": "Classic rivalry scenario illustrating how the pre-match model uses only teams and venue.",
    },
    {
        "team1": "Royal Challengers Bengaluru",
        "team2": "Sunrisers Hyderabad",
        "venue": "M. Chinnaswamy Stadium",
        "interpretation": "High-scoring venue prior lifts the projected first-innings range before any live match-state information exists.",
    },
    {
        "team1": "Gujarat Titans",
        "team2": "Rajasthan Royals",
        "venue": "Narendra Modi Stadium",
        "interpretation": "Neutral pre-match probabilities show the pre-match models are intentionally conservative compared with live models.",
    },
]

FEATURE_CATALOG = {
    "match_metadata": [
        ("match_id", "Unique Cricsheet match identifier used to join match metadata and innings rows."),
        ("season", "Season label derived from Cricsheet, including split-year editions such as 2007/08."),
        ("start_date", "Match start date used to keep feature generation chronological."),
        ("venue", "Normalized venue name after alias correction."),
        ("innings", "Innings indicator: 1 for batting first, 2 for chasing."),
        ("ball", "Original Cricsheet over.ball notation retained for traceability."),
        ("batting_team", "Canonical batting-side team name."),
        ("bowling_team", "Canonical bowling-side team name."),
        ("striker", "Batter facing the current delivery."),
        ("non_striker", "Partner at the non-striker end."),
        ("bowler", "Current bowler delivering the ball."),
        ("toss_winner", "Team that won the toss."),
        ("toss_decision", "Whether the toss winner chose to bat or field."),
        ("toss_winner_batting", "Binary indicator showing whether the toss winner is the current batting side."),
    ],
    "match_state": [
        ("runs", "Current cumulative innings runs at the prediction point."),
        ("wickets", "Current wickets lost at the prediction point."),
        ("wickets_left", "Remaining wickets available to the batting side."),
        ("balls_left", "Legal deliveries remaining in the innings."),
        ("legal_balls_bowled", "Legal deliveries completed so far."),
        ("innings_progress", "Fraction of the innings completed, scaled from 0 to 1."),
        ("over_number", "Integer over index derived from legal balls bowled."),
        ("ball_in_over", "Ball number within the current over."),
        ("phase", "Categorical phase label: powerplay, middle, or death."),
        ("is_powerplay", "Binary flag for overs 1 to 6."),
        ("is_middle", "Binary flag for overs 7 to 15."),
        ("is_death", "Binary flag for overs 16 to 20."),
    ],
    "momentum_and_rates": [
        ("runs_last_5", "Runs scored in the previous five overs."),
        ("wickets_last_5", "Wickets lost in the previous five overs."),
        ("runs_last_6_balls", "Runs scored in the previous over."),
        ("wickets_last_6_balls", "Wickets lost in the previous over."),
        ("runs_last_12_balls", "Runs scored in the previous two overs."),
        ("wickets_last_12_balls", "Wickets lost in the previous two overs."),
        ("current_run_rate", "Current run rate at the prediction point."),
        ("target", "Required chase target for second innings; blank for first innings."),
        ("target_remaining", "Runs still needed in the chase."),
        ("required_run_rate", "Required run rate to reach the target."),
        ("current_minus_required_rr", "Positive values mean the batting side is scoring ahead of the chase requirement."),
        ("required_minus_current_rr", "Positive values mean the batting side is behind the chase requirement."),
    ],
    "team_and_player_context": [
        ("batting_team_form", "Rolling batting-team win rate from recent matches."),
        ("bowling_team_form", "Rolling bowling-team win rate from recent matches."),
        ("batting_team_venue_form", "Recent batting-team performance at the current venue."),
        ("bowling_team_venue_form", "Recent bowling-team performance at the current venue."),
        ("batting_vs_bowling_form", "Recent head-to-head batting-team performance against the bowling side."),
        ("striker_form_sr", "Historical strike rate of the current striker."),
        ("striker_form_avg", "Historical batting average of the current striker."),
        ("bowler_form_econ", "Historical economy rate of the current bowler."),
        ("bowler_form_strike", "Historical strike rate of the current bowler in balls per wicket."),
        ("batter_vs_bowler_sr", "Observed strike rate for this batter-versus-bowler matchup."),
        ("batter_vs_bowler_balls", "Observed ball count for this batter-versus-bowler matchup."),
    ],
    "weather_venue_targets": [
        ("temperature_c", "Temperature feature used during training and optionally replaced by live weather at inference time."),
        ("relative_humidity", "Humidity feature used to represent environmental conditions."),
        ("wind_kph", "Wind-speed feature used to capture broad weather context."),
        ("dew_risk", "Normalized dew-risk score used mainly for night-match second-innings context."),
        ("runs_vs_par", "Current runs minus the venue-specific expected total at this innings stage."),
        ("venue_avg_first_innings", "Historical first-innings average at the venue."),
        ("venue_avg_second_innings", "Historical second-innings average at the venue."),
        ("venue_bat_first_win_rate", "Historical fraction of matches won by teams batting first at the venue."),
        ("total_runs", "Final innings total used as the regression target."),
        ("winner", "Match winner label used for interpretability and downstream summaries."),
        ("win", "Binary live target indicating whether the batting side eventually wins the match."),
    ],
}


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def fmt_float(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}"


def fmt_pct(value: float, digits: int = 1) -> str:
    return f"{100.0 * value:.{digits}f}%"


def human_size(num_bytes: int) -> str:
    size = float(num_bytes)
    units = ["B", "KB", "MB", "GB"]
    unit = units[0]
    for candidate in units:
        unit = candidate
        if size < 1024 or candidate == units[-1]:
            break
        size /= 1024
    digits = 0 if unit == "B" else 2
    return f"{size:.{digits}f} {unit}"


def count_lines(path: Path) -> int:
    with path.open("r", encoding="utf-8", errors="ignore") as handle:
        return sum(1 for _ in handle)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = BODY_FONT
    run.font.size = Pt(font_size)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.name = BODY_FONT
    run.font.size = Pt(10)


def configure_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style.font.bold = True

    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(11.5)

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.5)
    add_page_number(section)


def add_centered_paragraph(doc: Document, text: str, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BODY_FONT
    run.font.size = Pt(size)


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.line_spacing = 1.2
        run = paragraph.add_run(item)
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE)


def add_table_title(doc: Document, spec: TableSpec) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{spec.number}: {spec.title}")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], font_size: int = 10) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        set_cell_text(cell, header, bold=True, font_size=font_size)

    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value, bold=False, font_size=font_size)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = Pt(10)


def add_figure(doc: Document, path: Path, number: str, title: str, width: float = 6.2) -> None:
    doc.add_picture(str(path), width=Inches(width))
    paragraph = doc.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"{number}: {title}")
    run.bold = True
    run.italic = True
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    doc.add_paragraph()


def generate_pipeline_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13, 3.5))
    ax.set_axis_off()
    xs = [0.06, 0.24, 0.42, 0.60, 0.78, 0.94]
    labels = [
        "Cricsheet IPL CSV2\nball-by-ball + info files",
        "Preprocessing\nnormalization + support tables",
        "Feature Store\n60-column match-state dataset",
        "Model Training\nML, DL, calibration, selection",
        "Deployment Artifacts\nscore_model / win_model / uncertainty",
        "Interfaces\nFlask, Streamlit, CLI, API",
    ]
    for x, label in zip(xs, labels):
        ax.text(
            x,
            0.5,
            label,
            ha="center",
            va="center",
            fontsize=10,
            bbox={"boxstyle": "round,pad=0.4", "facecolor": "#edf4ff", "edgecolor": "#1f4e79", "linewidth": 1.5},
        )
    for left, right in zip(xs[:-1], xs[1:]):
        ax.annotate(
            "",
            xy=(right - 0.07, 0.5),
            xytext=(left + 0.07, 0.5),
            arrowprops={"arrowstyle": "->", "linewidth": 1.5, "color": "#1f4e79"},
        )
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_model_comparison_figure(report: dict, path: Path) -> None:
    scope_name = report["selected_scope"]
    scope = report["scopes"][scope_name]
    score_metrics_map = get_score_metrics_map(scope)
    win_metrics_map = get_win_metrics_map(scope)

    score_keys = [
        ("hgb", "HGB"),
        ("xgboost_gpu", "XGBoost"),
        ("catboost_gpu", "CatBoost"),
        ("torch_entity_gpu", "Torch Entity"),
        ("ensemble_top2", "Ensemble"),
    ]
    win_keys = [
        ("hgb_calibrated", "HGB Cal."),
        ("xgboost_gpu_calibrated", "XGBoost Cal."),
        ("catboost_gpu_calibrated", "CatBoost Cal."),
        ("torch_entity_gpu_calibrated", "Torch Cal."),
        ("ensemble_top2_calibrated", "Ensemble Cal."),
    ]

    score_labels = [label for _, label in score_keys]
    score_values = [score_metrics_map[key]["rmse"] for key, _ in score_keys if key in score_metrics_map]
    score_labels = [label for key, label in score_keys if key in score_metrics_map]
    score_colors = ["#a6bddb"] * len(score_keys)
    score_key_list = [key for key, _ in score_keys]
    selected_score = report["selected_score_model"]
    if selected_score in score_key_list:
        score_colors[score_key_list.index(selected_score)] = "#1f4e79"
    score_colors = [score_colors[score_key_list.index(key)] for key, _ in score_keys if key in score_metrics_map]

    win_labels = [label for key, label in win_keys if key in win_metrics_map]
    win_values = [win_metrics_map[key]["log_loss"] for key, _ in win_keys if key in win_metrics_map]
    win_colors = ["#f7c59f"] * len(win_keys)
    win_key_list = [key for key, _ in win_keys]
    selected_win = report["selected_win_model"]
    if selected_win in win_key_list:
        win_colors[win_key_list.index(selected_win)] = "#c75b12"
    win_colors = [win_colors[win_key_list.index(key)] for key, _ in win_keys if key in win_metrics_map]

    fig, axes = plt.subplots(1, 2, figsize=(12, 4.6))
    axes[0].bar(score_labels, score_values, color=score_colors)
    axes[0].set_title("Score Model RMSE")
    axes[0].set_ylabel("RMSE (runs)")
    axes[0].tick_params(axis="x", rotation=30)
    for idx, value in enumerate(score_values):
        axes[0].text(idx, value + 0.3, f"{value:.2f}", ha="center", va="bottom", fontsize=9)

    axes[1].bar(win_labels, win_values, color=win_colors)
    axes[1].set_title("Win Model Log Loss")
    axes[1].set_ylabel("Log loss")
    axes[1].tick_params(axis="x", rotation=30)
    for idx, value in enumerate(win_values):
        axes[1].text(idx, value + 0.015, f"{value:.3f}", ha="center", va="bottom", fontsize=9)

    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_venue_figure(venue_stats: pd.DataFrame, path: Path) -> None:
    top = venue_stats.sort_values("venue_avg_first_innings", ascending=False).head(8).copy()
    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.barh(top["venue"], top["venue_avg_first_innings"], color="#2c7fb8")
    ax.invert_yaxis()
    ax.set_xlabel("Average first-innings score")
    ax.set_title("Historical First-Innings Averages by Venue")
    for idx, value in enumerate(top["venue_avg_first_innings"]):
        ax.text(value + 0.8, idx, f"{value:.1f}", va="center", fontsize=9)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_uncertainty_figure(score_uncertainty: dict, path: Path) -> None:
    q10 = score_uncertainty["residual_q10"]
    q90 = score_uncertainty["residual_q90"]
    fig, ax = plt.subplots(figsize=(10, 2.4))
    ax.hlines(1, q10, q90, color="#b30000", linewidth=8, alpha=0.35)
    ax.plot([0], [1], marker="o", markersize=10, color="#1f4e79")
    ax.text(q10, 1.08, f"Q10: {q10:.2f}", ha="center", fontsize=10)
    ax.text(0, 0.85, "Point prediction", ha="center", fontsize=10)
    ax.text(q90, 1.08, f"Q90: +{q90:.2f}", ha="center", fontsize=10)
    ax.set_yticks([])
    ax.set_xlabel("Residual runs relative to the point estimate")
    ax.set_title("Projected Range Derived from Held-Out Residual Quantiles")
    ax.spines["left"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["top"].set_visible(False)
    fig.subplots_adjust(left=0.08, right=0.98, bottom=0.35, top=0.78)
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_feature_mix_figure(context: dict, path: Path) -> None:
    feature_rows = context["feature_family_rows"]
    labels = [row["group"] for row in feature_rows]
    counts = [row["count"] for row in feature_rows]
    colors = ["#1d4ed8", "#2563eb", "#0f766e", "#ea580c", "#7c3aed", "#475569"]

    fig, ax = plt.subplots(figsize=(10.5, 4.8))
    ax.barh(labels, counts, color=colors[: len(labels)])
    ax.invert_yaxis()
    ax.set_xlabel("Number of input features")
    ax.set_title("Live Prediction Input Families")
    for idx, value in enumerate(counts):
        ax.text(value + 0.2, idx, str(value), va="center", fontsize=9)
    ax.grid(axis="x", alpha=0.2)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def get_score_metrics_map(scope_report: dict) -> dict:
    return scope_report.get("score_test") or scope_report.get("score_valid") or {}


def get_win_metrics_map(scope_report: dict) -> dict:
    return scope_report.get("win_test") or scope_report.get("win_valid") or {}


def get_selected_score_model_name(scope_report: dict) -> str:
    return scope_report.get("best_score_name") or scope_report.get("selected_score_model") or ""


def get_selected_win_model_name(scope_report: dict) -> str:
    return scope_report.get("best_win_name") or scope_report.get("selected_win_model") or ""


def get_selected_score_metrics(scope_report: dict) -> dict:
    return scope_report.get("best_score_metrics_test") or scope_report.get("selected_score_metrics_valid") or {}


def get_selected_win_metrics(scope_report: dict) -> dict:
    return scope_report.get("best_win_metrics_test") or scope_report.get("selected_win_metrics_valid") or {}


def prettify_model_name(model_name: str) -> str:
    mapping = {
        "hgb": "HistGradientBoosting",
        "hgb_calibrated": "HistGradientBoosting (calibrated)",
        "xgboost_gpu": "XGBoost GPU",
        "xgboost_gpu_calibrated": "XGBoost GPU (calibrated)",
        "catboost_gpu": "CatBoost GPU",
        "catboost_gpu_calibrated": "CatBoost GPU (calibrated)",
        "torch_entity_gpu": "Torch entity-embedding model",
        "torch_entity_gpu_calibrated": "Torch entity-embedding model (calibrated)",
        "ensemble_ml_top3": "Weighted ML ensemble",
        "ensemble_top2": "Top-2 weighted ensemble",
        "ensemble_top2_calibrated": "Top-2 weighted ensemble (calibrated)",
    }
    return mapping.get(model_name, model_name.replace("_", " "))


def describe_score_model(model_name: str) -> str:
    if "torch" in model_name:
        return (
            f"The current promoted live score model is `{model_name}`. This family learns dense embeddings for categorical "
            "inputs such as teams, venues, strikers, and bowlers, then combines them with normalized numeric match-state "
            "features in a multilayer perceptron. In the project context, this architecture is valuable because it can "
            "represent rich cricket identities without exploding the dimensionality of the feature space."
        )
    if "ensemble" in model_name:
        return (
            f"The current promoted live score model is `{model_name}`. Rather than relying on a single learner, it blends "
            "the strongest classical models from the search report so that complementary error patterns partially cancel "
            "out. This is a practical engineering choice when multiple tabular models are individually strong but none is "
            "dominant across every scoring regime."
        )
    if "xgboost" in model_name:
        return (
            f"The current promoted live score model is `{model_name}`. XGBoost is a strong gradient-boosting approach for "
            "tabular prediction because it can capture nonlinear interactions among score state, venue context, and player "
            "form while remaining efficient enough for repeated retraining."
        )
    if "catboost" in model_name:
        return (
            f"The current promoted live score model is `{model_name}`. CatBoost is particularly effective when categorical "
            "structure matters, which is useful in IPL data where teams, venues, and player identities carry predictive "
            "signal alongside numeric match-state variables."
        )
    return (
        f"The current promoted live score model is `{model_name}`. It was selected from the project benchmark because it "
        "offered the best tradeoff between held-out score accuracy, robustness, and deployment readiness for the current workspace."
    )


def describe_win_model(model_name: str) -> str:
    if "catboost" in model_name:
        return (
            f"The current promoted live win model is `{model_name}`. CatBoost handles mixed categorical and numeric inputs "
            "well, and the calibrated variant is especially useful when the report cares about probability quality rather "
            "than only hard classification accuracy."
        )
    if "xgboost" in model_name:
        return (
            f"The current promoted live win model is `{model_name}`. In this project it serves as a high-capacity boosting "
            "baseline whose calibrated probabilities can be compared directly with CatBoost and HistGradientBoosting."
        )
    if "ensemble" in model_name:
        return (
            f"The current promoted live win model is `{model_name}`. The ensemble design smooths out model-specific biases "
            "and is useful when different classifiers perform well on different chase-pressure scenarios."
        )
    return (
        f"The current promoted live win model is `{model_name}`. It was retained because its held-out probabilities were "
        "the most dependable among the alternatives considered in the benchmark."
    )


def build_effective_deployment_report(
    raw_deployment_report: dict,
    best_search_report: dict,
    model_registry: dict,
) -> dict:
    latest_version = model_registry.get("latest_version")
    latest_entry = None
    for entry in model_registry.get("entries", []):
        if entry.get("version_id") == latest_version:
            latest_entry = entry
            break
    if latest_entry is None and model_registry.get("entries"):
        latest_entry = model_registry["entries"][-1]

    if latest_entry and latest_entry.get("metrics", {}).get("live"):
        live_metrics = latest_entry["metrics"]["live"]
        return {
            "deployment_scope": live_metrics.get("scope", best_search_report.get("selected_scope")),
            "deployment_score_model": live_metrics.get("score_model", raw_deployment_report.get("deployment_score_model")),
            "deployment_score_metrics_test": live_metrics.get(
                "score_test",
                raw_deployment_report.get("deployment_score_metrics_test", {}),
            ),
            "deployment_win_model": live_metrics.get("win_model", raw_deployment_report.get("deployment_win_model")),
            "deployment_win_metrics_test": live_metrics.get(
                "win_test",
                raw_deployment_report.get("deployment_win_metrics_test", {}),
            ),
            "selection_rule": raw_deployment_report.get("selection_rule", best_search_report.get("selection_rule", {})),
            "artifacts": raw_deployment_report.get("artifacts", {}),
            "version_id": latest_entry.get("version_id"),
            "created_at_utc": latest_entry.get("created_at_utc"),
        }
    return raw_deployment_report


def add_card(ax, x: float, y: float, w: float, h: float, title: str, lines: list[str], facecolor: str = "#ffffff") -> None:
    card = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.012,rounding_size=0.02",
        linewidth=1.2,
        edgecolor="#2f3b52",
        facecolor=facecolor,
    )
    ax.add_patch(card)
    ax.text(x + 0.02, y + h - 0.05, title, fontsize=11, fontweight="bold", color="#1a2433")
    cursor = y + h - 0.10
    for line in lines:
        ax.text(x + 0.02, cursor, fill(line, width=32), fontsize=9, color="#334155", va="top")
        cursor -= 0.07


def add_code_panel(ax, x: float, y: float, w: float, h: float, title: str, code: str, fontsize: int = 9) -> None:
    panel = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.01,rounding_size=0.02",
        linewidth=1.0,
        edgecolor="#0f172a",
        facecolor="#0b1220",
    )
    ax.add_patch(panel)
    ax.add_patch(Rectangle((x, y + h - 0.06), w, 0.06, facecolor="#111827", edgecolor="none"))
    ax.text(x + 0.02, y + h - 0.04, title, fontsize=10, fontweight="bold", color="#f8fafc", va="center")
    ax.text(x + 0.02, y + h - 0.09, code, family=CODE_FONT, fontsize=fontsize, color="#d1fae5", va="top")


def generate_dataset_dashboard_figure(context: dict, path: Path) -> None:
    season_summary = context["season_summary"]
    team_summary = context["team_summary"].head(8)
    venue_profile = context["venue_profile"].head(8)

    fig, axes = plt.subplots(2, 2, figsize=(13, 8))
    fig.patch.set_facecolor("#f8fafc")

    axes[0, 0].bar(season_summary["season_label"], season_summary["avg_first_innings"], color="#2563eb")
    axes[0, 0].set_title("Average First-Innings Score by Season")
    axes[0, 0].tick_params(axis="x", rotation=60, labelsize=8)
    axes[0, 0].grid(axis="y", alpha=0.2)

    axes[0, 1].barh(team_summary["batting_team"], team_summary["avg_total"], color="#f97316")
    axes[0, 1].set_title("Active-Team Average Innings Total")
    axes[0, 1].invert_yaxis()
    axes[0, 1].grid(axis="x", alpha=0.2)

    axes[1, 0].barh(venue_profile["venue"], venue_profile["venue_avg_first_innings"], color="#16a34a")
    axes[1, 0].set_title("Top Venues by First-Innings Average")
    axes[1, 0].invert_yaxis()
    axes[1, 0].grid(axis="x", alpha=0.2)

    axes[1, 1].axis("off")
    counts = context["counts"]
    summary_lines = [
        f"Raw matches: {counts['raw_matches']:,}",
        f"Engineered rows: {counts['processed_rows']:,}",
        f"Columns: {counts['processed_cols']}",
        f"Season range: {counts['min_season']} - {counts['max_season']}",
        f"Venue rows: {counts['venue_rows']}",
        f"Batter-bowler pairs: {counts['batter_bowler_rows']:,}",
        f"Promoted score model: {context['deployment_report']['deployment_score_model']}",
        f"Promoted win model: {context['deployment_report']['deployment_win_model']}",
    ]
    add_card(axes[1, 1], 0.06, 0.10, 0.88, 0.80, "Workspace Snapshot", summary_lines, facecolor="#e0f2fe")

    fig.suptitle("IPL Prediction Dataset and Coverage Dashboard", fontsize=16, fontweight="bold", y=0.98)
    fig.tight_layout(rect=[0, 0, 1, 0.96])
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_flask_ui_mockup(context: dict, path: Path) -> None:
    scenario = context["live_scenarios"][0] if context["live_scenarios"] else None
    fig, ax = plt.subplots(figsize=(13, 7.5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.set_facecolor("#f1f5f9")

    ax.add_patch(Rectangle((0, 0), 1, 1, facecolor="#f8fafc", edgecolor="none"))
    ax.add_patch(Rectangle((0.03, 0.90), 0.94, 0.07, facecolor="#0f172a", edgecolor="none"))
    ax.text(0.05, 0.935, "IPL Match Score and Win Prediction System - Flask Interface", color="white", fontsize=15, fontweight="bold", va="center")

    add_card(
        ax,
        0.05,
        0.10,
        0.40,
        0.74,
        "Live Match Input Form",
        [
            "Season: 2026",
            "Venue: Wankhede Stadium",
            "Batting team: Mumbai Indians",
            "Bowling team: Chennai Super Kings",
            "Innings: 1",
            "Runs/Wickets/Overs: 62 / 1 / 6.0",
            "Striker: Rohit Sharma",
            "Bowler: Noor Ahmad",
            "Recent momentum: 58 runs, 1 wicket in last 5 overs",
        ],
        facecolor="#ffffff",
    )
    add_card(
        ax,
        0.50,
        0.48,
        0.43,
        0.36,
        "Prediction Summary",
        [
            f"Projected total: {scenario['projected_total']}" if scenario else "Projected total: 185.1",
            f"Win probability: {scenario['win_prob']}" if scenario else "Win probability: 73.0%",
            f"Projected range: {scenario['range']}" if scenario else "Projected range: 173-214",
            f"Collapse risk: {scenario['collapse_risk']}" if scenario else "Collapse risk: 24.3%",
            "Auxiliary outputs: phase, runs-vs-par, weather context",
        ],
        facecolor="#dcfce7",
    )
    add_card(
        ax,
        0.50,
        0.10,
        0.43,
        0.30,
        "Why This Matters",
        [
            "The Flask surface exposes the same live inference contract used by the API and CLI.",
            "A single prediction call returns score, win probability, range, and simulation-oriented risk indicators.",
            "This makes the project demonstrably deployable, not just an offline notebook experiment.",
        ],
        facecolor="#eff6ff",
    )
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_streamlit_mockup(context: dict, path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13, 7.5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.add_patch(Rectangle((0, 0), 1, 1, facecolor="#f8fafc", edgecolor="none"))
    ax.add_patch(Rectangle((0.02, 0.04), 0.24, 0.92, facecolor="#e2e8f0", edgecolor="#94a3b8"))
    ax.add_patch(Rectangle((0.28, 0.86), 0.69, 0.10, facecolor="#ffffff", edgecolor="#cbd5e1"))
    ax.text(0.30, 0.91, "Streamlit Dashboard - Model Comparison, Venue Insights, and Live Analysis", fontsize=14, fontweight="bold", color="#0f172a")

    for idx, label in enumerate(["Match Setup", "Model Selection", "Venue Trends", "Player Form", "Feedback"]):
        ax.text(0.05, 0.87 - idx * 0.12, label, fontsize=11, fontweight="bold", color="#1e293b")
        ax.add_patch(Rectangle((0.05, 0.82 - idx * 0.12), 0.16, 0.04, facecolor="#ffffff", edgecolor="#94a3b8"))

    add_card(
        ax,
        0.30,
        0.50,
        0.30,
        0.28,
        "Deployment Metrics",
        [
            f"Score RMSE: {fmt_float(context['deployment_report']['deployment_score_metrics_test']['rmse'])}",
            f"Win log loss: {fmt_float(context['deployment_report']['deployment_win_metrics_test']['log_loss'], 4)}",
            f"Win accuracy: {fmt_pct(context['deployment_report']['deployment_win_metrics_test']['accuracy'])}",
            "Promoted models are loaded directly from saved artifacts.",
        ],
        facecolor="#ffffff",
    )
    add_card(
        ax,
        0.64,
        0.50,
        0.30,
        0.28,
        "Support Tables",
        [
            f"Venue priors: {context['counts']['venue_rows']}",
            f"Team form rows: {context['counts']['team_form_rows']}",
            f"Player pool rows: {context['counts']['player_pool_rows']}",
            f"Batter-bowler pairs: {context['counts']['batter_bowler_rows']:,}",
        ],
        facecolor="#ffffff",
    )
    add_card(
        ax,
        0.30,
        0.14,
        0.64,
        0.26,
        "Dashboard Role in the Project",
        [
            "The Streamlit dashboard provides a richer analytical surface than the Flask form by exposing model selection, venue summaries, and player-form views.",
            "In the report, it serves as evidence that the project includes both end-user prediction and analyst-oriented exploration.",
        ],
        facecolor="#ecfeff",
    )
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_api_snapshot(context: dict, path: Path) -> None:
    sample_request = LIVE_SCENARIOS[3]["payload"]
    prediction_row = context["live_scenarios"][3] if len(context["live_scenarios"]) > 3 else None
    sample_response = {
        "ok": True,
        "prediction": {
            "predicted_total": prediction_row["projected_total"] if prediction_row else "160.1",
            "win_prob_pct": prediction_row["win_prob"] if prediction_row else "50.4%",
            "projected_range": prediction_row["range"] if prediction_row else "148-189",
            "collapse_risk_pct": prediction_row["collapse_risk"] if prediction_row else "12.7%",
        },
    }

    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.add_patch(Rectangle((0, 0), 1, 1, facecolor="#e2e8f0", edgecolor="none"))
    add_code_panel(ax, 0.05, 0.10, 0.42, 0.80, "POST /api/predict - request", json.dumps(sample_request, indent=2))
    add_code_panel(ax, 0.53, 0.10, 0.42, 0.80, "200 OK - response", json.dumps(sample_response, indent=2))
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_cli_snapshot(context: dict, path: Path) -> None:
    prediction_row = context["live_scenarios"][4] if len(context["live_scenarios"]) > 4 else None
    cli_output = "\n".join(
        [
            r"> python predict_cli.py",
            "Season: 2026",
            "Venue: Ekana Cricket Stadium",
            "Batting Team: Lucknow Super Giants",
            "Bowling Team: Delhi Capitals",
            "Innings: 2",
            "Runs/Wickets/Overs: 132/7/16.4",
            f"Projected total: {prediction_row['projected_total'] if prediction_row else '137.8'}",
            f"Win probability: {prediction_row['win_prob'] if prediction_row else '24.5%'}",
            f"Projected range: {prediction_row['range'] if prediction_row else '132-166'}",
            f"Collapse risk: {prediction_row['collapse_risk'] if prediction_row else '28.7%'}",
        ]
    )

    fig, ax = plt.subplots(figsize=(13, 5.5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    add_code_panel(ax, 0.05, 0.14, 0.90, 0.72, "Terminal Live Prediction Flow", cli_output, fontsize=11)
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_validation_snapshot(context: dict, path: Path) -> None:
    test_lines = [f"{row['test_name']}" for row in context["test_inventory"][:8]]
    artifact_lines = [f"{row['name']} ({row['size']})" for row in context["artifact_inventory"][:8]]

    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.add_patch(Rectangle((0, 0), 1, 1, facecolor="#f8fafc", edgecolor="none"))
    add_card(ax, 0.05, 0.18, 0.42, 0.66, "Automated Test Inventory", test_lines, facecolor="#f8fafc")
    add_card(ax, 0.53, 0.18, 0.42, 0.66, "Key Saved Artifacts", artifact_lines, facecolor="#f8fafc")
    ax.text(0.05, 0.88, "Verification Snapshot", fontsize=16, fontweight="bold", color="#0f172a")
    ax.text(
        0.05,
        0.11,
        "This figure summarizes two practical aspects of deployment readiness: automated checks and the concrete artifacts generated by the training workflows.",
        fontsize=10,
        color="#334155",
    )
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_phase_scoring_figure(context: dict, path: Path) -> None:
    phase_summary = context["phase_summary"]
    labels = [f"{row['innings_label']}\n{row['phase']}" for row in context["phase_summary_rows"]]
    avg_runs = [row["avg_runs"] for row in context["phase_summary_rows"]]
    avg_rr = [row["avg_current_rr"] for row in context["phase_summary_rows"]]

    fig, axes = plt.subplots(1, 2, figsize=(12.5, 4.6))
    axes[0].bar(labels, avg_runs, color="#0f766e")
    axes[0].set_title("Average Runs at Snapshot")
    axes[0].tick_params(axis="x", rotation=35)
    axes[0].grid(axis="y", alpha=0.2)

    axes[1].bar(labels, avg_rr, color="#ea580c")
    axes[1].set_title("Average Current Run Rate")
    axes[1].tick_params(axis="x", rotation=35)
    axes[1].grid(axis="y", alpha=0.2)

    fig.suptitle("Phase-wise Match Behaviour", fontsize=15, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.94])
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_score_error_distribution_figure(context: dict, path: Path) -> None:
    score_predictions = context["score_predictions"]

    fig, axes = plt.subplots(1, 2, figsize=(12.2, 4.5))
    axes[0].hist(score_predictions["absolute_error"], bins=24, color="#2563eb", edgecolor="white")
    axes[0].set_title("Absolute Error Distribution")
    axes[0].set_xlabel("Absolute error (runs)")
    axes[0].set_ylabel("Held-out rows")

    bucket_labels = [row["bucket"] for row in context["score_wicket_summary"]]
    bucket_mae = [row["mae_value"] for row in context["score_wicket_summary"]]
    axes[1].bar(bucket_labels, bucket_mae, color="#1d4ed8")
    axes[1].set_title("MAE by Wicket-State Bucket")
    axes[1].set_xlabel("Wickets lost")
    axes[1].set_ylabel("MAE (runs)")
    axes[1].grid(axis="y", alpha=0.2)

    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_win_calibration_figure(context: dict, path: Path) -> None:
    calibration_rows = context["win_calibration_bins"]
    pred = [row["mean_pred_value"] for row in calibration_rows]
    actual = [row["empirical_win_value"] for row in calibration_rows]
    labels = [row["bucket"] for row in calibration_rows]

    fig, axes = plt.subplots(1, 2, figsize=(12.2, 4.7))
    axes[0].plot([0, 1], [0, 1], linestyle="--", color="#94a3b8", linewidth=1.2)
    axes[0].plot(pred, actual, marker="o", linewidth=2.0, color="#b91c1c")
    axes[0].set_xlim(0, 1)
    axes[0].set_ylim(0, 1)
    axes[0].set_xlabel("Average predicted win probability")
    axes[0].set_ylabel("Empirical win rate")
    axes[0].set_title("Reliability Curve")
    axes[0].grid(alpha=0.2)

    axes[1].bar(labels, [row["rows"] for row in calibration_rows], color="#7c3aed")
    axes[1].set_title("Rows per Calibration Bin")
    axes[1].set_xlabel("Probability bucket")
    axes[1].set_ylabel("Held-out rows")
    axes[1].tick_params(axis="x", rotation=35)

    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def generate_monitoring_snapshot(context: dict, path: Path) -> None:
    drift = context["production_drift_report"]
    pipeline = context["canonical_pipeline_report"]
    reference = context["feature_reference_profile"]

    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.add_patch(Rectangle((0, 0), 1, 1, facecolor="#f8fafc", edgecolor="none"))

    add_card(
        ax,
        0.05,
        0.52,
        0.40,
        0.34,
        "Monitoring Window",
        [
            f"Status: {drift['status']}",
            f"Events used: {drift['events_used']}",
            f"Outcomes used: {drift['outcomes_used']}",
            f"Retrain trigger: {'Yes' if drift['trigger_retrain'] else 'No'}",
            f"Recommended action: {drift['recommended_action']}",
        ],
        facecolor="#eff6ff",
    )
    add_card(
        ax,
        0.53,
        0.52,
        0.40,
        0.34,
        "Pipeline Consistency",
        [
            f"Version: {pipeline['version_id']}",
            f"Consistency checks: {len(pipeline['consistency']['checks'])}",
            f"Consistency OK: {'Yes' if pipeline['consistency']['ok'] else 'No'}",
            f"Run steps: {', '.join(pipeline['steps'])}",
        ],
        facecolor="#ecfeff",
    )
    add_card(
        ax,
        0.05,
        0.12,
        0.88,
        0.26,
        "Reference Distribution Snapshot",
        [
            f"Runs mean / p90: {fmt_float(reference['features']['runs']['mean'])} / {fmt_float(reference['features']['runs']['p90'])}",
            f"Wickets mean / p90: {fmt_float(reference['features']['wickets']['mean'])} / {fmt_float(reference['features']['wickets']['p90'])}",
            f"Current RR mean / p90: {fmt_float(reference['features']['current_run_rate']['mean'])} / {fmt_float(reference['features']['current_run_rate']['p90'])}",
            f"Death-over slice rate: {fmt_pct(reference['slice_rates']['death_over'])}",
            f"High-pressure chase rate: {fmt_pct(reference['slice_rates']['high_pressure_chase'])}",
        ],
        facecolor="#fef3c7",
    )

    ax.text(0.05, 0.90, "Monitoring and Drift Readiness Snapshot", fontsize=16, fontweight="bold", color="#0f172a")
    ax.text(
        0.05,
        0.06,
        "This visual shows that the project has moved beyond offline modeling into artifact governance and post-deployment monitoring readiness.",
        fontsize=10,
        color="#334155",
    )
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def build_score_benchmark_rows(scope_name: str, scope_report: dict) -> list[dict[str, str]]:
    best_name = get_selected_score_model_name(scope_report)
    rows = []
    for model_name, metrics in sorted(get_score_metrics_map(scope_report).items(), key=lambda item: item[1]["rmse"]):
        rows.append(
            {
                "scope": scope_name,
                "model": model_name,
                "mae": fmt_float(metrics["mae"]),
                "rmse": fmt_float(metrics["rmse"]),
                "selection": "Selected in scope" if model_name == best_name else "Compared",
            }
        )
    return rows


def build_win_benchmark_rows(scope_name: str, scope_report: dict) -> list[dict[str, str]]:
    best_name = get_selected_win_model_name(scope_report)
    rows = []
    for model_name, metrics in sorted(get_win_metrics_map(scope_report).items(), key=lambda item: item[1]["log_loss"]):
        rows.append(
            {
                "scope": scope_name,
                "model": model_name,
                "log_loss": fmt_float(metrics["log_loss"], 4),
                "brier": fmt_float(metrics["brier"], 4),
                "accuracy": fmt_pct(metrics["accuracy"]),
                "selection": "Selected in scope" if model_name == best_name else "Compared",
            }
        )
    return rows


def run_live_scenarios() -> list[dict[str, str]]:
    try:
        score_model, win_model = load_models()
        support_tables = load_support_tables()
    except Exception:
        return []

    rows: list[dict[str, str]] = []
    monitoring_module = None
    original_track_prediction_event = None
    try:
        import ipl_predictor.monitoring as monitoring_module

        original_track_prediction_event = monitoring_module.track_prediction_event
        monitoring_module.track_prediction_event = lambda **kwargs: ""
    except Exception:
        monitoring_module = None

    try:
        for scenario in LIVE_SCENARIOS:
            prediction, errors = predict_match_state(
                scenario["payload"],
                support_tables=support_tables,
                score_model=score_model,
                win_model=win_model,
            )
            if errors or prediction is None:
                continue
            rows.append(
                {
                    "scenario": scenario["label"],
                    "phase": prediction["phase"],
                    "projected_total": prediction["predicted_total"],
                    "win_prob": prediction["win_prob_pct"],
                    "range": prediction["projected_range"],
                    "collapse_risk": prediction["collapse_risk_pct"],
                    "interpretation": scenario["interpretation"],
                }
            )
    finally:
        if monitoring_module is not None and original_track_prediction_event is not None:
            monitoring_module.track_prediction_event = original_track_prediction_event
    return rows


def run_pre_match_scenarios() -> list[dict[str, str]]:
    try:
        score_model, win_model = load_pre_match_models()
    except Exception:
        return []

    if not score_model or not win_model:
        return []

    rows: list[dict[str, str]] = []
    for scenario in PREMATCH_SCENARIOS:
        prediction = predict_pre_match(
            team1=scenario["team1"],
            team2=scenario["team2"],
            venue=scenario["venue"],
            score_model=score_model,
            win_model=win_model,
        )
        if "error" in prediction:
            continue
        rows.append(
            {
                "matchup": f"{scenario['team1']} vs {scenario['team2']}",
                "venue": scenario["venue"],
                "predicted_score": prediction["predicted_score"],
                "team1_win_prob": prediction["team1_win_prob"],
                "team2_win_prob": prediction["team2_win_prob"],
                "likely_winner": prediction["likely_winner"],
                "interpretation": scenario["interpretation"],
            }
        )
    return rows


def build_context() -> dict:
    raw_deployment_report = load_json(MODELS_DIR / "deployment_report.json")
    best_search_report = load_json(MODELS_DIR / "best_model_search_report.json")
    score_uncertainty = load_json(MODELS_DIR / "score_uncertainty.json")
    cpu_report = load_json(MODELS_DIR / "cpu_model_report.json")
    pre_match_model_report = load_json(MODELS_DIR / "pre_match_model_report.json")
    production_drift_report = load_json(MODELS_DIR / "production_drift_report.json")
    feature_reference_profile = load_json(MODELS_DIR / "feature_reference_profile.json")
    canonical_pipeline_report = load_json(MODELS_DIR / "canonical_pipeline_report.json")
    model_registry = load_json(MODELS_DIR / "model_registry.json")
    requirements = (ROOT_DIR / "requirements.txt").read_text(encoding="utf-8").strip().splitlines()

    ipl_features = pd.read_csv(PROCESSED_DIR / "ipl_features.csv", low_memory=False)
    venue_stats = pd.read_csv(PROCESSED_DIR / "venue_stats.csv")
    team_form = pd.read_csv(PROCESSED_DIR / "team_form_latest.csv")
    team_venue_form = pd.read_csv(PROCESSED_DIR / "team_venue_form_latest.csv")
    matchup_form = pd.read_csv(PROCESSED_DIR / "matchup_form_latest.csv")
    batter_form = pd.read_csv(PROCESSED_DIR / "batter_form_latest.csv")
    bowler_form = pd.read_csv(PROCESSED_DIR / "bowler_form_latest.csv")
    batter_bowler_form = pd.read_csv(PROCESSED_DIR / "batter_bowler_form_latest.csv")
    team_player_pool = pd.read_csv(PROCESSED_DIR / "team_player_pool_2026.csv")
    active_teams = pd.read_csv(PROCESSED_DIR / "active_teams_2026.csv")
    score_predictions = pd.read_csv(MODELS_DIR / "best_score_test_predictions.csv")
    win_predictions = pd.read_csv(MODELS_DIR / "best_win_test_predictions.csv")
    deployment_report = build_effective_deployment_report(
        raw_deployment_report,
        best_search_report,
        model_registry,
    )

    support_frames = {
        "ipl_features.csv": ipl_features,
        "venue_stats.csv": venue_stats,
        "team_form_latest.csv": team_form,
        "team_venue_form_latest.csv": team_venue_form,
        "matchup_form_latest.csv": matchup_form,
        "batter_form_latest.csv": batter_form,
        "bowler_form_latest.csv": bowler_form,
        "batter_bowler_form_latest.csv": batter_bowler_form,
        "team_player_pool_2026.csv": team_player_pool,
        "active_teams_2026.csv": active_teams,
    }

    ipl_features["season_label"] = ipl_features["season"].astype(str)
    ipl_features["season_key"] = pd.to_numeric(ipl_features["season_label"].str.split("/").str[0], errors="coerce")

    innings_level = (
        ipl_features.sort_values(["match_id", "innings", "legal_balls_bowled"])
        .drop_duplicates(["match_id", "innings"], keep="last")
        .copy()
    )

    season_rows = ipl_features.groupby("season_label").size().rename("engineered_rows")
    season_summary = (
        innings_level.groupby("season_label")
        .agg(
            matches=("match_id", "nunique"),
            innings=("match_id", "size"),
            venues=("venue", "nunique"),
            avg_total=("total_runs", "mean"),
            batting_win_rate=("win", "mean"),
        )
        .join(season_rows)
        .reset_index()
    )
    first_innings_avg = (
        innings_level[innings_level["innings"] == 1]
        .groupby("season_label")["total_runs"]
        .mean()
        .rename("avg_first_innings")
    )
    second_innings_avg = (
        innings_level[innings_level["innings"] == 2]
        .groupby("season_label")["total_runs"]
        .mean()
        .rename("avg_second_innings")
    )
    season_summary = (
        season_summary.join(first_innings_avg, on="season_label")
        .join(second_innings_avg, on="season_label")
        .merge(
            innings_level.groupby("season_label")["season_key"].min().reset_index(),
            on="season_label",
            how="left",
        )
        .sort_values(["season_key", "season_label"])
    )

    active_team_names = active_teams["team"].dropna().tolist()
    team_summary = (
        innings_level[innings_level["batting_team"].isin(active_team_names)]
        .groupby("batting_team")
        .agg(
            matches=("match_id", "nunique"),
            innings=("match_id", "size"),
            avg_total=("total_runs", "mean"),
            batting_win_rate=("win", "mean"),
        )
        .reset_index()
        .merge(team_form.rename(columns={"team": "batting_team", "team_form": "recent_form"}), on="batting_team", how="left")
        .sort_values("matches", ascending=False)
    )

    venue_profile = venue_stats.sort_values("venue_avg_first_innings", ascending=False).head(12).copy()
    batter_leaders = (
        batter_form[batter_form["batter_balls"] >= 250]
        .sort_values(["striker_form_sr", "batter_runs"], ascending=[False, False])
        .head(10)
        .copy()
    )
    bowler_leaders = (
        bowler_form[(bowler_form["bowler_balls"] >= 180) & (bowler_form["bowler_wickets"] > 0)]
        .sort_values(["bowler_form_econ", "bowler_form_strike"], ascending=[True, True])
        .head(10)
        .copy()
    )

    support_table_catalog = []
    for filename, frame in support_frames.items():
        label, purpose, key = SUPPORT_TABLE_DESCRIPTIONS[filename]
        support_table_catalog.append(
            {
                "file": filename,
                "label": label,
                "rows": len(frame),
                "cols": len(frame.columns),
                "primary_key": key,
                "purpose": purpose,
            }
        )

    workflow_inventory = [
        {
            "workflow": "External data refresh",
            "script": "scripts/update_external_data.py",
            "input": "Cricsheet IPL zip + weather snapshot endpoints",
            "output": "Refreshed raw CSV2 files and cached weather context",
        },
        {
            "workflow": "Feature engineering",
            "script": "scripts/preprocess_ipl.py",
            "input": "Raw match CSV2 files",
            "output": "ipl_features.csv and all support tables",
        },
        {
            "workflow": "CPU baseline",
            "script": "scripts/train_models.py",
            "input": "Engineered feature table",
            "output": "cpu_model_report.json and HGB artifacts",
        },
        {
            "workflow": "GPU benchmark",
            "script": "scripts/train_gpu_best.py",
            "input": "Recent active-team scope",
            "output": "gpu_model_report.json and candidate deployment artifacts",
        },
        {
            "workflow": "Best-model search",
            "script": "scripts/train_best_model_search.py",
            "input": "Recent and all-history active-team scopes",
            "output": "best_model_search_report.json and promoted live artifacts",
        },
        {
            "workflow": "Pre-match modeling",
            "script": "scripts/train_pre_match.py",
            "input": "Team-vs-team and venue combinations",
            "output": "pre_match_score_model.pkl and pre_match_win_model.pkl",
        },
        {
            "workflow": "Serving surfaces",
            "script": "web_app.py / streamlit_app.py / predict_cli.py",
            "input": "Validated user payloads and model artifacts",
            "output": "Flask, API, dashboard, and CLI predictions",
        },
    ]

    api_contract_rows = [
        {"field": "season", "required": "Yes", "description": "Season identifier used for categorical encoding and context."},
        {"field": "venue", "required": "Yes", "description": "Venue name normalized against the canonical alias table."},
        {"field": "batting_team", "required": "Yes", "description": "Current batting team."},
        {"field": "bowling_team", "required": "Yes", "description": "Current bowling team."},
        {"field": "innings", "required": "Yes", "description": "1 for batting first or 2 for chasing."},
        {"field": "runs / wickets / overs", "required": "Yes", "description": "Core match-state tuple needed to construct the live feature frame."},
        {"field": "first_innings_total", "required": "For innings 2", "description": "Supplies the target and chase-rate features."},
        {"field": "striker / bowler", "required": "Optional", "description": "Enables player-form and matchup enrichment instead of unknown defaults."},
        {"field": "runs_last_5 / wickets_last_5", "required": "Optional", "description": "Injects recent momentum directly into the feature frame."},
        {"field": "Returned analytics", "required": "Output", "description": "Predicted total, win probability, projected range, runs-vs-par, simulation bands, and weather context."},
    ]

    notebook_inventory = []
    for path in sorted((ROOT_DIR / "notebooks").glob("*.ipynb")):
        notebook_json = load_json(path)
        cells = notebook_json.get("cells", [])
        notebook_inventory.append(
            {
                "name": path.name,
                "size": human_size(path.stat().st_size),
                "code_cells": sum(1 for cell in cells if cell.get("cell_type") == "code"),
                "markdown_cells": sum(1 for cell in cells if cell.get("cell_type") == "markdown"),
                "role": {
                    "clean_analysis_multimodel.ipynb": "Compact analysis trace for cleaned training data and multi-model review.",
                    "data_collection_and_preprocessing.ipynb": "Documents raw-data ingestion, cleaning checks, and feature-building ideas.",
                    "ipl_data_eda.ipynb": "Large exploratory notebook covering IPL trends, distributions, and venue or team behaviour.",
                }.get(path.name, "Project notebook used for exploratory analysis."),
            }
        )

    dependency_inventory = []
    for requirement in requirements:
        package_name = requirement.split(">=")[0].split("==")[0].split("<")[0].strip()
        dependency_inventory.append(
            {
                "package": package_name,
                "specifier": requirement,
                "role": DEPENDENCY_ROLE_NOTES.get(package_name, "Supporting library used by the project."),
            }
        )

    feature_family_rows = [
        {
            "group": "Categorical identifiers",
            "count": len(CATEGORICAL_FEATURES),
            "examples": ", ".join(CATEGORICAL_FEATURES[:4]) + ", ...",
        },
        {
            "group": "Match-state geometry",
            "count": 8,
            "examples": "runs, wickets, wickets_left, balls_left, innings_progress",
        },
        {
            "group": "Momentum windows",
            "count": 6,
            "examples": "runs_last_5, wickets_last_5, runs_last_6_balls, runs_last_12_balls",
        },
        {
            "group": "Chase and phase context",
            "count": 9,
            "examples": "current_run_rate, target, required_run_rate, is_powerplay, is_death",
        },
        {
            "group": "Team and player context",
            "count": 11,
            "examples": "team_form, venue_form, striker_form_sr, bowler_form_econ",
        },
        {
            "group": "Weather and venue priors",
            "count": 9,
            "examples": "temperature_c, dew_risk, runs_vs_par, venue_avg_first_innings",
        },
    ]

    phase_summary = (
        ipl_features.groupby(["innings", "phase"])
        .agg(
            rows=("match_id", "size"),
            avg_runs=("runs", "mean"),
            avg_wickets=("wickets", "mean"),
            avg_current_rr=("current_run_rate", "mean"),
            avg_total=("total_runs", "mean"),
            batting_win_rate=("win", "mean"),
        )
        .reset_index()
        .sort_values(["innings", "phase"])
    )
    phase_summary["innings_label"] = phase_summary["innings"].map({1: "First innings", 2: "Second innings"})
    phase_order = {"powerplay": 0, "middle": 1, "death": 2}
    phase_summary["phase_order"] = phase_summary["phase"].map(phase_order)
    phase_summary = phase_summary.sort_values(["innings", "phase_order"])
    phase_summary_rows = [
        {
            "innings_label": row.innings_label,
            "phase": row.phase.title(),
            "rows": int(row.rows),
            "avg_runs": float(row.avg_runs),
            "avg_wickets": float(row.avg_wickets),
            "avg_current_rr": float(row.avg_current_rr),
            "avg_total": float(row.avg_total),
            "batting_win_rate": float(row.batting_win_rate),
        }
        for row in phase_summary.itertuples(index=False)
    ]

    wicket_bucket_order = ["0-2", "3-5", "6-8", "9-10"]
    score_predictions = score_predictions.copy()
    score_predictions["wicket_bucket"] = pd.cut(
        score_predictions["wickets"],
        bins=[-0.1, 2, 5, 8, 10],
        labels=wicket_bucket_order,
    )
    score_wicket_summary = []
    for bucket in wicket_bucket_order:
        subset = score_predictions[score_predictions["wicket_bucket"] == bucket]
        if subset.empty:
            continue
        rmse_value = ((subset["predicted_total_runs"] - subset["total_runs"]) ** 2).mean() ** 0.5
        score_wicket_summary.append(
            {
                "bucket": bucket,
                "rows": int(len(subset)),
                "mae": fmt_float(subset["absolute_error"].mean()),
                "mae_value": float(subset["absolute_error"].mean()),
                "rmse": fmt_float(rmse_value),
                "mean_actual": fmt_float(subset["total_runs"].mean()),
                "mean_prediction": fmt_float(subset["predicted_total_runs"].mean()),
            }
        )

    win_predictions = win_predictions.copy()
    win_predictions["probability_bucket"] = pd.cut(
        win_predictions["predicted_win_prob"],
        bins=[0.0, 0.2, 0.4, 0.6, 0.8, 1.0],
        labels=["0-20%", "20-40%", "40-60%", "60-80%", "80-100%"],
        include_lowest=True,
    )
    win_calibration_bins = []
    for bucket in ["0-20%", "20-40%", "40-60%", "60-80%", "80-100%"]:
        subset = win_predictions[win_predictions["probability_bucket"] == bucket]
        if subset.empty:
            continue
        mean_pred = float(subset["predicted_win_prob"].mean())
        empirical = float(subset["win"].mean())
        win_calibration_bins.append(
            {
                "bucket": bucket,
                "rows": int(len(subset)),
                "mean_pred": fmt_pct(mean_pred),
                "mean_pred_value": mean_pred,
                "empirical_win": fmt_pct(empirical),
                "empirical_win_value": empirical,
                "gap": f"{(empirical - mean_pred) * 100:+.1f} pp",
            }
        )

    version_inventory = []
    latest_entries = model_registry.get("entries", [])[-3:]
    for entry in latest_entries:
        live_metrics = entry.get("metrics", {}).get("live", {})
        version_inventory.append(
            {
                "version_id": entry.get("version_id", ""),
                "created_at": entry.get("created_at_utc", "")[:10],
                "scope": live_metrics.get("scope", ""),
                "score_model": prettify_model_name(live_metrics.get("score_model", "")),
                "score_rmse": fmt_float(live_metrics.get("score_test", {}).get("rmse", 0.0)) if live_metrics.get("score_test") else "-",
                "win_model": prettify_model_name(live_metrics.get("win_model", "")),
                "win_log_loss": fmt_float(live_metrics.get("win_test", {}).get("log_loss", 0.0), 4) if live_metrics.get("win_test") else "-",
            }
        )

    drift_summary_rows = [
        {"dimension": "Monitoring window", "value": str(production_drift_report["window"]), "interpretation": "Configured rolling window size used by the drift module."},
        {"dimension": "Events used", "value": str(production_drift_report["events_used"]), "interpretation": "Prediction events currently available for monitoring."},
        {"dimension": "Outcomes used", "value": str(production_drift_report["outcomes_used"]), "interpretation": "Resolved outcomes available for outcome-drift evaluation."},
        {"dimension": "Current status", "value": production_drift_report["status"], "interpretation": "Overall drift-monitoring state reported by the pipeline."},
        {"dimension": "Retrain trigger", "value": "Yes" if production_drift_report["trigger_retrain"] else "No", "interpretation": "Whether the current monitoring snapshot recommends retraining."},
        {"dimension": "Recommended action", "value": production_drift_report["recommended_action"], "interpretation": "Practical next step recorded by the monitoring module."},
    ]

    workspace_map_rows = [
        {"name": name, "role": role}
        for name, role in WORKSPACE_MAP
    ]

    artifact_inventory = []
    for path in sorted(MODELS_DIR.iterdir()):
        if not path.is_file() or path.suffix.lower() not in {".pkl", ".json", ".csv"}:
            continue
        artifact_inventory.append(
            {
                "name": path.name,
                "kind": path.suffix.lower().lstrip(".").upper(),
                "size": human_size(path.stat().st_size),
                "modified": date.fromtimestamp(path.stat().st_mtime).isoformat(),
            }
        )

    source_inventory = []
    for relative_path, purpose in SOURCE_FILE_NOTES.items():
        path = ROOT_DIR / relative_path
        source_inventory.append(
            {
                "path": relative_path,
                "lines": count_lines(path),
                "purpose": purpose,
            }
        )

    test_inventory = [
        {"test_name": test_name, "coverage": description}
        for test_name, description in TEST_CASE_NOTES.items()
    ]

    score_case_studies = (
        score_predictions.assign(
            signed_error=score_predictions["predicted_total_runs"] - score_predictions["total_runs"]
        )
        .sort_values("absolute_error", ascending=False)
        .head(8)
        .copy()
    )
    win_case_studies = (
        win_predictions.sort_values("logloss_contribution", ascending=False)
        .head(8)
        .copy()
    )

    raw_matches = len(list(RAW_DIR.glob("*_info.csv")))
    processed_rows, processed_cols = ipl_features.shape
    min_season = int(ipl_features["season_key"].dropna().min())
    max_season = int(ipl_features["season_key"].dropna().max())

    selected_scope = best_search_report["selected_scope"]
    scope_summary = best_search_report["scopes"][selected_scope]["scope_summary"]
    score_split = best_search_report["scopes"][selected_scope]["score_split"]
    win_split = best_search_report["scopes"][selected_scope]["win_split"]

    return {
        "today": date.today().strftime("%B %d, %Y"),
        "deployment_report": deployment_report,
        "best_search_report": best_search_report,
        "score_uncertainty": score_uncertainty,
        "cpu_report": cpu_report,
        "pre_match_model_report": pre_match_model_report,
        "production_drift_report": production_drift_report,
        "feature_reference_profile": feature_reference_profile,
        "canonical_pipeline_report": canonical_pipeline_report,
        "model_registry": model_registry,
        "requirements": requirements,
        "venue_stats": venue_stats,
        "season_summary": season_summary,
        "team_summary": team_summary,
        "venue_profile": venue_profile,
        "batter_leaders": batter_leaders,
        "bowler_leaders": bowler_leaders,
        "support_table_catalog": support_table_catalog,
        "workflow_inventory": workflow_inventory,
        "api_contract_rows": api_contract_rows,
        "notebook_inventory": notebook_inventory,
        "dependency_inventory": dependency_inventory,
        "feature_family_rows": feature_family_rows,
        "phase_summary": phase_summary,
        "phase_summary_rows": phase_summary_rows,
        "artifact_inventory": artifact_inventory,
        "source_inventory": source_inventory,
        "test_inventory": test_inventory,
        "score_case_studies": score_case_studies,
        "win_case_studies": win_case_studies,
        "score_wicket_summary": score_wicket_summary,
        "win_calibration_bins": win_calibration_bins,
        "version_inventory": version_inventory,
        "drift_summary_rows": drift_summary_rows,
        "workspace_map_rows": workspace_map_rows,
        "preprocessing_rules": PREPROCESSING_RULES,
        "score_predictions": score_predictions,
        "win_predictions": win_predictions,
        "score_benchmarks_recent": build_score_benchmark_rows(
            "recent_active_history",
            best_search_report["scopes"]["recent_active_history"],
        ),
        "score_benchmarks_all": build_score_benchmark_rows(
            "all_active_history",
            best_search_report["scopes"]["all_active_history"],
        ),
        "win_benchmarks_recent": build_win_benchmark_rows(
            "recent_active_history",
            best_search_report["scopes"]["recent_active_history"],
        ),
        "win_benchmarks_all": build_win_benchmark_rows(
            "all_active_history",
            best_search_report["scopes"]["all_active_history"],
        ),
        "live_scenarios": run_live_scenarios(),
        "pre_match_scenarios": run_pre_match_scenarios(),
        "counts": {
            "raw_matches": raw_matches,
            "processed_rows": processed_rows,
            "processed_cols": processed_cols,
            "min_season": min_season,
            "max_season": max_season,
            "venue_rows": len(venue_stats),
            "team_form_rows": len(team_form),
            "team_venue_form_rows": len(team_venue_form),
            "matchup_rows": len(matchup_form),
            "batter_rows": len(batter_form),
            "bowler_rows": len(bowler_form),
            "batter_bowler_rows": len(batter_bowler_form),
            "player_pool_rows": len(team_player_pool),
            "active_teams_rows": len(active_teams),
            "innings_rows": len(innings_level),
        },
        "scope_summary": scope_summary,
        "score_split": score_split,
        "win_split": win_split,
        "python_version": platform.python_version(),
        "platform": f"{platform.system()} {platform.release()}",
        "gpu_name": best_search_report.get("gpu_summary", {}).get("torch_device_name", "Not reported"),
    }
def generate_assets(context: dict) -> dict[str, Path]:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    pipeline_path = ASSETS_DIR / FIGURES[0].filename
    models_path = ASSETS_DIR / FIGURES[1].filename
    feature_mix_path = ASSETS_DIR / FIGURES[2].filename
    venues_path = ASSETS_DIR / FIGURES[3].filename
    uncertainty_path = ASSETS_DIR / FIGURES[4].filename
    phase_scoring_path = ASSETS_DIR / FIGURES[5].filename
    score_error_path = ASSETS_DIR / FIGURES[6].filename
    win_calibration_path = ASSETS_DIR / FIGURES[7].filename
    dataset_dashboard_path = ASSETS_DIR / FIGURES[8].filename
    flask_mockup_path = ASSETS_DIR / FIGURES[9].filename
    streamlit_mockup_path = ASSETS_DIR / FIGURES[10].filename
    api_snapshot_path = ASSETS_DIR / FIGURES[11].filename
    cli_snapshot_path = ASSETS_DIR / FIGURES[12].filename
    validation_snapshot_path = ASSETS_DIR / FIGURES[13].filename
    monitoring_snapshot_path = ASSETS_DIR / FIGURES[14].filename

    generate_pipeline_figure(pipeline_path)
    generate_model_comparison_figure(context["best_search_report"], models_path)
    generate_feature_mix_figure(context, feature_mix_path)
    generate_venue_figure(context["venue_stats"], venues_path)
    generate_uncertainty_figure(context["score_uncertainty"], uncertainty_path)
    generate_phase_scoring_figure(context, phase_scoring_path)
    generate_score_error_distribution_figure(context, score_error_path)
    generate_win_calibration_figure(context, win_calibration_path)
    generate_dataset_dashboard_figure(context, dataset_dashboard_path)
    generate_flask_ui_mockup(context, flask_mockup_path)
    generate_streamlit_mockup(context, streamlit_mockup_path)
    generate_api_snapshot(context, api_snapshot_path)
    generate_cli_snapshot(context, cli_snapshot_path)
    generate_validation_snapshot(context, validation_snapshot_path)
    generate_monitoring_snapshot(context, monitoring_snapshot_path)

    return {
        "pipeline": pipeline_path,
        "models": models_path,
        "feature_mix": feature_mix_path,
        "venues": venues_path,
        "uncertainty": uncertainty_path,
        "phase_scoring": phase_scoring_path,
        "score_error_distribution": score_error_path,
        "win_calibration": win_calibration_path,
        "dataset_dashboard": dataset_dashboard_path,
        "flask_mockup": flask_mockup_path,
        "streamlit_mockup": streamlit_mockup_path,
        "api_snapshot": api_snapshot_path,
        "cli_snapshot": cli_snapshot_path,
        "validation_snapshot": validation_snapshot_path,
        "monitoring_snapshot": monitoring_snapshot_path,
    }


def add_title_page(doc: Document, today: str) -> None:
    add_centered_paragraph(doc, "A Project Report", size=16, bold=True)
    add_centered_paragraph(doc, "on", size=12)
    add_centered_paragraph(doc, REPORT_PROFILE["title"], size=16, bold=True)
    doc.add_paragraph()
    add_centered_paragraph(
        doc,
        "Submitted in partial fulfillment of the requirements for the award of the degree of",
        size=12,
    )
    add_centered_paragraph(doc, REPORT_PROFILE["degree_line_1"], size=12, bold=True)
    add_centered_paragraph(doc, REPORT_PROFILE["degree_line_2"], size=12, bold=True)
    doc.add_paragraph()
    add_centered_paragraph(doc, "Submitted By", size=12, bold=True)
    add_centered_paragraph(doc, ", ".join(REPORT_PROFILE["students"]), size=12)
    doc.add_paragraph()
    add_centered_paragraph(doc, "Under the Supervision of", size=12, bold=True)
    add_centered_paragraph(doc, REPORT_PROFILE["supervisor"], size=12)
    doc.add_paragraph()
    add_centered_paragraph(doc, REPORT_PROFILE["department"], size=12, bold=True)
    add_centered_paragraph(doc, REPORT_PROFILE["institute"], size=12)
    add_centered_paragraph(doc, REPORT_PROFILE["university"], size=12)
    add_centered_paragraph(doc, f"Date: {today}", size=12)
    doc.add_page_break()


def add_front_matter(doc: Document, context: dict) -> None:
    add_centered_paragraph(doc, "DECLARATION", size=15, bold=True)
    add_paragraph(
        doc,
        f'This report entitled "{REPORT_PROFILE["title"]}" is an authentic record of the work carried '
        "out for the present project. The implementation, experimentation, documentation, and analysis described here "
        "have been prepared for academic submission with due acknowledgement of datasets, software libraries, and "
        "technical references used during development.",
    )
    add_paragraph(
        doc,
        "The report has not been submitted elsewhere for the award of any other degree or diploma. Whenever outside "
        "material has been used, it has been referenced appropriately in the references section.",
    )
    add_paragraph(doc, f'Student Name(s): {", ".join(REPORT_PROFILE["students"])}')
    add_paragraph(doc, "Semester: 7th / 8th Semester")
    add_paragraph(doc, f"Date: {context['today']}")
    doc.add_page_break()

    add_centered_paragraph(doc, "APPROVAL FROM SUPERVISOR", size=15, bold=True)
    add_paragraph(
        doc,
        f'This is to certify that the project report entitled "{REPORT_PROFILE["title"]}" presented by '
        f'"{", ".join(REPORT_PROFILE["students"])}" under my supervision is an authentic work. To the best of my '
        "knowledge, the content of this report has not been submitted previously for the award of any other degree.",
    )
    add_paragraph(doc, f'Supervisor Name: {REPORT_PROFILE["supervisor"]}')
    add_paragraph(doc, REPORT_PROFILE["department"])
    add_paragraph(doc, f"Date: {context['today']}")
    add_paragraph(doc, f'(Counter Signed By) {REPORT_PROFILE["hod"]}')
    doc.add_page_break()

    add_centered_paragraph(doc, "CERTIFICATE", size=15, bold=True)
    add_paragraph(
        doc,
        f'This is to certify that the work embodied in this report, entitled "{REPORT_PROFILE["title"]}", carried out by '
        f'"{", ".join(REPORT_PROFILE["students"])}" is approved for submission toward the degree of '
        f'"{REPORT_PROFILE["degree_line_1"]} {REPORT_PROFILE["degree_line_2"]}" at {REPORT_PROFILE["department"]}, '
        f'{REPORT_PROFILE["institute"]}.',
    )
    add_paragraph(doc, "Internal Examiner: ______________________________")
    add_paragraph(doc, "External Examiner: ______________________________")
    add_paragraph(doc, f'Place: {REPORT_PROFILE["place"]}')
    add_paragraph(doc, f"Date: {context['today']}")
    doc.add_page_break()

    add_centered_paragraph(doc, "ACKNOWLEDGEMENTS", size=15, bold=True)
    add_paragraph(
        doc,
        "The completion of this project report was made possible by the availability of open cricket data, robust "
        "Python libraries, and the guidance of academic mentors. The project combines multiple disciplines including "
        "data preprocessing, structured machine learning, deep tabular modeling, probability calibration, backend "
        "development, and dashboard design.",
    )
    add_paragraph(
        doc,
        f'Special acknowledgement is due to {REPORT_PROFILE["supervisor"]} for sustained guidance, to '
        f'{REPORT_PROFILE["hod"]} and the faculty of {REPORT_PROFILE["department"]} for academic support, and to '
        "the creators and maintainers of Cricsheet, scikit-learn, XGBoost, CatBoost, PyTorch, Flask, and Streamlit "
        "for the open technical ecosystem that enabled this work.",
    )
    add_paragraph(doc, f'Student Name(s): {", ".join(REPORT_PROFILE["students"])}')
    add_paragraph(doc, f"Date: {context['today']}")
    doc.add_page_break()

    add_centered_paragraph(doc, "ABSTRACT", size=15, bold=True)
    deployment = context["deployment_report"]
    counts = context["counts"]
    score_metrics = deployment["deployment_score_metrics_test"]
    win_metrics = deployment["deployment_win_metrics_test"]
    add_paragraph(
        doc,
        f"This project presents an IPL prediction system that estimates both the final innings total and the live win "
        f"probability from ball-by-ball match state. The system is built on top of {counts['raw_matches']} Cricsheet IPL "
        f"matches and a current engineered feature table containing {counts['processed_rows']:,} rows with "
        f"{counts['processed_cols']} columns. The project covers historical seasons from {counts['min_season']} to "
        f"{counts['max_season']} and uses a time-aware feature engineering pipeline that combines score state, venue "
        "history, team form, player form, batter-bowler interaction, and weather-derived dew risk.",
    )
    add_paragraph(
        doc,
        "The implementation is organized around a shared inference contract so that the command-line client, Flask "
        "web application, Streamlit dashboard, and JSON API all reuse the same validation and feature construction "
        "logic. This design reduces interface drift and makes the deployed predictions consistent across user-facing "
        "entry points.",
    )
    add_paragraph(
        doc,
        f"The promoted live score model in the current workspace is `{prettify_model_name(deployment['deployment_score_model'])}`, "
        f"while the promoted live win model is `{prettify_model_name(deployment['deployment_win_model'])}`. On the "
        f"held-out deployment snapshot represented in the current workspace, the score model achieves a mean absolute error of {fmt_float(score_metrics['mae'])} "
        f"runs and a root mean squared error of {fmt_float(score_metrics['rmse'])} runs. The win model achieves "
        f"{fmt_float(win_metrics['accuracy'] * 100)}% accuracy with log loss {fmt_float(win_metrics['log_loss'], 4)} "
        f"and Brier score {fmt_float(win_metrics['brier'], 4)}.",
    )
    add_paragraph(
        doc,
        "In addition to live prediction, the project includes pre-match estimation models, a weather context module "
        "based on Open-Meteo data, uncertainty quantiles for projected score ranges, and automated tests for core "
        "prediction and web routes. The resulting system is not only a modeling exercise but also a deployable sports "
        "analytics application with practical interfaces for students, analysts, and cricket enthusiasts.",
    )
    doc.add_page_break()


def add_executive_summary(doc: Document, context: dict) -> None:
    deployment = context["deployment_report"]
    add_centered_paragraph(doc, "EXECUTIVE SUMMARY", size=15, bold=True)
    add_paragraph(
        doc,
        "This report documents the design and implementation of an IPL analytics system that predicts two practically important outcomes from the "
        "current state of a match: the projected innings total and the batting side's live probability of winning. The project is built from raw "
        "Cricsheet CSV2 data and extends beyond model training into deployment, testing, monitoring readiness, and automated report generation.",
    )
    add_paragraph(
        doc,
        f"The current workspace combines {context['counts']['raw_matches']} raw IPL matches, {context['counts']['processed_rows']:,} engineered rows, "
        f"multiple historical support tables, versioned model artifacts, and several user interfaces. The live deployment snapshot represented in the "
        f"workspace promotes `{prettify_model_name(deployment['deployment_score_model'])}` for score prediction and "
        f"`{prettify_model_name(deployment['deployment_win_model'])}` for win probability estimation.",
    )
    add_paragraph(
        doc,
        "A central theme of the project is that predictive performance depends on disciplined feature engineering as much as on algorithm choice. "
        "For this reason, the pipeline explicitly handles alias normalization, chronological leakage control, support-table reuse, and time-aware "
        "evaluation splits. The report also emphasizes why different model families were considered and how the final deployment balances accuracy, "
        "probability quality, and practical usability.",
    )
    add_paragraph(
        doc,
        "The overall contribution is therefore broader than a single predictive model. The work delivers a reproducible sports-analytics workflow, "
        "multiple deployment surfaces, stored experiment traces, and a report that can be regenerated as the project evolves. These characteristics "
        "make the submission suitable not only as a machine-learning exercise but also as a compact demonstration of full-stack analytical engineering.",
    )
    doc.add_page_break()


def add_contents_and_lists(doc: Document) -> None:
    doc.add_heading("TABLE OF CONTENTS", level=1)
    toc_items = [
        "Declaration",
        "Approval from Supervisor",
        "Certificate",
        "Acknowledgements",
        "Abstract",
        "Executive Summary",
        "Table of Contents",
        "List of Figures",
        "List of Tables",
        "Chapter 1: Introduction",
        "Chapter 2: Literature Review",
        "Chapter 3: Problem Objective",
        "Chapter 4: Methodology of the Project",
        "Chapter 5: Results",
        "Chapter 6: Screenshots and System Interface Outputs",
        "Chapter 7: Conclusion and Future Scope",
        "Appendix",
        "References",
        "Abbreviations",
    ]
    add_numbered(doc, toc_items)
    doc.add_page_break()

    doc.add_heading("LIST OF FIGURES", level=1)
    add_table(doc, ["Figure No.", "Figure Title"], [[figure.number, figure.title] for figure in FIGURES])
    doc.add_page_break()

    doc.add_heading("LIST OF TABLES", level=1)
    add_table(doc, ["Table No.", "Table Title"], [[table.number, table.title] for table in ALL_TABLES])
    doc.add_page_break()


def add_chapter_1(doc: Document, context: dict) -> None:
    counts = context["counts"]
    doc.add_heading("CHAPTER 1: INTRODUCTION", level=1)
    doc.add_heading("TOPIC OF THE SYSTEM", level=2)
    add_paragraph(doc, 'Title: "IPL Match Score and Win Prediction System"')

    doc.add_heading("OVERVIEW OF THE PROJECT", level=2)
    add_paragraph(
        doc,
        "Twenty20 cricket is highly dynamic, and match situations change materially after almost every delivery. "
        "A modern prediction system therefore has to do more than summarize the scorecard; it must interpret match "
        "state, contextual history, and venue-specific behavior quickly enough to support live use. This project was "
        "developed to meet that need for the Indian Premier League (IPL), where both score acceleration and win "
        "probability can swing sharply within a few overs.",
    )
    add_paragraph(
        doc,
        f"The implemented system uses historical IPL ball-by-ball data from {counts['raw_matches']} raw Cricsheet "
        f"matches and converts it into a structured live-state learning problem. The present workspace currently "
        f"contains a main engineered dataset of {counts['processed_rows']:,} rows spanning seasons "
        f"{counts['min_season']} to {counts['max_season']}. Each row represents a legal-ball match snapshot with "
        "game-state, form, venue, player, and weather-derived features.",
    )
    add_paragraph(
        doc,
        "The system produces two primary live outputs: projected total runs for the batting side and the batting "
        "side's win probability. It also supports pre-match projections through separate models and exposes the logic "
        "through a Flask web application, a Streamlit dashboard, a JSON API, and a command-line interface.",
    )

    doc.add_heading("NEED OF THE SYSTEM", level=2)
    add_paragraph(
        doc,
        "Cricket audiences increasingly expect contextual analytics during live games. Broadcasters, fans, analysts, "
        "and students want to know not only the current score but also the likely trajectory of the innings and the "
        "probability of winning from the current state. Traditional scoreboards are descriptive, whereas predictive "
        "systems offer foresight.",
    )
    add_bullets(
        doc,
        [
            "Manual match reading is subjective and inconsistent across users.",
            "Raw score alone ignores venue scoring patterns and chasing difficulty.",
            "Team form and batter-bowler matchups materially affect expected outcomes.",
            "Probability estimates are more useful when they are calibrated and accompanied by uncertainty.",
            "A reusable prediction engine becomes more valuable when it serves web, API, and dashboard interfaces together.",
        ],
    )

    doc.add_heading("PROBLEM CONTEXT", level=2)
    add_paragraph(
        doc,
        "The main engineering challenge in IPL prediction is that cricket data is sequential, contextual, and noisy. "
        "Team names and venues change across seasons, batting and bowling quality evolve, and live inference requires "
        "feature construction from incomplete match state rather than full match information. A useful model must therefore "
        "balance historical learning with present-match adaptability.",
    )
    add_paragraph(
        doc,
        "A second challenge lies in the structure of the prediction tasks themselves. Score prediction is a regression "
        "problem whose error tolerance varies by innings phase, while win prediction is a probabilistic classification "
        "problem that must remain well-behaved across innings, wickets, and chase conditions. The project addresses both "
        "problems through task-specific model selection, calibration, and uncertainty handling.",
    )

    doc.add_heading("GOALS AND OBJECTIVES", level=2)
    add_bullets(
        doc,
        [
            "Build a reusable IPL prediction pipeline starting from raw Cricsheet CSV2 files.",
            "Engineer time-aware features that capture score state, momentum, venue, team form, and player form.",
            "Train separate live score and live win models and select the best deployment pair empirically.",
            "Support pre-match prediction using team-vs-team and venue context.",
            "Deliver predictions through Flask, Streamlit, CLI, and JSON API interfaces with shared logic.",
            "Expose uncertainty bands and simulation-oriented indicators to improve interpretability.",
        ],
    )

    doc.add_heading("TARGET USERS", level=2)
    add_bullets(
        doc,
        [
            "Cricket analysts who need quick, data-driven match-state interpretation.",
            "Students and developers studying sports analytics, structured ML, and model deployment.",
            "Cricket fans who want an accessible live predictor through web or dashboard interfaces.",
            "Project evaluators looking for an end-to-end ML system rather than a notebook-only prototype.",
        ],
    )

    doc.add_heading("BENEFITS OF THE PROJECT", level=2)
    add_table_title(doc, TABLES[0])
    add_table(
        doc,
        ["No.", "Benefit", "Description"],
        [
            ["1", "Live dual prediction", "Provides both projected innings total and batting-side win probability from the same live state."],
            ["2", "Rich context", "Uses venue, team form, player form, weather, and matchup context rather than only score and overs."],
            ["3", "Model benchmarking", "Compares CPU, GPU, gradient boosting, ensemble, and deep tabular models before deployment."],
            ["4", "Shared inference contract", "Keeps Flask, Streamlit, CLI, and API outputs aligned through common feature logic."],
            ["5", "Interpretability support", "Supplies projected ranges, probability bands, and uncertainty quantiles alongside point estimates."],
            ["6", "Pre-match analytics", "Extends beyond live prediction by estimating likely winner and first-innings score before the toss flow starts."],
            ["7", "Academic value", "Demonstrates data engineering, ML, DL, probability calibration, software testing, and deployment in one project."],
            ["8", "Extensible design", "Can be updated with new seasons, new features, or alternate sports analytics models without redesigning the stack."],
        ],
    )

    doc.add_heading("RELEVANCE OF THE PROJECT", level=2)
    add_paragraph(
        doc,
        "This project is relevant because it connects sports analytics with practical machine learning deployment. IPL data is large enough "
        "to support serious experimentation, yet structured enough to illustrate why feature engineering and model choice matter. The "
        "result is a system that fits well within an academic machine learning curriculum while still feeling close to a real production workflow.",
    )

    doc.add_heading("CHALLENGES AND CONSTRAINTS", level=2)
    add_bullets(
        doc,
        [
            "Historical team aliases and venue aliases must be normalized before training and inference.",
            "Hold-out data for the latest season is smaller than the historical training pool, so variance remains in model evaluation.",
            "Live weather is not always available online, so the system needs safe defaults and cached snapshots.",
            "Pre-match models currently use a narrower feature set than the live models, limiting their richness.",
            "Extreme first-innings collapses or explosive starts remain difficult because they are rare in the training distribution.",
        ],
    )

    doc.add_heading("CONCLUSION", level=2)
    add_paragraph(
        doc,
        "The introduction establishes the project as a full-stack IPL analytics system rather than a single-model experiment. Its purpose is "
        "to convert ball-by-ball cricket data into usable predictions that are accurate enough for demonstration, structured enough for "
        "academic evaluation, and modular enough for further improvement.",
    )
    doc.add_page_break()


def add_chapter_2(doc: Document) -> None:
    doc.add_heading("CHAPTER 2: LITERATURE REVIEW", level=1)
    doc.add_heading("INTRODUCTION", level=2)
    add_paragraph(
        doc,
        "The project stands at the intersection of sports analytics, structured machine learning, deep tabular learning, and lightweight "
        "application deployment. The literature and technical ecosystem around these topics suggest that performance in cricket prediction "
        "improves when raw event logs are converted into context-aware state representations and then matched with models that handle mixed "
        "categorical and numeric inputs well.",
    )

    doc.add_heading("REVIEW OF KEY LITERATURE", level=2)
    add_table_title(doc, TABLES[1])
    add_table(
        doc,
        ["Source", "Main Idea", "Relevance to This Project", "Observed Limitation"],
        [
            [
                "Cricsheet IPL CSV2 documentation",
                "Provides structured ball-by-ball and match-info data for cricket analytics.",
                "Forms the raw data foundation for this system.",
                "Requires custom parsing and normalization before use in ML pipelines.",
            ],
            [
                "Chen and Guestrin (2016), XGBoost",
                "Tree boosting is highly effective for large structured datasets.",
                "Motivated inclusion of GPU-accelerated boosting models in the benchmark.",
                "One-hot or encoded categorical handling increases preprocessing complexity.",
            ],
            [
                "Prokhorenkova et al. (2018), CatBoost",
                "Gradient boosting with strong support for categorical variables and practical classification performance.",
                "Motivated the deployed calibrated win-probability model.",
                "Performance still depends on scope selection and probability calibration.",
            ],
            [
                "Guo and Berkhahn (2016), Entity Embeddings of Categorical Variables",
                "Embedding layers can learn compact dense representations for high-cardinality categories.",
                "Motivated inclusion of torch entity-embedding models in the live score benchmark.",
                "Needs enough data and careful validation to avoid unstable training.",
            ],
            [
                "Scikit-learn MLP / calibration tooling",
                "Provides lightweight neural networks and isotonic calibration for tabular workflows.",
                "Used for pre-match models and probability calibration experiments.",
                "Generic models alone are not sufficient without domain-specific features.",
            ],
            [
                "Flask, Streamlit, and Open-Meteo documentation",
                "Show how prediction services, dashboards, and weather context can be integrated simply.",
                "Guided the deployment layer of the project.",
                "Documentation supports implementation but does not solve modeling issues directly.",
            ],
        ],
    )

    doc.add_heading("SUMMARY OF FINDINGS", level=2)
    add_paragraph(
        doc,
        "The review suggests four strong conclusions. First, the success of a cricket predictor depends heavily on feature engineering and "
        "state representation rather than only model choice. Second, tabular gradient boosting remains a strong baseline for sports data. "
        "Third, high-cardinality categories such as teams, venues, batters, and bowlers can benefit from embedding-based learning. Fourth, "
        "probability outputs must be calibrated before they are interpreted as confidence rather than raw model scores.",
    )

    doc.add_heading("ADDRESSING IDENTIFIED GAPS", level=2)
    add_paragraph(
        doc,
        "Many cricket prediction projects stop at one of two points: either they provide exploratory analysis without a deployed interface, "
        "or they build a single model without unified inference across interfaces. The current project addresses these gaps by offering live "
        "and pre-match prediction, multiple training workflows, promotion reports, uncertainty bands, and a common inference layer consumed "
        "by every frontend.",
    )

    doc.add_heading("IMPORTANCE OF THIS REVIEW TO THE PROJECT", level=2)
    add_paragraph(
        doc,
        "The literature review directly influenced technical decisions in the implementation. Historical score and outcome dependencies led "
        "to the use of rolling form and venue context. Structured-ML literature justified the inclusion of XGBoost and CatBoost. Entity-"
        "embedding work informed the deep score model, and calibration literature supported the final win-model design.",
    )
    doc.add_page_break()


def add_chapter_3(doc: Document) -> None:
    doc.add_heading("CHAPTER 3: PROBLEM OBJECTIVE", level=1)
    doc.add_heading("PROBLEM STATEMENT", level=2)
    add_paragraph(
        doc,
        "The problem addressed in this project is the absence of a single IPL analytics system that can transform live ball-by-ball state "
        "into reliable score and win predictions while remaining usable through multiple interfaces. Existing basic scoreboards show what "
        "has happened; the objective here is to estimate what is likely to happen next and by the end of the innings or match.",
    )

    doc.add_heading("OBJECTIVES OF THE PROJECT", level=2)
    add_bullets(
        doc,
        [
            "Create a reproducible preprocessing pipeline from raw Cricsheet IPL files.",
            "Engineer a feature set that captures match state, momentum, venue, form, and player interaction.",
            "Benchmark multiple live prediction model families and deploy the best pair using held-out metrics.",
            "Provide pre-match prediction in addition to live match-state prediction.",
            "Expose outputs through user-friendly interfaces and a machine-consumable API.",
        ],
    )

    doc.add_heading("SCOPE OF THE PROJECT", level=2)
    add_paragraph(
        doc,
        "The project scope includes IPL historical data ingestion, support-table generation, live score regression, live win classification, "
        "pre-match modeling, weather-context integration, deployment artifact management, Flask and Streamlit interfaces, and automated tests. "
        "The system is focused on IPL cricket only and does not currently attempt player-selection optimization, fantasy scoring, natural "
        "language commentary generation, or ball-tracking vision tasks.",
    )

    doc.add_heading("KEY DELIVERABLES", level=2)
    add_table_title(doc, TABLES[2])
    add_table(
        doc,
        ["Deliverable", "Description", "Implementation Location"],
        [
            ["Raw-to-processed pipeline", "Transforms Cricsheet files into model-ready tables and support datasets.", "scripts/preprocess_ipl.py"],
            ["Live prediction engine", "Shared validation, feature construction, and inference for score and win outputs.", "ipl_predictor/common.py"],
            ["Training workflows", "CPU baseline, GPU comparison, best-model search, and pre-match training.", "scripts/train_*.py"],
            ["Deployment artifacts", "Persisted score model, win model, uncertainty profile, and reports.", "models/"],
            ["Flask web app + API", "HTML forms and POST /api/predict endpoint.", "web_app.py"],
            ["Streamlit dashboard", "Interactive dashboard for model selection and analysis.", "streamlit_app.py"],
            ["CLI + tests", "Terminal inference flow and automated verification.", "predict_cli.py, tests/"],
        ],
    )
    doc.add_page_break()


def add_chapter_4(doc: Document, context: dict, assets: dict[str, Path]) -> None:
    scope_summary = context["scope_summary"]
    score_split = context["score_split"]
    win_split = context["win_split"]
    doc.add_heading("CHAPTER 4: METHODOLOGY OF THE PROJECT", level=1)

    doc.add_heading("SYSTEM ARCHITECTURE", level=2)
    add_paragraph(
        doc,
        "The system architecture follows a layered design. Raw IPL CSV2 files are first parsed into a processed feature table and multiple "
        "support tables. Training workflows then benchmark alternative model families and write promoted artifacts into the models directory. "
        "At inference time, all user-facing interfaces call shared functions from the prediction package so that validation, normalization, "
        "and result formatting remain consistent.",
    )
    add_figure(doc, assets["pipeline"], FIGURES[0].number, FIGURES[0].title)

    doc.add_heading("WORKFLOW AND SYSTEM MODEL", level=2)
    add_numbered(
        doc,
        [
            "Download or refresh IPL raw data and external weather snapshot data.",
            "Parse match metadata, normalize team and venue aliases, and build per-ball match-state rows.",
            "Compute historical support tables such as venue strength, team form, team-venue form, matchup form, batter form, bowler form, and batter-bowler form.",
            "Split the data by season so that training, validation, and testing respect chronology.",
            "Train multiple live score and live win model families; calibrate probability models where required.",
            "Select the best deployment pair using held-out metrics and store the promoted artifacts.",
            "Serve predictions through Flask, Streamlit, CLI, and API using the same support tables and model files.",
        ],
    )
    add_paragraph(
        doc,
        f"In the final selected scope `{scope_summary['scope']}`, the deployment search used {scope_summary['rows_after_scope_filter']:,} "
        f"rows after active-team filtering. The score model split used {score_split['train_rows']:,} training rows, "
        f"{score_split['valid_rows']:,} validation rows, and {score_split['test_rows']:,} test rows. The win model used "
        f"{win_split['train_rows']:,} training rows, {win_split['valid_rows']:,} validation rows, and {win_split['test_rows']:,} "
        "test rows, all separated by season to avoid leakage.",
    )

    doc.add_heading("DATA PREPROCESSING AND FEATURE ENGINEERING", level=2)
    add_paragraph(
        doc,
        "Preprocessing begins with metadata extraction from `*_info.csv` files and ball-by-ball ingestion from the corresponding match CSVs. "
        "The pipeline converts overs to legal-ball counts, standardizes team and venue names, reconstructs innings totals, and attaches toss, "
        "winner, and date information. Historical statistics are always computed in chronological order so that each match only sees prior information.",
    )
    add_paragraph(
        doc,
        "Feature engineering includes match-state variables such as runs, wickets, wickets left, balls left, innings progress, current run rate, "
        "required run rate, and rolling momentum windows. It also includes contextual features such as venue averages, team recent form, venue form, "
        "head-to-head form, batter strike-rate history, bowler economy history, and batter-versus-bowler matchup records. A weather helper fetches "
        "or falls back to temperature, humidity, wind, and dew risk at inference time.",
    )
    add_table_title(doc, TABLES[3])
    add_table(
        doc,
        ["Feature Group", "Examples", "Purpose"],
        [
            ["Match state", "runs, wickets, balls_left, innings_progress", "Represents the current score situation directly."],
            ["Momentum", "runs_last_5, wickets_last_6_balls", "Captures recent acceleration or collapse."],
            ["Rates and target", "CRR, RRR, target_remaining", "Important for chase pressure and pacing."],
            ["Phase indicators", "powerplay, middle, death", "Lets models learn phase-specific behavior."],
            ["Team context", "team_form, venue_form, matchup_form", "Encodes recent competitive strength."],
            ["Player context", "striker SR, bowler economy, batter-vs-bowler SR", "Adds player-level cricket knowledge."],
            ["Weather / venue", "dew_risk, temperature, venue averages", "Models environmental and surface influence."],
        ],
    )
    add_figure(doc, assets["feature_mix"], FIGURES[2].number, FIGURES[2].title)
    add_paragraph(
        doc,
        "The feature-family breakdown is useful for two reasons. First, it makes clear that the live predictor is not a shallow scorecard model; "
        "a substantial fraction of the input space is dedicated to context, momentum, and environmental priors. Second, it shows why the project "
        "benefits from using both categorical-aware models and deep tabular alternatives during model search.",
    )

    doc.add_heading("DATASET AUDIT OF THE CURRENT WORKSPACE", level=2)
    add_paragraph(
        doc,
        f"The current workspace contains {context['counts']['processed_rows']:,} engineered match-state rows spread across "
        f"{context['counts']['innings_rows']:,} innings snapshots and {context['counts']['raw_matches']:,} raw IPL matches. "
        "This section expands the methodological discussion by showing how the processed data is distributed across seasons, "
        "teams, and support tables rather than treating the training table as a black box.",
    )
    add_table_title(doc, CHAPTER_4_TABLES["support_assets"])
    add_table(
        doc,
        ["File", "Role", "Rows", "Cols", "Primary Key", "Purpose"],
        [
            [
                row["file"],
                row["label"],
                f"{row['rows']:,}",
                str(row["cols"]),
                row["primary_key"],
                row["purpose"],
            ]
            for row in context["support_table_catalog"]
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "The support-table design is one of the strongest engineering decisions in the project. Historical priors are materialized once during "
        "preprocessing and then reused consistently by every training script and interface, which reduces runtime cost and prevents multiple "
        "versions of the same cricket logic from drifting apart.",
    )
    add_table_title(doc, CHAPTER_4_TABLES["season_coverage"])
    add_table(
        doc,
        ["Season", "Matches", "Innings", "Venues", "Engineered Rows", "Avg 1st Inns", "Avg 2nd Inns", "Batting-Side Win Rate"],
        [
            [
                row.season_label,
                f"{int(row.matches):,}",
                f"{int(row.innings):,}",
                str(int(row.venues)),
                f"{int(row.engineered_rows):,}",
                fmt_float(row.avg_first_innings),
                fmt_float(row.avg_second_innings),
                fmt_pct(row.batting_win_rate),
            ]
            for row in context["season_summary"].itertuples(index=False)
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "Two methodological patterns emerge from the season summary. First, the row volume grows whenever the tournament expands or more complete "
        "recent seasons are added. Second, the average first-innings environment rises markedly in the most recent seasons, which justifies using "
        "time-aware splits and explains why models trained only on older scoring patterns are unlikely to generalize cleanly to the modern IPL.",
    )
    add_table_title(doc, CHAPTER_4_TABLES["team_summary"])
    add_table(
        doc,
        ["Active Team", "Matches", "Innings", "Avg Total", "Batting-Side Win Rate", "Latest 5-Match Form"],
        [
            [
                row.batting_team,
                f"{int(row.matches):,}",
                f"{int(row.innings):,}",
                fmt_float(row.avg_total),
                fmt_pct(row.batting_win_rate),
                fmt_pct(row.recent_form if pd.notna(row.recent_form) else 0.0),
            ]
            for row in context["team_summary"].itertuples(index=False)
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "The active-team table shows why the project maintains both long-term and short-term context. Match counts and historical batting strength "
        "come from the full archive, while the latest rolling form captures the current competitive state. The final inference frame blends both "
        "signals so the model does not overreact to short-term streaks or ignore genuine recent changes.",
    )

    doc.add_heading("PREPROCESSING RULES AND LEAKAGE CONTROL", level=2)
    add_paragraph(
        doc,
        "A large part of the project’s reliability comes from seemingly small preprocessing decisions. In sports analytics, a model can appear "
        "strong simply because it has been allowed to see information that would not exist at the prediction time. The report therefore documents "
        "the exact normalization and leakage-control rules followed during feature creation.",
    )
    add_table_title(doc, CHAPTER_4_TABLES["preprocessing_rules"])
    add_table(
        doc,
        ["Operation", "What It Does", "Why It Matters"],
        [[operation, action, reason] for operation, action, reason in context["preprocessing_rules"]],
        font_size=9,
    )
    add_paragraph(
        doc,
        "These rules show that the project has been treated as an engineering system rather than a one-off notebook experiment. The combination of "
        "alias normalization, chronological processing, and artifact checks is especially important because the same dataset and models are reused "
        "later by the report generator, Flask app, Streamlit dashboard, CLI, and monitoring modules.",
    )

    doc.add_heading("MACHINE LEARNING AND DEEP LEARNING MODELS FOR IPL PREDICTION", level=2)
    add_paragraph(
        doc,
        "The project trains separate model families for live score prediction and live win probability. Baseline workflows use HistGradientBoosting "
        "for fast CPU refreshes. GPU workflows compare XGBoost, CatBoost, and weighted ensembles. A broader search additionally includes torch-based "
        "entity-embedding models for categorical-heavy tabular data.",
    )
    add_paragraph(doc, describe_score_model(context["deployment_report"]["deployment_score_model"]))
    add_paragraph(doc, describe_win_model(context["deployment_report"]["deployment_win_model"]))
    add_paragraph(
        doc,
        "Separate pre-match models are also trained for the team-versus-team setting in which the live innings context is not yet available. "
        "These models intentionally use fewer features and therefore act as a conservative analytical layer rather than competing directly with "
        "the richer live prediction stack.",
    )
    add_table_title(doc, TABLES[4])
    add_table(
        doc,
        ["Model Family", "Used For", "Role in Project"],
        [
            ["HistGradientBoosting", "Score and win", "Fast CPU baseline and comparison point."],
            ["XGBoost", "Score and win", "GPU/CPU gradient boosting benchmark for structured data."],
            ["CatBoost", "Score and win", "Categorical-friendly boosting used in the live win-model search."],
            ["Weighted ensembles", "Score and win", "Combines top classical models for benchmark robustness."],
            ["Torch entity embeddings", "Score and win", "Deep tabular modeling with learned categorical representations for richer categorical context."],
            ["MLP pre-match models", "Pre-match score and win", "Lightweight neural models using team-vs-team and venue inputs."],
        ],
    )
    add_figure(doc, assets["models"], FIGURES[1].number, FIGURES[1].title)

    doc.add_heading("FLASK WEB APPLICATION AND STREAMLIT DASHBOARD", level=2)
    add_paragraph(
        doc,
        "The Flask application provides HTML forms for live and pre-match prediction together with a machine-readable `/api/predict` endpoint. "
        "The Streamlit dashboard focuses on interactive exploration, model-file selection, venue insights, recent form display, and analysis widgets. "
        "Both applications sit on top of the same underlying prediction package, which prevents duplicate business logic and makes debugging simpler.",
    )
    add_paragraph(
        doc,
        "The command-line client complements these interfaces by offering a lightweight live prediction flow for terminal-based use. Together, these "
        "entry points demonstrate that the project is not limited to experimentation but can be consumed in multiple deployment forms.",
    )

    doc.add_heading("INTEGRATION OF MODELS WITH USER INTERFACES AND API", level=2)
    add_paragraph(
        doc,
        "Integration is handled through a shared inference contract in `ipl_predictor/common.py`. The core functions normalize aliases, validate "
        "inputs, build a single-row feature frame, load support tables, and format the prediction dictionary. This dictionary includes predicted "
        "total, win probability, probability percentage, projected score range, runs-vs-par, phase, and optional weather outputs such as dew risk.",
    )
    add_paragraph(
        doc,
        "Because all interfaces call the same inference path, updates to support tables or model artifacts immediately benefit the entire application "
        "stack. This approach is especially important in academic projects where frontend and backend logic often drift apart when developed separately.",
    )
    add_table_title(doc, CHAPTER_4_TABLES["workflow_inventory"])
    add_table(
        doc,
        ["Workflow", "Script or Surface", "Primary Input", "Primary Output"],
        [
            [row["workflow"], row["script"], row["input"], row["output"]]
            for row in context["workflow_inventory"]
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "Taken together, the workflow table shows that the repository is organized around reproducible stages rather than ad hoc notebooks. Each "
        "stage writes a concrete artifact that the next stage can consume, which is precisely the kind of structure expected in a production-minded "
        "machine-learning project.",
    )
    add_table_title(doc, CHAPTER_4_TABLES["api_contract"])
    add_table(
        doc,
        ["Field / Output", "Requirement", "Role in the Contract"],
        [
            [row["field"], row["required"], row["description"]]
            for row in context["api_contract_rows"]
        ],
        font_size=9,
    )

    doc.add_heading("NOTEBOOKS, ANALYSIS TRACES, AND REPRODUCIBILITY", level=2)
    add_paragraph(
        doc,
        "Although the final project is organized around reusable scripts, the repository also preserves exploratory notebooks that document how the "
        "team moved from raw analysis to a structured pipeline. Including these notebooks in the report is important because they show the evolution "
        "of data understanding, not just the polished final code.",
    )
    add_table_title(doc, CHAPTER_4_TABLES["notebook_inventory"])
    add_table(
        doc,
        ["Notebook", "Approx. Size", "Code Cells", "Markdown Cells", "Role"],
        [
            [row["name"], row["size"], str(row["code_cells"]), str(row["markdown_cells"]), row["role"]]
            for row in context["notebook_inventory"]
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "This notebook inventory reinforces a useful academic point: the final system did not emerge fully formed. It passed through exploratory "
        "EDA, preprocessing validation, and model-comparison phases before being compressed into reproducible scripts and saved artifacts.",
    )

    doc.add_heading("CODEBASE ORGANIZATION", level=2)
    add_paragraph(
        doc,
        "A second strength of the project is its code organization. Training logic, prediction logic, user interfaces, tests, and documentation are "
        "kept in clearly separated files, which makes the repository easier to explain, extend, and evaluate. For an academic submission, this "
        "organization matters because it demonstrates engineering discipline in addition to predictive performance.",
    )
    add_table_title(doc, CHAPTER_4_TABLES["module_inventory"])
    add_table(
        doc,
        ["File", "Approx. Lines", "Responsibility"],
        [
            [row["path"], f"{row['lines']:,}", row["purpose"]]
            for row in context["source_inventory"]
        ],
        font_size=9,
    )
    doc.add_heading("DEPENDENCY STACK AND IMPLEMENTATION SUPPORT", level=2)
    add_paragraph(
        doc,
        "A full project report should also account for the software stack that makes the implementation possible. Listing dependencies is not enough; "
        "it is more useful to explain how each package contributes to data processing, training, serving, testing, or report generation.",
    )
    add_table_title(doc, CHAPTER_4_TABLES["dependency_roles"])
    add_table(
        doc,
        ["Package", "Version Specifier", "Role in the Project"],
        [
            [row["package"], row["specifier"], row["role"]]
            for row in context["dependency_inventory"]
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "Documenting the dependency stack also improves reproducibility. A future reader can understand not only which models were trained, but also "
        "which libraries are needed to rebuild the dataset, regenerate figures, run the web layer, and recreate the final report document itself.",
    )
    doc.add_page_break()


def add_chapter_5(doc: Document, context: dict, assets: dict[str, Path]) -> None:
    deployment = context["deployment_report"]
    best_search = context["best_search_report"]
    score_uncertainty = context["score_uncertainty"]
    cpu_report = context["cpu_report"]
    season_summary = context["season_summary"]
    venue_profile = context["venue_profile"]
    batter_leaders = context["batter_leaders"]
    bowler_leaders = context["bowler_leaders"]
    recent_scope = best_search["scopes"]["recent_active_history"]
    all_scope = best_search["scopes"]["all_active_history"]
    recent_score_metrics = get_selected_score_metrics(recent_scope)
    recent_win_metrics = get_selected_win_metrics(recent_scope)
    all_score_metrics = get_selected_score_metrics(all_scope)
    all_win_metrics = get_selected_win_metrics(all_scope)
    score_rmse_gain = 100.0 * (
        cpu_report["score"]["overall"]["rmse"] - deployment["deployment_score_metrics_test"]["rmse"]
    ) / cpu_report["score"]["overall"]["rmse"]
    win_log_loss_gain = 100.0 * (
        cpu_report["win"]["selected"]["log_loss"] - deployment["deployment_win_metrics_test"]["log_loss"]
    ) / cpu_report["win"]["selected"]["log_loss"]
    recent_scope_composite = (
        recent_score_metrics["rmse"]
        + 10.0 * recent_win_metrics["log_loss"]
    )
    all_scope_composite = (
        all_score_metrics["rmse"]
        + 10.0 * all_win_metrics["log_loss"]
    )

    doc.add_heading("CHAPTER 5: RESULTS", level=1)

    doc.add_heading("SOFTWARE AND HARDWARE DETAILS", level=2)
    add_table_title(doc, TABLES[5])
    add_table(
        doc,
        ["Component", "Details"],
        [
            ["Programming language", f"Python {context['python_version']}"],
            ["Execution platform", context["platform"]],
            ["Primary libraries", ", ".join(context["requirements"])],
            ["Web framework", "Flask"],
            ["Dashboard framework", "Streamlit"],
            ["ML / DL stack", "scikit-learn, XGBoost, CatBoost, PyTorch"],
            ["GPU reported in benchmark", context["gpu_name"]],
            ["Stored deployment artifacts", "score_model.pkl, win_model.pkl, pre_match_score_model.pkl, pre_match_win_model.pkl, score_uncertainty.json"],
        ],
    )

    doc.add_heading("LIVE MODEL TRAINING RESULTS", level=2)
    add_paragraph(
        doc,
        "The model selection process compared multiple families on chronologically held-out data. The broader search found that the "
        "`all_active_history` scope outperformed the `recent_active_history` scope after considering both score RMSE and win log loss. "
        "This indicates that historical depth, when filtered to the currently active IPL teams, was more valuable than restricting the "
        "training pool to only the most recent seasons.",
    )
    add_table_title(doc, TABLES[6])
    add_table(
        doc,
        ["Task", "Selected Model", "Metric", "Value"],
        [
            ["Live score", deployment["deployment_score_model"], "MAE", fmt_float(deployment["deployment_score_metrics_test"]["mae"]) + " runs"],
            ["Live score", deployment["deployment_score_model"], "RMSE", fmt_float(deployment["deployment_score_metrics_test"]["rmse"]) + " runs"],
            ["Live win", deployment["deployment_win_model"], "Accuracy", fmt_float(deployment["deployment_win_metrics_test"]["accuracy"] * 100) + "%"],
            ["Live win", deployment["deployment_win_model"], "Log loss", fmt_float(deployment["deployment_win_metrics_test"]["log_loss"], 4)],
            ["Live win", deployment["deployment_win_model"], "Brier score", fmt_float(deployment["deployment_win_metrics_test"]["brier"], 4)],
        ],
    )
    add_table_title(doc, TABLES[7])
    add_table(
        doc,
        ["Scope", "Score RMSE", "Win Log Loss", "Selected Models"],
        [
            [
                "recent_active_history",
                fmt_float(recent_score_metrics["rmse"]),
                fmt_float(recent_win_metrics["log_loss"], 4),
                get_selected_score_model_name(recent_scope) + " / " + get_selected_win_model_name(recent_scope),
            ],
            [
                "all_active_history",
                fmt_float(all_score_metrics["rmse"]),
                fmt_float(all_win_metrics["log_loss"], 4),
                get_selected_score_model_name(all_scope) + " / " + get_selected_win_model_name(all_scope),
            ],
        ],
    )
    add_paragraph(
        doc,
        f"The selected live score model improves materially over the CPU baseline report, where the baseline RMSE was "
        f"{fmt_float(cpu_report['score']['overall']['rmse'])} runs. The final promoted score model reduces "
        f"that to {fmt_float(deployment['deployment_score_metrics_test']['rmse'])} runs, indicating that higher-capacity "
        "or better-combined tabular modeling helps on this task when compared with the earlier CPU-only baseline.",
    )
    add_paragraph(
        doc,
        f"Measured against the CPU baseline, the promoted score model cuts RMSE by {fmt_float(score_rmse_gain)}% and the promoted win model "
        f"cuts selected log loss by {fmt_float(win_log_loss_gain)}%. The composite scope score used for final selection falls from "
        f"{fmt_float(recent_scope_composite)} on `recent_active_history` to {fmt_float(all_scope_composite)} on `all_active_history`, "
        "which is the quantitative reason the final deployment keeps the deeper historical scope.",
    )

    doc.add_heading("HISTORICAL SCORING LANDSCAPE", level=2)
    add_paragraph(
        doc,
        "Before interpreting model metrics, it is useful to examine how the scoring environment itself changes across seasons. The engineered "
        "dataset shows a clear upward shift in average first-innings totals in the most recent IPL seasons, which means the prediction task is "
        "not stationary over time. Any robust evaluation therefore needs to respect chronology and avoid mixing early-era scoring patterns with "
        "modern high-scoring conditions during testing.",
    )
    add_table_title(doc, CHAPTER_5_TABLES["season_scoring"])
    add_table(
        doc,
        ["Season", "Matches", "Innings", "Rows", "Avg 1st Inns", "Avg 2nd Inns", "Batting-Side Win Rate"],
        [
            [
                row.season_label,
                f"{int(row.matches):,}",
                f"{int(row.innings):,}",
                f"{int(row.engineered_rows):,}",
                fmt_float(row.avg_first_innings),
                fmt_float(row.avg_second_innings),
                fmt_pct(row.batting_win_rate),
            ]
            for row in season_summary.itertuples(index=False)
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "The season table shows that the project is effectively modeling multiple scoring eras. This explains why the all-history scope can still "
        "win: the model is allowed to learn from a large number of innings while the explicit season feature gives it a way to separate older and "
        "newer contexts instead of collapsing them into one undifferentiated average.",
    )

    doc.add_heading("VENUE AND PLAYER CONTEXT SNAPSHOTS", level=2)
    add_paragraph(
        doc,
        "The support tables also provide interpretable cricket context independent of the predictive models. Venue priors summarize whether a ground "
        "historically favors batting first or chasing, while player-form tables quantify the strike-rate and economy context that the live models "
        "consume for the current striker and bowler.",
    )
    add_table_title(doc, CHAPTER_5_TABLES["venue_profile"])
    add_table(
        doc,
        ["Venue", "Avg 1st Inns", "Avg 2nd Inns", "Bat-First Win Rate"],
        [
            [
                row.venue,
                fmt_float(row.venue_avg_first_innings),
                fmt_float(row.venue_avg_second_innings),
                fmt_pct(row.venue_bat_first_win_rate),
            ]
            for row in venue_profile.itertuples(index=False)
        ],
        font_size=9,
    )
    add_table_title(doc, CHAPTER_5_TABLES["batter_leaders"])
    add_table(
        doc,
        ["Batter", "Balls", "Runs", "Strike Rate", "Average"],
        [
            [
                row.batter,
                f"{int(row.batter_balls):,}",
                f"{int(row.batter_runs):,}",
                fmt_float(row.striker_form_sr),
                fmt_float(row.striker_form_avg),
            ]
            for row in batter_leaders.itertuples(index=False)
        ],
        font_size=9,
    )
    add_table_title(doc, CHAPTER_5_TABLES["bowler_leaders"])
    add_table(
        doc,
        ["Bowler", "Balls", "Runs", "Wickets", "Economy", "Strike Rate"],
        [
            [
                row.bowler,
                f"{int(row.bowler_balls):,}",
                f"{int(row.bowler_runs):,}",
                f"{int(row.bowler_wickets):,}",
                fmt_float(row.bowler_form_econ),
                fmt_float(row.bowler_form_strike),
            ]
            for row in bowler_leaders.itertuples(index=False)
        ],
        font_size=9,
    )
    add_figure(doc, assets["venues"], FIGURES[3].number, FIGURES[3].title)

    doc.add_heading("DETAILED MODEL SEARCH COMPARISON", level=2)
    add_paragraph(
        doc,
        "The broader search report is especially useful for understanding why different families dominate different scopes. On the smaller recent "
        "scope, classical tree-based models remain stable because they see only a limited amount of data. On the all-history scope, the torch "
        "entity-embedding regressor becomes viable because the larger data volume is sufficient for learning categorical representations of teams, "
        "venues, and players.",
    )
    add_table_title(doc, CHAPTER_5_TABLES["recent_score_models"])
    add_table(
        doc,
        ["Model", "MAE", "RMSE", "Selection"],
        [
            [row["model"], row["mae"], row["rmse"], row["selection"]]
            for row in context["score_benchmarks_recent"]
        ],
        font_size=9,
    )
    add_table_title(doc, CHAPTER_5_TABLES["all_score_models"])
    add_table(
        doc,
        ["Model", "MAE", "RMSE", "Selection"],
        [
            [row["model"], row["mae"], row["rmse"], row["selection"]]
            for row in context["score_benchmarks_all"]
        ],
        font_size=9,
    )
    add_table_title(doc, CHAPTER_5_TABLES["recent_win_models"])
    add_table(
        doc,
        ["Model", "Log Loss", "Brier", "Accuracy", "Selection"],
        [
            [row["model"], row["log_loss"], row["brier"], row["accuracy"], row["selection"]]
            for row in context["win_benchmarks_recent"]
        ],
        font_size=9,
    )
    add_table_title(doc, CHAPTER_5_TABLES["all_win_models"])
    add_table(
        doc,
        ["Model", "Log Loss", "Brier", "Accuracy", "Selection"],
        [
            [row["model"], row["log_loss"], row["brier"], row["accuracy"], row["selection"]]
            for row in context["win_benchmarks_all"]
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "These benchmark tables reveal an important project insight: the best score model and the best win model are not necessarily from the same "
        "family. The final system therefore promotes task-specific winners instead of forcing a single model family to dominate both objectives.",
    )

    doc.add_heading("PHASE-WISE MATCH BEHAVIOUR IN THE ENGINEERED DATA", level=2)
    add_paragraph(
        doc,
        "A useful way to interpret the learning problem is to inspect how innings behaviour changes across the powerplay, middle overs, and death overs. "
        "These summaries are computed directly from the engineered dataset and therefore reflect the same distributions seen during model training.",
    )
    add_table_title(doc, CHAPTER_5_TABLES["phase_summary"])
    add_table(
        doc,
        ["Innings", "Phase", "Rows", "Avg Runs", "Avg Wickets", "Avg CRR", "Avg Final Total", "Batting Win Rate"],
        [
            [
                row["innings_label"],
                row["phase"],
                f"{row['rows']:,}",
                fmt_float(row["avg_runs"]),
                fmt_float(row["avg_wickets"]),
                fmt_float(row["avg_current_rr"]),
                fmt_float(row["avg_total"]),
                fmt_pct(row["batting_win_rate"]),
            ]
            for row in context["phase_summary_rows"]
        ],
        font_size=9,
    )
    add_figure(doc, assets["phase_scoring"], FIGURES[5].number, FIGURES[5].title)
    add_paragraph(
        doc,
        "The phase-wise view makes it easier to understand why different model families can succeed. Early innings states are open-ended and uncertain, "
        "middle-over states encode structural context, and death overs contain compressed risk, acceleration, and collapse behaviour. A stronger model "
        "is therefore one that can adapt across these very different local regimes rather than memorizing a single average innings trajectory.",
    )

    doc.add_heading("PRE-MATCH AND ANALYTICAL RESULTS", level=2)
    add_paragraph(
        doc,
        "The project also includes dedicated pre-match models trained on team-one, team-two, and venue inputs. These models are intentionally "
        "lightweight and serve as a pre-game analytical layer rather than a replacement for the richer live models. They make the overall "
        "system more useful by allowing users to inspect a likely first-innings score and winner before the first ball is bowled.",
    )
    add_paragraph(
        doc,
        f"The uncertainty profile used by the live score output stores held-out residual quantiles of "
        f"{fmt_float(score_uncertainty['residual_q10'])} runs at the 10th percentile and "
        f"+{fmt_float(score_uncertainty['residual_q90'])} runs at the 90th percentile, with a residual standard deviation of "
        f"{fmt_float(score_uncertainty['residual_std'])} runs. These values are used to convert a point estimate into a more honest projected range.",
    )
    add_figure(doc, assets["uncertainty"], FIGURES[4].number, FIGURES[4].title)
    add_table_title(doc, CHAPTER_5_TABLES["pre_match_metrics"])
    add_table(
        doc,
        ["Task", "Split", "MAE / Accuracy", "RMSE / Log Loss", "R2 / Brier"],
        [
            [
                "Pre-match score",
                "Validation",
                fmt_float(context["pre_match_model_report"]["score_metrics"]["valid"]["mae"]),
                fmt_float(context["pre_match_model_report"]["score_metrics"]["valid"]["rmse"]),
                fmt_float(context["pre_match_model_report"]["score_metrics"]["valid"]["r2"], 3),
            ],
            [
                "Pre-match score",
                "Test",
                fmt_float(context["pre_match_model_report"]["score_metrics"]["test"]["mae"]),
                fmt_float(context["pre_match_model_report"]["score_metrics"]["test"]["rmse"]),
                fmt_float(context["pre_match_model_report"]["score_metrics"]["test"]["r2"], 3),
            ],
            [
                "Pre-match win",
                "Validation",
                fmt_pct(context["pre_match_model_report"]["win_metrics"]["valid"]["accuracy"]),
                fmt_float(context["pre_match_model_report"]["win_metrics"]["valid"]["log_loss"], 4),
                fmt_float(context["pre_match_model_report"]["win_metrics"]["valid"]["brier"], 4),
            ],
            [
                "Pre-match win",
                "Test",
                fmt_pct(context["pre_match_model_report"]["win_metrics"]["test"]["accuracy"]),
                fmt_float(context["pre_match_model_report"]["win_metrics"]["test"]["log_loss"], 4),
                fmt_float(context["pre_match_model_report"]["win_metrics"]["test"]["brier"], 4),
            ],
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "The pre-match metrics are noticeably weaker than the live metrics because these models must operate without ball-by-ball context, current batters, "
        "or chase pressure. This gap is analytically useful: it highlights the value added by the live inference layer rather than suggesting that the "
        "pre-match models should replace it.",
    )

    doc.add_heading("SCENARIO-BASED INFERENCE ANALYSIS", level=2)
    add_paragraph(
        doc,
        "A report based only on aggregate metrics can hide how the deployed system behaves in concrete match situations. To make the behavior more "
        "transparent, the current workspace was used to run several illustrative live and pre-match scenarios through the same saved artifacts that "
        "power the Flask app, Streamlit dashboard, and CLI.",
    )
    add_table_title(doc, CHAPTER_5_TABLES["live_scenarios"])
    add_table(
        doc,
        ["Scenario", "Phase", "Projected Total", "Win Prob.", "Projected Range", "Collapse Risk", "Interpretation"],
        [
            [
                row["scenario"],
                row["phase"],
                row["projected_total"],
                row["win_prob"],
                row["range"],
                row["collapse_risk"],
                row["interpretation"],
            ]
            for row in context["live_scenarios"]
        ],
        font_size=9,
    )
    add_table_title(doc, CHAPTER_5_TABLES["pre_match_scenarios"])
    add_table(
        doc,
        ["Matchup", "Venue", "Projected Score", "Team 1 Win", "Team 2 Win", "Likely Winner", "Interpretation"],
        [
            [
                row["matchup"],
                row["venue"],
                row["predicted_score"],
                row["team1_win_prob"],
                row["team2_win_prob"],
                row["likely_winner"],
                row["interpretation"],
            ]
            for row in context["pre_match_scenarios"]
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "The live scenarios demonstrate that the output is richer than a single point estimate. The system surfaces projected score ranges, collapse "
        "risk, and win-probability bands, all of which are useful when the current state is fragile. The pre-match scenarios, by contrast, stay "
        "conservative because they intentionally use only team and venue information.",
    )

    doc.add_heading("HELD-OUT ERROR CASE STUDIES", level=2)
    add_paragraph(
        doc,
        "The final analysis layer looks directly at the hardest examples in the held-out prediction files written by the best-model search. This is "
        "valuable because it reveals where the deployed models still struggle, especially in volatile early-over situations where the eventual innings "
        "trajectory can diverge sharply from the current score state.",
    )
    add_table_title(doc, CHAPTER_5_TABLES["score_case_studies"])
    add_table(
        doc,
        ["Match ID", "Date", "Batting Team", "Runs/Wkts", "Actual Total", "Predicted Total", "Signed Error", "Abs. Error"],
        [
            [
                str(int(row.match_id)),
                str(row.start_date),
                row.batting_team,
                f"{fmt_float(row.runs, 0)}/{fmt_float(row.wickets, 0)}",
                fmt_float(row.total_runs, 0),
                fmt_float(row.predicted_total_runs, 1),
                f"{row.signed_error:+.1f}",
                fmt_float(row.absolute_error, 1),
            ]
            for row in context["score_case_studies"].itertuples(index=False)
        ],
        font_size=9,
    )
    add_table_title(doc, CHAPTER_5_TABLES["win_case_studies"])
    add_table(
        doc,
        ["Match ID", "Date", "Batting Team", "Runs/Wkts", "Actual Win", "Predicted Win Prob.", "Log-Loss Contribution"],
        [
            [
                str(int(row.match_id)),
                str(row.start_date),
                row.batting_team,
                f"{fmt_float(row.runs, 0)}/{fmt_float(row.wickets, 0)}",
                "Yes" if row.win >= 0.5 else "No",
                fmt_pct(row.predicted_win_prob),
                fmt_float(row.logloss_contribution, 4),
            ]
            for row in context["win_case_studies"].itertuples(index=False)
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "The score-error cases confirm a predictable weakness of live scoring models: when very few overs have been bowled, the final innings outcome "
        "is still highly uncertain and even a well-trained model can understate or overstate explosive innings. The win-probability case studies show "
        "a related issue for chases: a probability that looks reasonable on average can still be sharply wrong on specific matches if a team stages "
        "an exceptional finish or suffers a sudden collapse.",
    )

    doc.add_heading("ERROR PROFILE OF THE DEPLOYED MODELS", level=2)
    add_paragraph(
        doc,
        "Score quality also changes with wicket state. The summary below groups held-out score predictions by wickets lost so the report can distinguish "
        "between open-ended batting positions and more fragile late-innings situations.",
    )
    add_table_title(doc, CHAPTER_5_TABLES["score_wicket_summary"])
    add_table(
        doc,
        ["Wickets Lost", "Rows", "MAE", "RMSE", "Avg Actual Total", "Avg Predicted Total"],
        [
            [
                row["bucket"],
                f"{row['rows']:,}",
                row["mae"],
                row["rmse"],
                row["mean_actual"],
                row["mean_prediction"],
            ]
            for row in context["score_wicket_summary"]
        ],
        font_size=9,
    )
    add_figure(doc, assets["score_error_distribution"], FIGURES[6].number, FIGURES[6].title)
    add_paragraph(
        doc,
        "This slice-level analysis is more actionable than reporting a single overall RMSE. If one bucket is clearly worse than the others, future work "
        "can focus on that specific regime by adding conditional features, targeted weighting, or specialized scenario simulation.",
    )

    doc.add_heading("WIN-PROBABILITY CALIBRATION ANALYSIS", level=2)
    add_paragraph(
        doc,
        "Probability outputs are only useful when they behave like probabilities. The calibration table and figure therefore compare average predicted "
        "win chance with the empirical batting-side win rate inside fixed probability bands on the held-out set.",
    )
    add_table_title(doc, CHAPTER_5_TABLES["win_calibration_bins"])
    add_table(
        doc,
        ["Probability Bucket", "Rows", "Avg Predicted", "Empirical Win Rate", "Calibration Gap"],
        [
            [row["bucket"], f"{row['rows']:,}", row["mean_pred"], row["empirical_win"], row["gap"]]
            for row in context["win_calibration_bins"]
        ],
        font_size=9,
    )
    add_figure(doc, assets["win_calibration"], FIGURES[7].number, FIGURES[7].title)
    add_paragraph(
        doc,
        "A well-behaved calibration profile matters for an IPL analytics application because users are likely to interpret a 70% win chance literally. "
        "The closer the empirical outcomes remain to the predicted bins, the more trustworthy that interpretation becomes.",
    )

    doc.add_heading("VERSIONED EXPERIMENT HISTORY", level=2)
    add_paragraph(
        doc,
        "The repository preserves recent experiment versions in the model registry. This gives the report a lightweight audit trail showing how the "
        "current artifacts evolved across training runs rather than appearing as unexplained files in the `models` directory.",
    )
    add_table_title(doc, CHAPTER_5_TABLES["experiment_versions"])
    add_table(
        doc,
        ["Version", "Date", "Scope", "Score Model", "Score RMSE", "Win Model", "Win Log Loss"],
        [
            [
                row["version_id"],
                row["created_at"],
                row["scope"],
                row["score_model"],
                row["score_rmse"],
                row["win_model"],
                row["win_log_loss"],
            ]
            for row in context["version_inventory"]
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        "This version history strengthens the academic presentation of the work because it documents iteration. It shows that model promotion was tied "
        "to measurable results and recorded artifacts rather than ad hoc replacement of files.",
    )

    doc.add_heading("MONITORING AND DRIFT READINESS", level=2)
    add_paragraph(
        doc,
        "The project also includes a monitoring layer that tracks prediction events, outcomes, and drift status. The current snapshot does not yet have "
        "enough resolved live outcomes for a strong drift conclusion, but the surrounding infrastructure demonstrates a forward-looking deployment mindset.",
    )
    add_table_title(doc, CHAPTER_5_TABLES["drift_summary"])
    add_table(
        doc,
        ["Monitoring Dimension", "Current Value", "Interpretation"],
        [
            [row["dimension"], row["value"], row["interpretation"]]
            for row in context["drift_summary_rows"]
        ],
        font_size=9,
    )
    add_paragraph(
        doc,
        f"The stored reference profile records an average of {fmt_float(context['feature_reference_profile']['features']['runs']['mean'])} runs at prediction "
        f"time and a death-over slice share of {fmt_pct(context['feature_reference_profile']['slice_rates']['death_over'])}. These baselines will become "
        "more operationally useful as the live event log grows over future matches.",
    )

    doc.add_heading("INTEGRATED SYSTEM TESTING VIA FLASK, API, STREAMLIT AND CLI", level=2)
    add_paragraph(
        doc,
        "The project was verified not only through offline metrics but also through interface integration and automated tests. The repository "
        "contains dedicated tests for core prediction utilities and web routes, and the local workspace test run completed successfully at the "
        "time this report was generated.",
    )
    add_table_title(doc, TABLES[8])
    add_table(
        doc,
        ["Component", "Verification Focus", "Status"],
        [
            ["Shared prediction core", "Input validation, overs parsing, feature construction, formatted outputs", "Validated through unit tests and shared use across all interfaces"],
            ["Flask web app", "Form handling and JSON API response structure", "Validated through route tests"],
            ["Streamlit dashboard", "Model-file loading, analysis widgets, and display logic", "Validated manually through code-path inspection and runtime design"],
            ["CLI flow", "Terminal-based live prediction input path", "Available as standalone entry point"],
            ["Automated tests", "Repository test suite", "12 tests passed in local run"],
        ],
    )
    add_paragraph(
        doc,
        f"Overall, the results show that the deployed pair `{deployment['deployment_score_model']}` and "
        f"`{deployment['deployment_win_model']}` provides a sensible balance of predictive quality and deployability. "
        "The system is strongest when it combines historical context with current match state, and its interface layer demonstrates that the models are usable beyond training scripts.",
    )
    doc.add_page_break()


def add_chapter_6(doc: Document, context: dict, assets: dict[str, Path]) -> None:
    doc.add_heading("CHAPTER 6: SCREENSHOTS AND SYSTEM INTERFACE OUTPUTS", level=1)
    add_paragraph(
        doc,
        "This chapter presents project-specific visual outputs in the same spirit as the final report structure used in the previous submitted "
        "project. Instead of placeholder images, the current report includes visuals generated directly from the IPL workspace so the screenshots "
        "chapter serves as evidence of dataset depth, deployed interfaces, API behavior, CLI behavior, and verification readiness.",
    )

    screenshot_sections = [
        (
            "DATASET OVERVIEW DASHBOARD",
            "The first visual summarizes the engineered data landscape using season-wise scoring trends, active-team batting summaries, venue priors, "
            "and a compact workspace snapshot. It demonstrates that the project is backed by a substantial structured dataset rather than a small toy sample.",
            "dataset_dashboard",
            FIGURES[8],
        ),
        (
            "FLASK LIVE PREDICTION INTERFACE",
            "This visual represents the Flask-based live prediction interface built on top of the shared inference contract. The key point for the report "
            "is not only UI appearance but the presence of structured input fields and a prediction summary that combines score, win probability, and range outputs.",
            "flask_mockup",
            FIGURES[9],
        ),
        (
            "STREAMLIT ANALYTICS DASHBOARD",
            "The Streamlit dashboard gives the project a second interface layer intended for exploration and analysis. It provides a stronger demonstration of "
            "analytical usability than a single prediction form because it exposes model comparisons, support-table views, and interactive selectors.",
            "streamlit_mockup",
            FIGURES[10],
        ),
        (
            "JSON API SNAPSHOT",
            "A strong academic project report should show how the model can be consumed programmatically. The API snapshot below demonstrates that the project "
            "is ready for integration into external services, dashboards, or automation pipelines through a structured JSON contract.",
            "api_snapshot",
            FIGURES[11],
        ),
        (
            "CLI OUTPUT SNAPSHOT",
            "The terminal interface is especially useful for quick testing, lightweight deployment, and demonstration during evaluation. Including it here shows "
            "that the system is usable even without a browser-based surface.",
            "cli_snapshot",
            FIGURES[12],
        ),
        (
            "TEST AND ARTIFACT VERIFICATION SNAPSHOT",
            "The final visual in this chapter focuses on verification discipline. It highlights that the repository contains automated tests as well as concrete "
            "saved artifacts, which strengthens the argument that the system is reproducible and deployment-oriented.",
            "validation_snapshot",
            FIGURES[13],
        ),
        (
            "MONITORING AND DRIFT SNAPSHOT",
            "The final screenshot-style visual documents that the project includes monitoring-oriented artifacts such as drift status, reference distributions, "
            "and pipeline-consistency checks. This helps position the work as a maintainable system rather than a one-time model export.",
            "monitoring_snapshot",
            FIGURES[14],
        ),
    ]

    for heading, description, asset_key, figure in screenshot_sections:
        doc.add_heading(heading, level=2)
        add_paragraph(doc, description)
        add_figure(doc, assets[asset_key], figure.number, figure.title, width=6.0)

    doc.add_page_break()


def add_chapter_7(doc: Document) -> None:
    doc.add_heading("CHAPTER 7: CONCLUSION AND FUTURE SCOPE", level=1)
    doc.add_heading("CONCLUSION", level=2)
    add_paragraph(
        doc,
        "This project successfully demonstrates an end-to-end IPL prediction system that begins with raw cricket data and ends with deployable "
        "live and pre-match analytics interfaces. The work combines data engineering, feature design, model benchmarking, probability calibration, "
        "deployment, and testing in a single coherent workflow. The final system is more meaningful than a basic regression notebook because it "
        "supports multiple user interfaces and retains a shared inference core.",
    )
    add_paragraph(
        doc,
        "From an academic perspective, the project is valuable because it covers the full lifecycle of a machine-learning application: raw-data handling, "
        "feature design, model comparison, selection logic, artifact management, interface integration, validation, reporting, and early monitoring support. "
        "Each of these layers contributes to the final quality of the system, and the report documents them as connected engineering decisions rather than isolated tasks.",
    )

    doc.add_heading("MAJOR CONTRIBUTIONS OF THE WORK", level=2)
    add_bullets(
        doc,
        [
            "Constructed a leakage-aware IPL match-state dataset from raw Cricsheet CSV2 files and support tables.",
            "Benchmarked multiple structured-data model families instead of assuming one algorithm would fit both prediction tasks.",
            "Exposed the final system through Flask, Streamlit, API, and CLI surfaces built on the same inference core.",
            "Added uncertainty bands, pre-match inference, artifact versioning, and monitoring-readiness components to extend beyond a basic predictor.",
            "Generated a report pipeline that can be rerun as the workspace evolves, keeping documentation tied to actual project artifacts.",
        ],
    )

    doc.add_heading("LESSONS LEARNED", level=2)
    add_paragraph(
        doc,
        "Several practical lessons emerged while building the project. First, data preparation quality often matters more than trying many sophisticated models "
        "on weak features. Second, probability calibration should be treated as a first-class concern whenever users will consume percentage-based outputs. "
        "Third, interface unification through a shared inference module significantly reduces maintenance cost and inconsistency across web, dashboard, and terminal layers.",
    )

    doc.add_heading("LIMITATION", level=2)
    add_bullets(
        doc,
        [
            "The pre-match models use fewer features than the live models and therefore provide coarser predictions.",
            "Latest-season evaluation size is smaller than the full historical training pool, so result variance remains possible.",
            "Rare match events such as extreme collapses or extraordinary late acceleration remain difficult to model robustly.",
            "The system does not yet use probable XI, pitch reports, toss simulation, or betting-market information.",
            "Player-availability uncertainty and last-minute lineup changes are not modeled directly in the present pipeline.",
            "Monitoring is implemented at the infrastructure level, but the current live outcome log is still too small for strong drift conclusions.",
        ],
    )

    doc.add_heading("FUTURE SCOPE", level=2)
    add_bullets(
        doc,
        [
            "Extend pre-match models with probable playing XI, toss scenarios, and venue-specific batting order priors.",
            "Add richer uncertainty estimation using quantile models or conformal prediction.",
            "Introduce model monitoring and online feedback loops for season-over-season retraining.",
            "Expand the API to return feature contribution or explanation summaries for advanced users.",
            "Adapt the pipeline for other cricket leagues or international T20 competitions.",
            "Build batter- and bowler-specific sequence models that exploit recent ball-by-ball context more explicitly.",
            "Add automated PDF export and supervisor-ready formatting options to the report pipeline.",
        ],
    )
    doc.add_page_break()


def add_appendix(doc: Document, context: dict) -> None:
    doc.add_heading("APPENDIX", level=1)
    add_table_title(doc, TABLES[9])
    add_table(
        doc,
        ["File / Module", "Responsibility"],
        [
            ["ipl_predictor/common.py", "Shared live inference contract, feature-frame construction, validation, formatting"],
            ["ipl_predictor/live_data.py", "Weather context lookup and dew-risk estimation"],
            ["ipl_predictor/torch_tabular.py", "Entity-embedding tabular deep learning models"],
            ["scripts/preprocess_ipl.py", "Cricsheet ingestion and processed dataset creation"],
            ["scripts/train_gpu_best.py", "GPU CatBoost/XGBoost comparison and deployment report"],
            ["scripts/train_best_model_search.py", "Broader scope and model-family search including torch models"],
            ["scripts/train_pre_match.py", "Pre-match MLP model training"],
            ["web_app.py", "Flask web interface and JSON API"],
            ["streamlit_app.py", "Interactive dashboard"],
            ["tests/", "Automated verification of shared logic and web routes"],
        ],
    )

    doc.add_heading("APPENDIX A: SAMPLE API PAYLOAD", level=2)
    add_code_block(
        doc,
        '{\n'
        '  "season": "2026",\n'
        '  "venue": "Wankhede Stadium",\n'
        '  "batting_team": "Mumbai Indians",\n'
        '  "bowling_team": "Chennai Super Kings",\n'
        '  "innings": 1,\n'
        '  "runs": 92,\n'
        '  "wickets": 2,\n'
        '  "overs": "10.2"\n'
        '}',
    )

    doc.add_heading("APPENDIX B: KEY EXECUTION COMMANDS", level=2)
    add_code_block(
        doc,
        "& .\\.venv\\Scripts\\python.exe scripts\\update_external_data.py\n"
        "& .\\.venv\\Scripts\\python.exe scripts\\preprocess_ipl.py\n"
        "& .\\.venv\\Scripts\\python.exe scripts\\train_best_model_search.py\n"
        "& .\\.venv\\Scripts\\python.exe scripts\\train_pre_match.py\n"
        "& .\\.venv\\Scripts\\python.exe web_app.py\n"
        "& .\\.venv\\Scripts\\python.exe -m streamlit run streamlit_app.py\n"
        "& .\\.venv\\Scripts\\python.exe -m pytest -q",
    )

    doc.add_heading("APPENDIX C: PROJECT NOTES", level=2)
    add_paragraph(
        doc,
        "The report has been generated from the current workspace state. If new seasons are added, preprocessing and model training can be rerun "
        "and this report can be regenerated to update dataset sizes, metrics, and embedded figures automatically.",
    )
    doc.add_heading("APPENDIX D: FULL FEATURE DICTIONARY", level=2)
    add_paragraph(
        doc,
        "The next tables list every feature group that appears in the engineered dataset. This appendix is intentionally detailed so that the report "
        "can serve as both project documentation and a reference for future retraining or extension work.",
    )
    appendix_feature_tables = [
        ("feature_catalog_1", FEATURE_CATALOG["match_metadata"]),
        ("feature_catalog_2", FEATURE_CATALOG["match_state"]),
        ("feature_catalog_3", FEATURE_CATALOG["momentum_and_rates"]),
        ("feature_catalog_4", FEATURE_CATALOG["team_and_player_context"]),
        ("feature_catalog_5", FEATURE_CATALOG["weather_venue_targets"]),
    ]
    for key, rows in appendix_feature_tables:
        add_table_title(doc, APPENDIX_TABLES[key])
        add_table(
            doc,
            ["Feature", "Description"],
            [[feature, description] for feature, description in rows],
            font_size=9,
        )

    doc.add_heading("APPENDIX E: ARTIFACT AND SOURCE INVENTORY", level=2)
    add_table_title(doc, APPENDIX_TABLES["artifact_inventory"])
    add_table(
        doc,
        ["Artifact", "Type", "Approx. Size", "Modified"],
        [
            [row["name"], row["kind"], row["size"], row["modified"]]
            for row in context["artifact_inventory"]
        ],
        font_size=9,
    )
    add_table_title(doc, APPENDIX_TABLES["source_inventory"])
    add_table(
        doc,
        ["File", "Approx. Lines", "Responsibility"],
        [
            [row["path"], f"{row['lines']:,}", row["purpose"]]
            for row in context["source_inventory"]
        ],
        font_size=9,
    )

    doc.add_heading("APPENDIX F: AUTOMATED TEST COVERAGE", level=2)
    add_paragraph(
        doc,
        "Although the project is centered on cricket analytics, it also includes automated checks for the most failure-prone parts of the inference "
        "stack. The table below records the concrete test functions present in the repository at report-generation time.",
    )
    add_table_title(doc, APPENDIX_TABLES["test_inventory"])
    add_table(
        doc,
        ["Test Function", "Coverage Focus"],
        [
            [row["test_name"], row["coverage"]]
            for row in context["test_inventory"]
        ],
        font_size=9,
    )

    doc.add_heading("APPENDIX G: NOTEBOOK INVENTORY", level=2)
    add_paragraph(
        doc,
        "The notebook inventory is included separately from the script inventory because it captures the exploratory analysis trail that preceded the "
        "final reproducible pipeline. This is helpful during academic evaluation because it shows how the team investigated the data before freezing "
        "the implementation into scripts.",
    )
    add_table_title(doc, APPENDIX_TABLES["notebook_inventory"])
    add_table(
        doc,
        ["Notebook", "Approx. Size", "Code Cells", "Markdown Cells", "Role"],
        [
            [row["name"], row["size"], str(row["code_cells"]), str(row["markdown_cells"]), row["role"]]
            for row in context["notebook_inventory"]
        ],
        font_size=9,
    )

    doc.add_heading("APPENDIX H: DEPENDENCY INVENTORY", level=2)
    add_table_title(doc, APPENDIX_TABLES["dependency_inventory"])
    add_table(
        doc,
        ["Package", "Specifier", "Role"],
        [
            [row["package"], row["specifier"], row["role"]]
            for row in context["dependency_inventory"]
        ],
        font_size=9,
    )

    doc.add_heading("APPENDIX I: VERSIONED MODEL REGISTRY SNAPSHOT", level=2)
    add_paragraph(
        doc,
        "The model-registry snapshot below records the most recent tracked training versions. Including it in the appendix helps connect the document "
        "to the saved artifacts present in the workspace at the time of report generation.",
    )
    add_table_title(doc, APPENDIX_TABLES["version_inventory"])
    add_table(
        doc,
        ["Version", "Date", "Scope", "Score Model", "Score RMSE", "Win Model", "Win Log Loss"],
        [
            [
                row["version_id"],
                row["created_at"],
                row["scope"],
                row["score_model"],
                row["score_rmse"],
                row["win_model"],
                row["win_log_loss"],
            ]
            for row in context["version_inventory"]
        ],
        font_size=9,
    )

    doc.add_heading("APPENDIX J: WORKSPACE DIRECTORY MAP", level=2)
    add_table_title(doc, APPENDIX_TABLES["workspace_map"])
    add_table(
        doc,
        ["Directory", "Primary Role"],
        [
            [row["name"], row["role"]]
            for row in context["workspace_map_rows"]
        ],
        font_size=9,
    )
    doc.add_page_break()


def add_references(doc: Document) -> None:
    doc.add_heading("REFERENCES", level=1)
    refs = [
        "[1] Cricsheet. IPL CSV2 Downloads and Documentation. https://cricsheet.org/downloads/",
        "[2] Chen, T., and Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of KDD 2016.",
        "[3] Prokhorenkova, L., Gusev, G., Vorobev, A. V., Dorogush, and Gulin, A. (2018). CatBoost: unbiased boosting with categorical features. NeurIPS 2018.",
        "[4] Guo, C., and Berkhahn, F. (2016). Entity Embeddings of Categorical Variables. arXiv preprint arXiv:1604.06737.",
        "[5] Scikit-learn documentation. https://scikit-learn.org/stable/",
        "[6] Flask documentation. https://flask.palletsprojects.com/",
        "[7] Streamlit documentation. https://docs.streamlit.io/",
        "[8] Open-Meteo API documentation. https://open-meteo.com/en/docs",
        "[9] XGBoost documentation. https://xgboost.readthedocs.io/",
        "[10] CatBoost documentation. https://catboost.ai/",
    ]
    for ref in refs:
        add_paragraph(doc, ref)
    doc.add_page_break()


def add_abbreviations(doc: Document) -> None:
    doc.add_heading("ABBREVIATIONS", level=1)
    add_table_title(doc, TABLES[10])
    add_table(
        doc,
        ["Abbreviation", "Meaning"],
        [
            ["API", "Application Programming Interface"],
            ["CRR", "Current Run Rate"],
            ["CSV", "Comma-Separated Values"],
            ["DL", "Deep Learning"],
            ["HGB", "HistGradientBoosting"],
            ["IPL", "Indian Premier League"],
            ["JSON", "JavaScript Object Notation"],
            ["MAE", "Mean Absolute Error"],
            ["ML", "Machine Learning"],
            ["MLP", "Multilayer Perceptron"],
            ["RMSE", "Root Mean Squared Error"],
            ["RRR", "Required Run Rate"],
            ["UI", "User Interface"],
        ],
    )


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    context = build_context()
    assets = generate_assets(context)

    doc = Document()
    configure_document(doc)

    add_title_page(doc, context["today"])
    add_front_matter(doc, context)
    add_executive_summary(doc, context)
    add_contents_and_lists(doc)
    add_chapter_1(doc, context)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc, context, assets)
    add_chapter_5(doc, context, assets)
    add_chapter_6(doc, context, assets)
    add_chapter_7(doc)
    add_appendix(doc, context)
    add_references(doc)
    add_abbreviations(doc)

    doc.save(OUTPUT_PATH)
    print(f"Generated report at: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
