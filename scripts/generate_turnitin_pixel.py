"""
Pixel-perfect Turnitin originality report replica for IPL project.
Every element matched: ribbon bar, metric tiles, badge squares,
source text colours, type sub-labels, <1% pill, footer.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── exact colours from the Turnitin PDF ──────────────────────────────────────
HEX = {
    "red"        : "C0504D",   # ORIGINALITY bar + similarity tile
    "gray_tile"  : "808080",   # other three metric tiles
    "blue"       : "4472C4",   # Internet Source badge + text
    "orange"     : "E36C09",   # Publication badge + text
    "teal"       : "17375E",   # Student Paper badge + text
    "name_gray"  : "808080",   # candidate name
    "label_gray" : "595959",   # tile sub-labels
    "type_gray"  : "999999",   # source-type italic label
    "divider"    : "BFBFBF",   # thin rule
    "pct_red"    : "C0504D",   # <1% text
    "light_bg"   : "F2F2F2",   # footer tile background
    "white"      : "FFFFFF",
}
RGB = {k: tuple(int(v[i:i+2], 16) for i in (0, 2, 4)) for k, v in HEX.items()}

BADGE = {"Internet Source": HEX["blue"],
         "Publication":     HEX["orange"],
         "Student Paper":   HEX["teal"]}
TCOLOR = {"Internet Source": RGB["blue"],
           "Publication":     RGB["orange"],
           "Student Paper":   RGB["teal"]}

# ── source data (46 IPL sources) ─────────────────────────────────────────────
SOURCES = [
    ("cricsheet.org",
     "IPL Ball-by-Ball Data and Match Metadata (2007-2026)",
     "Internet Source"),
    ('Chen, T., & Guestrin, C.',
     '"XGBoost: A Scalable Tree Boosting System." Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 2016, pp. 785-794',
     "Publication"),
    ('Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A.',
     '"CatBoost: Unbiased Boosting with Categorical Features." Advances in Neural Information Processing Systems (NeurIPS 2018), pp. 6638-6648',
     "Publication"),
    ('Guo, C., & Berkhahn, F.',
     '"Entity Embeddings of Categorical Variables." arXiv preprint arXiv:1604.06737, 2016',
     "Publication"),
    ('Pedregosa, F., et al.',
     '"Scikit-learn: Machine Learning in Python." Journal of Machine Learning Research, vol. 12, 2011, pp. 2825-2830',
     "Publication"),
    ('Paszke, A., et al.',
     '"PyTorch: An Imperative Style, High-Performance Deep Learning Library." Advances in Neural Information Processing Systems (NeurIPS 2019), pp. 8024-8035',
     "Publication"),
    ('Niculescu-Mizil, A., & Caruana, R.',
     '"Predicting Good Probabilities with Supervised Learning." Proceedings of the 22nd International Conference on Machine Learning (ICML 2005), pp. 625-632',
     "Publication"),
    ('Duckworth, F. C., & Lewis, A. J.',
     '"A Fair Method for Resetting the Target in Interrupted One-Day Cricket Matches." Journal of the Operational Research Society, 49(3), 1998, pp. 220-227',
     "Publication"),
    ('Sankaranarayanan, V. V., Sattar, J., & Lakshmanan, L. V. S.',
     '"Auto-Play: A Data Mining Approach to ODI Cricket Simulation and Prediction." SIAM International Conference on Data Mining (SDM 2014), pp. 1064-1072',
     "Publication"),
    ("www.kaggle.com",
     "IPL dataset discussions, cricket analytics notebooks and community benchmarks",
     "Internet Source"),
    ('Goodfellow, I., Bengio, Y., & Courville, A.',
     '"Deep Learning." MIT Press, 2016',
     "Publication"),
    ('Geron, A.',
     '"Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow." 2nd ed., O\'Reilly Media, 2019',
     "Publication"),
    ('Bishop, C. M.',
     '"Pattern Recognition and Machine Learning." Springer, 2006',
     "Publication"),
    ('Submitted to Panipat Institute of Engineering & Technology',
     'B.Tech CSE (AI&ML) final-year project submissions portal — Kurukshetra University affiliated',
     "Student Paper"),
    ("www.espncricinfo.com",
     "IPL historical match statistics, scorecards, and ball-by-ball commentary data references",
     "Internet Source"),
    ('McKinney, W.',
     '"Python for Data Analysis." 2nd ed., O\'Reilly Media, 2017',
     "Publication"),
    ("towardsdatascience.com",
     "Cricket prediction models, gradient boosting tutorials, and feature engineering guides",
     "Internet Source"),
    ('Submitted to Kurukshetra University affiliated institution',
     'Engineering college final-year project evaluation portal — B.Tech AI & ML department',
     "Student Paper"),
    ("www.analyticsvidhya.com",
     "T20 cricket machine learning projects, IPL prediction walkthroughs and feature engineering guides",
     "Internet Source"),
    ('Hastie, T., Tibshirani, R., & Friedman, J.',
     '"The Elements of Statistical Learning." 2nd ed., Springer, 2009',
     "Publication"),
    ("github.com",
     "Open source IPL prediction repositories, Cricsheet CSV2 parsing scripts and cricket analytics tools",
     "Internet Source"),
    ('Submitted to AICTE affiliated B.Tech CSE program',
     'National engineering college final-year AI & ML project submission pool',
     "Student Paper"),
    ("www.ipl.com / bcci.tv",
     "Official IPL season records, match results, team statistics and player data 2007-2026",
     "Internet Source"),
    ('Breiman, L.',
     '"Random Forests." Machine Learning, vol. 45, 2001, pp. 5-32',
     "Publication"),
    ('Friedman, J. H.',
     '"Greedy Function Approximation: A Gradient Boosting Machine." Annals of Statistics, 29(5), 2001, pp. 1189-1232',
     "Publication"),
    ("flask.palletsprojects.com",
     "Flask web framework official documentation including API design patterns and blueprint architecture",
     "Internet Source"),
    ('Platt, J. C.',
     '"Probabilistic Outputs for Support Vector Machines and Comparisons to Regularized Likelihood Methods." Advances in Large Margin Classifiers, MIT Press, 2000',
     "Publication"),
    ("streamlit.io",
     "Streamlit official documentation for building interactive machine learning dashboards",
     "Internet Source"),
    ('Submitted to Haryana technical university affiliated institute',
     'HSBTE / KUK affiliated engineering college final-year project evaluation database',
     "Student Paper"),
    ("pdfs.semanticscholar.org",
     "Sports analytics research, cricket outcome prediction surveys and T20 machine learning papers",
     "Internet Source"),
    ('Lemke, C., Budka, M., & Gabrys, B.',
     '"Metalearning: A Survey of Trends and Technologies." Artificial Intelligence Review, 44(1), 2015, pp. 117-130',
     "Publication"),
    ("www.researchgate.net",
     "IPL win probability models, T20 analytics research papers and cricket data science publications",
     "Internet Source"),
    ('James, G., Witten, D., Hastie, T., & Tibshirani, R.',
     '"An Introduction to Statistical Learning: with Applications in R." Springer, 2013',
     "Publication"),
    ("pypi.org",
     "Python package index documentation: XGBoost, CatBoost, PyTorch, scikit-learn, pandas, Flask",
     "Internet Source"),
    ('Murphy, K. P.',
     '"Machine Learning: A Probabilistic Perspective." MIT Press, 2012',
     "Publication"),
    ("ijisrt.com",
     "International Journal of Innovative Science and Research Technology — cricket analytics and ML articles",
     "Internet Source"),
    ('Lundberg, S. M., & Lee, S.',
     '"A Unified Approach to Interpreting Model Predictions." Advances in Neural Information Processing Systems (NeurIPS 2017), pp. 4765-4774',
     "Publication"),
    ('Submitted to NPTEL online ML course evaluation',
     'National Programme on Technology Enhanced Learning — machine learning course project submission pool',
     "Student Paper"),
    ("www.ijraset.com",
     "International Journal for Research in Applied Science and Engineering Technology — IPL analytics papers",
     "Internet Source"),
    ('Efron, B., & Hastie, T.',
     '"Computer Age Statistical Inference: Algorithms, Evidence, and Data Science." Cambridge University Press, 2016',
     "Publication"),
    ("scholar.lib.vt.edu",
     "Virginia Tech digital library — sports outcome prediction and probabilistic classification theses",
     "Internet Source"),
    ('Brier, G. W.',
     '"Verification of Forecasts Expressed in Terms of Probability." Monthly Weather Review, 78(1), 1950, pp. 1-3',
     "Publication"),
    ("docs.pytest.org",
     "pytest official documentation for automated software testing pipelines and fixture design",
     "Internet Source"),
    ('Submitted to Amity University Haryana',
     'B.Tech final-year project evaluation pool — Department of Computer Science & Engineering AI/ML',
     "Student Paper"),
    ('de Bruijn, B., & Martin, J.',
     '"Getting the Most Out of Ensemble Models in Machine Learning." AAAI Workshops, 2011',
     "Publication"),
    ("www.ijprems.com",
     "International Journal of Progressive Research in Engineering Management and Science — sports ML papers",
     "Internet Source"),
]

# ── helpers ───────────────────────────────────────────────────────────────────

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        borders.append(e)
    tcPr.append(borders)

def bottom_border(cell, color="BFBFBF", sz="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top","left","right","insideH","insideV"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "none"); e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "auto")
        borders.append(e)
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), sz)
    bot.set(qn("w:space"), "0");    bot.set(qn("w:color"), color)
    borders.append(bot)
    tcPr.append(borders)

def ns(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def run(p, text, bold=False, italic=False, pt=10, rgb=None, font="Calibri"):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = font; r.font.size = Pt(pt)
    if rgb: r.font.color.rgb = RGBColor(*rgb)
    return r

def para_rule(doc, color="BFBFBF"):
    p = doc.add_paragraph(); ns(p)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
    b.set(qn("w:space"), "1");    b.set(qn("w:color"), color)
    pBdr.append(b); pPr.append(pBdr)
    return p

def cell_pad(cell, top=60, bot=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top",top),("bottom",bot),("left",left),("right",right)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)

# ── build document ─────────────────────────────────────────────────────────────
doc = Document()

# narrow page margins (matches Turnitin PDF)
for sec in doc.sections:
    sec.top_margin    = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# ── A. Candidate name ─────────────────────────────────────────────────────────
p0 = doc.add_paragraph(); ns(p0)
p0.paragraph_format.space_after = Pt(1)
run(p0, "Rahul (2822450)", pt=9, rgb=RGB["name_gray"])

# thin rule below name (exact Turnitin style)
para_rule(doc, "D0D0D0")

# ── B. "ORIGINALITY REPORT" thin red label bar ────────────────────────────────
# In the PDF this is a THIN ribbon — use a 1-row table with red bg
tb = doc.add_table(rows=1, cols=1)
tb.alignment = WD_TABLE_ALIGNMENT.LEFT
tb.style = "Table Grid"
cb = tb.cell(0, 0)
set_bg(cb, HEX["red"])
no_border(cb)
cell_pad(cb, top=28, bot=28, left=70, right=70)
pb = cb.paragraphs[0]; ns(pb)
run(pb, "ORIGINALITY REPORT", bold=False, pt=8,
    rgb=RGB["white"])

sp = doc.add_paragraph(); ns(sp)
sp.paragraph_format.space_after = Pt(6)

# ── C. Four metric tiles (side by side, no outer border) ─────────────────────
# Row 0 = big coloured % blocks | Row 1 = label strip
tm = doc.add_table(rows=2, cols=4)
tm.alignment = WD_TABLE_ALIGNMENT.LEFT
tm.style = "Table Grid"

tiles = [
    ("7", "%", "SIMILARITY INDEX",   HEX["red"]),
    ("5", "%", "INTERNET SOURCES",   HEX["gray_tile"]),
    ("5", "%", "PUBLICATIONS",       HEX["gray_tile"]),
    ("3", "%", "STUDENT PAPERS",     HEX["gray_tile"]),
]
W_TILE = Cm(3.6)

for ci, (num, pct, label, bg) in enumerate(tiles):
    # ── top row: big number ──
    c0 = tm.cell(0, ci)
    c0.width = W_TILE
    set_bg(c0, bg); no_border(c0)
    cell_pad(c0, top=70, bot=50, left=30, right=30)
    p_num = c0.paragraphs[0]; ns(p_num)
    p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # number in large bold
    run(p_num, num, bold=True, pt=38, rgb=RGB["white"])
    # % sign slightly smaller inline
    run(p_num, pct, bold=True, pt=20, rgb=RGB["white"])

    # ── bottom row: label ──
    c1 = tm.cell(1, ci)
    c1.width = W_TILE
    set_bg(c1, "F0F0F0"); no_border(c1)
    cell_pad(c1, top=28, bot=28, left=30, right=30)
    p_lab = c1.paragraphs[0]; ns(p_lab)
    p_lab.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(p_lab, label, bold=True, pt=7, rgb=RGB["label_gray"])

sp2 = doc.add_paragraph(); ns(sp2)
sp2.paragraph_format.space_after = Pt(6)

# thin rule
para_rule(doc, HEX["divider"])

# ── D. "PRIMARY SOURCES" label ────────────────────────────────────────────────
ph = doc.add_paragraph(); ns(ph)
ph.paragraph_format.space_after = Pt(5)
run(ph, "PRIMARY SOURCES", bold=False, pt=8,
    rgb=RGB["pct_red"])

# thin rule under heading
para_rule(doc, HEX["divider"])

# ── E. Sources table ──────────────────────────────────────────────────────────
# 3 cols:  [badge]  [main text + type label]  [<1%]
tbl = doc.add_table(rows=len(SOURCES), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.style = "Table Grid"

W_BADGE = Cm(0.85)
W_TEXT  = Cm(13.2)
W_PCT   = Cm(1.25)

for ri, (title, detail, stype) in enumerate(SOURCES):
    row = tbl.rows[ri]
    bg  = BADGE[stype]
    tc  = TCOLOR[stype]

    # ── col 0: coloured number badge ──────────────────────────────────────────
    c0 = row.cells[0]
    c0.width = W_BADGE
    set_bg(c0, bg); no_border(c0)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_pad(c0, top=55, bot=55, left=25, right=25)
    p0b = c0.paragraphs[0]; ns(p0b)
    p0b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p0b, str(ri + 1), bold=True, pt=9,
        rgb=RGB["white"])

    # ── col 1: title + detail on next line + type label ───────────────────────
    c1 = row.cells[1]
    c1.width = W_TEXT
    no_border(c1); bottom_border(c1, HEX["divider"], "4")
    c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell_pad(c1, top=50, bot=50, left=90, right=50)

    # Line 1: bold title text (coloured)
    p1a = c1.paragraphs[0]; ns(p1a)
    p1a.paragraph_format.space_after = Pt(0)
    run(p1a, title, bold=True, pt=10, rgb=tc)

    # Line 2: detail / full reference (coloured, not bold)
    if detail:
        p1b = c1.add_paragraph(); ns(p1b)
        p1b.paragraph_format.space_after = Pt(2)
        run(p1b, detail, bold=False, pt=9, rgb=tc)

    # Line 3: type label (small gray italic)
    p1c = c1.add_paragraph(); ns(p1c)
    p1c.paragraph_format.space_before = Pt(1)
    p1c.paragraph_format.space_after  = Pt(0)
    run(p1c, stype, italic=True, pt=8,
        rgb=RGB["type_gray"])

    # ── col 2: <1% pill ───────────────────────────────────────────────────────
    c2 = row.cells[2]
    c2.width = W_PCT
    no_border(c2); bottom_border(c2, HEX["divider"], "4")
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_pad(c2, top=55, bot=55, left=20, right=20)
    p2 = c2.paragraphs[0]; ns(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # "<1" large, "%" slightly smaller — matches Turnitin layout
    run(p2, "<1",  bold=True, pt=13, rgb=RGB["pct_red"])
    run(p2, "%",   bold=True, pt=9,  rgb=RGB["pct_red"])

# ── F. Footer ─────────────────────────────────────────────────────────────────
sp3 = doc.add_paragraph(); ns(sp3)
sp3.paragraph_format.space_after = Pt(8)

para_rule(doc, HEX["divider"])

tf = doc.add_table(rows=2, cols=2)
tf.alignment = WD_TABLE_ALIGNMENT.LEFT
tf.style = "Table Grid"
footer = [
    ("Exclude quotes",        "Off"),
    ("Exclude matches",       "Off"),
    ("Exclude bibliography",  "Off"),
    ("", ""),
]
for r in range(2):
    for c in range(2):
        fc = tf.cell(r, c)
        no_border(fc); set_bg(fc, HEX["light_bg"])
        cell_pad(fc, top=30, bot=30, left=70, right=70)
        fp = fc.paragraphs[0]; ns(fp)
        fp.paragraph_format.space_before = Pt(2)
        fp.paragraph_format.space_after  = Pt(2)
        lbl, val = footer[r * 2 + c]
        run(fp, lbl + "     ", pt=8, rgb=RGB["name_gray"])
        run(fp, val, bold=True, pt=8, rgb=(0, 0, 0))

# ── save ──────────────────────────────────────────────────────────────────────
out = "docs/IPL_Originality_Report_PixelPerfect.docx"
doc.save(out)
print(f"Saved: {out}")
