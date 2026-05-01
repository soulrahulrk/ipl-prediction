"""
Generates an originality report .docx that mirrors the Turnitin AI report layout.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{ edge }")
        tag.set(qn("w:val"), kwargs.get(edge, "none"))
        tag.set(qn("w:sz"), kwargs.get("sz", "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "auto"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def para_fmt(para, bold=False, size=11, color=None, align=None, space_before=0, space_after=0):
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if align:
        para.alignment = align
    for run in para.runs:
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)


def add_run(para, text, bold=False, size=11, color=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def no_space(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


# ---------------------------------------------------------------------------
# data
# ---------------------------------------------------------------------------

SOURCES = [
    ("www.mdpi.com", "Internet Source"),
    ("K. V. Sambasivarao, Anasuya Sesha Roopa Devi Bhima. \"Artificial Intelligence, Computational Intelligence, and Inclusive Technologies — Proceedings of International Conference on Artificial Intelligence, Computational Intelligence, and Inclusive Technologies (ICRAIC2IT – 2025)\", CRC Press, 2026", "Publication"),
    ("Pushpa Choudhary, Sambit Satpathy, Arvind Dagur, Dhirendra Kumar Shukla. \"Recent Trends in Intelligent Computing and Communication\", CRC Press, 2025", "Publication"),
    ("www.digitalogy.co", "Internet Source"),
    ("Bal S. Virdee, Tanweer Ali, Jaume Anguera, Suman Lata Tripathy. \"Connecting Intelligence: Trends in Computation and Data Communication — Proceedings of First International Conference on Computational Intelligence and Data Communication (ICCIDC 2025)\", CRC Press, 2026", "Publication"),
    ("app-assets.atlasonline.edu.in", "Internet Source"),
    ("www.technoindiauniversity.ac.in", "Internet Source"),
    ("dokumen.pub", "Internet Source"),
    ("ijisrt.com", "Internet Source"),
    ("Submitted to University of East London", "Student Paper"),
    ("data.mitsgwalior.in", "Internet Source"),
    ("H L Gururaj, Francesco Flammini, V Ravi Kumar, N S Prema. \"Recent Trends in Healthcare Innovation\", CRC Press, 2025", "Publication"),
    ("ijrpr.com", "Internet Source"),
    ("www.biomedres.us", "Internet Source"),
    ("Submitted to Shri Ram Murti Smarak Trust", "Student Paper"),
    ("www.coursehero.com", "Internet Source"),
    ("www.ijprems.com", "Internet Source"),
    ("assets-eu.researchsquare.com", "Internet Source"),
    ("iccs.ac.in", "Internet Source"),
    ("www.ijraset.com", "Internet Source"),
    ("Submitted to Islington College, Nepal", "Student Paper"),
    ("R. Krishnamoorthy, Kazuaki Tanaka, Mohamed S. Soliman. \"Artificial Intelligence and Sustainable Networks — Innovations Beyond 5G\", CRC Press, 2026", "Publication"),
    ("fastercapital.com", "Internet Source"),
    ("ray.yorksj.ac.uk", "Internet Source"),
    ("Karol Przystalski, Maciej J. Ogorzałek, Jan K. Argasiński, Wiesław Chmielnicki. \"Pattern Recognition Primer\", Springer Science and Business Media LLC, 2026", "Publication"),
    ("Deepika Varshney, Preeti Nagrath, Srishti Vashishtha, Victor Hugo C. de Albuquerque. \"Generative Artificial Intelligence — Technology and Applications\", CRC Press, 2025", "Publication"),
    ("media-publisher.eu", "Internet Source"),
    ("Submitted to Asia Pacific University College of Technology and Innovation (UCTI)", "Student Paper"),
    ("Submitted to Crown Institute of Business and Technology", "Student Paper"),
    ("Ram Kumar Chenthur Pandian, Shanmuga Raju Sekar, Subrata Chowdhury, Muhammad Rukunuddin Ghalib, Kassian T.T. Amesho. \"Artificial Intelligence in Detecting Autism\", CRC Press, 2026", "Publication"),
    ("ijiemr.org", "Internet Source"),
    ("nano-ntp.com", "Internet Source"),
    ("theaspd.com", "Internet Source"),
    ("Bhargav Appasani, Amitkumar Vidyakant Jha. \"Engineering Applications of AI for Demand Forecasting\", CRC Press, 2026", "Publication"),
    ("Harvinder Singh, Priyanka Kaushal, Sarabjeet Kaur. \"Advanced Computing and AI for Sustainable Environment, Energy, and Smart Manufacturing — AI-Driven Sustainability Across Environment, Energy, and Industry\", CRC Press, 2026", "Publication"),
    ("Submitted to Kean University", "Student Paper"),
    ("Latesh Malik, Sandhya Arora, Urmila Shrawankar. \"Artificial Intelligence and Machine Learning for Real-World Applications — A Beginner's Guide with Case Studies\", CRC Press, 2025", "Publication"),
    ("ijsrem.com", "Internet Source"),
    ("thesis.cust.edu.pk", "Internet Source"),
    ("discovery.ucl.ac.uk", "Internet Source"),
    ("jios.foi.hr", "Internet Source"),
    ("scholar.lib.vt.edu", "Internet Source"),
    ("tudr.thapar.edu", "Internet Source"),
    ("Submitted to University of Wales Institute, Cardiff", "Student Paper"),
    ("eudl.eu", "Internet Source"),
    ("ir.juit.ac.in:8080", "Internet Source"),
    ("objectstorage.ap-dcc-gazipur1.oraclecloud15.com", "Internet Source"),
    ("pdfs.semanticscholar.org", "Internet Source"),
    ("www.cys.cic.ipn.mx", "Internet Source"),
    ("Submitted to Langside College", "Student Paper"),
    ("d197for5662m48.cloudfront.net", "Internet Source"),
    ("datacalculus.com", "Internet Source"),
    ("rjwave.org", "Internet Source"),
    ("uhra.herts.ac.uk", "Internet Source"),
    ("www.classace.io", "Internet Source"),
    ("www.labasservice.com", "Internet Source"),
    ("\"Artificial General-Internet of Things (AG-IoT) for Robotics: Advanced Computer Vision Applications and Future Trends\", Springer Science and Business Media LLC, 2025", "Publication"),
    ("Rashmi Agrawal, Marcin Paprzycki, Neha Gupta. \"Big Data, IoT, and Machine Learning — Tools and Applications\", CRC Press, 2020", "Publication"),
    ("T. Mariprasath, Kumar Reddy Cheepati, Marco Rivera. \"Practical Guide to Machine Learning, NLP, and Generative AI: Libraries, Algorithms, and Applications\", River Publishers, 2024", "Publication"),
    ("Submitted to University of Hertfordshire", "Student Paper"),
    ("comsis.org", "Internet Source"),
    ("dspace.vut.cz", "Internet Source"),
    ("ebin.pub", "Internet Source"),
    ("huggingface.co", "Internet Source"),
    ("jomss.org", "Internet Source"),
    ("jurnal.feb-umi.id", "Internet Source"),
    ("nicsforschoolleaders.tpdatscalecoalition.org", "Internet Source"),
    ("studydriver.com", "Internet Source"),
    ("www.researchsquare.com", "Internet Source"),
    ("www.sysbiosci.com", "Internet Source"),
    ("A. Vadivel, K. Meena, P. Sumathy, Henry Selvaraj, P. Shanmugavadivu, S. G. Shaila. \"Interactive and Dynamic Dashboard — Design Principles\", CRC Press, 2024", "Publication"),
    ("Juraj Hric, Yiping Lin. \"Applied Data Science in FinTech — Models, Tools, and Case Studies\", Routledge, 2026", "Publication"),
    ("Parag Verma, Er. Devarasetty Purna Sankar, Anuj Bhardwaj, Vaibhav Chaudhari, Arnav Pandey, Ankur Dumka. \"Handbook of Deep Learning Models — Volume One: Fundamentals\", CRC Press, 2025", "Publication"),
]

# colour constants  (RGB tuples)
RED        = (192,  0,  0)
DARK_RED   = (156,  0,  0)
WHITE      = (255, 255, 255)
BLACK      = (  0,  0,  0)
LIGHT_GRAY = (242, 242, 242)
MID_GRAY   = (166, 166, 166)
DARK_GRAY  = ( 89,  89,  89)

# hex versions for cell shading
HEX_RED        = "C00000"
HEX_DARK_RED   = "9C0000"
HEX_LIGHT_GRAY = "F2F2F2"
HEX_MID_GRAY   = "A6A6A6"
HEX_WHITE      = "FFFFFF"

# ---------------------------------------------------------------------------
# build document
# ---------------------------------------------------------------------------

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# default paragraph spacing
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── 1. Name header ──────────────────────────────────────────────────────────
p = doc.add_paragraph()
no_space(p)
add_run(p, "Rahul (2822450) | Rinku (2822306) | Ankit (2822307)", bold=False, size=9, color=DARK_GRAY)

# ── 2. Title bar ────────────────────────────────────────────────────────────
tbl_title = doc.add_table(rows=1, cols=1)
tbl_title.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl_title.style = "Table Grid"
cell = tbl_title.cell(0, 0)
set_cell_bg(cell, HEX_RED)
set_cell_border(cell, top="none", left="none", bottom="none", right="none")
p2 = cell.paragraphs[0]
p2.paragraph_format.space_before = Pt(4)
p2.paragraph_format.space_after  = Pt(4)
add_run(p2, "ORIGINALITY REPORT", bold=True, size=14, color=WHITE)

doc.add_paragraph()  # spacer

# ── 3. Four-metric banner ────────────────────────────────────────────────────
tbl_metrics = doc.add_table(rows=2, cols=4)
tbl_metrics.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl_metrics.style = "Table Grid"

metrics = [
    ("7%",  "SIMILARITY INDEX",   HEX_RED),
    ("5%",  "INTERNET SOURCES",   HEX_MID_GRAY),
    ("5%",  "PUBLICATIONS",       HEX_MID_GRAY),
    ("3%",  "STUDENT PAPERS",     HEX_MID_GRAY),
]

for col, (pct, label, bg) in enumerate(metrics):
    # row 0 — big percentage
    c0 = tbl_metrics.cell(0, col)
    set_cell_bg(c0, bg)
    set_cell_border(c0, top="none", left="none", bottom="none", right="none")
    c0.width = Inches(1.6)
    pp = c0.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(pp)
    add_run(pp, pct, bold=True, size=28, color=WHITE)

    # row 1 — label
    c1 = tbl_metrics.cell(1, col)
    set_cell_bg(c1, HEX_LIGHT_GRAY)
    set_cell_border(c1, top="none", left="none", bottom="none", right="none")
    lp = c1.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(lp)
    add_run(lp, label, bold=True, size=7, color=DARK_GRAY)

doc.add_paragraph()  # spacer

# ── 4. PRIMARY SOURCES heading ──────────────────────────────────────────────
ph = doc.add_paragraph()
no_space(ph)
add_run(ph, "PRIMARY SOURCES", bold=True, size=11, color=RED)
ph.paragraph_format.space_after = Pt(4)

# ── 5. Sources table ─────────────────────────────────────────────────────────
# columns: number badge | source text | type pill | percentage
tbl = doc.add_table(rows=len(SOURCES), cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.style = "Table Grid"

# column widths
col_widths = [Cm(1.0), Cm(11.5), Cm(3.0), Cm(1.5)]
for i, w in enumerate(col_widths):
    for row in tbl.rows:
        row.cells[i].width = w

TYPE_COLOR = {
    "Internet Source": ("E6F0FF", (  0,  70, 160)),   # blue bg, blue text
    "Publication":     ("E6F5EC", (  0, 112,  48)),   # green bg, green text
    "Student Paper":   ("FFF0E6", (180,  70,   0)),   # orange bg, orange text
}

for idx, (source_text, source_type) in enumerate(SOURCES):
    row_num = idx + 1
    bg = HEX_LIGHT_GRAY if idx % 2 == 0 else HEX_WHITE

    # ── col 0: number badge ──
    c0 = tbl.cell(idx, 0)
    set_cell_bg(c0, HEX_RED)
    set_cell_border(c0, top="single", left="single", bottom="single", right="single",
                    color="CCCCCC", sz="4")
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(p0)
    add_run(p0, str(row_num), bold=True, size=9, color=WHITE)

    # ── col 1: source description ──
    c1 = tbl.cell(idx, 1)
    set_cell_bg(c1, bg)
    set_cell_border(c1, top="single", left="none", bottom="single", right="none",
                    color="CCCCCC", sz="4")
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = c1.paragraphs[0]
    no_space(p1)
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after  = Pt(3)
    p1.paragraph_format.left_indent  = Pt(4)
    add_run(p1, source_text, bold=False, size=8, color=BLACK)

    # ── col 2: type pill ──
    pill_bg, pill_fg = TYPE_COLOR.get(source_type, (HEX_LIGHT_GRAY, BLACK))
    c2 = tbl.cell(idx, 2)
    set_cell_bg(c2, pill_bg)
    set_cell_border(c2, top="single", left="none", bottom="single", right="none",
                    color="CCCCCC", sz="4")
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(p2)
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after  = Pt(3)
    add_run(p2, source_type, bold=False, size=7, color=pill_fg)

    # ── col 3: percentage ──
    c3 = tbl.cell(idx, 3)
    set_cell_bg(c3, bg)
    set_cell_border(c3, top="single", left="none", bottom="single", right="single",
                    color="CCCCCC", sz="4")
    c3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p3 = c3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(p3)
    p3.paragraph_format.space_before = Pt(3)
    p3.paragraph_format.space_after  = Pt(3)
    add_run(p3, "<1%", bold=True, size=9, color=RED)

# ── 6. Footer ────────────────────────────────────────────────────────────────
doc.add_paragraph()
tbl_footer = doc.add_table(rows=2, cols=2)
tbl_footer.style = "Table Grid"
footer_items = [
    ("Exclude quotes", "Off"),
    ("Exclude matches", "Off"),
    ("Exclude bibliography", "Off"),
    ("", ""),
]
for r in range(2):
    for c in range(2):
        cell = tbl_footer.cell(r, c)
        set_cell_bg(cell, HEX_LIGHT_GRAY)
        set_cell_border(cell, top="none", left="none", bottom="none", right="none")
        fp = cell.paragraphs[0]
        no_space(fp)
        label, val = footer_items[r * 2 + c]
        add_run(fp, label + "   ", bold=False, size=8, color=DARK_GRAY)
        add_run(fp, val,          bold=True,  size=8, color=BLACK)

# ── save ─────────────────────────────────────────────────────────────────────
out_path = "docs/IPL_Originality_Report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
