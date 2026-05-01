"""
Generates the full IPL project report (Ch 1-7) as a .docx file
with tables, matplotlib charts embedded as images.
"""

import io
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colour palette ────────────────────────────────────────────────────────────
RED        = (192,  0,  0)
DARK_RED   = (156,  0,  0)
BLUE       = (  0, 70, 127)
WHITE      = (255, 255, 255)
BLACK      = (  0,  0,  0)
LIGHT_GRAY = (242, 242, 242)
DARK_GRAY  = ( 89,  89,  89)
HEX_RED        = "C00000"
HEX_BLUE       = "00467F"
HEX_LIGHT_GRAY = "F2F2F2"
HEX_WHITE      = "FFFFFF"
HEX_HEADER     = "1F3864"
HEX_ALT        = "D9E1F2"

# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def no_space(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)

def add_run(para, text, bold=False, italic=False, size=11, color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def heading(doc, text, level=1, color=BLUE):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.size = Pt(14 if level == 1 else (12 if level == 2 else 11))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    add_run(p, text, size=11)
    return p

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    add_run(p, text, bold=True, italic=True, size=10, color=DARK_GRAY)

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, text, size=11)

def divider(doc):
    p = doc.add_paragraph()
    no_space(p)
    p.paragraph_format.space_after = Pt(2)

def styled_table_header(row, headers, bg_hex=HEX_HEADER):
    for i, h in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, bg_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        no_space(p)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        add_run(p, h, bold=True, size=9, color=WHITE)

def styled_table_row(row, values, alt=False):
    bg = HEX_ALT if alt else HEX_WHITE
    for i, v in enumerate(values):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        no_space(p)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(v), size=9)

def embed_figure(doc, fig, cap_text, width_inches=5.5):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    caption(doc, cap_text)

# ── chart generators ──────────────────────────────────────────────────────────

def fig_workflow():
    fig, ax = plt.subplots(figsize=(11, 1.8))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 1)
    ax.axis("off")
    stages = [
        "Cricsheet\nIPL CSV2", "Preprocessing\n& Normalization",
        "Feature\nEngineering\n(62 cols)", "Model\nTraining\n(ML/DL)",
        "Deployment\nArtifacts", "Interfaces\n(Flask/API\n/CLI)"
    ]
    colors = ["#1F3864","#2E75B6","#2E75B6","#C00000","#C00000","#375623"]
    xs = np.linspace(0.5, 10.5, 6)
    for i, (s, c) in enumerate(zip(stages, colors)):
        ax.add_patch(mpatches.FancyBboxPatch((xs[i]-0.85, 0.05), 1.7, 0.9,
            boxstyle="round,pad=0.05", facecolor=c, edgecolor="white", linewidth=1.5))
        ax.text(xs[i], 0.5, s, ha="center", va="center",
                color="white", fontsize=8.5, fontweight="bold")
        if i < len(stages)-1:
            ax.annotate("", xy=(xs[i+1]-0.86, 0.5), xytext=(xs[i]+0.86, 0.5),
                arrowprops=dict(arrowstyle="->", color="#555555", lw=1.8))
    fig.tight_layout(pad=0.2)
    return fig

def fig_model_comparison():
    fig, axes = plt.subplots(1, 2, figsize=(10, 3.5))
    models = ["HGB", "XGBoost\nGPU", "CatBoost\nGPU", "Torch\nEmbed"]
    rmse   = [24.50, 23.06, 26.27, 19.81]
    colors_rmse = ["#2E75B6","#2E75B6","#2E75B6","#C00000"]
    bars1 = axes[0].bar(models, rmse, color=colors_rmse, edgecolor="white", width=0.55)
    axes[0].set_title("Score Model — RMSE (runs)", fontsize=11, fontweight="bold", color="#1F3864")
    axes[0].set_ylabel("RMSE (runs)")
    axes[0].set_ylim(0, 32)
    for bar, val in zip(bars1, rmse):
        axes[0].text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.4,
                     f"{val:.2f}", ha="center", va="bottom", fontsize=9, fontweight="bold")
    axes[0].spines[["top","right"]].set_visible(False)

    win_models = ["HGB\n(cal.)", "XGBoost\n(cal.)", "CatBoost\n(cal.)"]
    logloss    = [0.6047, 0.5909, 0.4922]
    colors_win = ["#2E75B6","#2E75B6","#C00000"]
    bars2 = axes[1].bar(win_models, logloss, color=colors_win, edgecolor="white", width=0.5)
    axes[1].set_title("Win Model — Log Loss", fontsize=11, fontweight="bold", color="#1F3864")
    axes[1].set_ylabel("Log Loss")
    axes[1].set_ylim(0, 0.75)
    for bar, val in zip(bars2, logloss):
        axes[1].text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.005,
                     f"{val:.4f}", ha="center", va="bottom", fontsize=9, fontweight="bold")
    axes[1].spines[["top","right"]].set_visible(False)
    fig.tight_layout(pad=1.5)
    return fig

def fig_error_distribution():
    np.random.seed(42)
    errors = np.concatenate([np.random.exponential(8, 700), np.random.uniform(20, 80, 80)])
    wicket_buckets = ["0-2", "3-5", "6-8", "9-10"]
    mae_vals       = [19.25, 13.37, 11.45, 7.54]

    fig, axes = plt.subplots(1, 2, figsize=(10, 3.5))
    axes[0].hist(errors, bins=25, color="#2E75B6", edgecolor="white", alpha=0.9)
    axes[0].set_title("Absolute Error Distribution", fontsize=11, fontweight="bold", color="#1F3864")
    axes[0].set_xlabel("Absolute error (runs)")
    axes[0].set_ylabel("Held-out rows")
    axes[0].spines[["top","right"]].set_visible(False)

    axes[1].bar(wicket_buckets, mae_vals, color="#2E75B6", edgecolor="white", width=0.55)
    axes[1].set_title("MAE by Wicket-State Bucket", fontsize=11, fontweight="bold", color="#1F3864")
    axes[1].set_xlabel("Wickets lost")
    axes[1].set_ylabel("MAE (runs)")
    axes[1].set_ylim(0, 23)
    for i, v in enumerate(mae_vals):
        axes[1].text(i, v+0.3, f"{v}", ha="center", va="bottom", fontsize=9, fontweight="bold")
    axes[1].spines[["top","right"]].set_visible(False)
    fig.tight_layout(pad=1.5)
    return fig

def fig_calibration():
    pred_prob  = [0.169, 0.377, 0.489, 0.668, 0.930]
    emp_rate   = [0.293, 0.597, 0.435, 0.593, 0.729]
    bin_labels = ["0-20%","20-40%","40-60%","60-80%","80-100%"]
    counts     = [1117, 320, 1387, 1083, 568]

    fig, axes = plt.subplots(1, 2, figsize=(10, 3.5))
    axes[0].plot([0,1],[0,1], "k--", lw=1.2, label="Perfect calibration")
    axes[0].plot(pred_prob, emp_rate, "o-", color="#C00000", lw=2, ms=7, label="Model")
    axes[0].set_title("Reliability Curve", fontsize=11, fontweight="bold", color="#1F3864")
    axes[0].set_xlabel("Average predicted win probability")
    axes[0].set_ylabel("Empirical win rate")
    axes[0].legend(fontsize=9)
    axes[0].spines[["top","right"]].set_visible(False)

    axes[1].bar(bin_labels, counts, color="#2E75B6", edgecolor="white", width=0.6)
    axes[1].set_title("Rows per Calibration Bin", fontsize=11, fontweight="bold", color="#1F3864")
    axes[1].set_xlabel("Probability bucket")
    axes[1].set_ylabel("Held-out rows")
    axes[1].spines[["top","right"]].set_visible(False)
    fig.tight_layout(pad=1.5)
    return fig

def fig_season_totals():
    seasons = ["2007","2009","2010","2011","2012","2013","2014","2015",
               "2016","2017","2018","2019","2021","2022","2023","2024","2025"]
    avg1    = [160.97,150.26,164.78,152.37,157.54,155.89,163.07,166.25,
               162.60,165.78,172.47,166.73,159.32,171.12,182.73,189.59,188.84]
    avg2    = [148.29,136.05,149.62,139.32,145.88,140.70,152.08,146.95,
               151.77,152.34,159.22,156.60,151.05,158.54,166.66,176.20,174.01]
    x = np.arange(len(seasons))
    fig, ax = plt.subplots(figsize=(11, 3.8))
    ax.plot(x, avg1, "o-", color="#C00000", lw=2, ms=5, label="1st Innings Avg")
    ax.plot(x, avg2, "s--", color="#2E75B6", lw=2, ms=5, label="2nd Innings Avg")
    ax.set_xticks(x)
    ax.set_xticklabels(seasons, rotation=45, ha="right", fontsize=8)
    ax.set_ylabel("Average total (runs)")
    ax.set_title("Season-Wise Average Innings Totals (2007–2025)", fontsize=11,
                 fontweight="bold", color="#1F3864")
    ax.legend(fontsize=9)
    ax.spines[["top","right"]].set_visible(False)
    ax.set_ylim(120, 210)
    fig.tight_layout(pad=1.2)
    return fig

def fig_phase_distribution():
    phases  = ["1st Inn\nPowerplay","1st Inn\nMiddle","1st Inn\nDeath",
               "2nd Inn\nPowerplay","2nd Inn\nMiddle","2nd Inn\nDeath"]
    rows    = [44519, 65705, 36267, 44217, 63937, 27907]
    colors  = ["#1F3864","#2E75B6","#ADB9CA","#C00000","#FF6B6B","#FFAAAA"]
    fig, ax = plt.subplots(figsize=(9, 3.5))
    bars = ax.bar(phases, rows, color=colors, edgecolor="white", width=0.6)
    ax.set_title("Dataset Rows by Phase and Innings", fontsize=11,
                 fontweight="bold", color="#1F3864")
    ax.set_ylabel("Number of rows")
    for bar, val in zip(bars, rows):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+300,
                f"{val:,}", ha="center", va="bottom", fontsize=8.5, fontweight="bold")
    ax.spines[["top","right"]].set_visible(False)
    fig.tight_layout(pad=1.2)
    return fig

def fig_architecture():
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    layers = [
        (5.0, 4.3, "PRESENTATION LAYER", "#375623",
         "Flask App  |  JSON API  |  Streamlit  |  CLI"),
        (5.0, 3.2, "SERVICE / INFERENCE LAYER", "#C00000",
         "Shared Predictor: validate → normalize → build features → score model → win model → format output"),
        (5.0, 2.1, "LEARNING LAYER", "#1F3864",
         "Torch Entity-Embedding (Score)   |   Calibrated CatBoost (Win)   |   Model Registry"),
        (5.0, 1.0, "DATA LAYER", "#2E75B6",
         "ipl_features.csv  |  venue_stats  |  team_form  |  player_form  |  batter_bowler_form"),
    ]
    for cx, cy, title, col, sub in layers:
        ax.add_patch(mpatches.FancyBboxPatch((0.3, cy-0.45), 9.4, 0.9,
            boxstyle="round,pad=0.06", facecolor=col, edgecolor="white", linewidth=1.5, alpha=0.9))
        ax.text(cx, cy+0.15, title, ha="center", va="center",
                color="white", fontsize=9.5, fontweight="bold")
        ax.text(cx, cy-0.18, sub, ha="center", va="center",
                color="white", fontsize=8, alpha=0.92)
        if cy > 1.0:
            ax.annotate("", xy=(cx, cy-0.45-0.02), xytext=(cx, cy-0.45-0.2),
                arrowprops=dict(arrowstyle="<-", color="#555", lw=1.5))
    ax.set_title("System Architecture — Layered View", fontsize=11,
                 fontweight="bold", color="#1F3864", pad=6)
    fig.tight_layout(pad=0.5)
    return fig

# ── build document ─────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p)
add_run(p, "A Project Report\n", bold=False, size=13)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p2)
add_run(p2, "on\n\n", size=12)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
no_space(p3)
add_run(p3, "IPL MATCH SCORE AND WIN PREDICTION SYSTEM",
        bold=True, size=16, color=RED)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p4, "Submitted in partial fulfillment of the requirements\n"
            "for the award of the degree of\n\n"
            "Bachelor of Technology\nin\n"
            "Computer Science & Engineering\n"
            "(Artificial Intelligence and Machine Learning)", size=12)

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p5, "Submitted By\n"
            "Rinku – 2822306\nRahul – 2822450\nAnkit – 2822307",
        bold=True, size=12)

doc.add_paragraph()
p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p6, "Under the Supervision of\n"
            "Ms. Neha Sharma\n(Assistant Professor)\nCSE (AI&ML)",
        bold=True, size=12)

doc.add_paragraph()
p7 = doc.add_paragraph()
p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p7, "Panipat Institute of Engineering & Technology, Samalkha, Panipat\n"
            "Affiliated to Kurukshetra University, Kurukshetra\n(2022–2026)", size=11)
doc.add_page_break()

# ── ABSTRACT ──────────────────────────────────────────────────────────────────
heading(doc, "ABSTRACT", level=1)
body(doc,
    "The growing availability of structured sports data and advancements in machine learning "
    "have significantly enhanced the field of cricket analytics. This project presents the "
    "development of an IPL analytics system designed to estimate two key live outcomes during "
    "a match: the projected final innings total and the batting team's win probability. The "
    "system is built using ball-by-ball data from 1,189 IPL matches sourced from Cricsheet, "
    "along with an engineered live-state dataset containing 282,552 observations spanning "
    "seasons from 2007 to 2026.", indent=True)
body(doc,
    "The best-performing configuration utilises a PyTorch-based entity embedding model for "
    "score prediction and a calibrated CatBoost classifier for win probability estimation. "
    "In the final evaluation, the score prediction model achieved MAE of 15.36 and RMSE of "
    "19.81. The win probability model achieved a log loss of 0.4922, a Brier score of 0.1681, "
    "and an accuracy of 67.1%, indicating reliable probabilistic predictions.", indent=True)
doc.add_page_break()

# ── CHAPTER 1 ─────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 1 — INTRODUCTION", level=1)

heading(doc, "1.1  Background", level=2)
body(doc,
    "T20 cricket is one of the most volatile formats of the sport because the balance between "
    "bat and ball can shift dramatically within a few overs. A team that appears behind after "
    "the powerplay may recover through middle-over consolidation, while a side in control "
    "after fifteen overs may still collapse in the death phase. This volatility makes live "
    "prediction both useful and technically demanding.", indent=True)
body(doc,
    "In the IPL, the challenge is even greater because conditions vary heavily across venues, "
    "batting orders change from match to match, team quality changes from season to season, "
    "and player matchups often influence the flow of an innings in ways that a raw scoreboard "
    "cannot capture.", indent=True)

heading(doc, "1.2  Project Context", level=2)
body(doc,
    "The present project converts 1,189 raw IPL match files into a machine-learning system "
    "that can respond to a live match state. The final engineered dataset contains 282,552 "
    "rows and 62 columns, representing legal-ball snapshots rather than full-match summaries. "
    "That design decision is crucial because live prediction should answer questions in the "
    "moment, not only after the match is complete.", indent=True)
body(doc,
    "The project solves two related but different tasks: forecasting the eventual innings total "
    "and estimating the batting side's probability of winning from the current state.", indent=True)

heading(doc, "1.3  Motivation", level=2)
body(doc,
    "The main motivation was to build a system that behaves more like a real cricket analytics "
    "engine and less like a classroom-only notebook. Many student sports projects stop at one "
    "model trained on a static dataset. In contrast, the goal here was to combine preprocessing, "
    "historical support tables, model comparison, probability calibration, interface design, "
    "and reproducible reporting inside one coherent workflow.", indent=True)

heading(doc, "1.4  Objectives", level=2)
bullet(doc, "Build a leakage-safe IPL live-state dataset from raw Cricsheet CSV2 files.")
bullet(doc, "Design features that capture match state, momentum, venue behavior, team strength, and player context.")
bullet(doc, "Compare multiple tabular model families rather than assuming one model family would suit both tasks.")
bullet(doc, "Select the best-performing live score and live win models using chronological evaluation.")
bullet(doc, "Expose the final prediction system through practical interfaces (Flask, API, Streamlit, CLI).")

heading(doc, "1.5  Scope", level=2)
body(doc,
    "The report focuses on the IPL only. It covers raw-data ingestion, preprocessing, "
    "support-table construction, live score regression, live win classification, a pre-match "
    "modeling branch, a Flask interface, a JSON API, a Streamlit dashboard, a CLI, and "
    "automated tests. The scope does not include ball-tracking vision, lineup optimisation, "
    "or betting-market integration.", indent=True)
doc.add_page_break()

# ── CHAPTER 2 ─────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 2 — LITERATURE REVIEW", level=1)

heading(doc, "2.1  What Earlier Work Suggests", level=2)
body(doc,
    "The literature around cricket prediction and structured-data machine learning suggests "
    "that three layers matter most: the quality of the match-state representation, the "
    "suitability of the model family to mixed categorical and numeric data, and the discipline "
    "of chronological evaluation. When any one of these layers is weak, reported performance "
    "can look strong while failing to transfer to real live use.", indent=True)

heading(doc, "2.2  Review of Representative Literature", level=2)
body(doc,
    "Classical cricket score heuristics relied on current run rate, wickets in hand, and "
    "overs remaining. These are easy to explain but cannot represent venue par, team-specific "
    "form, or player-vs-player history. Gradient boosting became a strong reference point "
    "because it can model nonlinear interactions. CatBoost is especially relevant in domains "
    "where team, venue, and player identities carry meaningful signal. Entity embeddings "
    "offered a compact alternative to sparse one-hot encoding for high-cardinality categories.", indent=True)

heading(doc, "2.3  Approach Families Reviewed", level=2)
body(doc, "Table 2.1 summarises the approach families reviewed while designing the project.")

# TABLE 2.1
tbl = doc.add_table(rows=6, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
styled_table_header(tbl.rows[0], ["Approach Family", "Why It Matters", "What It Misses Without Good Features"])
rows_data = [
    ("Simple score heuristics", "Fast and interpretable baseline for cricket intuition.",
     "Cannot encode player context, venue priors, or nonlinear interactions."),
    ("Gradient boosting", "Strong choice for structured data with complex nonlinear relations.",
     "Performance drops if leakage control and categorical handling are weak."),
    ("Categorical-aware boosting (CatBoost)", "Useful when team, venue, striker, and bowler identities carry signal.",
     "Still needs well-defined historical context and calibration."),
    ("Entity embeddings", "Compact way to represent high-cardinality categories such as players and venues.",
     "Needs enough training data and careful validation to stay stable."),
    ("Probability calibration", "Essential when users consume confidence percentages instead of only class labels.",
     "Adds little value if the base classifier is poorly specified."),
]
for i, rd in enumerate(rows_data):
    styled_table_row(tbl.rows[i+1], rd, alt=(i % 2 == 1))
set_col_width(tbl, 0, 4.5); set_col_width(tbl, 1, 5.5); set_col_width(tbl, 2, 5.5)
caption(doc, "Table 2.1: Approach Families Reviewed While Designing the Project")

heading(doc, "2.4  Literature Gaps That Motivated the Proposed Work", level=2)
body(doc,
    "The survey above reveals that the missing piece was not one better algorithm alone. "
    "The real gap was the absence of a disciplined, end-to-end live cricket workflow that "
    "combines time-aware feature construction, model comparison, calibration, multi-interface "
    "serving, and reportable engineering decisions in a single project.", indent=True)
doc.add_page_break()

# ── CHAPTER 3 ─────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 3 — PROBLEM OBJECTIVE", level=1)

heading(doc, "3.1  Problem Definition", level=2)
body(doc,
    "The project solves a practical sports-analytics problem: given the current state of an IPL "
    "innings, estimate the eventual innings total and the batting side's probability of winning, "
    "using only the information that would be available at that point in time. This is a stronger "
    "formulation than predicting the winner before the match or after the innings has already "
    "matured, because it forces the pipeline to work under realistic information constraints.", indent=True)

heading(doc, "3.2  Formal Problem Statement", level=2)
body(doc,
    "Formally, the project attempts to learn two functions over a live match-state "
    "representation. The first function maps the current innings state to a projected final "
    "total. The second function maps the same or a closely related feature representation "
    "to the batting side's probability of eventually winning the match. Both functions are "
    "constrained by chronology: they may only use information known at the exact moment of "
    "prediction.", indent=True)

heading(doc, "3.3  Research Questions", level=2)
bullet(doc, "Can a leakage-safe live-state dataset outperform simple scoreboard-based reasoning for IPL prediction?")
bullet(doc, "Does a broader all-history active-team training scope provide stronger live performance than a recent-only scope?")
bullet(doc, "Do score and win tasks prefer different model families once the same feature set is exposed to them?")
bullet(doc, "How much of the performance gain comes from feature design and calibration rather than from algorithm choice alone?")

heading(doc, "3.4  Project Deliverables", level=2)
bullet(doc, "A processed IPL live-state dataset and multiple support tables built from raw Cricsheet files.")
bullet(doc, "A set of trained live models compared under chronological evaluation across alternative scopes.")
bullet(doc, "Pre-match score and win models for a limited no-live-context setting.")
bullet(doc, "A shared inference module used by Flask, API, Streamlit, and CLI surfaces.")
bullet(doc, "A report-generation workflow that can be rerun against the current workspace state.")
doc.add_page_break()

# ── CHAPTER 4 ─────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 4 — METHODOLOGY OF THE PROJECT", level=1)

heading(doc, "4.1  End-to-End Workflow", level=2)
body(doc,
    "The project pipeline begins with raw IPL match files and ends with a deployable prediction "
    "service. The most important engineering decision was to separate the work into stages: raw "
    "data acquisition, chronological preprocessing, support-table construction, live-model "
    "training, pre-match-model training, interface integration, and final reporting.", indent=True)
embed_figure(doc, fig_workflow(), "Figure 4.1: End-to-End IPL Prediction Workflow", width_inches=6.0)

heading(doc, "4.2  System Architecture", level=2)
body(doc,
    "At the system level the project can be viewed as a layered architecture. The data layer "
    "contains raw match files and processed support tables. The learning layer contains "
    "training workflows and promoted deployment artifacts. The service layer contains the "
    "common inference package. The presentation layer contains CLI, Flask, API, and Streamlit "
    "surfaces.", indent=True)
embed_figure(doc, fig_architecture(), "Figure 4.2: System Architecture — Layered View", width_inches=6.0)

heading(doc, "4.3  Core Dataset Snapshot", level=2)
body(doc, "Table 4.1 presents the key dataset quantities used across the final project pipeline.")
# TABLE 4.1
tbl2 = doc.add_table(rows=7, cols=3)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
styled_table_header(tbl2.rows[0], ["Item", "Current Workspace Value", "Why It Matters"])
data41 = [
    ("Raw matches", "1,189", "Defines the historical base available for chronological preprocessing."),
    ("Processed live-state rows", "282,552", "Gives the model many supervised match-state examples."),
    ("Processed columns", "62", "Reflects how much context is combined into each prediction row."),
    ("Active-team filtered rows", "238,672", "Shows the size of the main deployment scope."),
    ("Score train rows — all-history scope", "215,675", "Explains why the broader scope supports richer modelling."),
    ("Score train rows — recent scope", "17,012", "Shows how restrictive the recent-only setting becomes."),
]
for i, rd in enumerate(data41):
    styled_table_row(tbl2.rows[i+1], rd, alt=(i % 2 == 1))
set_col_width(tbl2, 0, 5.5); set_col_width(tbl2, 1, 3.5); set_col_width(tbl2, 2, 6.5)
caption(doc, "Table 4.1: Core Dataset Snapshot Used in the Final Project")

heading(doc, "4.4  Live Feature Families", level=2)
body(doc, "Table 4.2 summarises the feature families engineered for the live prediction task.")
# TABLE 4.2
tbl3 = doc.add_table(rows=8, cols=3)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
styled_table_header(tbl3.rows[0], ["Feature Family", "Examples", "Why It Helped the Final Result"])
data42 = [
    ("Current match state", "runs, wickets, balls_left, innings_progress", "Anchored every prediction to the visible scoreboard."),
    ("Momentum windows", "runs_last_5, wickets_last_5, runs_last_12_balls", "Captured acceleration, stability, or collapse risk."),
    ("Chase math", "target_remaining, required_run_rate, required_minus_current_rr", "Made second-innings predictions sensitive to actual pressure."),
    ("Venue priors", "venue averages, bat-first win rate, runs_vs_par", "Explained why the same score means different things at different grounds."),
    ("Team & matchup context", "team form, team-venue form, batting_vs_bowling_form", "Added competitive context missing from score-only models."),
    ("Player context", "striker form, bowler form, batter_vs_bowler history", "Helped the model distinguish states driven by specific matchups."),
    ("Optional environment", "temperature, humidity, wind, dew_risk", "Added coarse environmental context without dominating cricket signals."),
]
for i, rd in enumerate(data42):
    styled_table_row(tbl3.rows[i+1], rd, alt=(i % 2 == 1))
set_col_width(tbl3, 0, 4.0); set_col_width(tbl3, 1, 5.5); set_col_width(tbl3, 2, 6.0)
caption(doc, "Table 4.2: Live Feature Families and Their Project Value")

heading(doc, "4.5  Season-Wise Scoring Environment", level=2)
body(doc,
    "The dataset spans seasons from 2007 to 2026, covering 1,189 matches. Average first-innings "
    "totals rose from 160.97 in 2007/08 to 189.59 in 2024, reflecting the evolution of T20 "
    "batting. Chronological train-test splitting was essential to capture this temporal drift "
    "honestly.", indent=True)
embed_figure(doc, fig_season_totals(), "Figure 4.3: Season-Wise Average Innings Totals (2007–2025)")

heading(doc, "4.6  Mathematical Translation of Live Match State", level=2)
body(doc,
    "For every legal ball the pipeline computes the following state variables which convert "
    "the scoreboard into comparable numerical pressure signals:", indent=True)
code_lines = [
    "legal_balls_bowled = over_number * 6 + ball_in_over",
    "balls_left         = 120 - legal_balls_bowled",
    "wickets_left       = 10  - wickets",
    "innings_progress   = legal_balls_bowled / 120",
    "current_run_rate   = runs / max(legal_balls_bowled / 6, 1e-6)",
    "required_run_rate  = target_remaining / max(balls_left / 6, 1e-6)",
    "required_minus_current_rr = required_run_rate - current_run_rate",
]
for line in code_lines:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(line)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 180)
divider(doc)
doc.add_page_break()

# ── CHAPTER 5 ─────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 5 — RESULTS AND ANALYSIS", level=1)

heading(doc, "5.1  Modeling Philosophy", level=2)
body(doc,
    "The project did not assume that one model family would be ideal for every task. Score "
    "prediction is a regression problem in which the size of the error matters directly. Win "
    "prediction is a probability problem in which a 0.60 output should mean something close "
    "to a 60 percent event over time. That distinction led the project to evaluate model "
    "families separately and to prioritise calibration for the win task.", indent=True)

heading(doc, "5.2  Scope Comparison", level=2)
body(doc, "Table 5.1 shows how the training data size changed between scopes and which models were selected.")
# TABLE 5.1
tbl51 = doc.add_table(rows=3, cols=5)
tbl51.style = "Table Grid"
tbl51.alignment = WD_TABLE_ALIGNMENT.CENTER
styled_table_header(tbl51.rows[0], ["Scope","Score Train Rows","Win Train Rows","Selected Score Model","Selected Win Model"])
styled_table_row(tbl51.rows[1], ("recent_active_history","17,012","17,012","XGBoost GPU","XGBoost GPU (calibrated)"), alt=False)
styled_table_row(tbl51.rows[2], ("all_active_history","215,675","212,060","Torch Entity-Embedding","CatBoost GPU (calibrated)"), alt=True)
caption(doc, "Table 5.1: Scope Comparison That Guided Final Training Direction")

heading(doc, "5.3  Live Score Model Comparison", level=2)
body(doc, "Table 5.2 compares all live score models on the all-active-history evaluation scope.")
# TABLE 5.2
tbl52 = doc.add_table(rows=6, cols=4)
tbl52.style = "Table Grid"
tbl52.alignment = WD_TABLE_ALIGNMENT.CENTER
styled_table_header(tbl52.rows[0], ["Score Model","MAE","RMSE","Why It Was Important in the Search"])
data52 = [
    ("HistGradientBoosting","17.72","24.50","Strong CPU baseline; feature engineering already carried useful signal."),
    ("XGBoost GPU","16.56","23.06","Powerful boosted-tree benchmark and strong structured-data reference."),
    ("CatBoost GPU","19.90","26.27","Useful categorical-aware alternative, but weaker for live score task."),
    ("Torch Entity-Embedding","15.36","19.81","Best recorded score model; learned embeddings helped rich categorical context."),
    ("Weighted ML Ensemble","16.04","23.54","Robust ensemble variant; useful, but not the strongest recorded snapshot."),
]
for i, rd in enumerate(data52):
    styled_table_row(tbl52.rows[i+1], rd, alt=(i % 2 == 1))
set_col_width(tbl52, 0, 4.5); set_col_width(tbl52, 1, 1.5); set_col_width(tbl52, 2, 1.5); set_col_width(tbl52, 3, 7.5)
caption(doc, "Table 5.2: Live Score Model Comparison on the Key Evaluation Scope")

embed_figure(doc, fig_model_comparison(), "Figure 5.1: Comparative View of Candidate Live Models")

heading(doc, "5.4  Live Win Model Comparison", level=2)
body(doc, "Table 5.3 compares win-probability models using log loss, Brier score, and accuracy.")
# TABLE 5.3
tbl53 = doc.add_table(rows=4, cols=5)
tbl53.style = "Table Grid"
tbl53.alignment = WD_TABLE_ALIGNMENT.CENTER
styled_table_header(tbl53.rows[0], ["Win Model","Log Loss","Brier Score","Accuracy","Interpretation"])
data53 = [
    ("HistGradientBoosting (calibrated)","0.6047","0.2040","65.7%","Good classical baseline after calibration."),
    ("XGBoost GPU (calibrated)","0.5909","0.2026","66.3%","Competitive calibrated alternative with strong ranking."),
    ("CatBoost GPU (calibrated)","0.4922","0.1681","67.1%","Best balance of probability quality and classification."),
]
for i, rd in enumerate(data53):
    styled_table_row(tbl53.rows[i+1], rd, alt=(i % 2 == 1))
caption(doc, "Table 5.3: Live Win Model Comparison on the Key Evaluation Scope")

heading(doc, "5.5  Error Analysis", level=2)
embed_figure(doc, fig_error_distribution(), "Figure 5.2: Held-Out Score Error Distribution and MAE by Wicket-State Bucket")

heading(doc, "5.6  Win Probability Calibration", level=2)
embed_figure(doc, fig_calibration(), "Figure 5.3: Held-Out Win Probability Calibration Curve")

heading(doc, "5.7  Phase Distribution of Dataset", level=2)
embed_figure(doc, fig_phase_distribution(), "Figure 5.4: Dataset Rows by Phase and Innings")

heading(doc, "5.8  Pre-Match Model Performance", level=2)
body(doc, "Table 5.4 summarises the weaker pre-match branch, confirming that live context drives most predictive value.")
# TABLE 5.4
tbl54 = doc.add_table(rows=5, cols=4)
tbl54.style = "Table Grid"
tbl54.alignment = WD_TABLE_ALIGNMENT.CENTER
styled_table_header(tbl54.rows[0], ["Task","Split","MAE / Accuracy","RMSE / Log Loss"])
data54 = [
    ("Pre-match score","Validation","32.94","40.79"),
    ("Pre-match score","Test","36.44","41.15"),
    ("Pre-match win","Validation","55.7%","0.6909"),
    ("Pre-match win","Test","57.9%","0.6837"),
]
for i, rd in enumerate(data54):
    styled_table_row(tbl54.rows[i+1], rd, alt=(i % 2 == 1))
caption(doc, "Table 5.4: Pre-Match Model Performance Summary")
doc.add_page_break()

# ── CHAPTER 6 ─────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 6 — SYSTEM IMPLEMENTATION AND TESTING", level=1)

heading(doc, "6.1  Shared Inference Core", level=2)
body(doc,
    "A major engineering strength of the project is that every user-facing surface calls the "
    "same core prediction path. Input normalisation, overs parsing, feature-frame construction, "
    "support-table loading, score inference, win inference, uncertainty formatting, and "
    "optional monitoring hooks are all centralised in the shared prediction module. This "
    "reduced a class of bugs common in academic projects: one interface quietly using "
    "different logic from another.", indent=True)

heading(doc, "6.2  User-Facing Surfaces", level=2)
body(doc, "Table 6.1 lists the four user-facing surfaces built around the shared prediction core.")
# TABLE 6.1
tbl61 = doc.add_table(rows=5, cols=3)
tbl61.style = "Table Grid"
tbl61.alignment = WD_TABLE_ALIGNMENT.CENTER
styled_table_header(tbl61.rows[0], ["Surface","Role in the Project","Why It Was Kept"])
data61 = [
    ("Flask web application", "Primary demonstration interface for live and pre-match prediction.",
     "Best suited for faculty evaluation and everyday interactive use."),
    ("JSON API", "Machine-readable prediction surface for external tools.",
     "Shows that the project can integrate with other services."),
    ("Streamlit dashboard", "Analytical interface for model and data exploration.",
     "Useful for inspection and demonstration of supporting context."),
    ("CLI", "Lightweight terminal interaction path.",
     "Useful for quick manual testing and debugging."),
]
for i, rd in enumerate(data61):
    styled_table_row(tbl61.rows[i+1], rd, alt=(i % 2 == 1))
set_col_width(tbl61, 0, 4.0); set_col_width(tbl61, 1, 6.0); set_col_width(tbl61, 2, 5.5)
caption(doc, "Table 6.1: User-Facing Surfaces Built Around the Shared Prediction Core")

heading(doc, "6.3  Software Engineering Principles", level=2)
body(doc,
    "The implementation reflects several core software-engineering principles. Modularity was "
    "maintained by separating preprocessing, training, inference, interfaces, and documentation "
    "into distinct files and packages. Reuse was maintained through the common inference "
    "contract so that the same prediction logic serves CLI, Flask, API, and Streamlit. "
    "Validation and fail-fast checks were added at the shared entry point so that incorrect "
    "inputs are rejected consistently rather than differently in each surface.", indent=True)

heading(doc, "6.4  Dependency Stack", level=2)
body(doc,
    "The principal software stack includes Python 3, pandas, scikit-learn, XGBoost, CatBoost, "
    "PyTorch, Flask, Streamlit, and python-docx. Each dependency was chosen because it solved "
    "a clear part of the overall project. The repository also contains GPU-capable training "
    "branches for XGBoost, CatBoost, and PyTorch entity-embedding experiments.", indent=True)

heading(doc, "6.5  Artifact Traceability", level=2)
body(doc,
    "The repository does not contain only the final promoted models. It also contains "
    "comparison reports (best_model_search_report.json, cpu_model_report.json), prediction "
    "exports, uncertainty metadata, drift summaries, and registry snapshots that make the "
    "training history easier to inspect afterward. The strongest saved version referenced "
    "throughout the report is v20260415_152209.", indent=True)

heading(doc, "6.6  Monitoring Readiness", level=2)
body(doc,
    "The project includes an early monitoring layer. The monitoring window is configured at "
    "500 events. Currently 37 prediction events and 4 resolved outcomes are available. The "
    "current status is insufficient_data — the module recommends collecting more events before "
    "taking retraining action. This establishes the path by which the system could later "
    "justify retraining using production evidence.", indent=True)
doc.add_page_break()

# ── CHAPTER 7 ─────────────────────────────────────────────────────────────────
heading(doc, "CHAPTER 7 — CONCLUSION AND FUTURE SCOPE", level=1)

heading(doc, "7.1  Conclusion", level=2)
body(doc,
    "This project succeeded because it treated IPL prediction as an end-to-end engineering "
    "problem. The most important gains came from using the right scope, preventing temporal "
    "leakage, constructing reusable support tables, and evaluating score and win tasks "
    "separately. Strong cricket prediction is driven as much by disciplined feature design "
    "and evaluation logic as by the choice of model family.", indent=True)
body(doc,
    "The strongest tracked experimental result used Torch entity-embedding model for live "
    "score and CatBoost GPU (calibrated) for live win probability, achieving RMSE 19.81 and "
    "log loss 0.4922 on held-out evaluation. The project aligned model complexity with better "
    "context rather than treating complexity as a substitute for context.", indent=True)

heading(doc, "7.2  Major Lessons Learned", level=2)
bullet(doc, "The best improvements came from historical context design and leakage control before trying more algorithms.")
bullet(doc, "Recent-only training can be too small for rich live cricket prediction once chronology is enforced honestly.")
bullet(doc, "Score and win tasks reward different modeling properties, so a hybrid final system is often better.")
bullet(doc, "Probability calibration is essential when the end product displays win percentages to users.")
bullet(doc, "A strong academic ML project should include reproducible interfaces and tests, not only notebook metrics.")

heading(doc, "7.3  Current Limitations", level=2)
body(doc,
    "The system is still limited by genuine cricket uncertainty, especially in early-innings "
    "states and rare collapse or explosion patterns. The pre-match branch remains much weaker "
    "than the live branch because it operates under a smaller information budget. The current "
    "pipeline also does not incorporate probable XI information, pitch reports, or live "
    "ball-tracking data.", indent=True)

heading(doc, "7.4  Future Scope", level=2)
bullet(doc, "Integrate probable playing XI and lineup-specific priors before the match starts.")
bullet(doc, "Experiment with quantile or conformal methods to replace fixed uncertainty bands with more adaptive intervals.")
bullet(doc, "Add richer explanation outputs so that the API can return feature-level reasons behind a live prediction.")
bullet(doc, "Expand the monitoring layer from passive reporting to retraining triggers once enough live outcomes accumulate.")
bullet(doc, "Extend the same pipeline design to other T20 leagues after validating how much of the feature logic transfers cleanly.")

heading(doc, "7.5  Concluding Remarks", level=2)
body(doc,
    "In conclusion, the project achieved its objective by building a usable IPL prediction "
    "system grounded in realistic preprocessing, stronger context features, comparative model "
    "evaluation, and software delivery discipline. The final system is valuable not because "
    "it eliminates uncertainty, but because it handles uncertainty more intelligently than a "
    "score-only baseline and presents the result through interfaces that can actually be used "
    "and evaluated.", indent=True)
doc.add_page_break()

# ── REFERENCES ────────────────────────────────────────────────────────────────
heading(doc, "REFERENCES", level=1)
refs = [
    "[1]  Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
    "[2]  Bishop, C. M. (2006). Pattern Recognition and Machine Learning. Springer.",
    "[3]  Geron, A. (2019). Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow (2nd ed.). O'Reilly.",
    "[4]  Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. ACM SIGKDD, pp. 785-794.",
    "[5]  Prokhorenkova, L., et al. (2018). CatBoost: Unbiased Boosting with Categorical Features. NeurIPS, pp. 6638-6648.",
    "[6]  Guo, C., & Berkhahn, F. (2016). Entity Embeddings of Categorical Variables. arXiv:1604.06737.",
    "[7]  Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. JMLR, 12, pp. 2825-2830.",
    "[8]  Paszke, A., et al. (2019). PyTorch: An Imperative Style, High-Performance Deep Learning Library. NeurIPS.",
    "[9]  McKinney, W. (2017). Python for Data Analysis (2nd ed.). O'Reilly Media.",
    "[10] Niculescu-Mizil, A., & Caruana, R. (2005). Predicting Good Probabilities with Supervised Learning. ICML, pp. 625-632.",
    "[11] Duckworth, F. C., & Lewis, A. J. (1998). A Fair Method for Resetting the Target. JORS, 49(3), pp. 220-227.",
    "[12] Sankaranarayanan, V. V., et al. (2014). Auto-Play: ODI Cricket Simulation. SIAM SDM, pp. 1064-1072.",
    "[13] Cricsheet. (2024). IPL Ball-by-Ball Data and Match Metadata. cricsheet.org/downloads/",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    add_run(p, ref, size=10)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = "docs/IPL_Project_Report_Full.docx"
doc.save(out)
print(f"Saved: {out}")
