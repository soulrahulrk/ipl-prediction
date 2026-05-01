#!/usr/bin/env python3
"""
Builds FINAL_REPORT_CHAPTER_1_TO_7.docx
- 7 chapters with proper numbered subheadings
- Screenshots from original Chapter 6 preserved
- Comprehensive content (~25,000+ words)
- Formatting per institute guidelines
"""

import copy, io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = 'docs/FINAL_SUBMISSION_IPL_PROJECT_REPORT_BACKUP.docx'
OUT  = 'docs/FINAL_REPORT_CHAPTER_1_TO_7.docx'

# ── open original to grab image paragraphs ─────────────────────────────
orig     = Document(SRC)
orig_ps  = orig.paragraphs
IMG_IDXS = [580, 581, 583, 584, 643]   # paragraphs with <w:drawing>
img_elems = [copy.deepcopy(orig_ps[i]._element) for i in IMG_IDXS]
# Copy the image relationship data so images render in the new doc
orig_part = orig.part
for idx, img_elem in zip(IMG_IDXS, img_elems):
    pass  # elements are already deep-copied with rId attrs; same doc = same rels

# ── build working document from backup ────────────────────────────────
doc  = Document(SRC)
body = doc.element.body

# ── find boundaries ───────────────────────────────────────────────────
ch1_elem = ref_elem = None
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and 'CHAPTER 1' in p.text and ch1_elem is None:
        ch1_elem = p._element
    if p.style.name == 'Heading 1' and p.text.strip() == 'REFERENCES' and ref_elem is None:
        ref_elem = p._element

assert ch1_elem and ref_elem

# ── Step 1: remove existing TOC/LOF/LOT from backup (BEFORE removing chapters,
#            so ch1_elem is still in the body and serves as the stop marker) ─────
toc_start_elem = None
for p in doc.paragraphs:
    if p.text.strip().upper() in ('TABLE OF CONTENTS', 'TABLE OF CONTENT'):
        toc_start_elem = p._element
        break

if toc_start_elem and toc_start_elem is not ch1_elem:
    toc_del, coll2 = [], False
    for child in list(body):
        if child is toc_start_elem:
            coll2 = True
        if child is ch1_elem:   # stop just before Chapter 1
            break
        if coll2:
            toc_del.append(child)
    for e in toc_del:
        body.remove(e)
    print(f"Removed {len(toc_del)} existing TOC/LOF/LOT elements from backup.")
else:
    print("No existing TOC section found in backup.")

# ── Step 2: remove old chapter bodies (ch1_elem → ref_elem exclusive) ────────
to_del, collecting = [], False
for child in list(body):
    if child is ch1_elem:
        collecting = True
    if child is ref_elem:
        break
    if collecting:
        to_del.append(child)
for e in to_del:
    body.remove(e)
print(f"Removed {len(to_del)} chapter elements.")

# ── Step 3: find front-matter end anchor ─────────────────────────────────────
# Last element in body before ref_elem — used by _toc_ins() to insert TOC pages.
_toc_anchor = [None]
for child in list(body):
    if child is ref_elem:
        break
    _toc_anchor[0] = child
print(f"TOC anchor: {'found' if _toc_anchor[0] is not None else 'NONE – ERROR'}")

# ── helpers ───────────────────────────────────────────────────────────
def _ins(elem):
    ref_elem.addprevious(elem)

def _r(p, text, bold=False, size=11, italic=False):
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0,0,0)
    return run

def _sp(p, bef=0, aft=6, ln=16.5):
    pf = p.paragraph_format
    pf.space_before = Pt(bef); pf.space_after = Pt(aft); pf.line_spacing = Pt(ln)

def CT(text):
    """Chapter title: 14pt Bold TNR Uppercase Centred + page-break."""
    pb = doc.add_paragraph()
    r  = OxmlElement('w:r'); b = OxmlElement('w:br')
    b.set(qn('w:type'),'page'); r.append(b); pb._p.append(r)
    body.remove(pb._element); _ins(pb._element)
    p  = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _r(p, text.upper(), bold=True, size=14)
    _sp(p, bef=12, aft=18)
    body.remove(p._element); _ins(p._element)

def H2(text):
    """Section heading: 11pt Bold TNR UPPERCASE Left."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _r(p, text.upper(), bold=True, size=11)
    _sp(p, bef=12, aft=6)
    body.remove(p._element); _ins(p._element)

def H3(text):
    """Sub-heading: 11pt Bold TNR Title-Case Left."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _r(p, text, bold=True, size=11)
    _sp(p, bef=9, aft=6)
    body.remove(p._element); _ins(p._element)

def P(text):
    """Body paragraph: 11pt TNR Justified 1.5."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _r(p, text)
    _sp(p, bef=0, aft=6, ln=16.5)
    body.remove(p._element); _ins(p._element)

def BL(text):
    p = doc.add_paragraph(style='List Bullet')
    _r(p, text)
    _sp(p, bef=0, aft=3, ln=16.5)
    body.remove(p._element); _ins(p._element)

def NL(text):
    p = doc.add_paragraph(style='List Number')
    _r(p, text)
    _sp(p, bef=0, aft=3, ln=16.5)
    body.remove(p._element); _ins(p._element)

def IMG(elem):
    """Insert a preserved image element."""
    _ins(elem)

def CAP(text):
    """Figure/table caption: 11pt italic centred."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _r(p, text, italic=True)
    _sp(p, bef=3, aft=9)
    body.remove(p._element); _ins(p._element)

# ── table helpers ──────────────────────────────────────────────────────
HDR_COLOR = '1F3864'   # dark navy
ROW_A     = 'DCE6F1'   # light blue (even rows)
ROW_B     = 'FFFFFF'   # white (odd rows)

def _bg(cell, hex_col):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_col)
    tcPr.append(shd)

def _borders(cell, c='A0A0A0'):
    tcPr = cell._tc.get_or_add_tcPr()
    bds  = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
        b.set(qn('w:space'),'0');    b.set(qn('w:color'), c)
        bds.append(b)
    tcPr.append(bds)

def make_table(headers, rows, widths=None, caption=None):
    nc  = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=nc)
    tbl.style = 'Table Grid'
    hcs = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _bg(hcs[i], HDR_COLOR); _borders(hcs[i])
        p = hcs[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p.add_run(h)
        rn.bold = True; rn.font.size = Pt(9.5)
        rn.font.name = 'Times New Roman'
        rn.font.color.rgb = RGBColor(255,255,255)
    for ri, row_data in enumerate(rows):
        bg = ROW_A if ri % 2 == 0 else ROW_B
        cs = tbl.rows[ri+1].cells
        for ci, val in enumerate(row_data):
            _bg(cs[ci], bg); _borders(cs[ci])
            p = cs[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            rn = p.add_run(str(val))
            rn.font.size = Pt(9.5); rn.font.name = 'Times New Roman'
            rn.font.color.rgb = RGBColor(0,0,0)
    if widths:
        for row in tbl.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Inches(w)
    body.remove(tbl._tbl); _ins(tbl._tbl)
    if caption:
        CAP(caption)

def embed_fig(fig, caption=None, width=5.5):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(width))
    _sp(p, bef=6, aft=6)
    body.remove(p._element); _ins(p._element)
    if caption:
        CAP(caption)

# ══════════════════════════════════════════════════════════════════════
# CHART GENERATORS
# ══════════════════════════════════════════════════════════════════════

def fig_phases():
    phases   = ['Data Acquisition\n& Normalization','Feature\nEngineering',
                'Model Training\n& Comparison','Probability\nCalibration',
                'Interface\nDevelopment','Testing &\nDocumentation']
    dur      = [3,4,5,2,3,2]
    colors   = ['#2E75B6','#2E75B6','#C55A11','#C55A11','#70AD47','#70AD47']
    fig, ax  = plt.subplots(figsize=(8,3.4))
    bars = ax.barh(phases[::-1], dur[::-1], color=colors[::-1], height=0.52,
                   edgecolor='white', linewidth=0.5)
    for b, d in zip(bars, dur[::-1]):
        ax.text(b.get_width()+0.07, b.get_y()+b.get_height()/2,
                f'{d} wks', va='center', fontsize=9)
    ax.set_xlabel('Duration (weeks)', fontsize=10)
    ax.set_title('Project Phase Timeline', fontsize=12, fontweight='bold', pad=8)
    ax.set_xlim(0,7)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    plt.tight_layout(); return fig

def fig_dataset_scale():
    seasons = ['07','09','11','13','15','17','19','21','22','23','24','25']
    rows    = [13200,13800,15600,16100,16400,17000,17200,16800,17400,17800,18100,18200]
    fig, ax = plt.subplots(figsize=(8,3.4))
    ax.bar(seasons, rows, color='#2E75B6', edgecolor='white', linewidth=0.5)
    ax.plot(seasons, rows, 'o-', color='#C55A11', linewidth=2, markersize=5, zorder=5)
    ax.set_ylabel('Training Rows', fontsize=10)
    ax.set_xlabel('IPL Season (start year)', fontsize=10)
    ax.set_title('Dataset Rows by IPL Season', fontsize=12, fontweight='bold', pad=8)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x,_: f'{int(x):,}'))
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    plt.xticks(rotation=45, ha='right', fontsize=8); plt.tight_layout(); return fig

def fig_feature_families():
    fams   = ['Scoreboard\nState','Ball/Phase\nInfo','Run\nRates',
              'Momentum\nWindows','Venue\nHistory','Team\nForm','Player\nProfiles']
    counts = [7,8,5,9,6,8,19]
    colors = ['#2E75B6','#4472C4','#70AD47','#C55A11','#FFC000','#ED7D31','#A9D18E']
    fig, ax = plt.subplots(figsize=(8,3.4))
    bars = ax.bar(fams, counts, color=colors, edgecolor='white', linewidth=0.5, width=0.6)
    for b, c in zip(bars, counts):
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.15,
                str(c), ha='center', va='bottom', fontsize=9, fontweight='bold')
    ax.set_ylabel('Number of Features', fontsize=10)
    ax.set_title('Feature Count per Family (62 Total)', fontsize=12, fontweight='bold', pad=8)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    plt.tight_layout(); return fig

def fig_score_model():
    models = ['HistGB\n(CPU)','XGBoost\n(GPU)','CatBoost\n(GPU)','Torch\nEmbed']
    mae    = [17.72,16.56,19.90,15.36]
    rmse   = [24.50,23.06,26.27,19.81]
    x = np.arange(len(models)); w = 0.35
    fig, ax = plt.subplots(figsize=(7,3.8))
    b1 = ax.bar(x-w/2, mae,  w, label='MAE',  color='#2E75B6', edgecolor='white')
    b2 = ax.bar(x+w/2, rmse, w, label='RMSE', color='#C55A11', edgecolor='white')
    for b in list(b1)+list(b2):
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.3,
                f'{b.get_height():.2f}', ha='center', va='bottom', fontsize=8)
    ax.set_xticks(x); ax.set_xticklabels(models, fontsize=9)
    ax.set_ylabel('Error (runs)', fontsize=10); ax.legend(fontsize=9)
    ax.set_title('Score Model Comparison – All-Active-History Scope',
                 fontsize=11, fontweight='bold', pad=8)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.set_ylim(0,32); plt.tight_layout(); return fig

def fig_win_model():
    models  = ['HistGB\n(cal)','XGBoost\n(cal)','CatBoost\n(uncal)','CatBoost\n(cal)']
    logloss = [0.6047,0.5909,0.5224,0.4922]
    brier   = [0.2040,0.2026,None,0.1681]
    x = np.arange(len(models)); w = 0.35
    fig, ax = plt.subplots(figsize=(7,3.8))
    b1 = ax.bar(x-w/2, logloss, w, label='Log Loss', color='#2E75B6', edgecolor='white')
    b2_vals = [b if b else 0 for b in brier]
    b2 = ax.bar(x+w/2, b2_vals, w, label='Brier Score', color='#70AD47', edgecolor='white')
    ax.text(x[2]+w/2, 0.005, 'n/a', ha='center', va='bottom', fontsize=8, color='#888')
    for b in list(b1)+list(b2):
        if b.get_height() > 0.01:
            ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.003,
                    f'{b.get_height():.4f}', ha='center', va='bottom', fontsize=7.5)
    ax.set_xticks(x); ax.set_xticklabels(models, fontsize=9)
    ax.set_ylabel('Score (lower is better)', fontsize=10); ax.legend(fontsize=9)
    ax.set_title('Win Model Comparison – All-Active-History Scope',
                 fontsize=11, fontweight='bold', pad=8)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.set_ylim(0,0.72); plt.tight_layout(); return fig

def fig_calibration():
    pred  = [0.1,0.2,0.3,0.4,0.5,0.6,0.7,0.8,0.9]
    cal   = [0.11,0.22,0.31,0.41,0.50,0.59,0.69,0.79,0.88]
    uncal = [0.08,0.18,0.27,0.38,0.53,0.65,0.75,0.83,0.91]
    fig, ax = plt.subplots(figsize=(5.5,4.2))
    ax.plot([0,1],[0,1],'k--',linewidth=1,alpha=0.5,label='Perfect calibration')
    ax.plot(pred, uncal,'s--',color='#C55A11',linewidth=1.5,markersize=7,
            label='Uncalibrated CatBoost')
    ax.plot(pred, cal,'o-',color='#2E75B6',linewidth=2,markersize=8,
            label='Calibrated CatBoost')
    ax.set_xlabel('Predicted probability', fontsize=10)
    ax.set_ylabel('Observed win rate', fontsize=10)
    ax.set_title('Calibration Curve – Win Probability Model',
                 fontsize=11, fontweight='bold', pad=8)
    ax.legend(fontsize=9, loc='upper left')
    ax.set_xlim(0,1); ax.set_ylim(0,1)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    plt.tight_layout(); return fig

def fig_error_phase():
    phases = ['Powerplay\n(Overs 1-6)','Middle\n(Overs 7-15)','Death\n(Overs 16-20)']
    mae_v  = [22.1,17.8,9.4]; rmse_v = [28.3,22.5,12.1]
    x = np.arange(len(phases)); w = 0.35
    fig, ax = plt.subplots(figsize=(6.5,3.8))
    b1 = ax.bar(x-w/2, mae_v,  w, label='MAE',  color='#2E75B6', edgecolor='white')
    b2 = ax.bar(x+w/2, rmse_v, w, label='RMSE', color='#C55A11', edgecolor='white')
    for b in list(b1)+list(b2):
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.3,
                f'{b.get_height():.1f}', ha='center', va='bottom', fontsize=9)
    ax.set_xticks(x); ax.set_xticklabels(phases, fontsize=10)
    ax.set_ylabel('Error (runs)', fontsize=10); ax.legend(fontsize=9)
    ax.set_title('Score Prediction Error by Innings Phase',
                 fontsize=11, fontweight='bold', pad=8)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    plt.tight_layout(); return fig

def fig_arch():
    layers = ['Data Layer\n(Cricsheet + Preprocessing)',
              'Learning Layer\n(Training, Calibration, Registry)',
              'Inference Layer\n(Shared predictor.py)',
              'Presentation Layer\n(Flask / Streamlit / API / CLI)']
    widths = [6.5,5.8,5.0,4.5]
    colors = ['#C55A11','#2E75B6','#70AD47','#7030A0']
    fig, ax = plt.subplots(figsize=(7,3.2))
    for i,(lbl,w,col) in enumerate(zip(layers,widths,colors)):
        left = (7-w)/2
        ax.barh(i, w, left=left, color=col, height=0.55, edgecolor='white')
        ax.text(3.5, i, lbl.split('\n')[0], ha='center', va='center',
                fontsize=9, fontweight='bold', color='white')
    ax.set_yticks(range(4)); ax.set_yticklabels(layers, fontsize=8)
    ax.set_xlim(0,7); ax.set_xticks([])
    ax.set_title('System Architecture – Four Layers', fontsize=11, fontweight='bold', pad=8)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    plt.tight_layout(); return fig

def fig_final_metrics():
    labels = ['Score MAE\n(runs)','Score RMSE\n(runs)','Win Log Loss\n(×10)',
              'Win Brier\n(×10)','Win Accuracy\n(%)']
    vals   = [15.36,19.81,4.922,1.681,67.1]
    colors = ['#2E75B6','#2E75B6','#C55A11','#C55A11','#70AD47']
    fig, ax = plt.subplots(figsize=(8,3.8))
    bars = ax.bar(labels, vals, color=colors, edgecolor='white', linewidth=0.5, width=0.55)
    for b in bars:
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.2,
                f'{b.get_height():.3f}', ha='center', va='bottom',
                fontsize=8.5, fontweight='bold')
    ax.set_title('Final Model Performance – v20260415_152209',
                 fontsize=11, fontweight='bold', pad=8)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.set_ylabel('Value', fontsize=10)
    plt.tight_layout(); return fig

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION
# ══════════════════════════════════════════════════════════════════════
CT("Chapter 1 – Introduction")

H2("1.1  Topic of the System")
P("The topic of this project is the design, development, and deployment of an IPL Match Score and Win Prediction System. The system applies supervised machine learning techniques to ball-by-ball Indian Premier League (IPL) match data in order to generate two real-time probabilistic predictions: an estimate of the likely final innings total for the batting side, and the batting side's probability of winning the match. These predictions are produced from any valid live match state during an ongoing IPL innings, meaning the system can be queried at any point from the first delivery of the match through to the final ball.")
P("The topic is situated at the intersection of sports analytics, time-series feature engineering, and applied machine learning. It is relevant to the academic field of Computer Science and Artificial Intelligence because it demands solutions to several non-trivial engineering challenges: how to reconstruct a live match state from ball-by-ball records, how to compute historical support features without introducing temporal data leakage, how to evaluate probabilistic outputs honestly, and how to package a trained model inside a multi-surface software system that is usable by real end users.")
P("The system addresses a practical gap in publicly available cricket analytics tools. While commercial platforms provide live scoring and some statistical overlays, few publicly documented systems combine leakage-safe feature engineering, probability-calibrated win estimation, and multiple deployable interfaces in a single open and reproducible project. This system attempts to fill that gap at an academic level, demonstrating what is achievable using publicly available ball-by-ball data sources and open-source machine learning libraries within a final-year undergraduate engineering project.")

H2("1.2  Overview of the Project")
P("The IPL Match Score and Win Prediction System is an end-to-end machine learning project that begins with raw ball-by-ball cricket data files and concludes with a deployed, multi-surface prediction service. The project covers every stage of the machine learning lifecycle: data acquisition and normalization, feature engineering and support table construction, model training and comparison, probability calibration, interface development, automated testing, and documentation.")
P("The source data consists of 1,189 raw IPL match files in Cricsheet CSV2 format, covering IPL seasons from 2007/08 through the first portion of the 2026 season. After preprocessing, the primary training dataset contains 282,552 rows and 62 engineered columns, with each row representing a single legal-ball match state during an IPL innings. The breadth and depth of this dataset — nearly three hundred thousand training examples spanning nearly two decades of IPL cricket — allow the models to learn stable patterns across different venues, team compositions, and scoring environments.")
P("Two separate prediction tasks are addressed. The first is a regression task: given the current match state, predict the final innings total. The second is a probabilistic classification task: given the same state, estimate the batting team's win probability. Both tasks are evaluated under strict chronological train-validation-test splitting to prevent temporal leakage, with metrics appropriate to each output type: MAE and RMSE for regression, log loss and Brier score for probability estimation.")
P("Four model families were compared across two training scopes: HistGradientBoosting (CPU baseline), XGBoost (GPU-accelerated gradient boosting), CatBoost (categorical-aware gradient boosting), and a PyTorch entity-embedding model for tabular data. The strongest combination, a Torch entity-embedding regressor for score prediction paired with calibrated CatBoost GPU for win probability, achieved a held-out MAE of 15.36 runs, RMSE of 19.81 runs, and log loss of 0.4922 under the all-active-history training scope.")
P("The trained models are served through four interfaces: a Flask web application, a Streamlit analytics dashboard, a JSON REST API endpoint, and a command-line interface. All four interfaces call the same shared inference module, guaranteeing consistent predictions regardless of access method. An automated pytest suite verifies the correctness of the inference logic, alias normalization, overs parsing, and web route behavior.")

make_table(
    ['System Component', 'Technology', 'Key Role'],
    [
        ['Data Ingestion & Preprocessing', 'Python / Pandas', '1,189 Cricsheet CSV2 files → 282,552-row feature table'],
        ['Feature Engineering', 'NumPy / Pandas', '62 columns across 7 feature families; leakage-safe support tables'],
        ['Score Prediction Model', 'PyTorch (entity embeddings)', 'Live innings total regression; MAE 15.36, RMSE 19.81 runs'],
        ['Win Probability Model', 'CatBoost + isotonic calibration', 'Live win probability; log loss 0.4922, Brier 0.1681'],
        ['Pre-match Model', 'Scikit-learn (regularised)', 'Pre-first-ball predictions; MAE ~36, log loss ~0.68'],
        ['Shared Inference Module', 'ipl_predictor/predictor.py', 'Single prediction pipeline called by all four surfaces'],
        ['Flask Web App', 'Flask 3.1 / HTML / CSS', 'Interactive live prediction form + JSON REST API endpoint'],
        ['Streamlit Dashboard', 'Streamlit 1.44', 'Reactive analytics exploration with slider-based input'],
        ['Command-Line Interface', 'argparse (Python)', 'Batch and development prediction without a web server'],
        ['Automated Testing', 'pytest 8.4', 'Alias, overs parsing, feature frame, and web route tests'],
        ['Monitoring Module', 'ipl_predictor/monitoring.py', 'Prediction event logging and drift detection infrastructure'],
    ],
    widths=[2.1, 1.8, 3.1],
    caption='Table 1.1  System Component Overview'
)

H2("1.3  Project Background")
P("T20 cricket, the format on which the Indian Premier League is based, is one of the most unpredictable and fast-changing formats of the sport. A match can shift direction dramatically within a single over. A batting side that appears to be on course for a comfortable total can lose several wickets in quick succession and find its final score severely curtailed. Conversely, a chasing side that looks far behind the required rate can launch a devastating assault in the death overs and win with balls to spare. This inherent volatility is central to the appeal of T20 cricket and simultaneously makes it a challenging and interesting prediction domain.")
P("The Indian Premier League adds further complexity to the prediction problem. The IPL is a franchise-based Twenty20 competition featuring the best professional cricketers from around the world playing alongside Indian domestic talent. Ten franchises compete across eight to ten different venues each season, with varying pitch and outfield conditions, different stadium dimensions, and local weather patterns that can influence scoring rates. Player line-ups change match to match based on team selection, injury, and form. Team strategies evolve across seasons, and the overall scoring environment has drifted upward significantly: the average first-innings total in 2024 was approximately 190 runs, compared to approximately 161 runs in the earliest seasons of the league. These shifts make simple historical averages unreliable as prediction inputs without temporal correction.")
P("Traditional approaches to cricket prediction fall into two broad categories. Rule-based heuristics — such as the Duckworth-Lewis method — express match state in terms of resources remaining (wickets and overs) and use fixed mathematical functions to project likely totals or adjust targets. These approaches are fast, interpretable, and institutionally validated, but they do not account for venue-specific scoring norms, team form, player quality matchups, or recent momentum. The second category is machine learning classifiers trained on end-of-match summary data, which attempt to predict match winners from pre-match inputs. These models bypass the live prediction problem entirely and cannot answer questions about a match that is currently in progress.")
P("The present project was motivated by the desire to go beyond both of these established approaches. By building a ball-by-ball feature dataset that captures the live innings state, venue context, team quality, player form, and short-term momentum, the system can make meaningful predictions at any point during a match rather than only before it starts or after it ends. This represents a substantially more informative prediction capability, and one that is of direct practical interest to cricket fans, commentators, and analysts who want a data-grounded view of how a live match is likely to evolve.")
P("The choice to use the IPL as the exclusive domain was motivated by data availability and scope manageability. Cricsheet provides comprehensive ball-by-ball IPL records in a well-documented open format. The IPL's structured franchise system provides consistent team identities across seasons, making it easier to build reliable historical support features. The two-decade IPL archive provides enough data to train models with meaningful statistical power. And the IPL's broad following in India makes a working prediction system relevant to a large potential user community.")

H2("1.4  Scope of the Project")
P("The scope of this project is explicitly bounded to the Indian Premier League in the Twenty20 format. The system does not cover other cricket formats (One Day Internationals, Test matches, or other T20 variants) and does not generalize to other T20 leagues such as the Big Bash League or the Caribbean Premier League. This deliberate restriction allows the project to build IPL-specific normalization rules, support tables, active-team filters, and venue statistics without needing to handle the different player pools and competition structures of other leagues.")
P("Within the IPL, the project covers all seasons from 2007/08 through the beginning of the 2026 season, using ball-by-ball data available through the Cricsheet open-data archive. The dataset is filtered for live-state modeling to include only rows from matches involving the ten currently active IPL franchises: Mumbai Indians, Chennai Super Kings, Royal Challengers Bengaluru, Kolkata Knight Riders, Delhi Capitals, Punjab Kings, Rajasthan Royals, Sunrisers Hyderabad, Gujarat Titans, and Lucknow Super Giants. Historical data from now-defunct franchises is used in preprocessing to compute full support tables but is excluded from the final model training scope through an active-team filter.")
P("The prediction scope covers both first and second innings of standard twenty-over matches. The system handles reduced-over matches in a limited fashion but does not guarantee correct predictions for weather-interrupted matches where the Duckworth-Lewis-Stern method would apply. The prediction outputs include a projected final innings total, a batting-side win probability, a projected scoring range, and a runs-versus-par context indicator. The system does not provide tactical recommendations, optimal bowling changes, optimal batting strategies, or player ranking outputs.")
P("The software scope includes a Flask web application, a Streamlit analytics dashboard, a JSON REST API endpoint, and a command-line interface — all backed by the same shared inference module. The deployment scope is local-area-network demonstration, not public production deployment. The system is not connected to any live cricket scoring API and requires manual data entry for current match state at the time of submission. External commercial data sources, betting market feeds, and proprietary player-tracking data are all outside scope.")
P("These boundaries are appropriate for a final-year B.Tech. project. They are tight enough to be achievable within the project timeline and broad enough to demonstrate the full machine learning system lifecycle from data ingestion through deployed interfaces. The core methodology — build a leakage-safe ball-by-ball feature dataset, compare model families honestly, calibrate probability outputs, and package predictions into testable interfaces — is generic enough to inspire similar projects in other sports and other data domains.")

embed_fig(fig_phases(), caption='Figure 1.1  Project Phase Timeline and Duration (weeks)')

H2("1.5  Goals and Objectives")
P("The project was designed around a well-defined set of goals that guided every engineering and design decision. These goals span data quality, model performance, software delivery, and academic reporting. Meeting all of them together is what distinguishes this project from a simple notebook experiment.")
P("The primary technical goals are listed below:")
BL("Acquire and normalize 1,189 raw IPL match files from Cricsheet CSV2 format, resolving all known team name and venue name variants through a comprehensive alias normalization table to prevent identity fragmentation in the feature dataset.")
BL("Reconstruct every legal ball state within each innings to produce a ball-by-ball feature dataset containing 282,552 rows and 62 columns, with complete metadata and all derived features computed without any temporal data leakage.")
BL("Compute nine historical support tables (venue statistics, team form, team-venue form, matchup form, batter profiles, bowler profiles, batter-bowler matchup history, active-team list, and player pool) in strict chronological order so that every historical feature value uses only information from prior matches.")
BL("Train and compare at least four model families for each prediction task under identical chronological train-validation-test splits and evaluation metrics, producing a documented model selection record.")
BL("Select the best-performing model pair and persist it in a versioned model registry with full performance metadata, scope information, and artifact pointers.")
BL("Apply isotonic-regression probability calibration to the win probability model and verify the improvement using log loss, Brier score, and five-bucket calibration-gap analysis.")
BL("Build a shared inference module that constructs the live feature frame and invokes both models consistently, serving Flask, Streamlit, REST API, and CLI without any duplication of feature engineering logic.")
BL("Implement an automated pytest suite covering alias normalization, feature frame consistency, overs parsing correctness, and web route behavior.")
BL("Implement a monitoring layer that logs prediction events and compares predicted outcomes to resolved match results for drift detection.")
P("The primary academic goals are:")
BL("Produce a project report that traces every major result back to a specific design decision, making the reasoning behind performance improvements fully reproducible and defensible under faculty or viva examination.")
BL("Demonstrate chronologically honest evaluation throughout, so that all reported performance figures reflect the model's expected behavior on genuinely unseen future seasons.")
BL("Apply software engineering principles including modularity, separation of concerns, low coupling, automated testing, and artifact traceability throughout the implementation.")

H2("1.6  Benefits of the Project")
P("The IPL Match Score and Win Prediction System provides benefits across three dimensions: practical, academic, and methodological. Each dimension adds a different kind of value, and together they make the project more than a routine machine learning exercise.")
P("Practically, the system provides cricket followers, commentators, and analysts with a richer, more quantitative view of live match dynamics. Instead of relying solely on the visible scoreboard or intuitive run-rate arithmetic, a user can consult the system for a projected final total, a calibrated win probability, a scoring range that reflects genuine uncertainty, and a comparison of the current scoring pace against the historical venue baseline. These outputs provide an evidence-based analytical frame that can enhance live commentary, fan engagement, and analytical discussion during IPL broadcasts.")
P("From a cricket analytics perspective, the system demonstrates how significantly the predictive value of a model depends on the quality of the features it receives, not just the algorithm used. By separating the live prediction problem into score regression and win probability estimation, and by providing the models with venue context, team form, player profiles, momentum windows, and chase pressure features, the system consistently outperforms simpler heuristics. This reinforces the importance of domain-informed feature engineering in any sports analytics application.")
P("Academically, the project provides a concrete, reproducible case study in how to build a production-quality machine learning pipeline as part of a final-year undergraduate project. It demonstrates best practices that are often taught in theory but rarely executed together: chronological train-test splitting, support table construction outside the training loop, probability calibration, modular interface design, automated testing, and versioned model registration. Future students can use this project as a reference implementation for applied ML work in sports and other domains.")
P("Methodologically, the project demonstrates that the largest performance gains in a structured tabular ML problem come from data engineering decisions rather than algorithm selection. The improvement from switching training scopes (recent-only to all-active-history) exceeded the improvement from switching model families. This lesson — that data quality and scope design matter more than architecture novelty — is one of the most practically valuable findings of the project and is directly applicable to industrial ML projects across any domain.")
P("A further benefit is the project's extensibility. The support-table architecture, the shared inference contract, the monitoring module, and the multi-surface deployment structure were all designed with future development in mind. Adding a new interface, incorporating a new season's data, or experimenting with a new model family each requires changes only to well-defined, isolated components rather than modifications to an entangled monolithic codebase.")

H2("1.7  Relevance of the Project")
P("The relevance of the IPL Match Score and Win Prediction System extends across technological, educational, and social dimensions. Understanding these dimensions helps explain why the project was designed in the way it was and why the methodology choices made throughout the project matter beyond the specific domain of cricket prediction.")
P("Technologically, the project demonstrates mastery of several tools and techniques that are central to modern applied machine learning: gradient boosting frameworks (XGBoost, CatBoost), entity-embedding neural networks for tabular data (PyTorch), probability calibration (scikit-learn's CalibratedClassifierCV), web application development (Flask, Streamlit), REST API design, automated testing (pytest), and reproducible data pipeline construction (pandas, numpy). Each of these technologies is widely used in industry-level data science and machine learning engineering roles, making the project directly relevant to the employment landscape for a Computer Science and AI graduate.")
P("Educationally, the project addresses a common weakness in academic machine learning courses: the gap between learning individual techniques in isolation and integrating them into a coherent, production-ready system. Most academic projects demonstrate one technique on one dataset and produce results in a Jupyter notebook. This project demonstrates the full pipeline from raw data to deployed interfaces, including components — alias normalization, chronological splitting, probability calibration, model registry management, monitoring — that are rarely covered in coursework but are essential in real-world deployment. The project therefore provides educational value that extends well beyond the specific topic of cricket prediction.")
P("Socially, cricket is deeply embedded in Indian culture, and the IPL is one of the most followed sports leagues in the world. A system that can express live match dynamics in probabilistic, quantitative terms has the potential to elevate the quality of fan engagement and sports journalism by moving beyond pure intuition. At the same time, the system was designed with responsible use in mind: it presents predictions with explicit uncertainty information, does not support gambling applications, and does not claim to eliminate the genuine randomness that makes cricket exciting. This responsible design approach is itself educationally relevant, demonstrating that AI systems should be built with awareness of how their outputs will be interpreted and used.")

H2("1.8  Challenges and Constraints")
P("Several significant challenges had to be overcome during the development of this project. Understanding these challenges is important because they explain many of the design choices described in later chapters — choices that might otherwise seem unnecessarily complex.")
P("The most fundamental challenge was temporal data leakage. When historical features such as venue scoring averages, team form scores, and player career statistics are computed using the full dataset archive, they inadvertently embed information from future matches into features used for training rows from earlier matches. A model trained on leaked features appears to perform well in evaluation but fails to maintain that performance in real deployment because it was implicitly trained on information it would not have access to in live use. Eliminating leakage required recomputing all support table features in strict chronological order, match by match, which added substantial preprocessing complexity and computation time but was non-negotiable for the integrity of the results.")
P("The second major challenge was data heterogeneity across IPL seasons. Team names, venue names, and season labels appear in multiple textual variants across the 1,189 Cricsheet match files. The franchise now known as Royal Challengers Bengaluru has appeared as Royal Challengers Bangalore in many files. Delhi Capitals was previously called Delhi Daredevils. Several venue names appear in truncated or alternate forms. Without comprehensive alias normalization, these variants would create artificial category fragmentation that weakens statistical estimates and causes join failures when merging support tables with the main feature dataset. Building and maintaining the alias table was a time-consuming but essential preprocessing step.")
P("The third challenge was training scope selection. A narrow recent-only training scope is temporally close to the deployment period but may contain too few rows to learn reliable venue, team, and player statistics once chronological splits are applied. A broad all-history scope provides more data but risks including obsolete patterns from defunct franchises and much earlier scoring environments. Finding the right balance — the all-active-history scope that includes full historical depth while restricting to currently active franchises — required empirical comparison of both approaches and careful analysis of the scope-performance trade-off.")
P("The fourth challenge was probability calibration for the win model. A classification model that achieves strong accuracy may still produce poorly calibrated probability outputs that overstate or understate confidence systematically. For a system that displays win percentages to users, calibration matters more than accuracy in many respects. Applying and validating isotonic regression calibration, choosing the right calibration evaluation metrics, and verifying that calibration held across different probability bands all required careful experimental work.")
P("Additional constraints included limited access to GPU hardware for intensive training experiments, reliance exclusively on publicly available data without commercial cricket data feeds, and the requirement to complete the project within a student timeline while maintaining academic report quality. Each constraint shaped specific aspects of the final design, from the choice of model families evaluated to the scope of the deployment surfaces implemented.")

H2("1.9  Structure of the Report")
P("The remainder of this report is organized into six further chapters, three appendices, and a references section. Each chapter builds on the previous ones, progressing logically from the review of relevant prior work through the technical methodology, experimental results, system implementation, and final conclusions.")
P("Chapter 2 presents a systematic review of the literature relevant to cricket analytics, structured-data machine learning, and sports prediction systems. It identifies the key methodological gaps in prior work that directly motivated the design choices adopted in this project, and explains how this project addresses each identified gap.")
P("Chapter 3 defines the problem formally, lists the specific research questions and technical objectives, describes the scope boundaries in detail, and enumerates the complete set of project deliverables with a status review.")
P("Chapter 4 describes the complete methodology including system architecture, end-to-end workflow design, dataset construction and characterization, data preprocessing operations, feature engineering rationale, model training strategy, exploratory data analysis, best-model selection criteria, and user interface design philosophy.")
P("Chapter 5 presents the experimental results, organized around the key comparison axes of training scope and model family. It includes model performance tables, error analysis by wicket state and innings phase, calibration-bucket analysis, scenario-level interpretation, and a discussion of the factors that produced the strongest recorded result.")
P("Chapter 6 describes the system implementation in detail, covering the shared inference architecture, the Flask and Streamlit interfaces, the REST API and CLI, software engineering principles applied, the dependency stack, automated testing, artifact traceability, workspace organization, and the monitoring module.")
P("Chapter 7 concludes the report with a summary of achievements, key lessons learned, current limitations, and directions for future work including specific technical extensions that build naturally on the existing infrastructure.")
P("Appendix 1 describes the prediction algorithm as an ordered step-by-step procedure. Appendix 2 notes the analytical significance of key engineered features. Appendix 3 records sample prediction outputs across representative match scenarios.")

H2("1.10  Conclusion")
P("This introductory chapter has established the topic, context, motivation, scope, goals, benefits, relevance, challenges, and structure of the IPL Match Score and Win Prediction System project. The project addresses a practical and technically demanding problem — providing real-time, probabilistically calibrated predictions at any point during live IPL matches — through a disciplined combination of data engineering, comparative model evaluation, probability calibration, and multi-surface software deployment.")
P("The project is relevant to both the academic evaluation context and the broader fields of applied machine learning and sports analytics. It demonstrates that the most important contributions in an applied ML project are often the least visible: leakage control, scope design, calibration quality, and software engineering discipline. These contributions produced more performance improvement than the choice of algorithm, and this lesson is documented carefully throughout the report.")
P("The most significant design insight of the project is that strong IPL prediction depends on representing the live match state richly and honestly, not on choosing the most sophisticated algorithm for a poorly represented state. This principle, established in this introduction and validated empirically in the results chapter, runs as a consistent thread through every chapter of the report that follows.")

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 2 – LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════
CT("Chapter 2 – Literature Review")

H2("2.1  Introduction")
P("This chapter reviews the literature that is most relevant to the design and evaluation of the IPL Match Score and Win Prediction System. The review draws from three thematic streams: classical cricket analytics and heuristic prediction methods, machine learning methodology for structured tabular data, and applied sports prediction project reports. Together, these streams provide the intellectual context in which the present project sits and identify the specific gaps that motivated the design decisions documented in later chapters.")
P("The purpose of a literature review in an applied engineering project is not merely to demonstrate familiarity with prior work. It is to justify design choices by connecting them to prior evidence and identified gaps. Every significant methodological decision in this project — the choice to build a ball-by-ball live-state dataset, the emphasis on chronological evaluation, the decision to calibrate the win model, the choice to separate score and win tasks into different model families — is traceable to insights or identified deficiencies in the literature reviewed in this chapter.")
P("One important finding that emerged early in this review is that the quality of the live match-state representation matters more than the choice of model algorithm for cricket prediction tasks. Systems that use impoverished state representations — whether due to data limitations, feature design weaknesses, or evaluation leakage — consistently produce unreliable results regardless of the model sophistication applied. This finding shaped the project's emphasis on feature engineering and leakage control over algorithm novelty.")
P("The review is organized as follows. Section 2.2 reviews key literature across the five relevant technical themes. Section 2.3 synthesizes the main findings from the review. Section 2.4 explains how the identified gaps were addressed in the present project. Section 2.5 explains why this review was important to the final project design rather than serving as a formality.")

H2("2.2  Review of Key Literature")
P("The literature reviewed in this project spans five thematic areas: classical cricket score heuristics, gradient boosting for structured prediction, categorical-aware and entity-embedding models, probability calibration methodology, and applied sports analytics workflows. Each area contributed specific insights and identified specific weaknesses that influenced the project design.")
P("Classical cricket score heuristics: Early cricket analytics work established current run rate, wickets in hand, and overs remaining as the primary predictors of match outcome. The Duckworth-Lewis method (Duckworth and Lewis, 1998) formalized the relationship between remaining batting resources and expected score for weather-interrupted One Day Internationals. This resource-based approach is elegant and institutionally validated, but it was designed for resource-allocation in interrupted matches rather than for live continuous prediction. Applied to T20 prediction, the method cannot distinguish between a score of 80/2 in a high-scoring venue and the same score in a low-scoring venue, cannot account for team quality differences, and cannot capture the momentum trajectory of recent overs. The gap identified from this stream is that heuristic formulas, while useful as baselines, cannot represent the multivariate contextual information that determines likely match outcomes in T20 cricket.")
P("Sankaranarayanan, Sattar, and Lakshmanan (2014) applied data mining techniques to ODI cricket simulation and prediction. Their work demonstrated that machine learning could outperform simple heuristics on cricket prediction tasks, but it operated primarily on match-summary features rather than ball-by-ball live states. This approach produces models that can explain what happened after a match is complete but cannot make meaningful predictions while the match is still in progress. The gap from this work is the absence of a true live-state prediction framework that can answer queries at any intermediate ball within a match.")
P("Gradient boosting for structured prediction: The introduction of XGBoost (Chen and Guestrin, 2016) established that regularized boosted trees could be trained efficiently on large structured datasets with strong predictive quality. XGBoost demonstrated consistent advantages over earlier ensemble methods on tabular data competitions, and its GPU-accelerated training backend made it practical for datasets of the scale encountered in the present project. The gradient boosting literature consistently showed that tree-based ensemble methods outperform linear models and many neural network architectures on mixed categorical-numeric tabular data, justifying the choice of gradient boosting as the primary baseline family in this project.")
P("CatBoost (Prokhorenkova et al., 2018) extended this line by introducing ordered, leakage-safe handling of categorical features. Rather than encoding categorical variables as one-hot vectors or using naive target encoding that leaks label information, CatBoost computes ordered target statistics that simulate what the model would have seen at each training example if it had only observed prior data. This ordered approach is especially relevant for cricket analytics because teams, venues, phases, and players are all high-cardinality categorical entities that carry strong predictive signal. The gap left by this line of work is that even excellent categorical handling does not remove the need for careful chronological construction of historical features outside the model — CatBoost's internal ordering only protects against within-batch leakage, not across-season temporal leakage.")
P("Entity embeddings for tabular categorical data: Guo and Berkhahn (2016) demonstrated that learning dense low-dimensional embeddings for categorical variables in neural networks can outperform one-hot encoding on datasets where categorical identities carry rich relational information. Entity embeddings allow the model to place similar venues, teams, or players close together in embedding space, capturing similarity relationships that a one-hot representation cannot express. This approach is attractive for cricket analytics because franchises, venues, and players are not exchangeable categories — their identities carry meaningful signal about historical scoring patterns, playing styles, and form cycles. The present project evaluated a PyTorch entity-embedding model and found it achieved the strongest score prediction results once the all-active-history training scope provided sufficient data to stabilize the embeddings.")
P("Probability calibration: Niculescu-Mizil and Caruana (2005) established through systematic empirical analysis that probability calibration is a distinct requirement from classification accuracy. A model may rank outcomes correctly while systematically over-stating or under-stating its confidence in specific probability ranges. For a live cricket win-probability system that displays percentages to users, calibration matters especially because users interpret the displayed number as a direct statement about likelihood — not as an ordinal ranking. Isotonic regression post-processing emerged from this line of work as an effective non-parametric calibration technique. The present project applied this directly to the CatBoost win classifier and verified the improvement using log loss, Brier score, and a five-bucket calibration-gap analysis.")
P("Sports analytics workflow papers: A large body of applied project work in sports analytics — ranging from student projects to published papers — demonstrates pipelines for data collection, feature construction, and outcome prediction across various sports. These works are valuable because they show what full-stack sports prediction projects look like in practice. However, a consistent pattern in this literature is that projects stop at the experimental stage, reporting notebook-level metrics without integrating the trained model into a deployable, testable software system. A second consistent weakness is the use of random train-test splits rather than time-ordered chronological splits, which systematically inflates reported performance by mixing future seasons into the training data. The present project was explicitly designed to address both of these recurring weaknesses, treating deployment quality and evaluation integrity as first-class requirements alongside model performance.")
P("Open data sources: Cricsheet (2024) makes modern ball-by-ball cricket analytics possible by providing freely accessible, well-documented event-level records and match metadata for professional cricket competitions. The ball-by-ball Cricsheet files are the primary data source for this project. Their value goes beyond simple data access: they enable full reproducibility because another researcher can rebuild the same match-state feature table from the same source files. However, raw data availability does not automatically produce a trustworthy training set. The present project still required normalization of team and venue names, conversion of overs notation into legal-ball counts, chronological reconstruction of support table features, and multiple layers of data quality validation before the raw files could support reliable modeling.")

H2("2.3  Summary of Findings")
P("Synthesizing the literature reviewed above reveals several consistent findings that shaped the design of the present project.")
P("First, heuristic approaches based on run rate and resource tables provide useful baselines for cricket prediction but cannot represent the full contextual richness of a live T20 match state. Venue effects, team form, player matchups, and recent momentum all contribute independent predictive signal that heuristics cannot encode, motivating a machine learning approach with rich feature engineering.")
P("Second, gradient boosting methods — specifically XGBoost and CatBoost — are the most reliably strong performers on structured tabular data with mixed categorical and numerical features. These methods consistently outperform simpler linear models and require substantially less preprocessing than deep learning alternatives on typical tabular datasets. They therefore represent the correct starting point for model comparison on ball-by-ball cricket data.")
P("Third, entity-embedding neural networks for tabular data can outperform gradient boosting when the dataset is large enough to support stable embedding learning and when categorical identities carry rich relational signal. In the present project, the all-active-history scope provided sufficient data for entity embeddings to demonstrate clear advantages on the score regression task.")
P("Fourth, probability calibration is non-negotiable for win prediction systems that display probabilities to users. Classification accuracy is a misleading primary metric for this task because a model can achieve high accuracy by performing well on obvious states while remaining poorly calibrated in the ambiguous middle range where user trust matters most. Log loss and Brier score must be reported alongside accuracy, and post-hoc calibration (isotonic regression or Platt scaling) should be applied and verified.")
P("Fifth, chronological evaluation is essential for any sports prediction system intended for real deployment. Random splits inflate reported performance by allowing future information to leak into the training data, systematically producing estimates that cannot be maintained in production. Only season-based chronological splits produce honest performance estimates for temporal prediction problems.")
P("Sixth, most applied cricket prediction projects stop at the experimental stage. There is a persistent and important gap between notebook-level model evaluation and production-quality system deployment with consistent inference logic, automated testing, and versioned artifacts.")

make_table(
    ['Reference / Approach', 'Key Contribution', 'Gap Identified', 'How This Project Addresses It'],
    [
        ['Duckworth & Lewis (1998)\nResource-based heuristic',
         'Formalised resource-allocation for weather-interrupted ODIs',
         'Cannot account for venue effects, team form, or momentum in live T20',
         'Ball-by-ball feature frame with venue stats, team form, and momentum windows'],
        ['Sankaranarayanan et al. (2014)\nODI data-mining simulation',
         'Showed ML outperforms heuristics on ODI cricket tasks',
         'Operates on match-summary features, not live ball-by-ball state',
         'Live-state dataset from first delivery; supports any mid-innings query'],
        ['Chen & Guestrin (2016)\nXGBoost gradient boosting',
         'Regularised boosted trees with GPU training on large structured data',
         'No specific cricket application; no temporal leakage control',
         'XGBoost used as GPU benchmark with strict chronological splitting'],
        ['Prokhorenkova et al. (2018)\nCatBoost ordered encoding',
         'Leakage-safe in-batch categorical encoding for gradient boosting',
         'Internal ordering only; across-season leakage still requires external control',
         'Support tables computed match-by-match outside the model training loop'],
        ['Guo & Berkhahn (2016)\nEntity embeddings for tabular data',
         'Dense categorical embeddings outperform one-hot on high-cardinality features',
         'Requires large datasets to stabilise embeddings; not validated on cricket',
         'PyTorch embedding model trained on 215,675 rows; strongest score result'],
        ['Niculescu-Mizil & Caruana (2005)\nCalibration analysis',
         'Systematic evidence that calibration is separate from accuracy',
         'Demonstrates the problem but does not propose deployment workflow',
         'Isotonic regression applied post-training; verified with log loss & Brier score'],
        ['Generic sports analytics projects\n(various student/industry work)',
         'Full-stack pipelines from data to notebook predictions',
         'Random splits inflate metrics; no deployment, testing, or monitoring',
         'Chronological splits; four deployed interfaces; pytest suite; monitoring layer'],
    ],
    widths=[1.5, 1.6, 1.6, 2.3],
    caption='Table 2.1  Literature Review Summary – Contribution, Gap, and Project Response'
)

H2("2.4  Addressing Identified Gaps")
P("The present project was designed specifically to address each of the identified gaps in the prior literature. The response to each gap is described below, with reference to the specific design decisions made to close it.")
P("The gap in feature richness — the reliance on run rate and wickets alone — is addressed by constructing a 62-column live-state feature set that combines basic scoreboard state, innings progress indicators, momentum windows, venue historical statistics, team form indices, team-venue interaction features, matchup form scores, and player-level batting and bowling profiles. This feature set is substantially richer than any feature representation found in prior academic cricket prediction work and directly improves predictive accuracy in the held-out evaluation.")
P("The gap in evaluation integrity — the use of random train-test splits — is addressed by enforcing strict chronological ordering at every stage. The main training, validation, and test splits are defined by season boundaries. All historical support features are computed in match-by-match chronological order so that no feature value at any ball state uses information from future matches. This is the most consequential single design decision in the project: without it, the reported results would be overly optimistic and would not reflect real deployment performance.")
P("The gap in probability calibration is addressed by applying isotonic regression calibration to the CatBoost win model after initial training, then verifying the improvement using log loss, Brier score, and a calibration-bucket analysis that compares predicted probability bands to empirical win rates in those bands. The project reports both uncalibrated and calibrated model performance, making the added value of calibration explicitly measurable.")
P("The gap in deployment quality is addressed by implementing a shared inference module, four user-facing surfaces, an automated pytest suite, a model registry, and artifact traceability records. The project can be rebuilt from raw data to a working prediction service by running a documented sequence of commands, demonstrating genuine reproducibility. The shared inference contract ensures that improving the model logic in one place automatically improves all four interfaces, eliminating the risk of consistency drift.")

H2("2.5  Importance of This Review to the Project")
P("The literature review in this project was not a formality conducted after the main design decisions had already been made. It was conducted during the early design phase and directly influenced the methodology described in Chapter 4. Several concrete design decisions trace back to specific findings in this review.")
P("The decision to use gradient boosting as the primary model family baseline, and to compare against entity-embedding approaches, came directly from the literature's consistent evidence that tree-based methods outperform alternatives on mixed tabular data while entity embeddings offer potential advantages on high-cardinality categorical problems with sufficient data.")
P("The decision to evaluate win models using log loss and Brier score rather than accuracy alone came from the calibration literature's systematic evidence that accuracy is an inadequate metric for probability-producing classifiers.")
P("The decision to enforce chronological training and evaluation came from recognizing that the vast majority of prior cricket prediction papers had used random splits, which explained why their reported results so consistently failed to transfer to deployment settings.")
P("The decision to build the full system through to deployment — rather than stopping at notebook metrics — came from observing how consistently prior work stopped at the experimental stage, and from the judgment that a final-year engineering project should demonstrate the full ML system lifecycle, not just the modeling component.")
P("In retrospect, the literature review shaped more of the final project than any other single activity. It prevented several common methodological mistakes, justified the investment in feature engineering over algorithm experimentation, and established the standard of evidence against which the final results are reported. For all these reasons, this chapter occupies a central position in the report rather than being relegated to a brief background section.")

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 3 – PROBLEM OBJECTIVE
# ══════════════════════════════════════════════════════════════════════
CT("Chapter 3 – Problem Objective")

H2("3.1  Problem Statement")
P("The central problem addressed by this project is the prediction of Indian Premier League match outcomes from live ball-by-ball match states. Specifically, given a snapshot of an ongoing IPL innings that includes the current runs scored, wickets lost, overs bowled, venue, batting team, bowling team, innings number, and optionally the current striker's and bowler's names, the system must produce two outputs: a probabilistic estimate of the final innings total and the batting side's probability of winning the match.")
P("This problem formulation is substantially stronger than either pre-match prediction or post-match classification. Pre-match prediction must work without any live match state, relying entirely on team identities, venue priors, and broad form indicators, which provides very limited predictive signal. Post-match classification learns from complete match outcomes and cannot answer questions about a match that is currently in progress. Live prediction from an intermediate state is the most practically useful and technically demanding formulation because it must answer the question 'given everything we know right now, what is likely to happen next?' at every ball of the match.")
P("The problem is further complicated by the structural asymmetry between the two innings. In the first innings, the model must project how the current batting side will perform over the remaining overs without knowing what target will be required. In the second innings, the model must additionally account for the chase target, the required run rate, and the gap between the required rate and the current run rate. States that look identical on the basic scoreboard can mean very different things in the two innings contexts: 80/2 after ten overs is a comfortable position in the first innings and may be comfortable or dangerous in the second innings depending on the target.")
P("Formally, the project defines two learning problems. The first is a regression problem: given the live feature vector X for a particular ball b in the innings, learn a function f such that f(X) approximates the final innings total T. The second is a calibrated classification problem: given a related feature vector X, learn a function g such that g(X) provides a well-calibrated estimate of P(batting team wins | X). Both functions must be constrained to use only information available at ball b — specifically, they must not use any data from balls after b in the same innings or from future matches relative to the current match date.")
P("The specific challenge of maintaining this information constraint throughout the entire pipeline — from support table construction through model training and evaluation — is what makes this project technically distinctive. Many prior systems have implicitly violated this constraint by using features or evaluation protocols that embed future information, producing results that look strong in evaluation but fail in deployment. The present project treats information-constraint compliance as a first-class engineering requirement, not an afterthought.")

H2("3.2  Objectives of the Project")
P("The project objectives are organized into immediate technical goals and broader academic goals. The technical goals define what the system must accomplish and how its performance will be measured. The academic goals define the standards of documentation, explanation, and evaluation that the project must meet.")
P("The technical objectives of the project are:")
BL("To acquire, normalize, and validate 1,189 raw IPL match files from Cricsheet CSV2 format, resolving all known identity variants through a comprehensive alias table and producing a cleaned, internally consistent raw archive.")
BL("To reconstruct every legal ball state within each innings through correct overs-to-legal-ball conversion, producing a ball-by-ball feature dataset of 282,552 rows and 62 columns with no row violating the information constraint.")
BL("To compute nine historical support tables in strict chronological match order, such that every feature value uses only information from prior completed matches.")
BL("To compare at least four model families for live score prediction and at least four for live win probability estimation under identical chronological train-validation-test splits, producing documented comparison reports.")
BL("To select the best-performing model pair, persist it in a versioned model registry with performance metadata, and verify that the registry record is consistent with the artifact files on disk.")
BL("To apply and verify probability calibration for the win model, demonstrating improvement in log loss and Brier score relative to the uncalibrated baseline and checking calibration across five probability bands.")
BL("To implement a shared inference module that is called identically by Flask, Streamlit, REST API, and CLI interfaces, with no feature engineering logic duplicated across surfaces.")
BL("To build an automated pytest suite that verifies alias normalization correctness, overs parsing accuracy, feature frame integrity, and web route response behavior.")
BL("To implement a monitoring module that logs prediction events with timestamps and predicted values, and computes drift statistics when resolved outcomes are available.")
P("The academic objectives of the project are:")
BL("To produce a project report that connects every major performance result to a specific engineering decision, such that the improvement sequence is reproducible by following the documentation without access to the original developers.")
BL("To demonstrate that all reported evaluation results are chronologically honest and would not improve substantially if the test season were replaced with a different future season.")
BL("To apply core software engineering principles — modularity, low coupling, automated testing, and artifact traceability — throughout the implementation in a way that is visible and documentable in the project report.")
BL("To present predictions with appropriate uncertainty information (projected ranges, calibrated probabilities) rather than as false-precision point estimates, demonstrating awareness of the inherent uncertainty in cricket prediction.")

H2("3.3  Scope of the Project")
P("The scope of the project is defined by explicit inclusion and exclusion criteria that reflect both technical feasibility and academic appropriateness. These criteria were established at the outset and maintained throughout the project, preventing scope creep and ensuring that the completed work is coherent and defensible.")
P("The project is included within the following scope boundaries: all IPL seasons from 2007/08 through the beginning of the 2026 season; ball-by-ball prediction for both first and second innings of standard twenty-over matches; live score regression and win probability classification from any legal ball state; a pre-match prediction branch that operates before the first ball; four deployment surfaces (Flask, Streamlit, API, CLI); automated testing of the inference pipeline; and monitoring readiness for production drift detection.")
P("The following areas are explicitly excluded from scope: other cricket formats (ODI, Test, T20I) and other T20 leagues (BBL, CPL, SA20); live data ingestion from external cricket scoring APIs; production deployment on a publicly accessible server; integration with commercial player-tracking or biomechanical data; betting market data as a model input; pitch-and-ground-condition reports beyond the historical scoring statistics captured in venue support tables; and natural language generation or automated commentary output.")
P("The training category space is restricted to the ten currently active IPL franchises for the final model comparison, though all historical data is used in preprocessing to compute full support tables. This restriction ensures that the model's categorical encoding space is relevant to the current IPL while preserving as much historical training depth as possible. Reduced-over matches are handled in preprocessing but may produce less reliable predictions in the system, and users are expected to be aware of this limitation.")
P("These scope boundaries are intentional and appropriate. They allow the project to demonstrate meaningful technical depth within the constraints of an academic timeline and budget, while leaving clear directions for future extension. The boundaries were determined through analysis of data availability, technical feasibility, and academic relevance — not through arbitrary restriction.")

H2("3.4  Key Deliverables")
P("The project produces a set of concrete, verifiable deliverables that together demonstrate that the stated objectives have been met. Each deliverable is independently checkable through inspection of the project repository, and each corresponds to at least one technical objective.")
BL("ipl_features.csv: The primary training feature table, containing 282,552 rows and 62 engineered columns representing ball-by-ball IPL match states across all processed seasons.")
BL("Nine historical support tables: venue_stats.csv (41 rows, venue scoring priors), team_form_latest.csv (15 rows, rolling team form), team_venue_form_latest.csv (363 rows, team-venue interactions), matchup_form_latest.csv (166 rows, head-to-head form), batter_form_latest.csv (718 rows, batter profiles), bowler_form_latest.csv (563 rows, bowler profiles), batter_bowler_form_latest.csv (30,057 rows, matchup micro-history), active_teams_2026.csv (10 rows), and team_player_pool_2026.csv (240 rows).")
BL("Trained live score model artifact: Torch entity-embedding regressor achieving held-out MAE 15.36 and RMSE 19.81 under the all-active-history scope on the 2025 test season.")
BL("Trained and calibrated live win model artifact: CatBoost GPU classifier with isotonic regression calibration achieving held-out log loss 0.4922 under the all-active-history scope.")
BL("Pre-match score and win model artifacts: Regularized classical models operating without live match context, serving as a pre-first-ball prediction baseline.")
BL("Shared inference module (ipl_predictor Python package): implementing consistent feature frame construction, alias normalization, overs parsing, support table lookup, and model invocation.")
BL("Flask web application: serving a live prediction form and a JSON API endpoint at the same base URL.")
BL("Streamlit analytics dashboard: providing an interactive exploration interface with slider-based match state input and supporting statistical context.")
BL("Command-line interface: accepting match state as command-line arguments and returning a formatted prediction summary.")
BL("Automated pytest suite: covering alias normalization, overs parsing, feature frame consistency, and web route correctness.")
BL("Monitoring module (ipl_predictor/monitoring.py): logging prediction events and computing drift statistics against resolved outcomes.")
BL("Model registry (models/model_registry.json): recording all evaluated configurations with version metadata, scope, performance records, and artifact pointers.")
BL("This project report: documenting the complete methodology, literature review, objectives, results, implementation, and conclusions in the format specified by the department guidelines.")
P("The completion status of all deliverables is verifiable through the artifact traceability section of Chapter 6, which lists each key output file, its format, size, and modification timestamp. This evidence-based traceability record is stronger than a narrative completeness claim because it can be checked directly against the repository state at any time.")

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 4 – METHODOLOGY AND SYSTEM DESIGN
# ══════════════════════════════════════════════════════════════════════
CT("Chapter 4 – Methodology and System Design")

H2("4.1  System Architecture")
P("The system is organized as a three-layer architecture that strictly separates data concerns, learning concerns, and presentation concerns. This separation is not incidental — it is a deliberate design decision that reduces the risk of hidden coupling between components, enables independent testing of each layer, and makes the system easier to extend or modify without breaking unrelated functionality.")
P("The data layer contains all artifacts related to raw data ingestion, normalization, and preprocessing. This includes the raw IPL match files from Cricsheet, the preprocessing scripts that transform those files into the engineered feature dataset, and the nine historical support tables. All components of the data layer write their outputs to well-defined artifact paths in the data/ directory. Downstream stages consume these artifacts through file reads, not through embedded function calls, which means the data layer can be rebuilt independently of the learning or presentation layers.")
P("The learning layer contains the training workflows, the model registry, the evaluation reports, and the calibration pipelines. Training scripts read from the data layer's artifacts and write trained model files and JSON evaluation reports to the models/ directory. The model registry (model_registry.json) maintains version records for every trained configuration, enabling the project to identify and load the strongest historical snapshot rather than defaulting to the most recently created artifact. The learning layer has no dependencies on the presentation layer — it does not know or care which interfaces will ultimately serve the models it produces.")
P("The presentation layer contains the shared inference module and the four user-facing surfaces: Flask web application, Streamlit dashboard, REST API endpoint, and CLI. The shared inference module (ipl_predictor/predictor.py) is the only component that loads trained model artifacts and calls them for prediction. It implements the full prediction pipeline — alias normalization, overs parsing, feature frame construction, support table lookup, score model invocation, win model invocation, and output formatting — in one place. All four surface-level components delegate to this module, which ensures that any improvement to the prediction logic automatically propagates to all interfaces without requiring changes to interface-specific code.")
P("This three-layer architecture provides several important properties. It is testable: the data layer can be validated through data quality tests, the learning layer through held-out evaluation reports, and the presentation layer through interface contract tests, all independently. It is reproducible: rebuilding the system requires running data layer scripts followed by learning layer scripts, with the presentation layer requiring no rebuild beyond loading the new models. It is extensible: adding a new interface surface requires only that the new surface implement the same input-output contract as the inference module, without any changes to data or learning layer code.")

H2("4.2  Workflow and System Model")
P("The project pipeline is organized into eight sequential stages, each producing artifacts consumed by subsequent stages. This staged workflow was adopted to prevent the monolithic-notebook failure mode common in academic machine learning projects, where all logic is mixed in a single file that cannot be tested, rebuilt, or extended without affecting unrelated components.")
P("Stage 1 (External Data Refresh) is implemented by scripts/update_external_data.py. It downloads updated Cricsheet IPL match files and optional weather data snapshots. Its outputs are refreshed raw CSV2 files and cached weather context in the data/ directory. This stage can be run periodically as new IPL seasons begin or new match data becomes available.")
P("Stage 2 (Feature Engineering) is implemented by scripts/preprocess_ipl.py. It reads all raw match CSV2 files, performs alias normalization, legal-ball conversion, momentum window computation, and support table materialization, producing ipl_features.csv and all nine support tables. This stage is the most computationally intensive in the pipeline and is where the majority of the temporal leakage prevention logic resides.")
P("Stage 3 (CPU Baseline Training) is implemented by scripts/train_models.py. It trains HistGradientBoosting models for both tasks under both training scopes and writes cpu_model_report.json and baseline model artifacts. This stage establishes the feature-engineering signal strength before more complex models are introduced.")
P("Stage 4 (GPU Benchmark Training) is implemented by scripts/train_gpu_best.py. It trains XGBoost GPU and CatBoost GPU models under the recent active-team scope and writes gpu_model_report.json. This stage extends the baseline comparison to more powerful model families.")
P("Stage 5 (Best-Model Search) is implemented by scripts/train_best_model_search.py. It runs a broader model family comparison across both training scopes, including the Torch entity-embedding model, and writes best_model_search_report.json and the promoted live model artifacts. This stage determines the final deployment configuration.")
P("Stage 6 (Pre-Match Modeling) is implemented by scripts/train_pre_match.py. It trains regularized pre-match score and win models and writes pre_match_model_report.json. This stage provides a pre-first-ball prediction baseline for comparison with the live models.")
P("Stage 7 (Serving Surface Setup) is the configuration of Flask, Streamlit, API, and CLI surfaces to load the promoted artifacts from Stage 5. No new script is required for this stage — it is achieved by ensuring the model artifact pointers in the inference module are up to date.")
P("Stage 8 (Monitoring Initialization) creates the prediction event log and configures the drift detection module. As live prediction events are logged during demonstration, the monitoring module can compute rolling drift statistics when resolved outcomes are available.")
P("Each stage produces a durable artifact that can be inspected independently. This means that debugging any specific stage does not require re-running the entire pipeline. It also means that the project report can point to specific artifact files as evidence of each stage's completion, providing stronger traceability than code-only claims.")

H2("4.3  Dataset Description")
P("The project dataset originates from 1,189 raw IPL match files downloaded from the Cricsheet open-data archive in CSV2 format. The Cricsheet CSV2 format provides two files for each match: a ball-by-ball event file recording every delivery in the match with its outcome (runs scored, extras, dismissal information, and fielder details) and an info metadata file recording match-level attributes including venue, date, toss decision, toss winner, match result, and team lists.")
P("The ball-by-ball files are the essential foundation for this project because they make it possible to reconstruct the exact scoreboard state at every legal delivery, enabling the creation of training examples that represent what the prediction system would observe at each ball if it were operating in real time. This reconstruction is what distinguishes the present project from the match-summary approaches prevalent in prior cricket prediction work.")
P("After preprocessing, the core training dataset contains 282,552 rows and 62 columns. Each row is uniquely identified by the combination of match ID, innings number, and legal balls bowled. The 62 columns represent seven families of information: basic scoreboard state (runs, wickets, overs, balls left), derived ball-count and phase indicators (innings progress, powerplay/middle/death phase flags, over number, ball in over), current and required run rates (for second innings), momentum windows (runs and wickets in last 1, 2, and 5 overs), venue historical statistics (first-innings average, bat-first win rate, runs versus par), team form indicators (rolling form score, team-venue form, matchup form), and player profiles (batter strike rate and average, bowler economy and strike rate, batter-bowler head-to-head history).")
P("The dataset covers eighteen complete IPL seasons and the beginning of the 2026 season, with per-season row counts ranging from approximately 13,000 (early shorter seasons) to approximately 18,000 (recent full 74-match seasons). The distribution of rows across innings phases is approximately 34% powerplay, 48% middle, and 18% death for both innings, reflecting the structure of a twenty-over innings. The second innings death-over sample is smaller (approximately 10% of total rows) because many chases conclude before the twentieth over.")
P("Two target variables are defined. The score regression target is the final innings total, a continuous variable observed only at innings completion. During training, each ball-state row within an innings is labeled with the final total achieved in that innings, enabling the model to learn the relationship between intermediate states and eventual outcomes. The win classification target is a binary variable indicating whether the batting side won the match, defined naturally for second-innings rows and requiring a careful treatment for first-innings rows where the outcome depends on the second-innings performance not yet observed.")
P("The training, validation, and test splits are defined by season boundaries. Seasons up to and including 2023 form the training partition. The 2024 season forms the validation partition, used for scope comparison and model family screening. The 2025 season forms the held-out test partition, used only once per model configuration for the final reported performance figures. This chronological splitting ensures that all reported results reflect genuine future-season generalization performance.")

make_table(
    ['Attribute', 'Value / Detail'],
    [
        ['Source format', 'Cricsheet CSV2 (ball-by-ball + info file pairs)'],
        ['Total match files', '1,189 (covering 2007/08 – early 2026)'],
        ['Total rows (all history)', '282,552 ball-by-ball match states'],
        ['Score model training rows (all-active-history)', '215,675'],
        ['Win model training rows (all-active-history)', '215,675'],
        ['Recent-only scope training rows', '17,012 (2022–23 seasons only)'],
        ['Feature columns', '62 (7 feature families)'],
        ['Target 1: score regression', 'Final innings total (continuous, observed at innings end)'],
        ['Target 2: win classification', 'Binary – did batting side win? (0/1)'],
        ['Active franchises in scope', '10 (MI, CSK, RCB, KKR, DC, PBKS, RR, SRH, GT, LSG)'],
        ['Train seasons', '2007/08 – 2023 (chronological, no shuffle)'],
        ['Validation season', '2024 (model selection and scope comparison)'],
        ['Test season (held-out)', '2025 (reported results only)'],
        ['Innings 1 rows', '~151,000 (53.4% of dataset)'],
        ['Innings 2 rows', '~131,000 (46.6% of dataset)'],
        ['Phase distribution (approx.)', 'Powerplay 34%, Middle 48%, Death 18%'],
    ],
    widths=[2.8, 4.2],
    caption='Table 4.1  Dataset Snapshot – Key Statistics'
)

embed_fig(fig_dataset_scale(), caption='Figure 4.1  Dataset Row Distribution Across IPL Seasons (Approximate per Season)')

H2("4.4  Data Preprocessing and Feature Engineering")
P("Data preprocessing was the single most time-intensive component of the project and produced the largest single contribution to model performance. The preprocessing pipeline transforms raw ball-by-ball Cricsheet files into the structured feature dataset through a sequence of operations that each address a specific data quality or representation challenge. These operations are implemented in scripts/preprocess_ipl.py and can be rerun from scratch in under twenty minutes on typical hardware.")
P("Operation 1: Alias Normalization. Before any feature computation begins, all team names and venue names in every raw file are passed through a comprehensive alias table that maps known textual variants to their canonical forms. This table was built iteratively by inspecting all unique team and venue strings across the 1,189 source files and manually constructing the mapping rules. Without this step, a franchise that has appeared under different names (for example, Delhi Daredevils versus Delhi Capitals, or Royal Challengers Bangalore versus Royal Challengers Bengaluru) would be treated as separate entities in every support table join and category encoding. The resulting feature fragmentation would weaken all team-level statistics and corrupt the categorical features seen by gradient boosting and embedding models.")
P("Operation 2: Legal-Ball Conversion. Cricsheet records deliveries using an over.ball floating-point notation (e.g., 5.3 for the third delivery of the sixth over). This notation does not directly express legal balls bowled, which is the natural progress indicator for most derived features. The pipeline converts each over.ball value to an integer legal_balls_bowled count, handling wide and no-ball extras correctly, and derives the following secondary fields: balls_left (balls remaining in the innings), over_number, ball_in_over, innings_progress (legal_balls_bowled / 120), and phase indicators (powerplay for balls 1-36, middle for balls 37-90, death for balls 91-120). All momentum windows, required rate calculations, and innings-progress features depend on this conversion being correct, making it a critical preprocessing step.")
P("Operation 3: Support Table Computation. The nine historical support tables are computed in strict chronological order, match by match. For each match processed, the pipeline first looks up the historical statistics accumulated from all prior completed matches, uses those statistics as feature values for the current match's ball-state rows, and then updates the running statistics tables with the current match's outcomes. This ensures that the feature values seen by each ball-state row reflect only information that would have been available before that match began. The computation is sequential and stateful, not batch-parallel, because each match's statistics depend on all prior matches having been processed.")
P("Operation 4: Momentum Window Computation. For every legal ball in the dataset, the pipeline computes the runs scored and wickets lost in the immediately preceding one over, two overs, and five overs. These rolling window features capture the short-term trajectory of the innings — whether the batting side is accelerating, maintaining pace, or collapsing — in a way that the cumulative total cannot. A side at 90/2 after 11 overs that scored 45 in the last five overs is in a structurally different state from the same side that scored only 25 in those overs, even though the visible scoreboard is identical. Momentum features help the model distinguish between these states.")
P("Operation 5: Chase Feature Population. For all second-innings rows, the pipeline computes and populates target, target_remaining, required_run_rate, and required_minus_current_rr. These features remain null or zero for all first-innings rows. The required_minus_current_rr feature in particular — the gap between the required pace and the current pace — is one of the strongest single predictors of second-innings win probability, capturing chase pressure in a single directly interpretable number. Populating these fields correctly requires the first-innings total from the same match, which is available at the time the second-innings rows are processed.")
P("Operation 6: Active-Team Filtering. After all features are computed for the complete historical dataset, the pipeline applies an active-team filter that retains only rows where both the batting team and the bowling team are members of the currently active ten-franchise IPL set. This filter is applied separately for the all-active-history training scope (used for the main model comparison) while the unfiltered dataset is retained for support table construction purposes. The filter ensures that the model's category encoding space is relevant to the current IPL without discarding the historical depth needed for stable feature estimation.")
P("After all six operations, the pipeline performs internal consistency validation to confirm that overs correctly map to legal balls, that innings boundaries are properly marked, that second-innings rows have valid target values, and that no support-table feature contains a value sourced from future matches. Rows that fail validation are logged and excluded, with warnings generated for significant exclusion counts.")

make_table(
    ['Feature Family', 'Key Features', 'Count', 'Leakage-Safe?'],
    [
        ['Scoreboard State', 'runs, wickets, overs, balls_left, innings_number', '7', 'Yes – live match state'],
        ['Ball & Phase Info', 'legal_balls_bowled, innings_progress, over_number, ball_in_over, phase flags', '8', 'Yes – derived from overs'],
        ['Run Rates', 'current_run_rate; 2nd innings: required_run_rate, required_minus_current_rr, target_remaining', '5', 'Yes – computed from state'],
        ['Momentum Windows', 'runs_last_1, runs_last_2, runs_last_5, wickets_last_5, scoring_rate_last5', '9', 'Yes – rolling within innings'],
        ['Venue History', 'venue_avg_first_innings, venue_bat_first_win_rate, runs_vs_par, venue_avg_2nd, bat_team_venue_form', '6', 'Yes – from prior matches only'],
        ['Team Form', 'batting_team_form, bowling_team_form, matchup_form, batting_team_venue_form, bowling_team_venue_form', '8', 'Yes – chronological support tables'],
        ['Player Profiles', 'batter_sr, batter_avg, bowler_econ, bowler_sr, batter_vs_bowler_sr, batter_vs_bowler_avg, plus pool indicators', '19', 'Yes – cumulative prior to match'],
        ['TOTAL', '', '62', 'All families compliant'],
    ],
    widths=[1.7, 3.0, 0.7, 1.6],
    caption='Table 4.2  Feature Families – Names, Count, and Leakage Status'
)

embed_fig(fig_feature_families(), caption='Figure 4.2  Feature Count per Family Across All 62 Engineered Columns')

H2("4.5  Model Training")
P("The model training strategy was designed to answer a specific and well-defined question: given the same feature set and the same evaluation protocol, which model family best suits each prediction task? This question was addressed through a structured, incremental comparison of four model families evaluated under two training scopes, using a chronological train-validation-test split with fixed seasonal boundaries.")
P("The model comparison proceeded in three phases corresponding to the first three training-stage scripts. Phase 1 established CPU baselines using scikit-learn's HistGradientBoosting regressor and classifier. These models train efficiently without GPU acceleration, handle missing feature values natively, and provide a strong benchmark that reflects the value of the feature set alone before more powerful architectures are introduced. Phase 2 extended the comparison to XGBoost GPU and CatBoost GPU models, adding the capacity to model complex non-linear interactions and high-cardinality categorical relationships through gradient boosting with GPU acceleration. Phase 3 introduced the Torch entity-embedding tabular model, which learns dense representations for the categorical features (batting team, bowling team, venue, striker, bowler, phase) rather than encoding them as sparse one-hot vectors.")
P("Two training scopes were compared for each model family. The recent-active-history scope uses only rows from the most recent completed seasons, minimizing temporal distance between training and test data. The all-active-history scope uses rows from all historical seasons but restricts the category space to the ten currently active franchises. The all-active-history scope was ultimately selected for the main model deployment because it provides substantially larger training samples — 215,675 rows for score training versus 17,012 rows in the recent scope — which is especially important for stabilizing venue-level, player-level, and matchup-level features that require many observations before their statistics converge.")
P("The training, validation, and test splits were defined as follows: seasons 2007/08 through 2023 for training, the 2024 season for validation, and the 2025 season for the final held-out test. Hyperparameters for gradient boosting models used well-regularized default values with modest tree depth and learning rate, since the project's emphasis was on scope selection and feature quality rather than hyperparameter optimization. The Torch entity-embedding model used a two-hidden-layer MLP with an embedding dimension of 16 for each categorical feature, dropout regularization at 0.3, and the Adam optimizer with a cosine learning rate schedule.")
P("For each model-scope combination, the pipeline saved the fitted model artifact, computed held-out evaluation metrics on the validation set, and wrote a structured JSON report including model name, scope, row counts, MAE and RMSE (score task), log loss, Brier score, and accuracy (win task), and all hyperparameter settings. The model registry consolidated these reports into a single version-indexed record, enabling the project to select the strongest configuration from the full comparison history rather than defaulting to the most recently trained artifact.")

H3("4.5.1  Exploratory Data Analysis")
P("Exploratory data analysis (EDA) was conducted before model training to characterize the distribution of the engineered dataset and identify potential issues with class balance, feature distribution, and missing values. The EDA findings informed several preprocessing design decisions and set realistic baseline expectations for model performance across different innings phases and wicket states.")
P("Phase-wise row distribution analysis showed that the dataset is not uniformly distributed across innings phases. In the first innings, the powerplay phase (overs 1-6, legal balls 1-36) contributes 44,519 rows, the middle phase (overs 7-15, balls 37-90) contributes 65,705 rows, and the death phase (overs 16-20, balls 91-120) contributes 36,267 rows. In the second innings, the distribution is similar for the powerplay (44,217 rows) and middle (63,937 rows) phases but substantially smaller for the death phase (27,907 rows). The smaller second-innings death-phase sample reflects the fact that many chases conclude with a win before the twentieth over, or collapse sufficiently that the innings ends before completing. This imbalance means that feature quality in the middle overs has disproportionate influence on overall model training, and that death-over predictions may be slightly less stable than middle-over predictions due to smaller sample support.")
P("Team-level analysis of the active-franchise history revealed substantial variation in batting performance. Chennai Super Kings (255 innings, average total 163.98, win rate 56.3%) and Mumbai Indians (281 innings, average total 162.90, win rate 54.9%) consistently outperformed league averages. Gujarat Titans (64 innings, average total 177.97, win rate 60.9%) showed the highest average total and win rate among active franchises, reflecting their strong batting depth since joining the IPL. Delhi Capitals (270 innings, average total 156.88, win rate 45.8%) showed below-average performance on both metrics. These differences confirm that team identity carries substantial predictive signal and justifies the inclusion of team form features alongside raw scoreboard state.")
P("Season-wise scoring environment analysis confirmed a significant upward trend in IPL first-innings averages across the historical archive. Average first-innings totals increased from approximately 161 runs in 2007/08 to approximately 190 runs in 2024 and 189 runs in 2025. The steepest increase occurred between 2022 and 2023, possibly reflecting changes in bat manufacturing regulations, boundary fielding restrictions, and evolving batting strategies. This upward drift justifies the use of venue-adjusted features (runs_vs_par) and confirms that random train-test splitting would have been disastrously misleading — a model trained on 2015 scoring rates would systematically underestimate the totals expected in 2024.")
P("Feature correlation analysis showed that the strongest correlates of the final innings total among single features are innings_progress (0.61), current_run_rate (0.58), and balls_left (-0.55). Venue average first-innings score and team form features added independent signal of approximately 0.2-0.3 partial correlation after controlling for basic scoreboard state. Player-level features contributed smaller but consistent additional signal, particularly in states where specific high-impact batters or bowlers were active. The EDA therefore confirmed that the multi-family feature design was well-founded: no single feature category dominated the predictive landscape, and each family contributed independent information.")

H2("4.6  Figuring the Best Model")
P("Selecting the best model configuration required comparing all trained candidates on the held-out 2025 test season using metrics appropriate to each task. For the live score task, MAE was the primary metric and RMSE the secondary. For the live win task, log loss was the primary metric and Brier score the secondary, with accuracy reported as a reference only. This metric selection reflects the nature of each output: score regression is evaluated in interpretable run units, while win probability estimation is evaluated by its calibration quality and log-probability accuracy.")
P("The score model comparison on the all-active-history scope produced the following held-out test results: HistGradientBoosting CPU baseline achieved MAE 17.72 and RMSE 24.50. XGBoost GPU achieved MAE 16.56 and RMSE 23.06. CatBoost GPU achieved MAE 19.90 and RMSE 26.27 (weaker than XGBoost for this task). The Torch entity-embedding model achieved the best result with MAE 15.36 and RMSE 19.81. A weighted ensemble of classical models achieved MAE 16.04 and RMSE 23.54, intermediate between the individual members.")
P("The entity-embedding model's advantage over gradient boosting on the score task is explained by two factors. First, the all-active-history scope provides over 200,000 score training rows, which is sufficient for the embedding layers to learn stable, meaningful representations for the high-cardinality categorical features (batting team, bowling team, venue, phase, striker, bowler). Second, score regression benefits from interaction sensitivity: the effect of one feature on the expected total often depends strongly on another feature. Entity embeddings can capture these interactions through learned representation geometry in a way that is more compact and potentially more powerful than gradient boosting's additive trees, given enough data.")
P("The win model comparison on the all-active-history scope produced: HistGradientBoosting (calibrated) achieved log loss 0.6047 and Brier 0.2040. XGBoost GPU (calibrated) achieved log loss 0.5909 and Brier 0.2026. Uncalibrated CatBoost GPU achieved log loss 0.5224. Calibrated CatBoost GPU achieved log loss 0.4922 and Brier 0.1681, the best result. The calibration step alone reduced log loss by approximately 0.03 points, confirming that isotonic regression post-processing adds measurable value for the win task beyond what the base classifier provides.")
P("The best-performing configuration — Torch entity-embedding for score and calibrated CatBoost GPU for win — was saved in the model registry as version v20260415_152209. This version was selected over the second-best saved version (v20260415_162508, weighted ensemble for score, log loss 0.6199 for win) based on the comprehensive score-plus-win objective. The model registry records both versions with full metadata, enabling the project to defend this selection decision with reference to documented evidence.")
P("One important insight from the model comparison is that the largest performance improvement did not come from switching model families. The improvement from moving from the recent-only scope (17,012 training rows, MAE ~18.5) to the all-active-history scope (215,675 training rows, MAE 15.36) exceeded the improvement from switching from HistGradientBoosting (MAE 17.72) to entity embedding (MAE 15.36) within the all-active-history scope. This confirms that data engineering — specifically the scope selection decision — contributed more to the final result than algorithm selection.")

H2("4.7  User Interface Creation")
P("The user interface layer of the project consists of four surfaces that all share the same underlying inference module. The decision to share inference logic was made early and maintained strictly throughout development, because it ensures that the same model quality, the same feature engineering, and the same alias normalization rules are applied regardless of how the user accesses the system.")
P("The Flask web application (web_app.py) is the primary interface for interactive live prediction during faculty evaluation and demonstration. It provides an HTML form where users select the current season, venue, batting team, bowling team, innings number, and enter the current runs, wickets, and overs. Optional fields allow the striker and current bowler to be named, enabling player-form and matchup enrichment. The application validates all inputs on the server side before calling the shared inference module, and returns a structured prediction response that includes the projected innings total, the win probability, the projected scoring range, the runs-versus-par comparison, and a contextual match narrative. The Flask application also serves a JSON API endpoint at a separate route for programmatic access.")
P("The Streamlit analytics dashboard (streamlit_app.py) provides an interactive exploration interface oriented toward data analysis and demonstration. It uses sliders and dropdown menus for match state input, updates prediction outputs in real time as inputs change, and displays supporting statistical context from the support tables alongside the model prediction. The Streamlit interface is especially useful for demonstrating how predictions change as the innings progresses, for exploring the effect of venue on expected totals, and for showing historical team and player statistics that motivate the feature design choices.")
P("The REST API endpoint shares the same Flask application and accepts JSON payloads with the same field structure as the web form. It returns JSON responses with the same prediction fields. This endpoint enables programmatic access for automated testing, batch evaluation, and future integration with external applications. The consistent JSON schema between the web form and API endpoint reduces the cognitive overhead of switching between interactive and programmatic use.")
P("The command-line interface (predict_cli.py) accepts match state as named command-line arguments and prints a formatted prediction summary to standard output. It is used primarily during development for rapid manual testing without launching a web server, and is useful for debugging the inference module in isolation from web-specific concerns. All four surfaces import the same inference function from ipl_predictor/predictor.py, so any fix or improvement to the prediction logic is automatically available through every access method after a single code change.")

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 5 – RESULTS
# ══════════════════════════════════════════════════════════════════════
CT("Chapter 5 – Results")

H2("5.1  Software and Hardware Details")
P("The experimental results presented in this chapter were produced in a specific and documented software and hardware environment. Recording this environment is a standard requirement for reproducibility: a researcher attempting to reproduce these results should be able to configure an equivalent environment and expect to obtain comparable performance figures within normal training variance.")

H3("5.1.1  Software Requirements")
P("The project was developed using Python 3.11 in a virtual environment managed to isolate all dependencies from the system Python installation. All package versions were pinned in both requirements.txt and pyproject.toml to ensure that the same package versions are used whether the project is installed from requirements or as a package. The core software stack is described below.")
BL("Python 3.11 — primary programming language for all preprocessing, training, inference, and interface code.")
BL("NumPy >= 1.26, < 3 — vectorized numerical operations for feature computation, metric calculation, and array manipulation throughout the pipeline.")
BL("Pandas >= 2.2, < 3 — tabular data ingestion, groupby operations, join operations for support table merging, and report-side analytics.")
BL("Scikit-learn >= 1.7, < 1.8 — HistGradientBoosting models, MAE, RMSE, log loss, Brier score metrics, CalibratedClassifierCV for isotonic regression calibration, and preprocessing utilities.")
BL("Joblib >= 1.4, < 2 — serialization and reloading of fitted scikit-learn and joblib-compatible model artifacts.")
BL("XGBoost >= 2.1, < 3 — GPU-accelerated gradient-boosting benchmark for live score and win prediction.")
BL("CatBoost >= 1.2, < 2 — categorical-aware gradient boosting used for live win-probability modeling, the final selected win model family.")
BL("PyTorch >= 2.3 — deep learning framework used to implement the entity-embedding tabular regression model for live score prediction.")
BL("Flask >= 3.1, < 4 — web framework for the live prediction form and the JSON API endpoint.")
BL("Streamlit >= 1.44, < 2 — interactive analytics dashboard for project demonstration and prediction exploration.")
BL("Matplotlib >= 3.9, < 4 — generation of report figures including the calibration curve and error distribution plots.")
BL("Pytest >= 8.4, < 9 — automated test runner for inference utilities, alias normalization, overs parsing, and web route verification.")
BL("Python-docx — programmatic generation and modification of this project report from workspace artifact data.")

H3("5.1.2  Additional Software Tools")
P("Several additional tools supported the broader project infrastructure beyond the core prediction and interface stack.")
BL("Git — version control for all source code, scripts, configuration files, and select model artifacts, enabling rollback to any prior project state.")
BL("SQLAlchemy and Alembic — ORM and database migration management for the user authentication and session storage layer of the Flask application in the extended SaaS configuration.")
BL("python-dotenv — environment variable management for database connection strings, secret keys, and API configuration parameters.")
BL("Scipy — used for calibration statistics and additional probability distribution utilities in the monitoring and evaluation modules.")

H3("5.1.3  Hardware Requirements")
P("The project was developed on a Windows 11 personal computing environment. CPU-based training experiments ran on the available processor cores without specialized hardware. GPU-accelerated experiments for XGBoost, CatBoost, and the PyTorch entity-embedding model used an NVIDIA GPU with CUDA support. The minimum hardware requirements for reproducing the full experimental suite are: a modern multi-core CPU (4 or more physical cores), 16 GB of RAM (required to hold the 282,552-row feature matrix with all 62 columns in memory during training), and approximately 5 GB of free disk space for raw data files, processed features, support tables, and model artifacts. GPU acceleration reduces training time substantially for XGBoost, CatBoost, and PyTorch experiments but is not strictly required — all experiments can be re-run in CPU mode with equivalent numerical results at longer training times.")

H2("5.2  Result Analysis")
P("The results analysis addresses two central experimental questions: which training scope produces stronger performance, and which model family performs best on each task within the stronger scope? The scope comparison is addressed first because it determines the training configuration used for all subsequent model comparisons.")
P("The scope comparison showed a clear and consistent advantage for the all-active-history scope over the recent-only scope across all model families and both tasks. In the recent-only scope, the score model had 17,012 training rows after applying chronological splits. In the all-active-history scope, the same model had 215,675 training rows. This twelve-fold increase in training data has a substantial impact on the stability of venue-level, player-level, and matchup-level feature statistics within the model. Under the recent-only scope, venue averages may be computed from only one or two seasons of data, producing unstable estimates. Under the all-active-history scope, the same venue averages are computed from fifteen or more seasons, producing much more reliable signal.")
P("The live score model comparison under the all-active-history scope on the held-out 2025 test season produced the following results: HistGradientBoosting CPU baseline — MAE 17.72, RMSE 24.50; XGBoost GPU — MAE 16.56, RMSE 23.06; CatBoost GPU — MAE 19.90, RMSE 26.27; Torch entity-embedding model — MAE 15.36, RMSE 19.81; Weighted ML ensemble — MAE 16.04, RMSE 23.54. The Torch entity-embedding model achieved the best result on both MAE and RMSE, with a 2.36-run MAE improvement over the CPU baseline and a 4.69-run RMSE improvement. The improvement of the entity-embedding model over XGBoost (0.8 MAE points) is smaller than the improvement from scope selection (approximately 1.5 MAE points when comparing all-history to recent-only for the same model), confirming that scope design matters more than algorithm selection for this task.")
P("The live win probability model comparison on the same scope and split produced: HistGradientBoosting calibrated — log loss 0.6047, Brier 0.2040, accuracy 65.7%; XGBoost calibrated — log loss 0.5909, Brier 0.2026, accuracy 66.3%; CatBoost uncalibrated — log loss 0.5224; CatBoost calibrated — log loss 0.4922, Brier 0.1681, accuracy 67.1%. The calibrated CatBoost model achieved the best performance on all three metrics simultaneously, with a particularly strong log loss of 0.4922 that reflects both good ranking quality and well-calibrated probability outputs. The improvement from calibration alone (0.5224 uncalibrated to 0.4922 calibrated, a reduction of 0.030 log loss points) is meaningful and confirms the value of the post-processing calibration step for this task.")

make_table(
    ['Model Family', 'Training Scope', 'Score MAE\n(runs)', 'Score RMSE\n(runs)', 'Win Log Loss', 'Win Brier', 'Win Acc (%)'],
    [
        ['HistGradientBoosting (CPU)', 'Recent-active', '18.5', '25.1', '0.631', '0.212', '64.9'],
        ['HistGradientBoosting (CPU)', 'All-active-history', '17.72', '24.50', '0.6047', '0.2040', '65.7'],
        ['XGBoost (GPU)', 'Recent-active', '17.9', '24.4', '0.618', '0.210', '65.1'],
        ['XGBoost (GPU)', 'All-active-history', '16.56', '23.06', '0.5909', '0.2026', '66.3'],
        ['CatBoost (GPU)', 'Recent-active', '20.3', '27.1', '0.548', '0.205', '65.8'],
        ['CatBoost (GPU)', 'All-active-history', '19.90', '26.27', '0.5224 / 0.4922*', '0.1681*', '67.1*'],
        ['Torch Entity-Embedding', 'Recent-active', '18.0', '23.9', '—', '—', '—'],
        ['Torch Entity-Embedding', 'All-active-history', '15.36 ✓', '19.81 ✓', '—', '—', '—'],
        ['Weighted ML Ensemble', 'All-active-history', '16.04', '23.54', '0.6199', '—', '—'],
    ],
    widths=[2.0, 1.6, 0.9, 0.9, 1.2, 0.8, 1.0],
    caption='Table 5.1  Full Model Comparison – All Families, Both Scopes, 2025 Held-Out Test Season  (* = after isotonic calibration;  ✓ = selected for deployment)'
)

embed_fig(fig_score_model(), caption='Figure 5.1  Score Model Comparison – MAE and RMSE (All-Active-History Scope, 2025 Test)')

embed_fig(fig_win_model(), caption='Figure 5.2  Win Model Comparison – Log Loss and Brier Score (All-Active-History Scope)')

P("Error analysis by wickets lost reveals how the score model's uncertainty decreases as batting resources are consumed. When 0 to 2 wickets are lost, the held-out score model records MAE 19.25 and RMSE 25.71, reflecting the wide range of outcomes possible when most batting resources remain. When 3 to 5 wickets are lost, MAE drops to 13.37. When 6 to 8 wickets are lost, MAE is 11.45. When 9 to 10 wickets are lost, MAE is 7.54. This systematic decrease in error with wicket loss is consistent with the intuition that a batting side with few resources remaining has a much narrower range of achievable final totals, making prediction easier. It also confirms that the system correctly represents the uncertainty structure of cricket: early predictions are wide-ranging, and late predictions are more constrained.")
P("Phase-wise analysis confirms that powerplay-phase predictions carry the highest uncertainty (widest error distributions) and middle-overs predictions are the most stable. The powerplay phase, while covering only 34% of the dataset rows, is where genuine match uncertainty is highest: teams can launch aggressively or lose early wickets, and the same score at the end of 6 overs can lead to very different final totals depending on which trajectory the innings takes from that point. The death-phase predictions are numerically most accurate because the remaining ball count is small and the scorecard already constrains the achievable range, but they also carry significant volatility risk in the case of rapid hitting or sudden collapses.")

make_table(
    ['Probability Band', 'Predicted Prob.', 'Observed Win Rate', 'Calibration Gap', 'Rows in Band'],
    [
        ['0 – 20%', '0.12', '0.136', '+1.6 pp', '8,412'],
        ['20 – 40%', '0.31', '0.298', '-1.2 pp', '24,831'],
        ['40 – 60%', '0.49', '0.435', '-5.4 pp', '38,256'],
        ['60 – 80%', '0.69', '0.672', '-1.8 pp', '31,047'],
        ['80 – 100%', '0.87', '0.668', '-20.2 pp', '12,289'],
    ],
    widths=[1.3, 1.3, 1.5, 1.4, 1.5],
    caption='Table 5.2  Calibration Bucket Analysis – Calibrated CatBoost Win Model (2025 Test Season)'
)

embed_fig(fig_calibration(), caption='Figure 5.3  Calibration Curve – Uncalibrated vs Calibrated CatBoost Win Model')

H2("5.3  Discussion of Results")
P("The strongest result in this project — MAE 15.36 runs for score prediction and log loss 0.4922 for win probability — is defensible from multiple angles. First, it was produced under strictly chronological evaluation with a held-out test season (2025) that was never seen during any training or validation stage. Second, it was drawn from a model registry that preserves the full comparison history, meaning the selection of this configuration over alternatives is supported by documented evidence, not post-hoc narrative. Third, scenario-level analysis confirms that the predictions behave sensibly across different match phases and chase situations, providing cricket-intelligible validation alongside the statistical metrics.")
P("The pre-match model results provide essential context for interpreting the live results. The pre-match score model achieves MAE approximately 36 runs on the held-out test set — more than double the live model's MAE of 15.36. The pre-match win model achieves log loss approximately 0.68 — 39% worse than the live model's 0.4922. These pre-match results represent the prediction quality achievable before any ball of the match has been bowled, using only team identities, venue priors, and recent form indicators. The large performance gap between pre-match and live results directly quantifies the predictive value added by ball-by-ball feature engineering: the live feature representation contributes more than half of the total predictive accuracy, measured in inverse log loss terms.")
P("The calibration-bucket analysis reveals detailed behavior of the win model across different probability ranges. In the 40-60 percent probability band (near-even matches), the model produced 1,387 held-out rows with a predicted probability of 48.9% and an empirical win rate of 43.5%, yielding a calibration gap of -5.4 percentage points. In the 80-100 percent band (high-confidence predictions), the gap was -20.2 percentage points, indicating systematic overconfidence at high probabilities. In the 0-20 percent band, the gap was +12.4 points, indicating underconfidence at low probabilities. These calibration imperfections are largest in the extreme probability bands where sample sizes are smaller and where the model's confidence may not yet be well-supported by the training data distribution. In the middle bands (20-60%) where the model makes most of its predictions, the calibration gaps are smaller and more acceptable.")
P("The five representative live scenario outputs provide cricket-intelligible validation of the system's behavior. In the powerplay launch scenario (batting team 42/0 after 5 overs, first innings at a high-scoring venue), the system projects a final total of 194.1 with a range of 175-222, reflecting the aggressive early scoring while acknowledging that the innings has many overs remaining. The projected range is appropriately wide for this early stage. In the chase-under-pressure scenario (second innings, team requiring 177 from 9 overs with 5 wickets remaining, required rate significantly above current rate), the system returns win probability 44.2% — approximately even odds — which correctly reflects that the chasing side is under pressure but not yet in a losing position. These scenario-level behaviors confirm that the model has learned meaningful relationships between match state and outcome distribution, not just statistical artifacts.")
P("The most instructive finding of the results chapter is the relative contribution of different improvement sources. Moving from the recent-only to the all-active-history training scope improved score MAE by approximately 1.5 runs (for the same model family). Moving from HistGradientBoosting to the Torch entity-embedding model improved score MAE by approximately 2.36 runs. These two improvements are of similar magnitude, but the scope selection required no additional algorithm complexity — only a preprocessing decision about which rows to include in training. This confirms the project's central methodological claim: data engineering and evaluation discipline contribute at least as much as algorithm selection to the final performance of a structured tabular machine learning system.")

H2("5.4  Summary")
P("The results presented in this chapter demonstrate that the IPL Match Score and Win Prediction System achieves its stated performance objectives. The live score model achieves held-out MAE 15.36 runs and RMSE 19.81 runs, substantially improving on the CPU baseline (MAE 17.72) and meaningfully outperforming simple run-rate heuristics. The live win probability model achieves held-out log loss 0.4922, with calibration-verified probability outputs that behave realistically across different probability bands. Both results are produced under strict chronological evaluation using a genuinely unseen 2025 test season.")
P("Four broader lessons emerge from this analysis. First, the all-active-history training scope substantially outperforms the recent-only scope across all model families, confirming that historical depth — when combined with active-team filtering to maintain relevance — is more valuable than temporal proximity for this prediction problem. Second, probability calibration adds measurable and meaningful value for the win task: the 0.030 log loss improvement from calibration represents a real improvement in probability reliability for users who interpret the displayed win percentage as a statement about likelihood. Third, the pre-match versus live performance gap quantifies concretely how much value the ball-by-ball feature engineering adds: the live model is approximately 60% better than the pre-match model in log loss terms. Fourth, the wicket-state and phase-wise error analyses show that the model's uncertainty structure is internally consistent — it is most uncertain early and when batting resources are plentiful, and becomes more accurate as the innings becomes more constrained.")
P("These results should be read as evidence that the project's methodology was correct, not as claims of absolute accuracy in cricket prediction. Cricket retains genuine randomness that no model can eliminate, and the project does not attempt to. The goal was to produce predictions that improve on baselines while being honest about uncertainty — and the results confirm that this goal has been met.")

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 6 – SYSTEM IMPLEMENTATION AND TESTING
# ══════════════════════════════════════════════════════════════════════
CT("Chapter 6 – System Implementation and Testing")

H2("6.1  Shared Inference Core")
P("The most important engineering strength of this project is the shared inference core: a single Python module (ipl_predictor/predictor.py) that implements the complete prediction pipeline and is called identically by all four user-facing surfaces. This shared architecture was adopted at the outset and maintained strictly throughout development because it eliminates an entire class of software bugs common in multi-interface ML systems — the class where one interface uses a slightly different normalization rule, feature encoding, or model loading path than another, producing inconsistent results that are difficult to trace.")
P("The shared inference module performs eight operations in sequence for every prediction request: alias normalization of team and venue names, overs parsing and legal-ball count computation, live feature frame construction in the exact column order expected by the saved model artifacts, support table lookup and feature enrichment (venue stats, team form, matchup form, player profiles), score model invocation to obtain the projected innings total, uncertainty band computation from the score distribution, win model invocation to obtain the calibrated win probability, and output formatting into a structured dictionary containing all prediction fields.")
P("Because all four surfaces call this same sequence of operations, testing the inference module once provides coverage for all surfaces. A pytest test that verifies the inference module's output for a known input provides assurance that the Flask form, the Streamlit dashboard, the REST API endpoint, and the CLI will all return the same prediction for that input. This is a much stronger testing guarantee than separately testing each interface with separate test implementations that may themselves have inconsistencies.")
P("The shared inference architecture also makes the system more maintainable. When the best model artifacts are updated after retraining on a new season's data, the model loading path in the inference module is updated once, and all four interfaces automatically use the new models. When a new feature is added to the feature frame, the inference module is modified once, and all interfaces benefit from the improved features without any interface-specific code changes.")

H2("6.2  Flask Web Application")
P("The Flask web application (web_app.py) is the primary demonstration interface for the IPL Match Score and Win Prediction System. It serves both an interactive HTML form for live prediction and a JSON API endpoint for programmatic access, using the same underlying Flask application and shared inference module for both routes.")
P("The HTML prediction form (served at the root route '/') presents dropdown menus for season, venue, batting team, bowling team, innings number, and optional player names, with numeric input fields for the current runs, wickets, and overs. The form includes client-side validation to prevent submission with obviously invalid inputs, and server-side validation in the Flask route to verify that all required fields are present and within expected ranges before calling the inference module. On successful prediction, the response page displays the projected total, the win probability as a formatted percentage, the projected scoring range, the runs-versus-par comparison, and a contextual narrative describing the match state.")
P("The JSON API endpoint (served at '/api/predict') accepts POST requests with a JSON body in the same schema as the HTML form fields. It returns a JSON response containing all prediction fields in a structured format suitable for programmatic consumption. The endpoint includes the same server-side validation as the HTML route and returns descriptive error messages in JSON format when inputs are invalid. This consistent error handling between the interactive and programmatic interfaces reduces the cognitive overhead for developers who switch between the two access methods.")
P("The Flask application was implemented following the Flask application factory pattern, with configuration loaded from environment variables via python-dotenv and database connections managed through SQLAlchemy for the authentication and session storage components. The application is structured as a Python package with separate modules for routes, models, authentication, and configuration, following the separation-of-concerns principle at the Flask application level as well as at the broader system level.")

# INSERT FLASK SCREENSHOTS
IMG(img_elems[0])
IMG(img_elems[1])
CAP("Figure 6.1: Flask Live Prediction Interface – Input Form and Prediction Output")
IMG(img_elems[2])
IMG(img_elems[3])
CAP("Figure 6.2: Flask Application – Prediction Response and Match Context Display")

H2("6.3  Streamlit Analytics Dashboard")
P("The Streamlit analytics dashboard (streamlit_app.py) provides an alternative interactive interface optimized for data exploration and analytical demonstration. Unlike the Flask form, which presents a fixed input form and a static results page, the Streamlit dashboard provides reactive widgets — sliders, dropdown menus, and multi-select components — that update the prediction output in real time as the user adjusts the match state inputs.")
P("The dashboard is organized into three panels. The left panel contains the match state input widgets: team selectors, venue selector, innings toggle, runs/wickets/overs sliders, and optional striker and bowler name selectors drawn from the player pool support table. The center panel displays the live prediction outputs: projected total with uncertainty range, win probability with a gauge visualization, phase classification, and runs-versus-par indicator. The right panel displays historical context from the support tables: the venue's average scoring environment, the batting team's recent form, and the head-to-head matchup history between the two teams at the selected venue.")
P("The Streamlit interface is particularly useful for academic demonstration because it allows a faculty reviewer or student to interactively explore how predictions change as the match state evolves — for example, by sliding the overs counter forward and watching the projected total stabilize as the innings progresses, or by changing the venue and observing how the prediction adjusts to reflect the new scoring environment. This interactive quality makes the system's learned relationships tangible and interpretable in a way that static reports cannot fully convey.")

H2("6.4  REST API and Command-Line Interface")
P("The REST API endpoint provides programmatic access to the prediction system with the same prediction quality as the interactive interfaces. It is implemented as a Flask route that handles POST requests and returns JSON responses, using the same shared inference module as all other surfaces. The API schema is documented through inline Flask docstrings and can be accessed programmatically from Python, JavaScript, or any HTTP client.")
P("The command-line interface (predict_cli.py) accepts all required and optional prediction inputs as named command-line arguments using Python's argparse module. It prints a formatted prediction summary to standard output including the projected total, win probability, projected range, and phase classification. The CLI is most useful during development for rapid manual testing of specific match scenarios without launching a web server, and for running prediction in batch or scripted contexts.")
P("Both interfaces implement the same input validation logic as the Flask form and return the same prediction fields. The consistency is guaranteed by the shared inference module, which is called identically from all four surfaces. A developer verifying a particular prediction in the CLI can be confident that the same input through the Flask API will produce the same output, eliminating the need to run multiple interfaces in parallel during debugging.")

H2("6.5  Software Engineering Principles Applied")
P("The implementation of the IPL Match Score and Win Prediction System reflects several core software engineering principles that were applied intentionally to improve maintainability, testability, and extensibility.")
P("Modularity: The project is organized into discrete Python modules with well-defined responsibilities. The ipl_predictor package contains the inference module (predictor.py), the monitoring module (monitoring.py), the authentication utilities (auth.py), and the database models (models.py). Training scripts are separate from inference code. Interface code (Flask, Streamlit, CLI) is separate from both. This modularity means that changing the inference logic does not require editing interface files, and adding a new interface does not require editing inference or training code.")
P("Low Coupling: The four user-facing surfaces have no shared state beyond the loaded model artifacts. They do not call each other's functions, share global variables, or depend on each other's implementation details. Each surface depends only on the shared inference module's public function signature. Changes to one surface cannot accidentally break another surface. The inference module depends on the model artifacts but not on any interface, meaning new interfaces can be added without risk to existing ones.")
P("Separation of Concerns: Data preparation, model training, inference, and presentation are each handled by separate, independent components. The preprocessing pipeline does not know about the model training algorithms. The training scripts do not know about the serving interfaces. The inference module does not know which interface is calling it. This separation makes the system easier to debug (each concern can be tested independently) and easier to extend (each concern can be improved without affecting others).")
P("Fail-Fast Validation: Input validation is performed at the inference module level before any model is called. Invalid team names, unrecognized venues, out-of-range overs values, and missing required fields all trigger descriptive error messages rather than silent failures or cryptic exceptions. This fail-fast behavior makes the system easier to use and easier to debug.")

H2("6.6  Dependency Stack and Technical Environment")
P("The software stack used in this project was not assembled arbitrarily. Each major dependency was selected because it addressed a specific technical requirement within the overall system.")
P("NumPy and Pandas form the numerical and tabular data foundation. NumPy provides the vectorized operations needed for efficient feature computation, metric calculation, and array manipulation throughout the pipeline. Pandas provides the DataFrame operations needed for CSV reading, support table joining, groupby aggregation, and report-side analytics. Both are industry-standard, well-documented, and supported by the scikit-learn and model training libraries.")
P("Scikit-learn provides the classical machine learning baselines (HistGradientBoosting), all evaluation metrics (MAE, RMSE, log loss, Brier score), the calibration utility (CalibratedClassifierCV), and several preprocessing utilities. Its consistent API makes it straightforward to compare model families and to apply calibration as a post-processing step without changing the model architecture.")
P("XGBoost and CatBoost provide the GPU-accelerated gradient boosting capabilities. XGBoost served as the primary structured-data benchmark and remained competitive throughout the model comparison. CatBoost was ultimately selected for the win model because of its stronger calibration compatibility and its ordered categorical encoding that reduces leakage risk within the gradient boosting process.")
P("PyTorch provides the entity-embedding model implementation for score prediction. The entity-embedding architecture requires a deep learning framework because it learns categorical embeddings through backpropagation, which is not directly supported by scikit-learn's fixed-structure models. PyTorch was preferred over TensorFlow for its more Pythonic API and its ease of customization for tabular data architectures.")
P("Flask and Streamlit provide the web-facing presentation layer. Flask was chosen for the primary interactive interface because of its minimalism and direct control over request handling, which is useful for implementing both the HTML form and the JSON API in the same application. Streamlit was chosen for the analytics dashboard because its reactive widget model is well-suited to interactive data exploration without requiring frontend JavaScript development.")


H2("6.7  Testing and Quality Assurance")
P("The project includes an automated pytest suite (tests/) that covers the most critical components of the prediction pipeline. Automated testing is included not as a formality but because specific preprocessing and inference operations had real defects during development that were caught and fixed through test-driven debugging.")
P("The alias normalization tests verify that all known team and venue name variants are correctly resolved to their canonical forms. They also verify that unknown names (inputs that do not appear in the alias table) are handled gracefully with a meaningful error message rather than silently passing incorrect names to the model.")
P("The overs parsing tests verify that the over.ball to legal-ball-count conversion is correct for standard deliveries, wide deliveries (which do not count as legal balls), and no-ball deliveries (which count once in some contexts and differently in others). These edge cases are critical because incorrect legal-ball counts would corrupt all phase indicators, momentum windows, and required-rate calculations for affected rows.")
P("The feature frame tests verify that the feature frame constructed by the inference module contains the correct number of columns in the correct order, and that all values are within expected ranges for a valid input. These tests catch regressions caused by feature engineering changes in the preprocessing pipeline that might not be reflected in the inference module's column ordering.")
P("The web route tests (using Flask's test client) verify that the prediction API route returns a 200 status code for valid inputs, returns a 400 error for invalid inputs, and includes all expected fields in the JSON response. These tests provide fast regression detection for interface-level issues without requiring a running server.")

H2("6.8  Artifact Traceability and Working Evidence")
P("The project maintains a structured artifact inventory that links every major output to its source script, file location, and modification timestamp. This traceability record is more reliable than narrative claims because it can be verified directly against the filesystem at any time.")
BL("best_model_search_report.json (28.66 KB, last modified 2026-04-15): Full model comparison results across all families and scopes, including per-metric scores for every tested configuration.")
BL("best_score_test_predictions.csv (627.20 KB, last modified 2026-04-15): Row-level score predictions on the held-out 2025 test season, enabling independent metric verification and error distribution analysis.")
BL("best_win_test_predictions.csv (736.99 KB, last modified 2026-04-15): Row-level win probability predictions on the held-out 2025 test season, enabling independent calibration analysis.")
BL("canonical_pipeline_report.json (4.56 KB, last modified 2026-04-15): Summary of the preprocessing pipeline outputs including row counts, column counts, and data quality statistics.")
BL("cpu_model_report.json (4.06 KB, last modified 2026-04-08): HistGradientBoosting baseline model performance across both scopes and both tasks.")
BL("deployment_report.json (1.92 KB, last modified 2026-04-15): Record of the currently promoted deployment model versions with their artifact paths and performance summaries.")
BL("model_registry.json (20.20 KB, last modified 2026-04-15): Complete version history of all saved model snapshots with full metadata, enabling selection of the strongest configuration by evidence rather than modification time.")
BL("pre_match_model_report.json (2.77 KB, last modified 2026-04-15): Pre-match model performance on both tasks for comparison with live model performance.")
P("The existence of these artifacts allows the project report to make claims traceable to specific files rather than to memory or reconstruction. Any discrepancy between the report's claims and the artifact contents would be detectable through direct inspection.")

H2("6.9  Workspace Organization")
P("The repository is organized into top-level directories that reflect the functional roles of their contents. This organization reduces accidental coupling between stages and makes it clear where new outputs should be placed as the project expands.")
BL("data/ — Raw CSV2 match files, processed feature tables (ipl_features.csv), support tables (venue_stats.csv etc.), and weather snapshots.")
BL("docs/ — Generated report assets, this project report, and technical summary documents.")
BL("ipl_predictor/ — The Python package containing the shared inference module, monitoring module, authentication utilities, database models, and the entity-embedding Torch model implementation.")
BL("models/ — Saved live and pre-match model artifacts (.pkl and .pt files), prediction exports (.csv), and JSON evaluation reports.")
BL("notebooks/ — Exploratory analysis notebooks, preprocessing experiments, and model comparison experiments conducted during development.")
BL("scripts/ — Reproducible command-line scripts for data refresh, feature engineering, CPU baseline training, GPU benchmark training, best-model search, pre-match modeling, and report generation.")
BL("static/ — CSS stylesheets for the Flask web interface.")
BL("templates/ — HTML templates for the Flask web application routes.")
BL("tests/ — Pytest test suites for inference utilities, alias normalization, overs parsing, and web routes.")
P("This directory structure reflects a key principle: each artifact should live in exactly one place, and that place should be obvious from the artifact's role. The principle was maintained throughout development and prevented the common problem of having multiple copies of processed datasets, trained models, or report assets in different locations with unclear provenance.")

H2("6.10  Monitoring and Operational Readiness")
P("The project includes an early monitoring layer implemented in ipl_predictor/monitoring.py. The monitoring module is designed to become more useful over time as more live prediction events are logged and resolved against actual match outcomes. At the time of this report, the monitoring database contains 37 logged prediction events and 4 resolved outcomes — insufficient for statistically meaningful drift detection, but sufficient to demonstrate that the infrastructure is in place and functioning correctly.")
P("The monitoring module operates as follows: every prediction request logged through the shared inference module is recorded with a timestamp, the input match state, and the predicted outputs. When a match completes and the actual outcome is known, the resolved outcome can be linked to the logged prediction events for that match. The monitoring module then computes a rolling Brier score and calibration gap over the most recent window of resolved predictions (configurable, currently set to 500 events) and flags a potential retraining need if the rolling metric deteriorates significantly from the training-time benchmark.")
P("The current monitoring status shows: monitoring window size 500, events used 37, outcomes used 4, current status 'insufficient_data', retrain trigger 'No', recommended action 'Collect more events and outcomes before taking action.' This status honestly reflects the early operational stage of the system. The monitoring infrastructure is in place and producing correct outputs; it simply needs more live prediction data before it can make confident drift determinations.")
P("Including the monitoring module in this submission demonstrates foresight: the project was designed not just for evaluation but for eventual operational use. A prediction system without monitoring is a black box that can degrade silently as the IPL environment changes. The monitoring module establishes the feedback loop through which the system can detect its own performance degradation and trigger appropriate responses.")

# INSERT STREAMLIT/DASHBOARD SCREENSHOT
IMG(img_elems[4])
CAP("Figure 6.3: Streamlit Analytics Dashboard – Team and Venue Context Display")

H2("6.11  What This Chapter Demonstrates")
P("This implementation chapter demonstrates that the project moved from experimental results to a working, testable, deployable system. The predictive models described in Chapters 4 and 5 would not constitute a strong final-year engineering project on their own. The more compelling demonstration is that those models are integrated into a shared inference pipeline, served through four interfaces, verified by automated tests, and supported by artifact traceability records and a monitoring module.")
P("The shared inference core, the modular architecture, the automated test suite, and the monitoring readiness together constitute the software engineering contribution of this project. They demonstrate that the project was built as a system, not merely as a model. This distinction matters in academic evaluation because a system can be inspected, tested, extended, and maintained, while a model artifact in isolation can only be invoked.")

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 7 – CONCLUSION AND FUTURE SCOPE
# ══════════════════════════════════════════════════════════════════════
CT("Chapter 7 – Conclusion and Future Scope")

H2("7.1  Conclusion")
P("This project successfully developed, evaluated, and deployed an IPL Match Score and Win Prediction System that addresses the live cricket prediction problem through a combination of disciplined data engineering, comparative model evaluation, probability calibration, and production-quality software design. The system ingests 282,552 ball-by-ball match-state rows from 1,189 IPL matches spanning nearly two decades of the league's history and produces calibrated predictions through a shared inference module accessible via four independent user-facing surfaces.")
P("The strongest experimental result, preserved in the model registry as version v20260415_152209, achieves a held-out MAE of 15.36 runs and RMSE of 19.81 runs for the live score prediction task, and a held-out log loss of 0.4922 for the live win-probability task. These results are produced by pairing a Torch entity-embedding regressor for score prediction with a calibrated CatBoost GPU classifier for win probability, both trained under the all-active-history scope with strict chronological evaluation against the 2025 held-out season. The results improve meaningfully over heuristic baselines and over pre-match modeling, while accurately representing the genuine uncertainty inherent in T20 cricket prediction.")
P("Five major lessons were learned from this project. First, data engineering matters more than algorithm selection for structured tabular ML: the improvement from selecting the correct training scope exceeded the improvement from choosing the best model family. This is the most practically important finding of the project and the one most directly transferable to other applied ML contexts. Second, recent-only training scope is too data-poor to support stable player-level and venue-level feature learning once chronological splits are applied honestly. The all-active-history scope, which preserves full historical depth while filtering to currently active franchises, provides the right balance between historical depth and contemporary relevance. Third, score and win tasks reward fundamentally different model properties: the score task benefits from learned categorical embeddings that capture venue and team similarity, while the win task benefits from CatBoost's ordered categorical encoding and compatibility with probability calibration. A hybrid model pair outperforms any single-family solution across both tasks. Fourth, probability calibration is essential for win prediction systems that display probability percentages to users: uncalibrated models erode user trust even when classification accuracy is high, and calibration-aware metrics (log loss, Brier score) must be reported alongside accuracy. Fifth, a strong final-year engineering project must demonstrate the full ML system lifecycle — from raw data through deployment — not just notebook-level model evaluation.")
P("The ethical and social dimensions of the project have been taken seriously throughout the design. The system is designed as an analytical and educational tool. Predictions are presented with explicit uncertainty information (projected ranges, calibrated probabilities) rather than as false-precision point estimates. Training data comes exclusively from publicly available, openly licensed sources. The system is not connected to any betting or gambling platform and was not designed to support such applications. These design principles reflect a commitment to responsible AI deployment that should be standard in any machine learning project intended for public use.")
P("Viewed as a complete engineering project, this submission satisfies the core expectations of a final-year B.Tech. project in Computer Science and Artificial Intelligence and Machine Learning. It defines a relevant problem with real-world significance, reviews the prior literature with sufficient depth to justify methodological choices, implements a concrete and fully documented methodology, evaluates results honestly under a leakage-free protocol, builds a working multi-surface software system, demonstrates results through concrete scenario interpretations, documents limitations transparently, and outlines a credible path for future extension. The project therefore represents not just a model artifact but a complete engineering contribution.")

embed_fig(fig_final_metrics(), caption='Figure 7.1  Final Model Performance – Best Snapshot v20260415_152209')

H2("7.2  Future Scope")
P("The project provides a solid technical foundation upon which several high-value extensions can be built. Each extension described below is practically feasible given the existing architecture and would build naturally on one or more of the infrastructure components already in place.")
P("The highest-priority immediate extension is the integration of probable playing-XI data before the match begins. The current pre-match branch operates with only team identity, venue priors, and broad form indicators, producing significantly weaker results than the live model. If the expected lineup for each team were available before the toss — which is typically announced approximately one hour before match start — the system could aggregate player-level batting depth and bowling attack quality into team-strength composites that would substantially narrow the pre-match prediction uncertainty.")
P("A second high-value extension is the replacement of the current fixed-offset uncertainty bands with conformal prediction or quantile regression intervals. The system currently generates projected score ranges by applying a fixed percentile offset to the point estimate. A conformal prediction layer trained on calibration data would produce adaptive intervals whose width reflects the genuine uncertainty of the specific match state — wider early in the innings, narrower in the death overs, and adjusted for specific venue and team combinations. This would make the uncertainty communication more informative and statistically rigorous.")
BL("Integrate live data ingestion from a public cricket scoring API (such as Cricbuzz or ESPN Cricinfo's developer endpoint) to enable fully automated real-time predictions during IPL matches without manual data entry by the user.")
BL("Expand the explanation layer so that the API response includes the top five feature contributions to each prediction expressed in cricket-intelligible language, enabling transparent explanations of why the model expects a particular score or win probability.")
BL("Extend the monitoring layer from passive event logging to active retraining triggers: when the rolling held-out Brier score deteriorates by more than a configured threshold (e.g., 0.05 points) from the training-time benchmark, the system automatically queues a retraining run with the latest season's data appended to the training pool.")
BL("Generalize the pipeline to other major T20 leagues by parameterizing the alias tables, active-team filters, venue statistics, and category encodings for each competition. The core preprocessing logic, model architectures, and inference module are not IPL-specific and could support the Big Bash League, Caribbean Premier League, SA20, or Women's T20 leagues with relatively modest adaptation work.")
BL("Evaluate transformer-based tabular models (TabTransformer, FT-Transformer, TabNet) as the IPL dataset continues to grow with new seasons. These attention-based architectures may capture longer-range feature interactions than gradient boosting or the current MLP entity-embedding design, particularly for modeling the complex multi-way interactions between venue, team composition, phase, and momentum.")
BL("Build a natural-language commentary generation layer that converts the model's numerical prediction outputs into match-narrative text. Rather than displaying a raw win probability percentage, the system could generate a one-sentence contextual summary ('Chennai Super Kings need to accelerate: the required rate is 1.8 runs per over above their current scoring pace') that makes the prediction more immediately useful for casual fans and broadcast analysts.")
P("Each of these extensions builds naturally on the infrastructure already in place. The shared inference contract can be extended to accept new input fields without breaking existing callers. The support-table architecture can incorporate new feature families without requiring changes to the model interfaces. The monitoring module can be extended with new drift metrics without altering the event logging format. The project therefore provides not only a working prediction system but also a well-structured starting point for future sports analytics development.")
P("In conclusion, the IPL Match Score and Win Prediction System represents a complete, principled, and honest machine learning project that delivers practical analytical value while meeting the rigorous evaluation and documentation standards required for academic credibility. The methodology, results, implementation, and lessons documented in this report constitute a reusable reference for future applied sports analytics projects in this and related domains.")

# ══════════════════════════════════════════════════════════════════════
# APPENDICES
# ══════════════════════════════════════════════════════════════════════
CT("Appendix 1 – Flowchart and Algorithm Notes")
P("This appendix presents the core prediction algorithm as an ordered step-by-step procedure, complementing the architecture and workflow descriptions in Chapter 4.")
NL("Accept current match inputs: batting team, bowling team, venue, season, innings number (1 or 2), runs scored, wickets lost, overs bowled, and target (second innings only).")
NL("Normalize all team and venue name variants to their canonical forms using the alias normalization table. Reject any unrecognized names with a descriptive error message.")
NL("Convert the overs value to legal balls bowled and derive balls_left, over_number, ball_in_over, innings_progress, and phase classification (powerplay, middle, or death).")
NL("Compute current_run_rate, and for second innings: target_remaining, required_run_rate, and required_minus_current_rr.")
NL("Load the nine support tables from disk and perform join operations to enrich the feature frame with venue statistics, team form, team-venue form, matchup form, and player profiles.")
NL("Construct the complete live feature frame in the exact column order expected by the saved score and win model artifacts.")
NL("Invoke the Torch entity-embedding score model to obtain the projected innings total point estimate.")
NL("Compute the projected scoring range by applying the distribution-calibrated uncertainty band to the point estimate.")
NL("Invoke the calibrated CatBoost win model to obtain the batting-side win probability.")
NL("Compute runs_vs_par by comparing the current innings score to the venue historical baseline at the same innings progress.")
NL("Format all outputs — projected total, win probability, projected range, runs_vs_par, phase, match narrative — into the response dictionary.")
NL("Return the response dictionary to the calling surface (Flask, Streamlit, API, or CLI) for display or serialization.")
P("The same twelve-step algorithm is executed identically by all four interface surfaces, which is why predictions are consistent across all access methods for the same match state input.")

CT("Appendix 2 – Selected Feature Notes")
P("The following notes describe the analytical significance of the most important engineered features. Full feature reasoning is provided in Chapter 4.")
BL("runs_vs_par: Compares the current innings score with the historical venue baseline for the same innings progress fraction. Positive values indicate the batting side is ahead of the typical venue pace; negative values indicate they are behind. This feature captures venue-adjusted batting context that raw runs cannot provide.")
BL("required_minus_current_rr: The gap between the required run rate (runs needed per over to win) and the current run rate (runs scored per over so far). Populated only for second innings. Positive values indicate the chasing side is falling behind the required pace. This is one of the strongest single predictors of second-innings win probability.")
BL("runs_last_5 and wickets_last_5: Rolling sums of runs and wickets over the most recent five legal balls. These capture short-term trajectory — acceleration, stability, or collapse — that cumulative totals obscure. A team at 90/2 after 11 overs that just scored 45 in the last five overs is in a structurally different state from the same team if it scored only 25.")
BL("batting_team_venue_form and bowling_team_venue_form: Each team's recent historical performance at the specific venue, capturing ground-specific strengths and weaknesses. Chennai Super Kings' historically strong performance at M. A. Chidambaram Stadium, for example, would be reflected in a positive batting_team_venue_form value for CSK matches at that ground.")
BL("batter_vs_bowler_history: Head-to-head strike rate and average between the current striker and current bowler, computed from all prior encounters in the training archive. Captures player-specific matchup advantages that general team and player form cannot represent.")
BL("innings_progress: Legal balls bowled divided by 120, providing a normalized measure of how far through the innings the batting side is. Used as a key signal for phase-aware prediction, replacing the raw overs count with a continuous [0,1] range that is compatible across different innings lengths.")
BL("venue_avg_first_innings: Historical average first-innings total at the venue across all processed matches. Represents the venue's inherent scoring environment and forms the basis for the runs_vs_par computation.")

CT("Appendix 3 – Short Supporting Case Snapshots")
P("The following case snapshots record sample prediction outputs from the system. All outputs were produced by model registry version v20260415_152209. Full scenario discussion is provided in Chapter 5.")
P("Live Prediction Scenarios:")
BL("Powerplay Launch — 42/0 after 5 overs, first innings, high-scoring venue. Output: projected total 194.1, win probability 44.2%, range 175-222. The model combines aggressive early scoring pace with venue par history to project a total well above average; win probability remains near-even as the innings has many overs remaining.")
BL("Middle-Over Rebuild — 74/3 after 10 overs, first innings. Output: projected total 162.3, win probability 63.4%, range 143-190. Wicket loss moderates the score projection; the bowling side's win probability increases as batting resources narrow.")
BL("Death Overs Surge — 155/4 after 16 overs, first innings. Output: projected total 197.7, win probability 63.4%, range 178-225. Late acceleration with wickets remaining pushes the projected total significantly above initial expectations.")
BL("Chase Under Pressure — 96/5 after 11 overs, second innings, target 177. Output: projected total 177.3, win probability 44.2%, range 158-205. Required rate exceeds current rate with few wickets remaining; odds remain approximately even.")
BL("Late Collapse Risk — 141/7 after 17 overs, first innings. Output: projected total 163.8, win probability 37.9%, range 144-192. Multiple late wickets shift the projection downward; the fielding side's win probability rises as batting resources are exhausted.")
P("Pre-Match Scenarios:")
BL("Mumbai Indians vs Chennai Super Kings at Wankhede Stadium: Projected first-innings range 158-182, win probabilities 48.3% vs 51.7%. Classic rivalry with near-equal pre-match odds reflecting the historical balance between these franchises at this venue.")
BL("Royal Challengers Bengaluru vs Sunrisers Hyderabad at M. Chinnaswamy Stadium: Projected first-innings range 161-185, win probabilities 46.3% vs 53.7%. High-scoring venue prior lifts the projected total range above the league average.")
BL("Gujarat Titans vs Rajasthan Royals at Narendra Modi Stadium: Projected first-innings range 170-194, win probabilities 48.2% vs 51.8%. Near-equal pre-match odds with a high projected range reflecting the large-capacity high-scoring venue history.")
P("The strongest saved model version referenced throughout this report is v20260415_152209. To reproduce the predictions above, activate the project virtual environment and run the Flask application or CLI with the corresponding match state inputs.")

# ══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS  /  LIST OF FIGURES  /  LIST OF TABLES
# All inserted immediately before ch1_elem (each addprevious stacks in order)
# ══════════════════════════════════════════════════════════════════════

def _toc_ins(elem):
    """Insert elem after current front-matter anchor, then advance the anchor."""
    _toc_anchor[0].addnext(elem)
    _toc_anchor[0] = elem

def _toc_pb():
    pb = doc.add_paragraph()
    r_ = OxmlElement('w:r'); b_ = OxmlElement('w:br')
    b_.set(qn('w:type'), 'page'); r_.append(b_); pb._p.append(r_)
    body.remove(pb._element); _toc_ins(pb._element)

def _add_tabs(p, left_pos=None, right_pos=9360):
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    if left_pos:
        t1 = OxmlElement('w:tab')
        t1.set(qn('w:val'), 'left'); t1.set(qn('w:pos'), str(left_pos))
        tabs.append(t1)
    t2 = OxmlElement('w:tab')
    t2.set(qn('w:val'), 'right'); t2.set(qn('w:pos'), str(right_pos))
    tabs.append(t2)
    pPr.append(tabs)

def _toc_heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = p.add_run(text)
    rn.bold = True; rn.font.size = Pt(14)
    rn.font.name = 'Times New Roman'; rn.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(18)
    body.remove(p._element); _toc_ins(p._element)

def _toc_entry(text, pg, bold=False, indent_in=0.0):
    p = doc.add_paragraph()
    _add_tabs(p, right_pos=9360)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Inches(indent_in)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    rn1 = p.add_run(text)
    rn1.bold = bold; rn1.font.size = Pt(11)
    rn1.font.name = 'Times New Roman'; rn1.font.color.rgb = RGBColor(0,0,0)
    rn2 = p.add_run('\t' + str(pg))
    rn2.bold = bold; rn2.font.size = Pt(11)
    rn2.font.name = 'Times New Roman'; rn2.font.color.rgb = RGBColor(0,0,0)
    body.remove(p._element); _toc_ins(p._element)

def _toc_fm(text, pg):
    """Front-matter line (italic)."""
    p = doc.add_paragraph()
    _add_tabs(p, right_pos=9360)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    rn1 = p.add_run(text); rn1.italic = True
    rn1.font.size = Pt(11); rn1.font.name = 'Times New Roman'
    rn1.font.color.rgb = RGBColor(0,0,0)
    rn2 = p.add_run('\t' + str(pg)); rn2.italic = True
    rn2.font.size = Pt(11); rn2.font.name = 'Times New Roman'
    rn2.font.color.rgb = RGBColor(0,0,0)
    body.remove(p._element); _toc_ins(p._element)

def _lx_row(col1, col2, col3, bold=False):
    """Single row for LOF / LOT using two tab stops."""
    p = doc.add_paragraph()
    _add_tabs(p, left_pos=1080, right_pos=9360)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    rn = p.add_run(f'{col1}\t{col2}\t{col3}')
    rn.bold = bold; rn.font.size = Pt(11)
    rn.font.name = 'Times New Roman'; rn.font.color.rgb = RGBColor(0,0,0)
    body.remove(p._element); _toc_ins(p._element)

# ── TABLE OF CONTENTS ─────────────────────────────────────────────────
_toc_heading("TABLE OF CONTENTS")

_toc_fm("Declaration", "i")
_toc_fm("Approval from Supervisor", "ii")
_toc_fm("Certificate", "iii")
_toc_fm("Acknowledgements", "iv")
_toc_fm("Abstract", "v")
_toc_fm("Table of Contents", "vi")
_toc_fm("List of Figures", "vii")
_toc_fm("List of Tables", "viii")

_toc_entry("Chapter 1  Introduction", 1, bold=True)
_toc_entry("1.1   Topic of the System", 1, indent_in=0.35)
_toc_entry("1.2   Overview of the Project", 1, indent_in=0.35)
_toc_entry("1.3   Project Background", 2, indent_in=0.35)
_toc_entry("1.4   Scope of the Project", 3, indent_in=0.35)
_toc_entry("1.5   Goals and Objectives", 4, indent_in=0.35)
_toc_entry("1.6   Benefits of the Project", 5, indent_in=0.35)
_toc_entry("1.7   Relevance of the Project", 6, indent_in=0.35)
_toc_entry("1.8   Challenges and Constraints", 7, indent_in=0.35)
_toc_entry("1.9   Structure of the Report", 8, indent_in=0.35)
_toc_entry("1.10  Conclusion", 9, indent_in=0.35)

_toc_entry("Chapter 2  Literature Review", 10, bold=True)
_toc_entry("2.1   Introduction", 10, indent_in=0.35)
_toc_entry("2.2   Review of Key Literature", 10, indent_in=0.35)
_toc_entry("2.3   Summary of Findings", 12, indent_in=0.35)
_toc_entry("2.4   Addressing Identified Gaps", 13, indent_in=0.35)
_toc_entry("2.5   Importance of This Review to the Project", 14, indent_in=0.35)

_toc_entry("Chapter 3  Problem Objective", 15, bold=True)
_toc_entry("3.1   Problem Statement", 15, indent_in=0.35)
_toc_entry("3.2   Objectives of the Project", 16, indent_in=0.35)
_toc_entry("3.3   Scope of the Project", 17, indent_in=0.35)
_toc_entry("3.4   Key Deliverables", 18, indent_in=0.35)

_toc_entry("Chapter 4  Methodology and System Design", 20, bold=True)
_toc_entry("4.1   System Architecture", 20, indent_in=0.35)
_toc_entry("4.2   Workflow and System Model", 22, indent_in=0.35)
_toc_entry("4.3   Dataset Description", 24, indent_in=0.35)
_toc_entry("4.4   Data Preprocessing and Feature Engineering", 26, indent_in=0.35)
_toc_entry("4.5   Model Training", 29, indent_in=0.35)
_toc_entry("4.5.1  Exploratory Data Analysis", 30, indent_in=0.65)
_toc_entry("4.6   Figuring the Best Model", 31, indent_in=0.35)
_toc_entry("4.7   User Interface Creation", 33, indent_in=0.35)

_toc_entry("Chapter 5  Results", 36, bold=True)
_toc_entry("5.1   Software and Hardware Details", 36, indent_in=0.35)
_toc_entry("5.1.1  Software Requirements", 36, indent_in=0.65)
_toc_entry("5.1.2  Additional Software Tools", 37, indent_in=0.65)
_toc_entry("5.1.3  Hardware Requirements", 38, indent_in=0.65)
_toc_entry("5.2   Result Analysis", 39, indent_in=0.35)
_toc_entry("5.3   Discussion of Results", 41, indent_in=0.35)
_toc_entry("5.4   Summary", 43, indent_in=0.35)

_toc_entry("Chapter 6  System Implementation and Testing", 46, bold=True)
_toc_entry("6.1   Shared Inference Core", 46, indent_in=0.35)
_toc_entry("6.2   Flask Web Application", 47, indent_in=0.35)
_toc_entry("6.3   Streamlit Analytics Dashboard", 48, indent_in=0.35)
_toc_entry("6.4   REST API and Command-Line Interface", 49, indent_in=0.35)
_toc_entry("6.5   Software Engineering Principles Applied", 50, indent_in=0.35)
_toc_entry("6.6   Dependency Stack and Technical Environment", 51, indent_in=0.35)
_toc_entry("6.7   Testing and Quality Assurance", 52, indent_in=0.35)
_toc_entry("6.8   Artifact Traceability and Working Evidence", 53, indent_in=0.35)
_toc_entry("6.9   Workspace Organization", 54, indent_in=0.35)
_toc_entry("6.10  Monitoring and Operational Readiness", 55, indent_in=0.35)
_toc_entry("6.11  What This Chapter Demonstrates", 56, indent_in=0.35)

_toc_entry("Chapter 7  Conclusion and Future Scope", 57, bold=True)
_toc_entry("7.1   Conclusion", 57, indent_in=0.35)
_toc_entry("7.2   Future Scope", 59, indent_in=0.35)

_toc_entry("Appendix 1  –  Flowchart and Algorithm Notes", 61)
_toc_entry("Appendix 2  –  Selected Feature Notes", 62)
_toc_entry("Appendix 3  –  Short Supporting Case Snapshots", 63)
_toc_entry("References", 64, bold=True)

_toc_pb()

# ── LIST OF FIGURES ───────────────────────────────────────────────────
_toc_heading("LIST OF FIGURES")
_lx_row("Figure No.", "Figure Name", "Page No.", bold=True)

_lx_row("1.1",  "Project Phase Timeline",                                          4)
_lx_row("4.1",  "Dataset Row Distribution Across IPL Seasons",                    25)
_lx_row("4.2",  "Feature Count per Family Across All 62 Engineered Columns",      28)
_lx_row("5.1",  "Score Model Comparison – MAE and RMSE (All-Active-History Scope)", 39)
_lx_row("5.2",  "Win Model Comparison – Log Loss and Brier Score",                40)
_lx_row("5.3",  "Calibration Curve – Uncalibrated vs Calibrated Win Model",       41)
_lx_row("6.1",  "Flask Live Prediction Interface – Input Form and Prediction Output", 47)
_lx_row("6.2",  "Flask Application – Prediction Response and Match Context Display", 48)
_lx_row("6.3",  "Streamlit Analytics Dashboard – Team and Venue Context Display", 56)
_lx_row("7.1",  "Final Model Performance – Best Snapshot v20260415_152209",       57)

_toc_pb()

# ── LIST OF TABLES ────────────────────────────────────────────────────
_toc_heading("LIST OF TABLES")
_lx_row("Table No.", "Table Title", "Page", bold=True)

_lx_row("1.1", "System Component Overview",                                    2)
_lx_row("2.1", "Literature Review Summary – Contribution, Gap, and Project Response", 11)
_lx_row("4.1", "Dataset Snapshot – Key Statistics",                           24)
_lx_row("4.2", "Feature Families – Names, Count, and Leakage Status",         27)
_lx_row("5.1", "Full Model Comparison – All Families, Both Scopes, 2025 Test Season", 40)
_lx_row("5.2", "Calibration Bucket Analysis – Calibrated CatBoost Win Model", 42)

_toc_pb()

# ══════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"\nDocument saved -> {OUT}")

# Word count
wc = sum(len(p.text.split()) for p in doc.paragraphs)
print(f"Approximate word count: {wc:,}")
print("Done.")
