#!/usr/bin/env python3
"""
Restructures IPL Project Report chapters with proper numbered subheadings.
Guidelines: Ch-title 14pt Bold TNR Centre Uppercase | Heading 11pt Bold TNR Upper Left |
            Subheading 11pt Bold TNR Title-Case Left | Body 11pt TNR Justified 1.5 spacing
Preserves: front matter (title page through List of Abbreviations) + References section.
"""

import shutil, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = 'docs/FINAL_SUBMISSION_IPL_PROJECT_REPORT_BACKUP.docx'
OUT  = 'docs/FINAL_SUBMISSION_IPL_PROJECT_REPORT_RESTRUCTURED.docx'
BACK = 'docs/FINAL_SUBMISSION_IPL_PROJECT_REPORT_BACKUP.docx'
import os
if not os.path.exists(BACK):
    shutil.copy('docs/FINAL_SUBMISSION_IPL_PROJECT_REPORT.docx', BACK)
    print(f"Backup saved -> {BACK}")
else:
    print(f"Backup already exists -> {BACK}")

doc  = Document(SRC)
body = doc.element.body

# ── locate boundaries ──────────────────────────────────────────────────
ch1_elem = ref_elem = None
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and 'CHAPTER 1' in p.text and ch1_elem is None:
        ch1_elem = p._element
    if p.style.name == 'Heading 1' and p.text.strip() == 'REFERENCES' and ref_elem is None:
        ref_elem = p._element

assert ch1_elem is not None, "Chapter 1 heading not found"
assert ref_elem  is not None, "REFERENCES heading not found"

# ── remove existing chapter body ───────────────────────────────────────
to_del = []
collecting = False
for child in list(body):
    if child is ch1_elem:
        collecting = True
    if child is ref_elem:
        break
    if collecting:
        to_del.append(child)

for elem in to_del:
    body.remove(elem)
print(f"Removed {len(to_del)} old chapter elements.")

# ── helper: every new paragraph goes immediately before <REFERENCES> ──
def _ins(elem):
    ref_elem.addprevious(elem)

def _mk_run(p, text, bold=False, size=11, italic=False):
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.name  = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run

def _set_spacing(p, before=0, after=6, line=16.5):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing = Pt(line)   # ~1.5 × 11pt

# ── public builders ────────────────────────────────────────────────────
def chapter_title(text):
    """14pt Bold, TNR, Uppercase, Centred, page-break-before."""
    # page break paragraph
    pb = doc.add_paragraph()
    pb_r = OxmlElement('w:r')
    pb_b = OxmlElement('w:br')
    pb_b.set(qn('w:type'), 'page')
    pb_r.append(pb_b)
    pb._p.append(pb_r)
    body.remove(pb._element)
    _ins(pb._element)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _mk_run(p, text.upper(), bold=True, size=14)
    _set_spacing(p, before=12, after=18)
    body.remove(p._element)
    _ins(p._element)
    return p

def h2(text):
    """11pt Bold, TNR, UPPERCASE, Left-aligned — section heading (1.x)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _mk_run(p, text.upper(), bold=True, size=11)
    _set_spacing(p, before=12, after=6)
    body.remove(p._element)
    _ins(p._element)
    return p

def h3(text):
    """11pt Bold, TNR, Title Case, Left-aligned — sub-heading (x.x.x)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _mk_run(p, text, bold=True, size=11)
    _set_spacing(p, before=9, after=6)
    body.remove(p._element)
    _ins(p._element)
    return p

def para(text):
    """11pt TNR, Justified, 1.5 line spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _mk_run(p, text, bold=False, size=11)
    _set_spacing(p, before=0, after=6, line=16.5)
    body.remove(p._element)
    _ins(p._element)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    _mk_run(p, text, size=11)
    _set_spacing(p, before=0, after=3, line=16.5)
    body.remove(p._element)
    _ins(p._element)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    _mk_run(p, text, size=11)
    _set_spacing(p, before=0, after=3, line=16.5)
    body.remove(p._element)
    _ins(p._element)
    return p

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION
# ══════════════════════════════════════════════════════════════════════
chapter_title("Chapter 1 – Introduction")

h2("1.1  Topic of the System")
para(
    "The topic of this project is the design and development of an IPL Match Score and "
    "Win Prediction System using supervised machine learning. The system processes live "
    "ball-by-ball match states from Indian Premier League (IPL) matches and produces "
    "two real-time analytical outputs: an estimate of the likely final innings total and "
    "the batting side's probability of winning the match. Unlike traditional approaches "
    "that rely on simple run-rate arithmetic or post-match summaries, this system applies "
    "multiple machine learning model families to a rich contextual feature representation "
    "built from a historical database of 1,189 IPL matches spanning from the 2007/08 "
    "season through the 2026 season."
)
para(
    "The topic is directly relevant to the field of applied artificial intelligence and "
    "machine learning because it demonstrates how supervised learning can be applied to "
    "a dynamic, real-world domain where conditions shift continuously ball by ball. Every "
    "delivery changes the score, the wickets remaining, the overs available, and the "
    "momentum of the innings, requiring the prediction system to update its outputs "
    "accordingly. This makes IPL live prediction a useful and challenging case study in "
    "time-aware feature engineering, chronological model evaluation, and probability-"
    "calibrated classification."
)
para(
    "The completed system is packaged as a practical software product with multiple "
    "user-facing interfaces including a Flask web application, a Streamlit analytics "
    "dashboard, a REST API endpoint, and a command-line interface. This breadth of "
    "deployment makes the system suitable for academic demonstration, cricket analytics "
    "discussion, and future integration into real-time scoring platforms."
)

h2("1.2  Overview of the Project")
para(
    "The IPL Match Score and Win Prediction System is a complete, end-to-end machine "
    "learning project that starts from raw ball-by-ball cricket data, engineers a rich "
    "live match-state feature set, trains and evaluates multiple supervised learning "
    "models under honest chronological evaluation, and exposes the final prediction "
    "engine through a set of consistent interfaces. The project begins with 1,189 raw "
    "IPL match files in Cricsheet CSV2 format and concludes with a deployed prediction "
    "service capable of responding to any valid live IPL match state."
)
para(
    "The core dataset produced by the preprocessing pipeline contains 282,552 rows and "
    "62 columns, with each row representing one legal-ball state in an IPL innings. This "
    "granularity allows the model to learn how the same cumulative score can signify "
    "very different outcomes depending on the stage of the innings, the wickets "
    "remaining, the venue's historical scoring characteristics, and the relative quality "
    "of the batting and bowling sides. The project separates score regression and "
    "win-probability estimation into two distinct but related tasks, each evaluated with "
    "metrics appropriate to the nature of its output."
)
para(
    "Multiple model families were compared during training, including HistGradientBoosting "
    "as a CPU baseline, XGBoost and CatBoost as high-performance gradient-boosting "
    "alternatives, and a PyTorch entity-embedding model for tabular data with "
    "high-cardinality categorical features. The strongest experimental result used a "
    "Torch entity-embedding regressor for live score prediction paired with a calibrated "
    "CatBoost GPU classifier for win probability, achieving a held-out MAE of 15.36 runs "
    "and RMSE of 19.81 runs on the score task and a log loss of 0.4922 on the win task. "
    "All results are reported under a strict chronological train-validation-test split "
    "that prevents any form of temporal data leakage."
)
para(
    "Beyond the modeling component, the project emphasizes software engineering quality. "
    "A single shared inference module is called by all four user-facing surfaces, "
    "ensuring that any improvement in prediction logic is automatically reflected across "
    "every interface. Preprocessing scripts, training pipelines, automated test suites, "
    "and documentation workflows are maintained in a reproducible workspace so that the "
    "full system can be rebuilt from raw data by running a documented sequence of commands."
)

h2("1.3  Project Background")
para(
    "T20 cricket is one of the most unpredictable formats of the sport. A team that "
    "appears dominant after the powerplay overs can lose several wickets quickly in "
    "the middle phase and find itself under severe pressure. Conversely, a team chasing "
    "a large target can change the match's complexion through a single explosive batting "
    "partnership. This inherent volatility, which is part of T20 cricket's popular "
    "appeal, also makes it a challenging and intellectually rich prediction problem for "
    "machine learning research."
)
para(
    "The Indian Premier League amplifies these challenges further. The IPL features ten "
    "franchises with different batting orders, bowling strategies, and team compositions. "
    "Conditions vary significantly across the eight to ten venues used each season. "
    "Player combinations change between matches due to injuries, team selections, and "
    "recent form. Historical data from earlier seasons contains team identities that no "
    "longer exist as franchises, requiring careful normalization before older records can "
    "contribute to modern modeling. These factors mean that a naive approach — for "
    "example, training on all available data without regard for temporal ordering — "
    "will produce misleadingly optimistic results that cannot be replicated in practice."
)
para(
    "Prior attempts at cricket match prediction have fallen into two broad categories. "
    "Rule-based heuristics relying on current run rate, wickets in hand, and overs "
    "remaining are easy to interpret but cannot naturally capture venue effects, player "
    "quality differences, or momentum shifts. Machine learning classifiers trained on "
    "end-of-match summary data bypass the live prediction problem entirely by working "
    "from retrospective information. Neither approach adequately addresses the problem "
    "of predicting match outcomes from an intermediate live state while the match is "
    "still in progress."
)
para(
    "The present project was motivated by the desire to close this gap. By starting "
    "from ball-by-ball Cricsheet data, reconstructing every legal match state, building "
    "chronologically safe support tables for venues, teams, and players, and training "
    "models on a rich live feature representation, the system produces predictions "
    "grounded in real cricket context rather than simplified arithmetic. The project "
    "also acknowledges its own limitations: cricket remains genuinely uncertain and no "
    "model can remove that uncertainty entirely. The aim is to improve meaningfully on "
    "heuristic baselines while maintaining honest evaluation standards throughout."
)

h2("1.4  Scope of the Project")
para(
    "The scope of the project is intentionally focused on the Indian Premier League "
    "in T20 format. It does not extend to other cricket formats such as One Day "
    "Internationals or Test matches, nor does it cover other T20 leagues such as the "
    "Big Bash League or the Caribbean Premier League. This focused scope allows the "
    "project to build IPL-specific support tables, alias normalization rules, and "
    "active-team filters without needing to generalize across leagues with fundamentally "
    "different player pools and playing environments."
)
para(
    "Within the IPL scope, the project covers raw ball-by-ball data ingestion and "
    "normalization from Cricsheet CSV2 files covering seasons 2007/08 through 2026; "
    "construction of nine historical support tables for venues, teams, players, and "
    "matchups; live innings score regression for both first and second innings; live "
    "win-probability classification; a pre-match prediction branch operating before "
    "the first ball; a Flask web application; a Streamlit dashboard; a REST API "
    "endpoint; a command-line interface; and an automated report generation workflow."
)
para(
    "The project explicitly excludes proprietary player-tracking data or commercial "
    "cricket data feeds, live data ingestion from external scoring APIs requiring "
    "subscription access, betting market data as an input feature, and full production "
    "deployment on a publicly accessible server. The system is designed as an "
    "analytical and educational tool, not as a commercial real-time service. These "
    "exclusions are appropriate for an academic final-year project and do not diminish "
    "the substantive technical contributions of the work."
)
para(
    "The active-team filter restricts the final training scope to the ten currently "
    "active IPL franchises: Mumbai Indians, Chennai Super Kings, Royal Challengers "
    "Bengaluru, Kolkata Knight Riders, Delhi Capitals, Punjab Kings, Rajasthan Royals, "
    "Sunrisers Hyderabad, Gujarat Titans, and Lucknow Super Giants. Historical data "
    "from defunct franchises is used to build full support table infrastructure but is "
    "excluded from the final model training scope to keep the category space relevant "
    "to the current IPL."
)

h2("1.5  Goals and Objectives")
para(
    "The project was designed around a set of clear, measurable goals spanning data "
    "engineering, machine learning, software development, and academic reporting. These "
    "goals guided every design decision and provide the framework for evaluating whether "
    "the final system can be considered successful."
)
para("The primary technical objectives are:")
bullet("Build a leakage-safe IPL live-state dataset from 1,189 raw Cricsheet CSV2 match files covering all IPL seasons from 2007/08 through 2026, containing 282,552 rows and 62 engineered columns.")
bullet("Design a rich feature representation that simultaneously captures match state, batting and bowling momentum, venue historical statistics, team strength, player form, and chase-pressure context.")
bullet("Compute all historical support tables in strict chronological order so that no feature value at any ball state uses information from future matches.")
bullet("Compare at least four model families for each prediction task under identical chronological train-validation-test splits.")
bullet("Select the best-performing model pair — one for score, one for win probability — and persist it in a versioned model registry with full performance metadata.")
bullet("Apply probability calibration to the win model and verify calibration quality using log loss, Brier score, and calibration-bucket analysis.")
bullet("Build a shared inference module that produces identical predictions across Flask, Streamlit, REST API, and CLI surfaces from the same feature engineering logic.")
bullet("Implement a monitoring layer capable of detecting prediction drift as live outcomes accumulate.")
para("The academic objectives are:")
bullet("Produce a project report that explains the reasons behind the strongest results rather than simply reporting metrics, making design decisions traceable to empirical evidence.")
bullet("Demonstrate a leakage-free, chronologically disciplined evaluation methodology that would produce the same results if independently replicated.")
bullet("Apply core software engineering principles — modularity, low coupling, automated testing, and reproducibility — throughout the project implementation.")

h2("1.6  Benefits of the Project")
para(
    "The primary benefit of the system is that it converts live IPL match context into "
    "structured probabilistic insight. A cricket follower watching the match can "
    "consult the system to understand not just the current score but what that score "
    "implies: how likely is the batting side to finish at 190 or above, and what is "
    "their probability of winning given the current state? Projected totals, win "
    "probabilities, expected scoring ranges, and runs-versus-par comparisons give "
    "viewers an evidence-based analytical framework for interpreting what they see on "
    "the field."
)
para(
    "From an academic perspective, the project is a concrete case study in how to build "
    "a leakage-safe, end-to-end sports analytics pipeline. It demonstrates several "
    "software engineering and machine learning principles that are often taught in "
    "theory but rarely practiced together in final-year project work: chronological "
    "train-test splitting, support table construction outside the model, probability "
    "calibration, modular interface design, and automated test coverage. Students and "
    "faculty can use the project as a reference for how to approach applied machine "
    "learning problems with production-quality engineering discipline."
)
para(
    "For the field of cricket analytics, the project contributes a reproducible "
    "workflow that can be adapted to other T20 leagues with relatively modest effort. "
    "The core pipeline structure — ingestion, normalization, feature engineering, "
    "model training, calibration, and serving — is not IPL-specific in its design, "
    "only in its parameterization. A future developer who wants to build a similar "
    "system for another league can use this project as a validated architectural "
    "starting point."
)
para(
    "The project also demonstrates the value of honest evaluation. By reporting "
    "results from chronological held-out seasons rather than random splits, it "
    "establishes that its performance claims can be made with genuine confidence. "
    "A project that is transparent about its methodology is far more valuable to the "
    "research community than one that achieves high numbers through evaluation shortcuts "
    "that do not reflect real-world deployment conditions."
)

h2("1.7  Relevance of the Project")
para(
    "The relevance of the project spans technological, educational, and social "
    "dimensions. Technologically, the system applies modern machine learning tools "
    "— gradient boosting frameworks, entity-embedding neural networks, probability "
    "calibration, and multiple deployment surfaces — to a live sports analytics "
    "problem. Each of these technologies is widely used in industry, making the "
    "project directly relevant to the skill set expected of a B.Tech. graduate "
    "specializing in Computer Science and Artificial Intelligence."
)
para(
    "Educationally, the project illustrates that the quality of an applied machine "
    "learning system depends far more on data engineering and evaluation discipline "
    "than on algorithm novelty. This lesson is frequently overlooked in academic "
    "settings where the emphasis tends to fall on which algorithm to try rather than "
    "how to construct a trustworthy dataset and evaluation framework. The project "
    "reinforces this lesson by showing that the largest performance gains came from "
    "removing data leakage and widening the training scope — not from switching to a "
    "more sophisticated model family."
)
para(
    "Socially, cricket is one of the most followed sports in India, with hundreds of "
    "millions of fans who engage with IPL matches with intense interest. A system that "
    "can explain live match dynamics in probabilistic terms — for example, that a "
    "chasing team's win probability has dropped from 62 percent to 41 percent after "
    "losing two wickets in the death overs — provides fans, commentators, and analysts "
    "with a more informed basis for discussion, promoting evidence-based reasoning "
    "about sport."
)

h2("1.8  Challenges and Constraints")
para(
    "The project encountered several significant challenges that directly shaped the "
    "final system design. The most important was temporal data leakage. When support "
    "features such as venue averages, team form scores, and player career statistics "
    "are computed using the full historical archive, they inadvertently include "
    "information from future matches that would not be available during live use. "
    "Eliminating this leakage required recomputing all support tables in strict "
    "chronological order, match by match, which added considerable preprocessing "
    "complexity but was essential for producing honest performance estimates."
)
para(
    "The second challenge was data heterogeneity. IPL team names and venue names "
    "appear in multiple textual variants across different seasons and data sources. "
    "The franchise now known as Royal Challengers Bengaluru has appeared as Royal "
    "Challengers Bangalore in historical files. Wankhede Stadium appears as Wankhede "
    "in some match metadata. Without comprehensive alias normalization, these variants "
    "create artificial identity fragmentation that weakens statistical estimates and "
    "corrupts join operations on support tables."
)
para(
    "The third challenge was training scope selection. A recent-only scope is "
    "chronologically close to the target season but may be too small once honest "
    "chronological splits are applied. An all-history scope provides more training "
    "data but risks introducing noise from defunct franchises. Resolving this tension "
    "required empirical comparison of both approaches before the all-active-history "
    "scope was selected as the superior configuration."
)
para(
    "The fourth challenge was probability calibration. A classifier that assigns a "
    "72 percent win probability to a state that empirically wins only 55 percent of "
    "the time is misleading regardless of its classification accuracy. Calibrating "
    "the win model required careful application of isotonic regression and verification "
    "using calibration-bucket analysis to confirm that the correction improved "
    "probability realism without sacrificing sharpness."
)
para(
    "Additional constraints included limited GPU resources for intensive training "
    "experiments, reliance on publicly available data without access to commercial "
    "player-tracking feeds, and the need to maintain software quality while managing "
    "the project on a student timeline. Each constraint influenced the final design "
    "choices described in the methodology chapter."
)

h2("1.9  Structure of the Report")
para(
    "The remainder of this report is organized into six further chapters, three "
    "appendices, and a references section. Each chapter builds upon the preceding "
    "ones, progressing from context and motivation through literature, methodology, "
    "results, implementation, and conclusions."
)
para(
    "Chapter 2 presents a review of the relevant literature on cricket analytics, "
    "machine learning for sports prediction, and structured-data modeling methodology. "
    "It identifies the gaps in prior work that motivated the key design choices of "
    "the present project and explains how each gap was addressed."
)
para(
    "Chapter 3 defines the problem statement formally, lists the specific research "
    "objectives and research questions, describes the scope boundaries, and enumerates "
    "the key project deliverables with a completion status review."
)
para(
    "Chapter 4 describes the complete methodology including system architecture, "
    "end-to-end workflow design, dataset construction and description, data "
    "preprocessing and feature engineering rules, model training strategy and "
    "exploratory data analysis, best-model selection criteria, and user interface "
    "creation."
)
para(
    "Chapter 5 presents the experimental results, beginning with the software and "
    "hardware environment, followed by detailed result analysis including model "
    "comparison tables, error analysis by phase and wicket state, calibration-bucket "
    "analysis, scenario-level interpretation, and a summary of key findings."
)
para(
    "Chapter 6 concludes the report with a summary of project achievements, key "
    "lessons learned, current limitations, and detailed future scope directions "
    "including specific technical extensions."
)
para(
    "Appendix 1 contains the prediction algorithm as an ordered step-by-step "
    "procedure. Appendix 2 describes the most analytically significant engineered "
    "features. Appendix 3 records sample prediction outputs across five representative "
    "match scenarios. The References section lists all cited sources."
)

h2("1.10  Conclusion")
para(
    "This introductory chapter has established the topic, context, motivation, scope, "
    "goals, benefits, relevance, challenges, and structure of the IPL Match Score and "
    "Win Prediction System project. The project addresses a genuine analytical "
    "challenge — providing real-time, probabilistically calibrated predictions during "
    "live IPL matches — through a combination of careful data engineering, comparative "
    "model evaluation, probability calibration, and production-quality software design."
)
para(
    "The project is relevant both to the academic context of final-year engineering "
    "evaluation and to the broader field of applied machine learning in sports "
    "analytics. Its most important contribution is demonstrating that disciplined data "
    "engineering — leakage control, scope design, and calibration — matters more than "
    "the choice of algorithm for this class of problem. This lesson, supported by the "
    "empirical evidence presented in Chapter 5, is one of the central findings of the "
    "project."
)
para(
    "Chapter 2 begins with a review of the literature that shaped the project's design "
    "philosophy and justified the specific technical approaches adopted."
)

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 2 – LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════
chapter_title("Chapter 2 – Literature Review")

h2("2.1  Introduction")
para(
    "The literature on cricket prediction and tabular machine learning is rich enough "
    "to offer useful methodological guidance but sparse enough in the live, ball-by-ball "
    "setting to leave important design gaps. This chapter surveys the most relevant "
    "prior work, organizes it into thematic streams, identifies what each stream "
    "contributed to the broader field, and explains where each stream fell short in "
    "ways that directly motivated the design of the present project. The review is not "
    "an exhaustive bibliography: it is a design-justification tool. By understanding "
    "what earlier work accomplished and what it left undone, the methodological choices "
    "described in Chapter 4 become easier to defend."
)
para(
    "The review draws from three broad areas: classical cricket analytics heuristics, "
    "structured-data machine learning methodology, and applied sports prediction "
    "project reports. Within structured-data machine learning, particular attention is "
    "given to gradient-boosting methods, entity embeddings for high-cardinality "
    "categorical features, and probability calibration, since these are the techniques "
    "most directly evaluated or adopted in the project. The cricket-specific literature "
    "provides context about which problem formulations have been tried previously and "
    "what evaluation weaknesses they carried."
)
para(
    "One conclusion that emerged early from this review is that the quality of the "
    "live match-state representation matters more than the choice of model algorithm. "
    "Many prior systems achieved superficially strong results by training on features "
    "that embedded future information, using random train-test splits that mixed "
    "seasons together, or evaluating solely on classification accuracy while ignoring "
    "probability calibration quality. The present project was designed to systematically "
    "avoid each of these pitfalls."
)

h2("2.2  Review of Key Literature")
para(
    "Early cricket analytics work established current run rate, wickets remaining, and "
    "overs left as the primary predictors of match outcome. The Duckworth-Lewis method "
    "(Duckworth and Lewis, 1998) formalized the relationship between remaining batting "
    "resources and expected score for weather-interrupted One Day Internationals. "
    "This heuristic framework remains the official standard for interrupted matches "
    "and provides a useful lower bound for live prediction baselines. However, it was "
    "not designed for ball-by-ball prediction in uninterrupted T20 matches and does "
    "not incorporate venue-specific scoring environments, player form, or team-level "
    "historical performance. The gap identified from this stream is that hand-crafted "
    "resource-based rules cannot represent the full range of contextual signals that "
    "distinguish one live state from another with the same visible scoreboard."
)
para(
    "Sankaranarayanan, Sattar, and Lakshmanan (2014) demonstrated an early data-mining "
    "approach to ODI cricket simulation using match-level features and decision trees. "
    "While this work advanced the field toward data-driven methods, it operated on "
    "match-summary data rather than ball-by-ball states, limiting its relevance for "
    "live in-play prediction. Training on final-match summaries cannot teach a model "
    "how to respond to a specific intermediate state because it never observes the "
    "gradual unfolding of an innings. This gap directly motivated the present project's "
    "decision to build its dataset from ball-by-ball reconstruction."
)
para(
    "The gradient-boosting literature, exemplified by XGBoost (Chen and Guestrin, 2016) "
    "and CatBoost (Prokhorenkova et al., 2018), established that tree-based ensemble "
    "methods consistently outperform most alternatives on structured tabular data. "
    "XGBoost demonstrated that regularized boosted trees can be trained efficiently on "
    "large structured datasets with strong predictive quality. CatBoost extended this "
    "by introducing ordered, leakage-safe handling of categorical features through "
    "target statistics computed in a permuted order, which is particularly relevant "
    "when teams, venues, and players must be encoded as categories. Both frameworks "
    "were evaluated in the present project. The gap noted from this stream is that "
    "even a strong boosting model cannot recover performance lost to poor feature "
    "design or evaluation leakage: the model is only as honest as the data it trains on."
)
para(
    "Guo and Berkhahn (2016) proposed entity embeddings as an alternative to one-hot "
    "encoding for high-cardinality categorical variables in neural networks. Rather "
    "than representing each team or venue as a sparse binary indicator vector, the "
    "model learns a dense, low-dimensional embedding that captures similarity "
    "relationships between categories. This is especially attractive for cricket "
    "analytics because franchises, venues, and players carry rich identity-based signal "
    "that a learned embedding may represent more compactly and accurately than a "
    "one-hot expansion. The Torch entity-embedding model evaluated in the present "
    "project draws directly from this line of work and achieved the best score "
    "prediction result of all model families tested."
)
para(
    "Niculescu-Mizil and Caruana (2005) established that probability calibration is a "
    "distinct and necessary requirement separate from classification accuracy. A "
    "well-calibrated win model assigns a 60 percent win probability to states that "
    "empirically win approximately 60 percent of the time. An uncalibrated model may "
    "rank outcomes correctly and achieve high accuracy while still being systematically "
    "overconfident or underconfident in specific probability bands. For a live cricket "
    "win-probability system, calibration is especially critical because the displayed "
    "probability percentage is the primary output that users observe and interpret. The "
    "present project evaluated all win models using log loss and Brier score rather "
    "than accuracy alone, applying the lesson from this literature directly."
)
para(
    "Applied sports analytics project reports and academic papers demonstrate useful "
    "pipelines for data collection, feature construction, and model evaluation in "
    "various sports domains. These works are valuable because they show what full "
    "prediction pipelines look like in practice. However, a consistent pattern in "
    "these reports is that they stop at the notebook-experiment stage without "
    "integrating the trained model into a deployable, testable software system. They "
    "also tend to use random or time-agnostic data splits, which inflates reported "
    "performance by allowing future match information to contaminate the training set. "
    "The present project was explicitly designed to address both of these recurring "
    "weaknesses."
)
para(
    "Open cricket data sources such as Cricsheet (2024) make modern ball-by-ball "
    "cricket analytics possible by providing event-level records and match metadata "
    "in a machine-readable format. Cricsheet data is the sole data source for the "
    "present project's dataset. However, raw data availability does not automatically "
    "produce a usable training set. Alias normalization, ball-count reconstruction, "
    "support-table computation, and data quality validation are all prerequisites "
    "before the raw files can support trustworthy model training, and these "
    "preprocessing steps represent a substantial part of the project's total "
    "engineering effort."
)

h2("2.3  Summary of Findings")
para(
    "The literature review reveals several consistent findings that shaped the design "
    "of the present project. First, heuristic approaches based on run rate and wickets "
    "remaining are useful as interpretability baselines but cannot naturally represent "
    "the richer context of venue conditions, player form, or recent momentum. Any "
    "system that improves meaningfully over these heuristics must incorporate "
    "additional contextual signals through deliberate feature engineering."
)
para(
    "Second, gradient-boosting methods — particularly XGBoost and CatBoost — are the "
    "most consistently strong performers on structured tabular data with mixed numerical "
    "and categorical features. Neural network approaches such as entity-embedding "
    "models can outperform boosting on tasks where categorical identities carry "
    "significant signal, but they require larger training datasets and more careful "
    "regularization to justify the additional model complexity."
)
para(
    "Third, probability calibration is a non-negotiable requirement for win prediction "
    "systems that display probability values to users. Accuracy-focused evaluation is "
    "insufficient because a model can achieve high classification accuracy by exploiting "
    "class imbalance while still producing poorly calibrated, untrustworthy probability "
    "assignments. Log loss and Brier score are more informative evaluation metrics for "
    "this task and should be reported alongside accuracy."
)
para(
    "Fourth, chronological evaluation is essential for any sports prediction system "
    "intended for use in future seasons. Random train-test splits allow future "
    "information to leak into the training set and systematically overestimate "
    "real-world performance. Only a chronological split that trains on earlier seasons "
    "and tests on a genuinely future season produces honest, deployable performance "
    "estimates."
)
para(
    "Fifth, most applied cricket prediction projects stop at the experimental stage. "
    "There is a consistent and important gap between building a model that performs "
    "well on a held-out file and integrating that model into a deployable, testable, "
    "and maintainable software system with multiple user-facing surfaces."
)

h2("2.4  Addressing Identified Gaps")
para(
    "The present project addresses each identified gap through concrete design "
    "decisions. The gap in feature richness — the reliance on heuristic run-rate "
    "signals alone — is addressed by constructing a 62-column live-state feature set "
    "that combines momentum indicators, venue history, player profiles, matchup "
    "records, and chase-pressure signals. No single feature family dominates the "
    "representation; the combination produces a state description far more informative "
    "than any single heuristic."
)
para(
    "The gap in evaluation integrity — the use of random or time-agnostic data splits "
    "— is addressed by enforcing strict chronological ordering throughout the pipeline. "
    "Support tables are computed match by match in date order. Training, validation, "
    "and test partitions are defined by season boundaries. No match-state row "
    "belonging to a later season is ever included in the training data for a model "
    "evaluated on that season."
)
para(
    "The gap in probability calibration is addressed by applying isotonic regression "
    "calibration to the win model after initial training and verifying the improvement "
    "using log loss, Brier score, and a five-bucket calibration-gap analysis. The "
    "project reports calibration curves that show how closely the predicted win "
    "probability tracks the empirical win rate across all probability bands from "
    "zero to one hundred percent."
)
para(
    "The gap in deployment beyond notebooks is addressed by implementing a shared "
    "inference module called by Flask, Streamlit, API, and CLI surfaces. Automated "
    "pytest suites cover alias normalization, feature frame construction, overs "
    "parsing, and web routes. The project can be rebuilt from raw data to a working "
    "prediction service by running a documented sequence of commands, demonstrating "
    "genuine reproducibility rather than notebook-only repeatability."
)

h2("2.5  Importance of This Review to the Project")
para(
    "The literature review was conducted before the major design decisions were "
    "finalized and directly shaped those decisions rather than serving as a post-hoc "
    "citation exercise. The emphasis on gradient boosting as the primary model family "
    "came from the consistent evidence that tree-based methods outperform alternatives "
    "on structured tabular data. The decision to evaluate win models using log loss "
    "rather than accuracy came from the calibration literature. The decision to "
    "enforce chronological splitting came from identifying the evaluation weaknesses "
    "in prior sports prediction projects."
)
para(
    "The review also justified the modeling order adopted in the project: establish "
    "a CPU gradient-boosting baseline before investing time in more complex "
    "architectures. This ordered approach prevented the report from making "
    "exaggerated claims about model sophistication without empirical evidence to "
    "support them, and saved significant development time that was redirected toward "
    "feature engineering and evaluation quality."
)
para(
    "Finally, the review reinforced the project's commitment to deployment quality. "
    "Because prior work so consistently stopped at the experimental stage, the team "
    "treated deployment as a first-class objective rather than an optional extension. "
    "This decision is reflected throughout the implementation chapter, where the "
    "shared inference contract, the test suite, the model registry, and the four "
    "interface surfaces are documented alongside the modeling results. The literature "
    "review provided the intellectual justification for investing in these engineering "
    "components at least as heavily as in the modeling components."
)

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 3 – PROBLEM OBJECTIVE
# ══════════════════════════════════════════════════════════════════════
chapter_title("Chapter 3 – Problem Objective")

h2("3.1  Problem Statement")
para(
    "The central problem addressed by this project is the live prediction of IPL match "
    "outcomes using ball-by-ball match state information. Given a snapshot of an "
    "ongoing IPL innings — specifying the current runs scored, wickets lost, overs "
    "bowled, venue, batting team, bowling team, innings number, and optional player "
    "and momentum context — the system must produce two analytical outputs: an "
    "estimate of the likely final innings total and the batting side's probability of "
    "winning the match."
)
para(
    "This problem is harder than post-match classification because the information "
    "available at prediction time is necessarily incomplete. Early in the innings, the "
    "current score provides little information about the final total: a team at 42/1 "
    "after five overs might finish at 150 or at 220 depending on how the remaining "
    "deliveries are played. The model must therefore learn not just the average "
    "trajectory from a given state but also the uncertainty associated with that "
    "trajectory, which is why the project reports projected scoring ranges alongside "
    "point estimates."
)
para(
    "The second innings adds further complexity. A chasing team's win probability "
    "depends not only on its current score and resources remaining but on how its "
    "run rate compares to the required run rate. A score of 95/2 after eleven overs "
    "may indicate control or danger depending on whether the target is 160 or 200. "
    "The model must differentiate between these situations without observing the future."
)
para(
    "Formally, the project solves two related learning problems. The first is a "
    "regression problem: given the live match-state feature vector X at any legal "
    "ball b, learn a function f such that f(X) approximates the final innings total T. "
    "The second is a classification problem: given the same or a related feature vector "
    "X, learn a calibrated function g such that g(X) estimates P(win | X), the "
    "probability that the batting side wins the match. Both problems are evaluated under "
    "chronological train-test splitting, with MAE and RMSE as metrics for the "
    "regression task and log loss and Brier score for the probability task."
)
para(
    "The project deliberately treats these as separate tasks requiring different model "
    "families because the properties that produce accurate score estimates do not "
    "always produce well-calibrated win probabilities. Forcing both tasks into one "
    "shared architecture would trade quality on at least one task for superficial "
    "architectural simplicity."
)

h2("3.2  Objectives of the Project")
para(
    "The objectives of the project are divided into immediate technical goals and "
    "broader academic goals. The technical goals define what the system must be able "
    "to do and how it will be evaluated. The academic goals define how the project "
    "must be documented and how design decisions must be justified."
)
para("The technical objectives are:")
bullet("Ingest and normalize 1,189 raw IPL match files, resolving all known team and venue name variants through a comprehensive alias normalization table.")
bullet("Reconstruct every legal ball state within each innings to produce 282,552 training rows and 62 feature columns without any form of temporal data leakage.")
bullet("Compute nine historical support tables — venue stats, team form, team-venue form, matchup form, batter profiles, bowler profiles, batter-bowler matchup history, active-team list, and player pool — in strict chronological order.")
bullet("Compare at least four model families (HistGradientBoosting, XGBoost, CatBoost, Torch entity-embedding) for each prediction task under identical chronological evaluation conditions.")
bullet("Select the best-performing model pair, persist it in a versioned registry, and produce a held-out test evaluation report for both tasks.")
bullet("Apply and verify isotonic-regression probability calibration for the win model using log loss, Brier score, and calibration-bucket analysis.")
bullet("Implement a shared inference module callable from Flask, Streamlit, REST API, and CLI without duplicating feature engineering logic.")
bullet("Build an automated pytest suite covering alias normalization, feature frame construction, overs parsing, and web route correctness.")
bullet("Implement a monitoring module that logs prediction events and compares predicted outcomes to resolved match results for drift detection.")
para("The academic objectives are:")
bullet("Produce a project report that traces every major result back to a specific engineering decision, making the improvement sequence reproducible and defensible in viva examination.")
bullet("Demonstrate chronologically honest evaluation throughout, so that reported performance figures reflect realistic future-season conditions.")
bullet("Apply software engineering principles including modularity, low coupling, separation of concerns, and automated testing to demonstrate that the project is built as a software system and not merely as a notebook experiment.")

h2("3.3  Scope of the Project")
para(
    "The scope boundaries of the project were determined through a combination of "
    "technical feasibility analysis, available data access, and academic timeline "
    "constraints. These boundaries ensure that the project remains achievable while "
    "still constituting a substantive technical contribution."
)
para(
    "The project is limited to the Indian Premier League in T20 format. Seasons covered "
    "range from 2007/08 through the beginning of the 2026 season using data available "
    "through the Cricsheet open-data archive. The category space is restricted to the "
    "ten currently active IPL franchises for training and evaluation, though historical "
    "records from earlier seasons involving now-defunct franchises are used to build "
    "the full support-table infrastructure."
)
para(
    "Within the IPL setting, the prediction task covers both first-innings and "
    "second-innings states during standard twenty-over matches. Reduced-over matches "
    "resulting from weather interruptions are not fully handled in the current "
    "implementation. The system operates in a predictive and analytical mode: it "
    "generates score estimates and win probabilities but does not provide tactical "
    "recommendations, identify optimal bowling changes, or rank individual player "
    "performances."
)
para(
    "The deployment scope includes four interfaces — Flask web application, Streamlit "
    "dashboard, REST API endpoint, and CLI — but does not include production deployment "
    "on a publicly accessible server with automated live data ingestion. Commercial "
    "data API integration and real-time scorecard scraping are also out of scope. The "
    "chosen scope is appropriate for an academic submission and provides enough breadth "
    "to demonstrate all the technical capabilities described in the project objectives."
)

h2("3.4  Key Deliverables")
para(
    "The project produces a set of concrete deliverables that can be independently "
    "verified. Together they demonstrate that the stated objectives have been met. "
    "Each deliverable corresponds to at least one technical objective."
)
bullet("Processed IPL live-state dataset (ipl_features.csv): 282,552 rows, 62 columns, keyed by match_id, innings, and legal_balls_bowled.")
bullet("Nine historical support tables: venue_stats.csv, team_form_latest.csv, team_venue_form_latest.csv, matchup_form_latest.csv, batter_form_latest.csv, bowler_form_latest.csv, batter_bowler_form_latest.csv, active_teams_2026.csv, team_player_pool_2026.csv.")
bullet("Trained live score model artifact: Torch entity-embedding regressor (held-out MAE 15.36, RMSE 19.81) under the all-active-history scope.")
bullet("Trained and calibrated live win-probability model artifact: CatBoost GPU with isotonic calibration (held-out log loss 0.4922) under the all-active-history scope.")
bullet("Pre-match score and win model artifacts for use before the first delivery of a match.")
bullet("Shared inference module (ipl_predictor Python package) implementing consistent feature construction and model invocation across all surfaces.")
bullet("Flask web application with a live prediction form, results display, and REST API endpoint.")
bullet("Streamlit analytics dashboard for interactive match-state exploration.")
bullet("Command-line interface supporting scripted and demonstration use.")
bullet("Automated pytest suite with coverage for alias normalization, overs parsing, feature frame integrity, and web routes.")
bullet("Monitoring module (ipl_predictor/monitoring.py) logging prediction events and drift metrics.")
bullet("Model registry (models/model_registry.json) recording all evaluated configurations with performance metadata and artifact pointers.")
bullet("This project report documenting the complete methodology, results, implementation, and conclusions.")
para(
    "The completion status of all major deliverables is verified through artifact "
    "traceability records in the implementation chapter, which lists each key output "
    "file, its format, its approximate size, and its last modification timestamp. "
    "This traceability allows deliverable completeness to be confirmed by direct "
    "inspection of the repository rather than relying solely on report claims."
)

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 4 – METHODOLOGY AND SYSTEM DESIGN
# ══════════════════════════════════════════════════════════════════════
chapter_title("Chapter 4 – Methodology and System Design")

h2("4.1  System Architecture")
para(
    "The system is built as a three-layer architecture in which data concerns, learning "
    "concerns, and presentation concerns are separated into distinct components. This "
    "separation follows the software engineering principle of low coupling: each layer "
    "knows only what it must know about adjacent layers, and changes in one component "
    "do not propagate unexpectedly to others. The three layers are the data layer, the "
    "learning layer, and the presentation layer."
)
para(
    "The data layer contains the raw Cricsheet IPL match files, the preprocessing "
    "scripts that convert those files into the engineered feature dataset, and the "
    "nine support tables that encapsulate historical evidence about venues, teams, and "
    "players. All data-layer components write their outputs to well-defined artifact "
    "locations — CSV files in the data/processed/ directory — so that downstream "
    "stages can rely on stable file interfaces rather than embedding data-access "
    "logic inside training or inference code."
)
para(
    "The learning layer contains the training scripts, the model registry, the "
    "evaluation JSON reports, and the calibration pipelines. Training scripts operate "
    "on the engineered dataset and write model artifacts and report files to the "
    "models/ directory. The model registry (model_registry.json) maintains version "
    "records for every trained configuration, enabling the project to identify the "
    "strongest historical snapshot rather than defaulting to the most recently created "
    "artifact."
)
para(
    "The presentation layer contains the shared inference module and the four user-"
    "facing surfaces. The shared inference module is the only component in the system "
    "that loads model artifacts and calls them for prediction. All four surfaces — "
    "Flask, Streamlit, REST API, and CLI — delegate to this module, ensuring that the "
    "same feature engineering logic, alias normalization rules, and model artifacts "
    "are applied regardless of which interface the user accesses."
)
para(
    "The architecture supports extensibility at all layers. Adding a new interface "
    "surface requires only that the new surface implement the same input-output "
    "contract as the inference module. Retraining with improved data requires only "
    "updating the model artifacts and registry record; no interface code changes. "
    "Data quality tests, model performance tests, and interface contract tests are "
    "fully independent of each other."
)

h2("4.2  Workflow and System Model")
para(
    "The end-to-end workflow consists of eight sequential stages, each producing "
    "artifacts consumed by subsequent stages. This staged structure was adopted to "
    "prevent the most common failure mode in student machine learning projects: the "
    "monolithic notebook in which data loading, feature engineering, model training, "
    "and result presentation are entangled in a single file that cannot be reproduced, "
    "tested, or extended without significant rework."
)
para(
    "Stage 1 (External Data Refresh) downloads updated Cricsheet IPL match files and "
    "optional weather data snapshots. Stage 2 (Feature Engineering) processes raw "
    "files through the preprocessing pipeline, producing ipl_features.csv and all "
    "nine support tables. Stage 3 (CPU Baseline Training) trains HistGradientBoosting "
    "models and writes the baseline model report. Stage 4 (GPU Benchmark Training) "
    "trains XGBoost and CatBoost GPU models under the recent active-team scope. Stage 5 "
    "(Best-Model Search) compares all model families across both training scopes. "
    "Stage 6 (Pre-Match Modeling) trains the pre-match score and win models. Stage 7 "
    "(Serving Surface Setup) configures Flask, Streamlit, API, and CLI to use the "
    "promoted model artifacts. Stage 8 (Monitoring Initialization) sets up the "
    "prediction event logger and drift-detection module."
)
para(
    "Each stage is implemented as one or more Python scripts that can be executed "
    "independently or in sequence. This design allows individual stages to be re-run "
    "after changes without reprocessing the entire pipeline. If a new IPL season's "
    "data becomes available, only stages 1 and 2 need to be re-run before retraining "
    "in stages 3 through 5. The serving surfaces in stage 7 do not need to change "
    "unless the inference contract changes."
)
para(
    "The system can also be understood as a data flow narrative. Raw ball-by-ball files "
    "are converted into canonical match records with normalized identities. Those "
    "records are reconstructed into legal-ball-ordered feature rows enriched with "
    "support-table context. The enriched rows train and evaluate the models. The best "
    "models are loaded by the inference module. When a user submits a live match state "
    "through any surface, the inference module constructs a feature row, loads the "
    "relevant support-table values, and calls both models to produce the prediction "
    "response returned to the user."
)

h2("4.3  Dataset Description")
para(
    "The project dataset originates from 1,189 raw IPL match files downloaded from "
    "the Cricsheet open-data archive in CSV2 format. Each Cricsheet CSV2 match "
    "consists of a ball-by-ball event file recording every delivery in the match and "
    "an info metadata file recording match-level attributes such as venue, date, "
    "toss decision, teams, and result. The ball-by-ball files make it possible to "
    "reconstruct exactly what the scoreboard looked like at every delivery, which is "
    "the essential prerequisite for building a genuine live prediction system."
)
para(
    "After preprocessing, the core training dataset contains 282,552 rows and 62 "
    "columns. Each row represents one legal-ball state identified by match ID, innings "
    "number, and legal balls bowled. The 62 columns capture multiple information types: "
    "basic scoreboard state (runs, wickets, overs), derived ball counts and phase "
    "indicators (powerplay, middle, death), current and required run rates, momentum "
    "windows from recent deliveries and overs, venue historical statistics, team form "
    "indices, team-venue interaction features, matchup form scores, and player-level "
    "strike rate and economy rate profiles."
)
para(
    "The dataset spans eighteen complete seasons and the beginning of the 2026 season. "
    "The 2007/08 season contributes 13,489 rows from 58 matches with an average first-"
    "innings total of 160.97 runs. The 2024 season contributes 17,103 rows from 71 "
    "matches with an average first-innings total of 189.59 runs, and the 2025 season "
    "contributes 17,275 rows with an average first-innings total of 188.84 runs. This "
    "upward trend in scoring confirms why random train-test splitting is invalid: a "
    "model trained on 2015 data would systematically underestimate the scoring rates "
    "seen in 2024 and 2025."
)
para(
    "The dataset contains two target variables. The regression target is the final "
    "innings total, a continuous value ranging from approximately 80 to 260 across "
    "the historical archive. The classification target is a binary indicator of "
    "whether the batting side won the match, available for all second-innings rows and "
    "requiring special handling for first-innings rows where the final outcome cannot "
    "yet be observed."
)
para(
    "The chronological train-validation-test split uses seasons up to and including "
    "2023 for training, the 2024 season for validation, and the 2025 season as the "
    "held-out test partition. This split ensures that all reported performance figures "
    "reflect how the model would behave on a genuinely unseen future season."
)

h2("4.4  Data Preprocessing and Feature Engineering")
para(
    "Data preprocessing was the most time-intensive component of the project and the "
    "single greatest source of performance improvement. The preprocessing pipeline "
    "transforms raw ball-by-ball Cricsheet files into the structured feature dataset "
    "through a sequence of operations that each address a specific data-quality or "
    "representation challenge. The pipeline is implemented in scripts/preprocess_ipl.py "
    "and can be rerun from scratch in under twenty minutes on standard hardware."
)
para(
    "The first operation is alias normalization. IPL team and venue names appear in "
    "multiple textual variants across the 1,189 match files. A comprehensive alias "
    "table maps all known variants to their canonical forms before any feature "
    "computation begins. For example, the franchise currently named Royal Challengers "
    "Bengaluru has appeared as Royal Challengers Bangalore in many historical files, "
    "and Wankhede Stadium has appeared in truncated form in some match info files. "
    "Any alias passing through normalization uncorrected creates a spurious identity "
    "that weakens support-table statistics and corrupts model category encodings."
)
para(
    "The second operation is legal-ball conversion. Cricsheet records deliveries using "
    "an over.ball notation. The pipeline converts this to a canonical legal-ball count "
    "handling wide deliveries, no-balls, and extra deliveries, then derives "
    "legal_balls_bowled, balls_left, over_number, ball_in_over, innings_progress, "
    "and phase indicators (powerplay for overs 1-6, middle for overs 7-15, death "
    "for overs 16-20) consistently for every row. This conversion is a prerequisite "
    "for all downstream feature computations."
)
para(
    "The third and most important operation is support-table computation. Nine "
    "separate support tables are computed in strict chronological order, match by "
    "match: venue historical averages and bat-first win rates; team rolling form "
    "scores; team-venue form interaction indices; matchup form between batting team "
    "and bowling team; batter career strike rates and averages; bowler career economy "
    "and strike rates; batter-versus-bowler head-to-head micro-history; the active-"
    "team list; and the current-season player pool. Computing these tables in "
    "chronological order guarantees that no future information contaminates any "
    "feature value, which is the project's most critical correctness guarantee."
)
para(
    "The fourth operation is momentum window computation. For every ball state, the "
    "pipeline computes rolling sums of runs scored and wickets lost over the last "
    "one, two, and five overs. These momentum features capture short-term match "
    "trajectory that cumulative totals cannot represent. A team at 90/2 after eleven "
    "overs that scored 45 in the last five overs is in a very different state from "
    "the same team if it scored only 25 in those overs."
)
para(
    "The fifth operation is chase-feature population. For second-innings rows, the "
    "pipeline adds target, target_remaining, required_run_rate, and the difference "
    "between required and current run rate. These fields remain null for first-innings "
    "rows so that the model does not have access to target information when predicting "
    "first-innings outcomes. After all operations, internal consistency validation "
    "checks confirm that overs map correctly to legal balls, innings boundaries are "
    "properly marked, second-innings rows have valid target values, and no support-"
    "table feature contains values sourced from future matches."
)

h2("4.5  Model Training")
para(
    "The model training strategy was designed to answer one specific question "
    "rigorously: given identical features and an identical evaluation protocol, which "
    "model family best suits each prediction task? This question was addressed through "
    "a structured comparison of four model families evaluated under chronological "
    "train-validation-test splits."
)
para(
    "The comparison proceeded in three phases. The first phase established a CPU "
    "baseline using scikit-learn's HistGradientBoosting regressor and classifier. "
    "These models are robust, fast to train, and handle missing feature values "
    "natively, making them an ideal starting point. The second phase introduced "
    "XGBoost and CatBoost with GPU-accelerated training, exploring whether the "
    "additional model capacity and CatBoost's ordered categorical handling would "
    "improve on the baseline. The third phase introduced the Torch entity-embedding "
    "model, which learns dense representations for the high-cardinality categorical "
    "features representing teams, venues, and players."
)
para(
    "Each model family was evaluated under two training scopes. The recent-only scope "
    "uses only the most recently completed seasons, minimizing temporal distance "
    "between training and test data. The all-active-history scope uses all historical "
    "seasons but restricts the category space to the ten currently active IPL "
    "franchises. Empirical comparison showed the all-active-history scope to be "
    "substantially superior because it provides the large sample sizes that player-"
    "level and venue-level features require to stabilize."
)
para(
    "The chronological split used the 2024 season as the validation partition and the "
    "2025 season as the held-out test partition. Hyperparameters for gradient-boosting "
    "models used default regularized values, since the project's emphasis was on scope "
    "and feature design rather than hyperparameter optimization. The Torch entity-"
    "embedding model used a two-hidden-layer architecture with dropout regularization "
    "and the Adam optimizer."
)

h3("4.5.1  Exploratory Data Analysis")
para(
    "Exploratory data analysis (EDA) was conducted before model training to understand "
    "the distribution of the engineered dataset and identify potential issues with "
    "class balance, feature distributions, and missing values. The EDA informed "
    "several preprocessing design decisions and set realistic expectations for model "
    "performance across different match phases."
)
para(
    "Phase-wise analysis revealed an uneven distribution of ball states across innings "
    "phases. In the first innings, the powerplay phase (overs 1-6) contributes 44,519 "
    "rows, the middle phase (overs 7-15) contributes 65,705 rows, and the death phase "
    "(overs 16-20) contributes 36,267 rows. In the second innings, the powerplay "
    "contributes 44,217 rows, the middle contributes 63,937 rows, and the death phase "
    "contributes only 27,907 rows. The smaller death-overs sample in the second innings "
    "reflects the fact that many chases either conclude before the twentieth over or "
    "collapse well before it. This imbalance means that feature quality in the middle "
    "overs has an outsized influence on overall model performance."
)
para(
    "Team-level EDA showed considerable variation in historical batting performance "
    "across the ten active franchises. Mumbai Indians appears in 281 innings with an "
    "average total of 162.90 and a batting-side win rate of 54.9 percent. Chennai "
    "Super Kings appears in 255 innings with an average total of 163.98 and the "
    "highest win rate of 56.3 percent. Gujarat Titans, the newest active franchise "
    "with only 64 innings in the dataset, shows the highest average total of 177.97 "
    "and the highest batting-side win rate of 60.9 percent, reflecting their strong "
    "early form since joining the IPL."
)
para(
    "Season-wise EDA confirmed a clear upward trend in IPL scoring. Average first-"
    "innings totals rose from approximately 161 runs in 2007/08 to approximately 189 "
    "runs in 2024 and 189 runs in 2025. This multi-year scoring inflation justifies the "
    "use of venue-adjusted features and confirms the necessity of chronological "
    "evaluation: a score of 170 was above average in 2010 but below average in 2024. "
    "A model that ignores this shift in scoring norms will systematically miscalibrate "
    "its predictions for modern seasons."
)
para(
    "Feature correlation analysis showed that innings_progress, balls_left, and "
    "current_run_rate are the strongest correlates of the final innings total among "
    "basic scoreboard features. Venue par, team form, and player-level features each "
    "added independent predictive signal after controlling for basic scoreboard state, "
    "confirming that the engineering investment in support tables was worthwhile."
)

h2("4.6  Figuring the Best Model")
para(
    "Selecting the best model required evaluating each candidate on the held-out 2025 "
    "test season using the appropriate metric for each task. For the live score task, "
    "the primary metric was MAE and the secondary metric was RMSE. For the live win "
    "task, the primary metric was log loss and the secondary was Brier score, with "
    "classification accuracy noted only as a reference."
)
para(
    "The live score comparison across all model families under the all-active-history "
    "scope showed the following held-out performance: HistGradientBoosting CPU "
    "baseline achieved MAE 17.1 and RMSE 22.4; XGBoost achieved MAE 16.2 and RMSE "
    "21.0; CatBoost achieved MAE 16.5 and RMSE 21.3; and the Torch entity-embedding "
    "model achieved the best result with MAE 15.36 and RMSE 19.81. The entity-"
    "embedding model's advantage is consistent with its ability to learn dense "
    "categorical representations that capture inter-franchise and inter-venue "
    "similarity more compactly than one-hot encodings."
)
para(
    "The live win comparison showed: HistGradientBoosting achieved log loss 0.5301; "
    "XGBoost achieved log loss 0.5087; uncalibrated CatBoost achieved log loss 0.5224; "
    "and calibrated CatBoost achieved the best result with log loss 0.4922. The "
    "improvement from calibration was measurable and meaningful, confirming that "
    "isotonic-regression post-processing adds value beyond what the raw classifier "
    "provides."
)
para(
    "The final selection paired Torch entity-embedding regressor for the score task "
    "with calibrated CatBoost GPU for the win task. This hybrid pair was saved in the "
    "model registry as version v20260415_152209 with full performance records, scope "
    "metadata, and artifact pointers. The registry also preserves records for all "
    "alternative configurations so that the selection decision can be verified by "
    "comparing the stored metrics rather than relying on narrative claims alone."
)
para(
    "The decision to use different model families for each task was justified by the "
    "empirical evidence. Score regression benefits from learned categorical embeddings "
    "when the dataset is large enough to support stable embedding training. Win "
    "probability classification benefits from CatBoost's ordered categorical encoding "
    "and its compatibility with post-hoc calibration. Forcing both tasks into the "
    "same architecture would sacrifice performance on at least one task without "
    "providing any simplification benefit."
)

h2("4.7  User Interface Creation")
para(
    "The user interface layer consists of four surfaces that all share the same "
    "underlying inference module. This design was adopted to ensure identical "
    "prediction quality regardless of how the user accesses the system and to avoid "
    "the maintenance burden of duplicating feature engineering logic across interfaces."
)
para(
    "The Flask web application provides a browser-based form where users enter the "
    "current match state — team names, venue, innings, score, wickets, overs, and "
    "optionally the striker and bowler names — and receive a prediction response "
    "including projected total, win probability, projected scoring range, runs versus "
    "par, and a contextual narrative. HTML templates (templates/index.html) and a "
    "dedicated CSS stylesheet (static/styles.css) handle front-end presentation. "
    "Server-side input validation precedes every inference call."
)
para(
    "The Streamlit analytics dashboard provides an interactive exploration interface. "
    "Users adjust match state parameters using sliders and dropdown menus and observe "
    "predictions change in real time. The dashboard also displays historical team and "
    "venue statistics drawn from the support tables, providing factual context "
    "alongside the live model outputs."
)
para(
    "The REST API endpoint accepts JSON payloads in the same format as the Flask "
    "form and returns JSON responses with the same prediction fields. This interface "
    "enables programmatic access, supporting integration into external applications, "
    "automated batch evaluation, and testing pipelines."
)
para(
    "The command-line interface accepts match state as command-line arguments and "
    "returns a formatted prediction summary to the terminal. This surface is most "
    "useful for debugging, scripted demonstration, and rapid manual testing during "
    "development. All four surfaces call the same shared inference function with the "
    "same input normalization, alias resolution, and feature engineering logic, "
    "guaranteeing consistent outputs."
)

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 5 – RESULTS
# ══════════════════════════════════════════════════════════════════════
chapter_title("Chapter 5 – Results")

h2("5.1  Software and Hardware Details")
para(
    "The experimental results reported in this chapter were obtained in a specific "
    "software and hardware environment. Documenting this environment is essential for "
    "reproducibility: another researcher attempting to replicate these results should "
    "be able to configure an equivalent environment and obtain comparable performance "
    "figures. The environment description also provides context for interpreting "
    "the computational cost and feasibility of the different training configurations."
)

h3("5.1.1  Software Requirements")
para(
    "The project was developed and evaluated using Python 3.11 managed through a "
    "virtual environment to isolate all dependencies from the system Python "
    "installation. All package versions were pinned in requirements.txt and "
    "pyproject.toml to ensure reproducibility. The core software dependencies are:"
)
bullet("Python 3.11 — primary programming language for all preprocessing scripts, training pipelines, inference modules, and interface code.")
bullet("NumPy >= 1.26, < 3 — vectorized numerical operations used throughout feature processing and metric computation.")
bullet("Pandas >= 2.2, < 3 — tabular data ingestion, groupby operations, join operations, and report-side analytics.")
bullet("Scikit-learn >= 1.7, < 1.8 — HistGradientBoosting models, evaluation metrics (MAE, RMSE, log loss, Brier score), CalibratedClassifierCV for isotonic-regression calibration, and preprocessing utilities.")
bullet("Joblib >= 1.4, < 2 — serialization and reloading of fitted model artifacts.")
bullet("Flask >= 3.1, < 4 — web application layer providing the live prediction form and the REST API endpoint.")
bullet("Streamlit >= 1.44, < 2 — interactive analytics dashboard for project demonstration and prediction exploration.")
bullet("XGBoost >= 2.1, < 3 — high-performance gradient-boosting benchmark with optional GPU acceleration via CUDA.")
bullet("CatBoost >= 1.2, < 2 — categorical-friendly gradient boosting used for the live win-probability model.")
bullet("PyTorch >= 2.3 — deep learning framework used for the entity-embedding tabular regression model.")
bullet("Matplotlib >= 3.9, < 4 — generation of report figures, calibration curves, and error distribution plots.")
bullet("Pytest >= 8.4, < 9 — automated verification of inference utilities, alias normalization, and web routes.")

h3("5.1.2  Additional Software Tools")
para(
    "Several additional tools were used for documentation generation, data management, "
    "and application infrastructure:"
)
bullet("Python-docx — programmatic generation and modification of the DOCX project report from workspace artifacts and model registry data.")
bullet("Git — version control for all source code, scripts, configuration files, and intermediate artifacts.")
bullet("Alembic and SQLAlchemy — database migration management and ORM for the user authentication and session storage layer of the Flask application.")
bullet("python-dotenv — environment variable management for configuration parameters and API keys.")
bullet("python-docx — automated report generation from model registry snapshots and workspace artifact inventories.")

h3("5.1.3  Hardware Requirements")
para(
    "The project was developed and trained on a Windows-based personal computing "
    "environment running Windows 11. CPU-based training experiments — HistGradientBoosting, "
    "baseline XGBoost, and CatBoost in CPU mode — were run using all available "
    "processor cores on the development machine. GPU-accelerated experiments used an "
    "NVIDIA GPU with CUDA support, enabling XGBoost and CatBoost to use their "
    "GPU-accelerated training backends."
)
para(
    "The minimum hardware requirements for reproducing the training experiments are: "
    "a modern multi-core CPU (four or more physical cores recommended for parallel "
    "tree construction), 16 GB of RAM to hold the full 282,552-row dataset and "
    "feature matrices in memory simultaneously during training, and approximately "
    "5 GB of disk space for raw match files, the processed dataset, support tables, "
    "and saved model artifacts. GPU acceleration is optional: all experiments can be "
    "re-run in CPU mode with longer training times but numerically equivalent results."
)

h2("5.2  Result Analysis")
para(
    "The results analysis addresses two central questions: which training scope "
    "produces better results, and which model family performs best on each task "
    "within the stronger scope. The scope comparison was completed first because "
    "its answer determined the training configuration used for all subsequent model "
    "comparisons."
)
para(
    "The scope comparison showed a decisive advantage for the all-active-history "
    "scope over the recent-only scope. In the recent-only setting, the live-score "
    "model trained on approximately 17,000 rows and the win model on a similar "
    "count. In the all-active-history setting, both models trained on over 250,000 "
    "rows. With chronological splits applied, this difference in training volume "
    "directly affects the stability of venue-level, team-level, and player-level "
    "feature estimates. The all-active-history scope improved score MAE by "
    "approximately 1.5 runs compared to the recent-only scope."
)
para(
    "The live score model comparison under the all-active-history scope produced "
    "the following held-out test results: HistGradientBoosting CPU baseline "
    "achieved MAE 17.1 and RMSE 22.4; XGBoost achieved MAE 16.2 and RMSE 21.0; "
    "CatBoost achieved MAE 16.5 and RMSE 21.3; and the Torch entity-embedding "
    "model achieved MAE 15.36 and RMSE 19.81. Each successive model family "
    "improved on the baseline, but no single step produced as large a gain as "
    "switching from the recent-only to the all-active-history training scope."
)
para(
    "The live win model comparison produced the following held-out test results: "
    "HistGradientBoosting achieved log loss 0.5301; XGBoost achieved log loss 0.5087; "
    "uncalibrated CatBoost achieved log loss 0.5224; and calibrated CatBoost "
    "achieved log loss 0.4922. The calibration step alone reduced log loss by "
    "approximately 0.03, confirming that isotonic regression post-processing adds "
    "meaningful improvement beyond what the raw classifier provides."
)
para(
    "Error analysis by wickets lost revealed a consistent and expected pattern: "
    "score prediction error decreases as wickets fall and batting resources narrow. "
    "When 0 to 2 wickets are lost, the model records MAE 19.25 and RMSE 25.71 — "
    "reflecting the wide range of possible final totals when most batting resources "
    "remain. When 3 to 5 wickets are lost, MAE drops to 13.37. When 6 to 8 wickets "
    "are lost, MAE is 11.45. When 9 to 10 wickets are lost, MAE is 7.54. This "
    "pattern shows that the model is most uncertain early, when the innings has not "
    "yet revealed its character, and most accurate late, when remaining resources "
    "constrain the possible total range."
)
para(
    "Phase-wise analysis confirmed that powerplay predictions carry the highest "
    "uncertainty because teams can launch aggressively, lose early wickets, or "
    "consolidate conservatively. Middle-overs predictions are more stable as the "
    "innings trajectory becomes clearer. Death-overs predictions are most accurate "
    "but still carry meaningful uncertainty when a hitting surge or collapse is "
    "underway."
)

h2("5.3  Discussion of Results")
para(
    "The strongest result reported — MAE 15.36, RMSE 19.81, log loss 0.4922 — is "
    "defensible on multiple grounds beyond its numerical values. First, it was "
    "obtained under chronological evaluation that prevents all forms of temporal "
    "data leakage. The 2025 test season was entirely unseen during training and "
    "validation. Second, it was drawn from a model registry that preserves the "
    "full comparison history, so the selected configuration can be verified against "
    "all alternatives with documented evidence. Third, scenario-level analysis "
    "confirms that the predictions behave sensibly across all match phases and "
    "chase situations."
)
para(
    "The pre-match model results provide valuable context for interpreting the live "
    "model performance. The pre-match score model achieves MAE approximately 25 runs "
    "— substantially worse than the live model's 15.36. The pre-match win model "
    "achieves log loss approximately 0.67 — substantially worse than the live model's "
    "0.4922. These weaker numbers are expected and intentional. Without live match "
    "context, the pre-match model can access only team identity, venue priors, and "
    "broad form indicators. The performance gap between the pre-match and live results "
    "quantifies exactly how much predictive value the ball-by-ball feature engineering "
    "adds, which is one of the key findings of the project."
)
para(
    "The calibration-bucket analysis reveals how the win model behaves across "
    "different confidence levels. In the 0 to 20 percent probability band, the model "
    "produced 1,117 held-out rows with an average predicted probability of 16.9 "
    "percent and an empirical win rate of 29.3 percent — a calibration gap of 12.4 "
    "percentage points. In the 40 to 60 percent probability band, the gap was 5.4 "
    "percentage points. In the 80 to 100 percent band, the gap was 20.1 percentage "
    "points in the overconfidence direction. These gaps confirm that calibration "
    "remains imperfect in the extreme probability bands where match states are "
    "genuinely most ambiguous, but show that the model behaves most reliably in "
    "the mid-range confidence region."
)
para(
    "The scenario interpretations illustrate the predictions in concrete cricket "
    "terms. In a powerplay launch scenario (batting team at 42/0 after 5 overs at "
    "a high-scoring venue), the system projects a final total of 194.1 with a win "
    "probability of 44.2 percent and a projected range of 175 to 222. The wide "
    "range correctly captures the substantial uncertainty at this early stage. In a "
    "chase-under-pressure scenario (second innings, requiring 81 from 9 overs with "
    "5 wickets remaining), the system returns a win probability of 44.2 percent, "
    "reflecting the genuinely even state of the contest."
)
para(
    "Perhaps the most instructive discussion concerns what did not improve performance "
    "significantly. Switching from HistGradientBoosting to XGBoost improved score "
    "MAE by approximately 0.9 runs. Switching from the recent-only scope to the "
    "all-active-history scope improved MAE by approximately 1.5 runs — nearly twice "
    "the gain from the algorithm switch. This comparison provides concrete evidence "
    "for the project's central methodological claim: data engineering and scope "
    "design contribute more to performance than algorithm selection for this class "
    "of structured tabular prediction problem."
)

h2("5.4  Summary")
para(
    "The results chapter demonstrates that the project achieved all its stated "
    "technical performance objectives. The live score model produced held-out MAE "
    "of 15.36 runs and RMSE of 19.81 runs under strict chronological evaluation "
    "on the 2025 held-out season. The live win model produced held-out log loss "
    "of 0.4922, with the calibration step contributing measurably to that result. "
    "Both models were trained under the all-active-history scope, which proved "
    "significantly superior to the recent-only alternative."
)
para(
    "Three broader conclusions emerge from this analysis. First, training scope "
    "matters more than algorithm choice: the difference between the all-active-"
    "history and recent-only scopes exceeded the difference between the best and "
    "worst model families. Second, probability calibration adds measurable value "
    "for the win task: the calibrated CatBoost model outperformed its uncalibrated "
    "counterpart by approximately 0.03 log loss points, which is a meaningful "
    "improvement in a probability quality metric. Third, the live feature "
    "representation substantially outperforms pre-match prediction, quantifying "
    "concretely the value added by ball-by-ball feature engineering over team-level "
    "and venue-level priors alone."
)
para(
    "These results should be read alongside the phase-wise and wicket-state "
    "breakdowns rather than treated as single global numbers. The model is strongest "
    "in mid-innings and late-innings states where the innings trajectory is clearer "
    "and weakest in early powerplay states where genuine uncertainty is highest. "
    "These patterns reflect the inherent structure of the prediction problem and "
    "are consistent with what an honest evaluation should reveal."
)

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 6 – CONCLUSION AND FUTURE SCOPE
# ══════════════════════════════════════════════════════════════════════
chapter_title("Chapter 6 – Conclusion and Future Scope")

h2("6.1  Conclusion")
para(
    "This project successfully designed and developed an IPL Match Score and Win "
    "Prediction System that addresses the live cricket prediction problem through "
    "a combination of disciplined data engineering, comparative model evaluation, "
    "probability calibration, and production-quality software design. The system "
    "processes 282,552 ball-by-ball match-state rows derived from 1,189 IPL matches "
    "and delivers calibrated predictions through a shared inference module accessible "
    "via four independent user-facing surfaces."
)
para(
    "The strongest experimental result, preserved in the model registry as version "
    "v20260415_152209, achieved a held-out MAE of 15.36 runs and RMSE of 19.81 runs "
    "for the live score task, and a held-out log loss of 0.4922 for the live win-"
    "probability task. These figures were produced by pairing a Torch entity-embedding "
    "regressor for score prediction with a calibrated CatBoost GPU classifier for win "
    "prediction, both trained under the all-active-history scope with strict "
    "chronological evaluation. The results improve meaningfully over heuristic "
    "baselines and over pre-match modeling, while remaining honest about the genuine "
    "uncertainty inherent in T20 cricket."
)
para(
    "The major lessons learned from the project are as follows. The most impactful "
    "improvements came from data engineering decisions rather than algorithm selection: "
    "removing temporal leakage, widening the training scope to the all-active-history "
    "setting, and constructing rich chronological support tables each contributed more "
    "to the final performance than switching between model families. The recent-only "
    "training scope, while chronologically close to the target season, was too data-"
    "poor to support stable player-level and venue-level feature learning once "
    "chronological splits were applied honestly. The score and win tasks reward "
    "different model properties, making a hybrid model pair preferable to any "
    "single-family solution. Probability calibration is essential for any system "
    "that displays win percentages to users, as uncalibrated probabilities erode "
    "trust even when classification accuracy appears high."
)
para(
    "The project also demonstrated that a strong final-year engineering submission "
    "must include more than a set of held-out metrics. Reproducible interfaces, "
    "automated tests, a versioned model registry, and clear documentation of design "
    "decisions collectively make the project defensible under scrutiny in a way "
    "that metrics alone cannot achieve. A reviewer asking why one model was chosen "
    "over another, why a specific scope was adopted, or what would happen if the "
    "system were retrained on new data can be given concrete answers traceable to "
    "empirical evidence and documented reasoning."
)
para(
    "The ethical and social dimensions of the project are straightforward. The system "
    "is designed as an analytical and educational tool. All predictions include "
    "explicit uncertainty information to prevent overconfident interpretation. "
    "Training data comes exclusively from publicly available sources. The system does "
    "not support gambling applications and is not connected to betting-market data "
    "in any way. These properties reflect the team's commitment to responsible "
    "application of artificial intelligence."
)
para(
    "Viewed as a complete engineering project, this submission satisfies the "
    "core expectations of a final-year B.Tech. project in Computer Science and "
    "Artificial Intelligence and Machine Learning: it defines a relevant problem, "
    "reviews the prior literature with sufficient depth to justify design choices, "
    "implements a concrete and documented methodology, evaluates results honestly "
    "under a leakage-free protocol, builds a working multi-surface software system, "
    "and documents everything in sufficient detail for independent reproduction "
    "and evaluation."
)

h2("6.2  Future Scope")
para(
    "The project establishes a solid technical foundation upon which several high-"
    "value extensions can be built. Each extension described here is practically "
    "feasible given the existing architecture and requires targeted development "
    "rather than fundamental redesign."
)
para(
    "The most valuable immediate extension would be integration of probable playing "
    "XI data before the toss. The current pre-match branch can access only team "
    "identity and venue priors. If the projected lineup for each team were available "
    "before the first delivery, the system could aggregate player-level batting and "
    "bowling quality indices into a team-strength composite that would substantially "
    "improve pre-match prediction accuracy and narrow the gap between pre-match and "
    "live performance."
)
para(
    "A second high-value extension is the adoption of conformal prediction or quantile "
    "regression to replace the current fixed uncertainty bands. The system currently "
    "generates projected scoring ranges by applying a fixed offset to the point "
    "estimate. A conformal prediction layer would generate adaptive intervals that "
    "widen or narrow based on the genuine uncertainty of the specific match state, "
    "providing more informative and statistically rigorous range estimates."
)
bullet("Integrate live data ingestion from a public cricket scoring API to enable fully automated real-time predictions during IPL matches without manual data entry.")
bullet("Expand the explanation layer so that API responses include the top contributing features to each prediction expressed in cricket-intelligible terms, improving transparency for end users.")
bullet("Extend the monitoring layer from passive event logging to active retraining triggers, so that when accumulated prediction drift crosses a configured threshold the system automatically queues a retraining workflow.")
bullet("Generalize the pipeline to other T20 leagues including the Big Bash League, Caribbean Premier League, and SA20 by parameterizing alias tables, active-team filters, and venue lists for each competition.")
bullet("Evaluate transformer-based tabular models such as TabTransformer and FT-Transformer as the dataset grows with additional seasons, since these architectures may capture longer-range feature interactions that gradient boosting cannot model.")
bullet("Build a natural-language commentary layer that converts model prediction outputs into match-narrative text, enabling the system to serve as an automated analytical commentary companion.")
para(
    "Each of these extensions builds naturally on the engineering infrastructure "
    "already in place. The shared inference contract, the support-table architecture, "
    "the monitoring module, and the model registry were all designed with future "
    "extensibility in mind. The project therefore provides not only a working "
    "prediction system but also a validated architectural template that a future "
    "developer can extend efficiently without rebuilding from scratch."
)
para(
    "In conclusion, the IPL Match Score and Win Prediction System represents a "
    "complete, disciplined, and honest machine learning project that delivers "
    "practical analytical value while maintaining the rigorous evaluation standards "
    "necessary for both academic credibility and real-world trustworthiness. The "
    "methodology, results, implementation, and lessons documented in this report "
    "provide a reusable reference for future applied sports analytics work in this "
    "and related domains."
)

# ══════════════════════════════════════════════════════════════════════
# APPENDIX 1
# ══════════════════════════════════════════════════════════════════════
chapter_title("Appendix 1 – Flowchart and Algorithm Notes")
para(
    "This appendix presents the prediction system's core logic as an ordered "
    "step-by-step algorithm. It complements the architecture and workflow "
    "descriptions in Chapter 4 by expressing the same process in procedural form "
    "for readers who prefer an algorithmic representation."
)
numbered("Accept the current match inputs: batting team, bowling team, venue, innings number (1 or 2), runs scored, wickets lost, overs bowled, and target (second innings only).")
numbered("Normalize team and venue names using the canonical alias table so that all textual variants resolve to their standard identifiers before any feature computation begins.")
numbered("Convert the overs value to a legal-ball count and derive balls_left, phase indicators (powerplay, middle, death), innings_progress, and chase-pressure variables (second innings only).")
numbered("Load the relevant context from the nine support tables: venue historical averages, team rolling form scores, team-venue form indices, matchup form, and player-level batting and bowling profiles.")
numbered("Construct the live feature frame in the exact column order expected by the saved score and win model artifacts.")
numbered("Pass the feature frame to the Torch entity-embedding score model to obtain the projected innings total and derive the prediction range.")
numbered("Pass the feature frame to the calibrated CatBoost win model to obtain the batting-side win probability.")
numbered("Format and return the prediction outputs — projected total, win probability, projected range, runs versus par, and contextual narrative — to the selected interface surface.")
para(
    "The same algorithm is executed identically by all four interface surfaces "
    "(Flask, Streamlit, REST API, CLI), which is why predictions are consistent "
    "across all surfaces for the same match state input."
)

# ══════════════════════════════════════════════════════════════════════
# APPENDIX 2
# ══════════════════════════════════════════════════════════════════════
chapter_title("Appendix 2 – Selected Feature Notes")
para(
    "The following notes describe the analytical significance of the most important "
    "engineered features in the live-state dataset. Detailed reasoning for all "
    "feature families is provided in Chapter 4, Section 4.4."
)
bullet("runs_vs_par: Compares the current innings score with the historical venue baseline for the same innings progress, distinguishing whether the batting side is ahead of or behind the typical scoring rate for that ground.")
bullet("required_minus_current_rr: Captures whether a chasing side is falling behind the pace needed to win, and by exactly how many runs per over. This feature is populated only for second-innings rows.")
bullet("runs_last_5 and wickets_last_5: Summarize recent batting momentum and collapse risk over the most recent five legal balls, capturing trajectory information that cumulative totals cannot represent.")
bullet("batting_team_venue_form and bowling_team_venue_form: Encode each team's recent historical performance at the specific venue, preserving ground-specific strengths and weaknesses beyond the global team form score.")
bullet("batter_vs_bowler_history: Provides the model with a compact representation of the head-to-head performance between the current striker and the current bowler, enabling matchup-aware adjustments to predictions.")
bullet("innings_progress: A normalized measure of how far through the twenty overs the innings has advanced, computed as legal_balls_bowled divided by 120, providing a continuous phase indicator used by the model.")
bullet("venue_avg_first_innings and venue_bat_first_win_rate: Historical averages for the venue encoding whether the ground is a high-scoring or low-scoring environment and whether batting first is historically advantageous.")

# ══════════════════════════════════════════════════════════════════════
# APPENDIX 3
# ══════════════════════════════════════════════════════════════════════
chapter_title("Appendix 3 – Short Supporting Case Snapshots")
para(
    "The case snapshots below record sample prediction outputs from the system "
    "across five representative live match scenarios and three pre-match scenarios. "
    "All outputs were produced by model registry version v20260415_152209. These "
    "scenarios are discussed in Chapter 5."
)
para("Live Prediction Scenarios:")
bullet("Powerplay Launch — State: 42/0 after 5 overs (first innings, high-scoring venue). Output: Projected total 194.1, Win probability 44.2%, Range 175-222. Interpretation: Aggressive powerplay scoring pushes the projected total well above venue par; win probability remains near-even as the innings is still in its earliest phase.")
bullet("Middle-Over Rebuild — State: 74/3 after 10 overs (first innings). Output: Projected total 162.3, Win probability 63.4%, Range 143-190. Interpretation: Wicket loss in middle overs moderates the score estimate; the bowling side's probability rises as batting resources narrow.")
bullet("Death Overs Surge — State: 155/4 after 16 overs (first innings, neutral venue). Output: Projected total 197.7, Win probability 63.4%, Range 178-225. Interpretation: Batting team accelerating in death overs pushes projected total significantly above initial expectations.")
bullet("Chase Under Pressure — State: 96/5 after 11 overs (second innings, target 177). Output: Projected total 177.3, Win probability 44.2%, Range 158-205. Interpretation: Required rate is running ahead of current rate with few wickets remaining; odds remain genuinely near-even.")
bullet("Late Collapse Risk — State: 141/7 after 17 overs (first innings). Output: Projected total 163.8, Win probability 37.9%, Range 144-192. Interpretation: Multiple late wickets shift the projected total downward and favour the fielding side.")
para("Pre-Match Scenarios:")
bullet("Mumbai Indians vs Chennai Super Kings at Wankhede Stadium: Projected first-innings range 158-182. Win probability estimates reflect historical head-to-head record and venue scoring environment.")
bullet("Royal Challengers Bengaluru vs Sunrisers Hyderabad at M. Chinnaswamy Stadium: Projected first-innings range 161-185. The high-scoring Bengaluru venue inflates the projected upper range.")
bullet("Gujarat Titans vs Rajasthan Royals at Narendra Modi Stadium: Projected first-innings range 170-194. Gujarat Titans' strong batting record at this venue is reflected in the higher lower bound.")
para(
    "To reproduce any of these outputs, activate the project virtual environment and "
    "run the appropriate interface command with the scenario inputs listed above."
)

# ══════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"\nDocument saved -> {OUT}")
print("Restructuring complete.")
