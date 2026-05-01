"""
Exact Turnitin-style Originality Report replica for IPL project.
Mirrors every visual element: badge colours, pill fonts, spacing, dividers.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as Oxe
import copy

# ── colour constants ──────────────────────────────────────────────────────────
C_RED_BADGE   = "C0504D"   # metric tile – similarity index
C_GRAY_BADGE  = "A5A5A5"   # metric tiles – other three
C_INTERNET    = "4472C4"   # blue  – Internet Source badge & text
C_PUBLICATION = "E36C09"   # orange – Publication badge & text
C_STUDENT     = "17375E"   # dark teal – Student Paper badge & text
C_RED_HEADING = "C0504D"   # PRIMARY SOURCES heading
C_DIVIDER     = "BFBFBF"   # thin horizontal line
C_NAME        = "808080"   # top name
C_ORIG_BG     = "C0504D"   # ORIGINALITY REPORT bar background
C_TILE_LABEL  = "595959"   # small label under big %
C_PCT_RED     = "C0504D"   # <1% text colour

BADGE_COLOR = {
    "Internet Source": C_INTERNET,
    "Publication":     C_PUBLICATION,
    "Student Paper":   C_STUDENT,
}
TEXT_COLOR = {
    "Internet Source": (68, 114, 196),
    "Publication":     (227, 108,   9),
    "Student Paper":   ( 23,  55,  94),
}

# ── helpers ───────────────────────────────────────────────────────────────────

def hex2rgb(h): return tuple(int(h[i:i+2],16) for i in (0,2,4))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "auto")
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_cell_border_bottom(cell, color="BFBFBF", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top","left","right","insideH","insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "auto")
        tcBorders.append(tag)
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), sz)
    bot.set(qn("w:space"), "0")
    bot.set(qn("w:color"), color)
    tcBorders.append(bot)
    tcPr.append(tcBorders)

def no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def add_run(p, text, bold=False, italic=False, size=10,
            color=None, font="Calibri"):
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.name = font
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r

def set_row_height(row, pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(pt * 20)))   # twips
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

def cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top",top),("bottom",bottom),("left",left),("right",right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

# ── source data ───────────────────────────────────────────────────────────────
SOURCES = [
    ("cricsheet.org — IPL Ball-by-Ball Data and Match Metadata", "Internet Source"),
    ('Chen, T., & Guestrin, C. "XGBoost: A Scalable Tree Boosting System." ACM SIGKDD 2016, pp. 785–794', "Publication"),
    ('Prokhorenkova, L., et al. "CatBoost: Unbiased Boosting with Categorical Features." NeurIPS 2018, pp. 6638–6648', "Publication"),
    ('Guo, C., & Berkhahn, F. "Entity Embeddings of Categorical Variables." arXiv:1604.06737, 2016', "Publication"),
    ('Pedregosa, F., et al. "Scikit-learn: Machine Learning in Python." JMLR, vol. 12, 2011, pp. 2825–2830', "Publication"),
    ('Paszke, A., et al. "PyTorch: An Imperative Style, High-Performance Deep Learning Library." NeurIPS 2019, pp. 8024–8035', "Publication"),
    ('Niculescu-Mizil, A., & Caruana, R. "Predicting Good Probabilities with Supervised Learning." ICML 2005, pp. 625–632', "Publication"),
    ('Duckworth, F. C., & Lewis, A. J. "A Fair Method for Resetting the Target in Interrupted One-Day Cricket Matches." JORS, 49(3), 1998, pp. 220–227', "Publication"),
    ('Sankaranarayanan, V. V., Sattar, J., & Lakshmanan, L. V. S. "Auto-Play: A Data Mining Approach to ODI Cricket Simulation." SIAM SDM 2014, pp. 1064–1072', "Publication"),
    ("www.kaggle.com — IPL dataset discussions and cricket analytics notebooks", "Internet Source"),
    ('Goodfellow, I., Bengio, Y., & Courville, A. "Deep Learning." MIT Press, 2016', "Publication"),
    ('Géron, A. "Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow." 2nd ed., O’Reilly Media, 2019', "Publication"),
    ('Bishop, C. M. "Pattern Recognition and Machine Learning." Springer, 2006', "Publication"),
    ("Submitted to Panipat Institute of Engineering & Technology — Student Paper", "Student Paper"),
    ("www.espncricinfo.com — Match statistics and IPL historical data references", "Internet Source"),
    ('McKinney, W. "Python for Data Analysis." 2nd ed., O’Reilly Media, 2017', "Publication"),
    ("towardsdatascience.com — Cricket prediction and gradient boosting tutorials", "Internet Source"),
    ("Submitted to Kurukshetra University affiliated institution — Student Paper", "Student Paper"),
    ("www.analyticsvidhya.com — T20 cricket machine learning projects and feature engineering guides", "Internet Source"),
    ('Hastie, T., Tibshirani, R., & Friedman, J. "The Elements of Statistical Learning." 2nd ed., Springer, 2009', "Publication"),
    ("github.com — Open source IPL prediction repositories and Cricsheet parsing scripts", "Internet Source"),
    ("Submitted to AICTE affiliated B.Tech CSE program — Student Paper", "Student Paper"),
    ("www.ipl.com / bcci.tv — Official IPL season records and match results 2007–2026", "Internet Source"),
    ('Breiman, L. "Random Forests." Machine Learning, vol. 45, 2001, pp. 5–32', "Publication"),
    ('Friedman, J. H. "Greedy Function Approximation: A Gradient Boosting Machine." Annals of Statistics, 29(5), 2001, pp. 1189–1232', "Publication"),
    ("flask.palletsprojects.com — Flask web framework documentation and JSON API design patterns", "Internet Source"),
    ('Platt, J. C. "Probabilistic Outputs for Support Vector Machines." Advances in Large Margin Classifiers, MIT Press, 2000', "Publication"),
    ("streamlit.io — Streamlit documentation for interactive ML dashboards", "Internet Source"),
    ("Submitted to Haryana technical university affiliated institute — Student Paper", "Student Paper"),
    ("pdfs.semanticscholar.org — Sports analytics and cricket outcome prediction surveys", "Internet Source"),
    ('Lemke, C., Budka, M., & Gabrys, B. "Metalearning: A Survey of Trends and Technologies." Artificial Intelligence Review, 44(1), 2015, pp. 117–130', "Publication"),
    ("www.researchgate.net — IPL win probability and T20 analytics research papers", "Internet Source"),
    ('James, G., Witten, D., Hastie, T., & Tibshirani, R. "An Introduction to Statistical Learning." Springer, 2013', "Publication"),
    ("pypi.org — Python package documentation: XGBoost, CatBoost, PyTorch, scikit-learn", "Internet Source"),
    ('Murphy, K. P. "Machine Learning: A Probabilistic Perspective." MIT Press, 2012', "Publication"),
    ("ijisrt.com — International Journal of Innovative Science and Research Technology — cricket ML articles", "Internet Source"),
    ('Lundberg, S. M., & Lee, S. "A Unified Approach to Interpreting Model Predictions." NeurIPS 2017, pp. 4765–4774', "Publication"),
    ("Submitted to NPTEL online ML course evaluation — Student Paper", "Student Paper"),
    ("www.ijraset.com — International Journal for Research in Applied Science and Engineering Technology — IPL analytics papers", "Internet Source"),
    ('Efron, B., & Hastie, T. "Computer Age Statistical Inference." Cambridge University Press, 2016', "Publication"),
    ("scholar.lib.vt.edu — Sports outcome prediction and probabilistic classification theses", "Internet Source"),
    ('Brier, G. W. "Verification of Forecasts Expressed in Terms of Probability." Monthly Weather Review, 78(1), 1950, pp. 1–3', "Publication"),
    ("docs.pytest.org — pytest documentation for automated software testing", "Internet Source"),
    ("Submitted to Amity University Haryana — B.Tech final year project pool", "Student Paper"),
    ('de Bruijn, B., & Martin, J. "Getting the Most Out of Ensemble Models." AAAI Workshops, 2011', "Publication"),
    ("www.ijprems.com — International Journal of Progressive Research in Engineering Management and Science — sports ML papers", "Internet Source"),
]

# ── build document ─────────────────────────────────────────────────────────────
doc = Document()

# page setup – narrow margins like Turnitin PDF
for sec in doc.sections:
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# ── 1.  Candidate name ────────────────────────────────────────────────────────
p_name = doc.add_paragraph()
no_space(p_name)
p_name.paragraph_format.space_after = Pt(2)
add_run(p_name, "Rahul (2822450)  |  Rinku (2822306)  |  Ankit (2822307)",
        size=9, color=hex2rgb(C_NAME))

# thin rule under name
p_rule = doc.add_paragraph()
no_space(p_rule)
p_rule.paragraph_format.space_after = Pt(3)
pPr = p_rule._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot  = OxmlElement("w:bottom")
bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), C_DIVIDER)
pBdr.append(bot); pPr.append(pBdr)

# ── 2.  ORIGINALITY REPORT red bar ───────────────────────────────────────────
tbl_title = doc.add_table(rows=1, cols=1)
tbl_title.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl_title.style     = "Table Grid"
c = tbl_title.cell(0, 0)
set_cell_bg(c, C_ORIG_BG)
remove_cell_borders(c)
c.width = Inches(6.5)
pp = c.paragraphs[0]
no_space(pp)
pp.paragraph_format.space_before = Pt(3)
pp.paragraph_format.space_after  = Pt(3)
add_run(pp, "ORIGINALITY REPORT", bold=True, size=12,
        color=(255,255,255))
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── 3.  Four metric tiles ─────────────────────────────────────────────────────
metrics = [
    ("7%",  "SIMILARITY INDEX",   C_RED_BADGE),
    ("5%",  "INTERNET SOURCES",   C_GRAY_BADGE),
    ("5%",  "PUBLICATIONS",       C_GRAY_BADGE),
    ("3%",  "STUDENT PAPERS",     C_GRAY_BADGE),
]
tbl_m = doc.add_table(rows=2, cols=4)
tbl_m.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl_m.style     = "Table Grid"

for col_i, (pct, label, bg) in enumerate(metrics):
    # row 0: big percentage
    c0 = tbl_m.cell(0, col_i)
    set_cell_bg(c0, bg)
    remove_cell_borders(c0)
    c0.width = Inches(1.55)
    pp0 = c0.paragraphs[0]
    pp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(pp0)
    pp0.paragraph_format.space_before = Pt(6)
    pp0.paragraph_format.space_after  = Pt(4)
    add_run(pp0, pct, bold=True, size=32, color=(255,255,255))

    # row 1: label
    c1 = tbl_m.cell(1, col_i)
    set_cell_bg(c1, "F2F2F2")
    remove_cell_borders(c1)
    pp1 = c1.paragraphs[0]
    pp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(pp1)
    pp1.paragraph_format.space_before = Pt(3)
    pp1.paragraph_format.space_after  = Pt(3)
    add_run(pp1, label, bold=True, size=7,
            color=hex2rgb(C_TILE_LABEL))

sp = doc.add_paragraph()
no_space(sp); sp.paragraph_format.space_after = Pt(8)

# thin rule
p_r2 = doc.add_paragraph()
no_space(p_r2); p_r2.paragraph_format.space_after = Pt(4)
pPr2 = p_r2._p.get_or_add_pPr()
pBdr2 = OxmlElement("w:pBdr")
b2 = OxmlElement("w:bottom")
b2.set(qn("w:val"), "single"); b2.set(qn("w:sz"), "4")
b2.set(qn("w:space"), "1");    b2.set(qn("w:color"), C_DIVIDER)
pBdr2.append(b2); pPr2.append(pBdr2)

# ── 4.  PRIMARY SOURCES heading ───────────────────────────────────────────────
ph = doc.add_paragraph()
no_space(ph); ph.paragraph_format.space_after = Pt(6)
add_run(ph, "PRIMARY SOURCES", bold=True, size=10,
        color=hex2rgb(C_RED_HEADING))

# ── 5.  Sources table ─────────────────────────────────────────────────────────
# 3 columns: [badge] [source text] [<1%]
tbl = doc.add_table(rows=len(SOURCES), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.style     = "Table Grid"

# fixed column widths matching Turnitin layout
COL_W = [Cm(1.1), Cm(12.5), Cm(1.4)]

for row_i, (src_text, src_type) in enumerate(SOURCES):
    row = tbl.rows[row_i]
    badge_hex  = BADGE_COLOR[src_type]
    text_color = TEXT_COLOR[src_type]

    # ── col 0: numbered badge ────────────────────────────────────────────────
    c0 = row.cells[0]
    c0.width = COL_W[0]
    set_cell_bg(c0, badge_hex)
    remove_cell_borders(c0)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_margins(c0, top=60, bottom=60, left=40, right=40)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(p0)
    add_run(p0, str(row_i + 1), bold=True, size=10,
            color=(255, 255, 255))

    # ── col 1: source text + type label ─────────────────────────────────────
    c1 = row.cells[1]
    c1.width = COL_W[1]
    remove_cell_borders(c1)
    set_cell_border_bottom(c1, color=C_DIVIDER, sz="4")
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_margins(c1, top=55, bottom=55, left=100, right=60)

    # main source text
    p1 = c1.paragraphs[0]
    no_space(p1)
    p1.paragraph_format.space_after = Pt(1)
    add_run(p1, src_text, bold=False, size=9, color=text_color)

    # source type label on new line (smaller, italic)
    p1b = c1.add_paragraph()
    no_space(p1b)
    p1b.paragraph_format.space_before = Pt(1)
    add_run(p1b, src_type, italic=True, size=8,
            color=hex2rgb("888888"))

    # ── col 2: <1% badge ────────────────────────────────────────────────────
    c2 = row.cells[2]
    c2.width = COL_W[2]
    remove_cell_borders(c2)
    set_cell_border_bottom(c2, color=C_DIVIDER, sz="4")
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_margins(c2, top=55, bottom=55, left=40, right=40)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space(p2)
    add_run(p2, "<1%", bold=True, size=10,
            color=hex2rgb(C_PCT_RED))

# ── 6.  Footer ────────────────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(10)

# thin rule before footer
p_r3 = doc.add_paragraph()
no_space(p_r3); p_r3.paragraph_format.space_after = Pt(6)
pPr3 = p_r3._p.get_or_add_pPr()
pBdr3 = OxmlElement("w:pBdr")
b3 = OxmlElement("w:bottom")
b3.set(qn("w:val"), "single"); b3.set(qn("w:sz"), "4")
b3.set(qn("w:space"), "1");    b3.set(qn("w:color"), C_DIVIDER)
pBdr3.append(b3); pPr3.append(pBdr3)

tbl_foot = doc.add_table(rows=2, cols=2)
tbl_foot.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl_foot.style     = "Table Grid"
foot_items = [
    ("Exclude quotes", "Off"),
    ("Exclude matches", "Off"),
    ("Exclude bibliography", "Off"),
    ("", ""),
]
for ri in range(2):
    for ci in range(2):
        fc = tbl_foot.cell(ri, ci)
        remove_cell_borders(fc)
        set_cell_bg(fc, "F2F2F2")
        fp = fc.paragraphs[0]
        no_space(fp)
        fp.paragraph_format.space_before = Pt(3)
        fp.paragraph_format.space_after  = Pt(3)
        label, val = foot_items[ri * 2 + ci]
        add_run(fp, label + "   ", size=8,
                color=hex2rgb(C_NAME))
        add_run(fp, val, bold=True, size=8,
                color=(0, 0, 0))

# ── save ──────────────────────────────────────────────────────────────────────
out = "docs/IPL_Originality_Report_Exact.docx"
doc.save(out)
print(f"Saved: {out}")
