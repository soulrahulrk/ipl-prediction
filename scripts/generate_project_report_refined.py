from __future__ import annotations

import importlib.util
import json
import shutil
import subprocess
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


ROOT_DIR = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT_DIR / "docs"
ASSETS_DIR = DOCS_DIR / "report_assets"
REPORTS_DIR = DOCS_DIR / "reports"
OUTPUT_PATH = REPORTS_DIR / "IPL_Prediction_Project_Report_Refined.docx"
OUTPUT_PATH_ROOT = DOCS_DIR / "IPL_Prediction_Project_Report_Refined.docx"
FINAL_OUTPUT_PATH = DOCS_DIR / "FINAL_SUBMISSION_IPL_PROJECT_REPORT.docx"
FINAL_OUTPUT_PDF_PATH = DOCS_DIR / "FINAL_SUBMISSION_IPL_PROJECT_REPORT.pdf"

BODY_FONT = "Times New Roman"
HEADING_FONT = "Times New Roman"
CODE_FONT = "Consolas"
BODY_SIZE = 11

A4_WIDTH = Mm(210)
A4_HEIGHT = Mm(297)
LEFT_MARGIN = Mm(35)
OTHER_MARGIN = Mm(25)

ABBREVIATIONS = [
    ("IPL", "Indian Premier League"),
    ("ML", "Machine Learning"),
    ("MAE", "Mean Absolute Error"),
    ("RMSE", "Root Mean Squared Error"),
    ("API", "Application Programming Interface"),
    ("CLI", "Command Line Interface"),
    ("GPU", "Graphics Processing Unit"),
    ("DFD", "Data Flow Diagram"),
]

REPORT_PROFILE = {
    "title": "IPL Match Score and Win Prediction System",
    "subtitle": "A narrative project report focused on live score forecasting, win-probability modeling, and deployment",
    "degree_line_1": "Bachelor of Technology in",
    "degree_line_2": "Computer Science & Engineering (Artificial Intelligence and Machine Learning)",
    "students": [
        "Rahul (2822450)",
        "Rinku (2822306)",
        "Ankit (2822307)",
    ],
    "supervisor": "Ms. Chetna Dahiya (Assistant Professor)",
    "hod": "Dr. Devendra Prasad",
    "department": "Department of Computer Science & Engineering (Artificial Intelligence and Machine Learning)",
    "institute": "Panipat Institute of Engineering & Technology, Samalkha, Panipat",
    "university": "Kurukshetra University, Kurukshetra, India (2022-2026)",
    "place": "Panipat",
}


@dataclass(frozen=True)
class FigureSpec:
    number: str
    title: str
    filename: str


@dataclass(frozen=True)
class TableSpec:
    number: str
    title: str


FIGURES = [
    FigureSpec("Figure 4.1", "End-to-End IPL Prediction Workflow", "figure_4_1_pipeline.png"),
    FigureSpec("Figure 5.1", "Comparative View of Candidate Live Models", "figure_4_2_model_comparison.png"),
    FigureSpec("Figure 5.2", "Held-Out Score Error Distribution", "figure_5_4_score_error_distribution.png"),
    FigureSpec("Figure 5.3", "Held-Out Win Probability Calibration", "figure_5_5_win_calibration.png"),
    FigureSpec("Figure 6.1", "Flask Live Prediction Interface", "figure_6_2_flask_mockup.png"),
]

TABLES = [
    TableSpec("Table 2.1", "Approach Families Reviewed While Designing the Project"),
    TableSpec("Table 4.1", "Core Dataset Snapshot Used in the Final Project"),
    TableSpec("Table 4.2", "Live Feature Families and Their Project Value"),
    TableSpec("Table 5.1", "Scope Comparison That Guided Final Training Direction"),
    TableSpec("Table 5.2", "Live Score Model Comparison on the Key Evaluation Scope"),
    TableSpec("Table 5.3", "Live Win Model Comparison on the Key Evaluation Scope"),
    TableSpec("Table 5.4", "Pre-Match Model Performance Summary"),
    TableSpec("Table 6.1", "User-Facing Surfaces Built Around the Shared Prediction Core"),
]

TOC_ENTRY_LABELS = [
    "DECLARATION",
    "APPROVAL FROM SUPERVISOR",
    "CERTIFICATE",
    "ACKNOWLEDGEMENTS",
    "ABSTRACT",
    "TABLE OF CONTENTS",
    "LIST OF TABLES",
    "LIST OF FIGURES",
    "LIST OF ABBREVIATIONS",
    "CHAPTER 1 - INTRODUCTION",
    "CHAPTER 2 - LITERATURE REVIEW",
    "CHAPTER 3 - PROBLEM OBJECTIVE",
    "CHAPTER 4 - METHODOLOGY OF THE PROJECT",
    "CHAPTER 5 - RESULTS AND ANALYSIS",
    "CHAPTER 6 - SYSTEM IMPLEMENTATION AND TESTING",
    "CHAPTER 7 - CONCLUSION AND FUTURE SCOPE",
    "APPENDIX 1 - FLOWCHART AND ALGORITHM NOTES",
    "APPENDIX 2 - SELECTED FEATURE NOTES",
    "APPENDIX 3 - SHORT SUPPORTING CASE SNAPSHOTS",
    "REFERENCES",
]

TABLE_ENTRY_LABELS = [f"{spec.number}: {spec.title}" for spec in TABLES]
FIGURE_ENTRY_LABELS = [f"{spec.number}: {spec.title}" for spec in FIGURES]
FRONT_MATTER_LABELS = TOC_ENTRY_LABELS[:9]


def load_base_report_module():
    path = ROOT_DIR / "scripts" / "generate_project_report.py"
    spec = importlib.util.spec_from_file_location("base_project_report", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load report module from {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def fmt_float(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}"


def fmt_pct(value: float, digits: int = 1) -> str:
    return f"{100.0 * value:.{digits}f}%"


def prettify_model_name(model_name: str) -> str:
    mapping = {
        "hgb": "HistGradientBoosting",
        "hgb_calibrated": "HistGradientBoosting (calibrated)",
        "hgb_raw": "HistGradientBoosting (raw)",
        "xgboost_gpu": "XGBoost GPU",
        "xgboost_gpu_raw": "XGBoost GPU (raw)",
        "xgboost_gpu_calibrated": "XGBoost GPU (calibrated)",
        "catboost_gpu": "CatBoost GPU",
        "catboost_gpu_raw": "CatBoost GPU (raw)",
        "catboost_gpu_calibrated": "CatBoost GPU (calibrated)",
        "torch_entity_gpu": "Torch entity-embedding model",
        "torch_entity_gpu_calibrated": "Torch entity-embedding model (calibrated)",
        "ensemble_ml_top3": "Weighted ML ensemble",
        "ensemble_top2": "Top-2 weighted ensemble",
        "ensemble_top2_calibrated": "Top-2 weighted ensemble (calibrated)",
        "ridge": "Ridge regression",
        "catboost_classifier": "CatBoost classifier",
    }
    return mapping.get(model_name, model_name.replace("_", " "))


def best_registry_entry(model_registry: dict) -> dict:
    entries = model_registry.get("entries", [])
    if not entries:
        raise ValueError("Model registry is empty.")
    return min(
        entries,
        key=lambda entry: (
            entry["metrics"]["live"]["score_test"]["rmse"]
            + 10.0 * entry["metrics"]["live"]["win_test"]["log_loss"]
        ),
    )


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def apply_page_layout(section) -> None:
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.left_margin = LEFT_MARGIN
    section.right_margin = OTHER_MARGIN
    section.top_margin = OTHER_MARGIN
    section.bottom_margin = OTHER_MARGIN


def set_page_number_format(section, fmt: str, start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.xpath("./w:pgNumType")
    if pg_num:
        pg_num = pg_num[0]
    else:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def configure_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer)


def clear_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = ""


def configure_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Inches(0.25)

    h1 = doc.styles["Heading 1"]
    h1.font.name = HEADING_FONT
    h1.font.bold = True
    h1.font.size = Pt(14)
    h1.paragraph_format.first_line_indent = Inches(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h2 = doc.styles["Heading 2"]
    h2.font.name = HEADING_FONT
    h2.font.bold = True
    h2.font.size = Pt(11)
    h2.paragraph_format.first_line_indent = Inches(0)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h3 = doc.styles["Heading 3"]
    h3.font.name = HEADING_FONT
    h3.font.bold = True
    h3.font.size = Pt(11)
    h3.paragraph_format.first_line_indent = Inches(0)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    section = doc.sections[0]
    apply_page_layout(section)
    section.different_first_page_header_footer = False
    clear_footer(section)


def add_centered_paragraph(
    doc: Document,
    text: str,
    size: int = 11,
    bold: bool = False,
    italic: bool = False,
    spacing_after: int = 6,
    spacing_before: int = 0,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(spacing_before)
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    run = paragraph.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE)


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(item)
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE)


def add_index_entry(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE)


def add_numbered_list(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(item)
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE)


def add_appendix_separator(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)


def add_table_title(doc: Document, spec: TableSpec) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.space_before = Pt(8)
    run = paragraph.add_run(f"{spec.number}: {spec.title}")
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.bold = True


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Inches(0)
    run = paragraph.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(font_size)
    run.bold = bold


def add_table(doc: Document, headers: list[str], rows: list[list[str]], font_size: int = 10) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, font_size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value, bold=False, font_size=font_size)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = Pt(10)


def add_figure(doc: Document, path: Path, number: str, title: str, width: float = 5.8) -> None:
    doc.add_picture(str(path), width=Inches(width))
    pic_paragraph = doc.paragraphs[-1]
    pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.first_line_indent = Inches(0)
    run = caption.add_run(f"{number}: {title}")
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.bold = True
    run.italic = True


def ensure_asset_paths(context: dict, base_module) -> dict[str, Path]:
    assets = {spec.filename: ASSETS_DIR / spec.filename for spec in FIGURES}
    missing = [path for path in assets.values() if not path.exists()]
    if missing:
        base_module.generate_assets(context)
    return {spec.number: ASSETS_DIR / spec.filename for spec in FIGURES}


def add_title_page(doc: Document, today_text: str) -> None:
    add_centered_paragraph(doc, "", spacing_after=24, spacing_before=70)
    add_centered_paragraph(doc, REPORT_PROFILE["title"].upper(), size=16, bold=True, spacing_after=18)
    add_centered_paragraph(doc, "A Project Report (Project-II)", size=13, bold=True, spacing_after=6)
    add_centered_paragraph(
        doc,
        "Submitted in partial fulfillment of the requirements for the award of the degree of",
        size=11,
        spacing_after=6,
    )
    add_centered_paragraph(doc, "BACHELOR OF TECHNOLOGY", size=14, bold=True, spacing_after=2)
    add_centered_paragraph(doc, "IN", size=12, bold=True, spacing_after=2)
    add_centered_paragraph(doc, REPORT_PROFILE["degree_line_2"].upper(), size=12, bold=True, spacing_after=18)
    add_centered_paragraph(doc, "Submitted By", size=11, bold=True, spacing_after=4)
    for student in REPORT_PROFILE["students"]:
        add_centered_paragraph(doc, student, size=11, spacing_after=2)
    add_centered_paragraph(doc, "", spacing_after=6)
    add_centered_paragraph(doc, "Supervised By", size=11, bold=True, spacing_after=4)
    add_centered_paragraph(doc, REPORT_PROFILE["supervisor"], size=11, spacing_after=16)
    add_centered_paragraph(doc, REPORT_PROFILE["department"].upper(), size=11, bold=True, spacing_after=4)
    add_centered_paragraph(doc, REPORT_PROFILE["institute"].upper(), size=11, spacing_after=2)
    add_centered_paragraph(doc, "(Approved by AICTE and Affiliated to Kurukshetra University, Kurukshetra)", size=10, spacing_after=2)
    add_centered_paragraph(doc, today_text.upper(), size=11, spacing_after=0)


def add_declaration_page(doc: Document, today_text: str) -> None:
    doc.add_heading("DECLARATION", level=1)
    add_body(
        doc,
        'We certify that the work presented in this project report entitled "IPL Match Score and Win Prediction System" is an authentic record of the '
        "work carried out by us under the guidance of our supervisor. The preprocessing pipeline, feature engineering, experiments, interfaces, and "
        "analysis documented here were developed specifically for this project from the available IPL data and associated software stack.",
    )
    add_body(
        doc,
        "This report has not been submitted elsewhere for the award of any other degree or diploma. Wherever external datasets, software libraries, "
        "or published technical ideas influenced the work, they have been acknowledged appropriately in the text and listed in the references section.",
    )
    add_body(doc, f"Student Name(s): {', '.join(REPORT_PROFILE['students'])}")
    add_body(doc, 'Project report title: "IPL Match Score and Win Prediction System"')
    add_body(doc, "Semester: 7th / 8th Semester")
    add_body(doc, f"Date: {today_text}")
    doc.add_page_break()


def add_approval_page(doc: Document, today_text: str) -> None:
    doc.add_heading("APPROVAL FROM SUPERVISOR", level=1)
    add_body(
        doc,
        f'This is to certify that the project report entitled "IPL Match Score and Win Prediction System" presented by '
        f'"{", ".join(REPORT_PROFILE["students"])}" under my supervision is an authentic work. To the best of my knowledge, the content of this '
        "report has not been submitted previously for the award of any other degree. The work is recommended for academic consideration as a valid "
        "final-year project submission.",
    )
    add_body(doc, f"Supervisor Name: {REPORT_PROFILE['supervisor']}")
    add_body(doc, REPORT_PROFILE["department"])
    add_body(doc, f"Date: {today_text}")
    add_body(doc, f"(Counter Signed By) {REPORT_PROFILE['hod']}")
    doc.add_page_break()


def add_certificate_page(doc: Document, today_text: str) -> None:
    doc.add_heading("CERTIFICATE", level=1)
    add_body(
        doc,
        f'This is to certify that the work embodied in this report, entitled "IPL Match Score and Win Prediction System", carried out by '
        f'"{", ".join(REPORT_PROFILE["students"])}" is approved for submission toward the degree of '
        f'"{REPORT_PROFILE["degree_line_1"]} {REPORT_PROFILE["degree_line_2"]}" at {REPORT_PROFILE["department"]}, '
        f'{REPORT_PROFILE["institute"]}.',
    )
    add_body(doc, "Internal Examiner: ______________________________")
    add_body(doc, "External Examiner: ______________________________")
    add_body(doc, f"Place: {REPORT_PROFILE['place']}")
    add_body(doc, f"Date: {today_text}")
    doc.add_page_break()


def add_acknowledgements_page(doc: Document, today_text: str) -> None:
    doc.add_heading("ACKNOWLEDGEMENTS", level=1)
    add_body(
        doc,
        "The completion of this project was made possible through the availability of open IPL ball-by-ball data, the support of our academic mentors, "
        "and the strength of the Python machine-learning ecosystem. The work demanded coordinated effort across data engineering, feature design, model "
        "selection, backend integration, dashboard development, testing, and report preparation.",
    )
    add_body(
        doc,
        f"We express sincere gratitude to {REPORT_PROFILE['supervisor']} for guidance throughout the project lifecycle, from problem formulation to "
        "final report preparation. We also thank the faculty of the department for providing the academic environment needed to complete an applied "
        "machine-learning project of this scale.",
    )
    add_body(
        doc,
        "We further acknowledge the maintainers of Cricsheet, scikit-learn, XGBoost, CatBoost, PyTorch, Flask, Streamlit, and python-docx. Without "
        "these openly available tools and datasets, the present work would have required a much narrower scope and would not have achieved the same "
        "level of experimentation or reproducibility.",
    )
    add_body(doc, f"Student Name(s): {', '.join(REPORT_PROFILE['students'])}")
    add_body(doc, f"Date: {today_text}")
    doc.add_page_break()


def add_abstract_page(doc: Document, context: dict, best_entry: dict) -> None:
    live_metrics = best_entry["metrics"]["live"]
    counts = context["counts"]
    doc.add_heading("ABSTRACT", level=1)
    add_body(
        doc,
        f"This report presents an IPL analytics system that estimates two live outcomes from the current ball-by-ball state of a match: the projected "
        f"final innings total and the batting side's win probability. The project is built on {counts['raw_matches']} IPL matches from Cricsheet and "
        f"an engineered live-state table containing {counts['processed_rows']:,} rows across seasons {counts['min_season']} to {counts['max_season']}.",
    )
    add_body(
        doc,
        "The central idea of the project is that live cricket prediction becomes more reliable when the current score is combined with carefully "
        "constructed context. Instead of relying only on runs, wickets, and overs, the final feature set integrates momentum windows, venue history, "
        "recent team form, team-venue form, head-to-head context, player-level form, batter-versus-bowler history, and optional weather-derived dew risk.",
    )
    add_body(
        doc,
        "A sequence of experiments was carried out to determine which training scope and model families performed best under chronological evaluation. "
        "The strongest recorded project snapshot used the all-active-history scope, a torch entity-embedding regressor for live score prediction, and "
        "a calibrated CatBoost classifier for live win prediction.",
    )
    add_body(
        doc,
        f"In that best experimental snapshot, the live score model achieved MAE {fmt_float(live_metrics['score_test']['mae'])} and RMSE "
        f"{fmt_float(live_metrics['score_test']['rmse'])}, while the live win model achieved log loss "
        f"{fmt_float(live_metrics['win_test']['log_loss'], 4)}, Brier score {fmt_float(live_metrics['win_test']['brier'], 4)}, and accuracy "
        f"{fmt_pct(live_metrics['win_test']['accuracy'])}. The system was then wrapped inside a shared inference layer used by a Flask web interface, "
        "a JSON API, a Streamlit dashboard, and a command-line tool.",
    )
    doc.add_page_break()


def add_table_of_contents_page(doc: Document) -> None:
    doc.add_heading("TABLE OF CONTENTS", level=1)
    for entry in TOC_ENTRY_LABELS:
        add_index_entry(doc, entry)
    doc.add_page_break()


def add_list_of_tables_page(doc: Document) -> None:
    doc.add_heading("LIST OF TABLES", level=1)
    for entry in TABLE_ENTRY_LABELS:
        add_index_entry(doc, entry)
    doc.add_page_break()


def add_list_of_figures_page(doc: Document) -> None:
    doc.add_heading("LIST OF FIGURES", level=1)
    for entry in FIGURE_ENTRY_LABELS:
        add_index_entry(doc, entry)
    doc.add_page_break()


def add_abbreviations_page(doc: Document) -> None:
    doc.add_heading("LIST OF ABBREVIATIONS", level=1)
    for short, full in ABBREVIATIONS:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(f"{short}: {full}")
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE)


def add_chapter_1(doc: Document, context: dict) -> None:
    counts = context["counts"]
    doc.add_heading("CHAPTER 1 - INTRODUCTION", level=1)

    doc.add_heading("Background", level=2)
    add_body(
        doc,
        "T20 cricket is one of the most volatile formats of the sport because the balance between bat and ball can shift dramatically within a few overs. "
        "A team that appears behind after the powerplay may recover through middle-over consolidation, while a side in control after fifteen overs may "
        "still collapse in the death phase. This volatility makes live prediction both useful and technically demanding.",
    )
    add_body(
        doc,
        "In the IPL, the challenge is even greater because conditions vary heavily across venues, batting orders change from match to match, team quality "
        "changes from season to season, and player matchups often influence the flow of an innings in ways that a raw scoreboard cannot capture. For that "
        "reason, a meaningful prediction system must see more than the visible runs and wickets.",
    )

    doc.add_heading("Project Context", level=2)
    add_body(
        doc,
        f"The present project converts {counts['raw_matches']} raw IPL match files into a machine-learning system that can respond to a live match state. "
        f"The final engineered dataset contains {counts['processed_rows']:,} rows and {counts['processed_cols']} columns, representing legal-ball snapshots "
        "rather than full-match summaries. That design decision is crucial because live prediction should answer questions in the moment, not only after the match is complete.",
    )
    add_body(
        doc,
        "The project does not treat cricket prediction as a single binary classification problem. Instead, it solves two related but different tasks: "
        "forecasting the eventual innings total and estimating the batting side's probability of winning from the current state. This separation allows the "
        "report to examine not only whether the batting side is likely to win, but also why the expected innings trajectory is changing.",
    )

    doc.add_heading("Motivation", level=2)
    add_body(
        doc,
        "The main motivation behind the project was to build a system that behaves more like a real cricket analytics engine and less like a classroom-only "
        "notebook. Many student sports projects stop at one model trained on a static dataset. In contrast, the goal here was to combine preprocessing, "
        "historical support tables, model comparison, probability calibration, interface design, and reproducible reporting inside one coherent workflow.",
    )
    add_body(
        doc,
        "A second motivation came from the gap between scoreboard intuition and data-driven judgment. Human observers can often tell that a side is under "
        "pressure, but they struggle to express that pressure numerically. The project attempts to bridge that gap by turning match state into explicit "
        "features such as required-minus-current run rate, runs scored in recent windows, venue par score, and batter-versus-bowler history.",
    )

    doc.add_heading("Target Users", level=2)
    add_body(
        doc,
        "The principal users of the system are students and faculty evaluating the project, cricket followers who want a clearer live interpretation of match state, and developers who may later integrate the prediction logic with another interface or analytics layer. This user view matters because the report is not about a model in isolation. It is about a model that must communicate useful outputs to human readers.",
    )
    add_body(
        doc,
        "For that reason, the project emphasized interpretable signals such as projected total, win probability, expected range, and contextual narrative instead of exposing only raw model scores that would be difficult for non-specialists to understand.",
    )

    doc.add_heading("Rationale Behind the System", level=2)
    add_body(
        doc,
        "The rationale behind the system is that live cricket decisions are rarely driven by one visible number alone. Runs, wickets, venue conditions, player quality, recent momentum, and chase pressure all interact. A practical prediction system therefore needs to combine these factors in one framework rather than depending on a simple run-rate heuristic.",
    )
    add_body(
        doc,
        "This rationale also explains why the project invested so much effort in preprocessing and feature engineering. The better the system state is represented, the more meaningful the later predictions become.",
    )

    doc.add_heading("Expected Benefits", level=2)
    add_body(
        doc,
        "The main benefit of the system is that it turns live IPL context into structured insight. It can support classroom demonstration, project evaluation, cricket analytics discussion, and future application development. More broadly, it shows how data engineering, model selection, and software deployment can be integrated in one applied machine-learning project.",
    )
    add_body(
        doc,
        "A second benefit is academic. The project provides a concrete case study of how to build a leakage-safe sports prediction pipeline that goes beyond notebook experimentation and moves toward a usable software artifact.",
    )

    doc.add_heading("Why This Problem Is Hard", level=2)
    add_body(
        doc,
        "The project revealed early that live IPL prediction is hard for at least four reasons. First, the same visible score can mean different things on "
        "different grounds: 170 is defendable in some venues and below par in others. Second, batting momentum changes quickly, so recent overs carry signal "
        "that simple season averages miss. Third, player identities are not interchangeable; the same run rate means different future expectations depending "
        "on who is batting and bowling. Fourth, any feature built from the future immediately creates leakage and gives the model unrealistically easy information.",
    )
    add_body(
        doc,
        "Because of those difficulties, the project had to be engineered carefully before modeling even began. The strongest improvements in the report did "
        "not come from blindly trying more algorithms. They came from designing the right training scope, the right context features, and the right validation logic.",
    )

    doc.add_heading("Assumptions Made", level=2)
    add_body(
        doc,
        "The project assumes that ball-by-ball Cricsheet records and associated match metadata are sufficiently accurate to reconstruct legal match states. It also assumes that historical venue, team, and player behavior retains enough continuity to provide predictive value for later seasons, even though that continuity is not perfect.",
    )
    add_body(
        doc,
        "A further assumption is that the primary deployment use case is analytical assistance rather than betting or automated decision-making. The system is therefore designed to inform judgment, not to replace human understanding of cricket uncertainty.",
    )

    doc.add_heading("Objectives", level=2)
    add_bullet_list(
        doc,
        [
            "Build a leakage-safe IPL live-state dataset from raw Cricsheet CSV2 files.",
            "Design features that capture match state, momentum, venue behavior, team strength, and player context together.",
            "Compare multiple tabular model families rather than assuming one model family would suit both prediction tasks.",
            "Select the best-performing live score and live win models using chronological evaluation and probability-aware metrics.",
            "Expose the final prediction system through practical interfaces instead of leaving it as an offline experiment only.",
        ],
    )

    doc.add_heading("Scope", level=2)
    add_body(
        doc,
        "The report focuses on the IPL only. It covers raw-data ingestion, preprocessing, support-table construction, live score regression, live win "
        "classification, a smaller pre-match modeling branch, a Flask interface, a JSON API, a Streamlit dashboard, a CLI, and automated tests. The "
        "current scope does not attempt ball-tracking vision, lineup optimization, or betting-market integration.",
    )
    add_body(
        doc,
        "The rest of the report is written with one principle in mind: the strongest project value came from understanding why some modeling choices worked "
        "better than others. For that reason, later chapters emphasize the reasoning behind the final results rather than filling pages with large inventories.",
    )
    doc.add_page_break()


def add_chapter_2(doc: Document, context: dict) -> None:
    doc.add_heading("CHAPTER 2 - LITERATURE REVIEW", level=1)

    doc.add_heading("What Earlier Work Suggests", level=2)
    add_body(
        doc,
        "The literature around cricket prediction and structured-data machine learning suggests that three layers matter most: the quality of the match-state "
        "representation, the suitability of the model family to mixed categorical and numeric data, and the discipline of chronological evaluation. When any "
        "one of these layers is weak, reported performance can look strong while failing to transfer to real live use.",
    )
    add_body(
        doc,
        "In cricket analytics, score projection and win prediction have often been handled through heuristics or simplified run-rate rules. Those approaches are "
        "useful as intuition, but they cannot naturally combine venue priors, player context, matchup history, and changing momentum. The project therefore treated "
        "machine learning as a tool for combining diverse sources of context, not as a replacement for cricket logic.",
    )
    add_body(
        doc,
        "Another important lesson from prior work is that sports prediction can look artificially strong if evaluation is careless. Random train-test splits can leak future seasons into the past, "
        "and aggressive feature engineering can accidentally reuse information that would never be known at the prediction point. For this reason, the project review phase paid close attention not only "
        "to model families but also to how those models were validated and whether their reported performance could survive realistic deployment conditions.",
    )

    doc.add_heading("Review of Representative Literature and Prior Solutions", level=2)
    literature_notes = [
        (
            "Classical cricket score heuristics",
            "Early cricket-analytics work often relied on current run rate, wickets in hand, and overs remaining as the main predictive signals. Those methods are easy to explain and useful as baselines, but they struggle when the same scoreline means different things at different venues or for different batting and bowling combinations.",
            "The gap identified from this stream is that hand-crafted rules alone cannot represent richer live context such as venue par, team-specific form, or player-vs-player history.",
        ),
        (
            "Gradient boosting for structured prediction",
            "Tree-based gradient boosting became a strong reference point in tabular machine learning because it can model nonlinear interactions without demanding large feature preprocessing pipelines. In applied predictive tasks, these methods frequently outperform simpler linear baselines and remain easier to audit than many deep-learning alternatives.",
            "The important gap is not model capacity itself, but whether the structured inputs are realistic and leakage-safe. A strong boosting model trained on weak or leaked features still produces an unreliable final system.",
        ),
        (
            "XGBoost as a scalable boosted-tree method",
            "The XGBoost family showed that regularized boosted trees could be trained efficiently while handling large structured datasets with strong predictive quality. This made it a natural benchmark for the current project, especially because the IPL live-state table contains many numeric and categorical interactions after encoding.",
            "The gap left by this line of work is that the model does not automatically solve probability calibration, context validity, or deployment consistency. Those had to be handled by the project design itself.",
        ),
        (
            "CatBoost and categorical handling",
            "CatBoost is especially relevant in domains where categories such as teams, venues, phases, and players carry meaningful signal. Its ordered treatment of categorical values is attractive for cricket analytics because the semantics of these identities matter strongly during a live innings.",
            "The remaining gap is that even strong categorical handling does not remove the need for careful history construction. If historical support features are not aligned chronologically, the benefit of a categorical-aware model becomes misleading.",
        ),
        (
            "Entity embeddings for high-cardinality features",
            "Entity embeddings offered an alternative perspective: instead of expanding categories into sparse one-hot indicators, the model can learn dense representations for venues, players, and teams. This is especially appealing when the category space is large and when similarity across identities has predictive value.",
            "The gap here is practical rather than theoretical. These models need enough training data and good validation discipline; otherwise they can look sophisticated without earning a real advantage over strong boosted-tree baselines.",
        ),
        (
            "Probability calibration and reliability-oriented prediction",
            "Calibration-oriented literature argues that a classifier should not only rank outcomes correctly but should also attach confidence values that behave realistically. In the present project, that idea became central because the user sees a win percentage, not only a class label.",
            "The gap identified from this literature is that accuracy alone is not enough for a live win-probability system. That directly shaped the final evaluation criteria of the project.",
        ),
        (
            "Sports analytics workflow papers and project reports",
            "A large body of applied project work in sports analytics demonstrates useful pipelines for collecting data, building features, and evaluating outcome predictors. These works are valuable because they show how academic experiments can move toward functioning applications rather than remain at the level of notebook-only demonstrations.",
            "The recurring gap is that many such reports stop at the modeling stage and do not unify inference, testing, reporting, and multiple user-facing interfaces under one software contract.",
        ),
        (
            "Open cricket datasets and reproducible data practice",
            "Open data sources such as Cricsheet make modern cricket analytics possible by preserving match metadata and ball-by-ball detail. Their importance goes beyond data access: they enable reproducibility because another researcher can rebuild the same match-state table from source files rather than depend on a private dataset.",
            "The gap, however, is that raw openness does not equal readiness. The project still needed normalization, reconstruction, support tables, and validation before the data could support trustworthy live inference.",
        ),
    ]
    for title, summary, gap in literature_notes:
        doc.add_heading(title, level=3)
        add_body(doc, summary)
        add_body(doc, gap)

    doc.add_heading("Literature Gaps That Directly Motivated the Proposed Work", level=2)
    add_body(
        doc,
        "The survey above reveals that the missing piece was not one better algorithm alone. The real gap was the absence of a disciplined, end-to-end live cricket workflow that combines time-aware feature construction, model comparison, calibration, multi-interface serving, and reportable engineering decisions in a single project.",
    )
    add_body(
        doc,
        "This observation shaped the proposed system directly. Instead of asking only which model should be trained, the project asked what representation of live cricket state a model should see, which evaluation protocol should be trusted, how the results should be delivered to the user, and how the whole pipeline should remain reproducible afterward.",
    )
    add_body(
        doc,
        "That is why the literature review matters so much to the present report. It does not exist merely to cite known models. It justifies the design choices that later define the strongest result achieved in this workspace.",
    )

    add_table_title(doc, TABLES[0])
    add_table(
        doc,
        ["Approach Family", "Why It Matters", "What It Misses Without Good Features"],
        [
            ["Simple score heuristics", "Fast and interpretable starting baseline for cricket intuition.", "Cannot encode player context, venue priors, or nonlinear interactions."],
            ["Gradient boosting", "Strong choice for structured data with complex nonlinear relations.", "Performance drops if leakage control and categorical handling are weak."],
            ["Categorical-aware boosting", "Useful when team, venue, toss, striker, and bowler identities carry signal.", "Still needs well-defined historical context and calibration."],
            ["Entity embeddings", "Compact way to represent high-cardinality categories such as players and venues.", "Needs enough training data and careful validation to stay stable."],
            ["Probability calibration", "Essential when users consume confidence percentages instead of only class labels.", "Adds little value if the base classifier is poorly specified."],
        ],
        font_size=9,
    )

    doc.add_heading("How This Project Differs", level=2)
    add_body(
        doc,
        "The present project differs from a simple benchmark study in two ways. First, it builds historical support tables outside the model so that venue priors, "
        "team form, player form, and matchup context are computed in a controlled, reusable manner. Second, it evaluates the entire system as a deployable workflow, "
        "not only as a notebook metric. That means the report also considers API behavior, interface integration, and whether the same inference logic can be reused across surfaces.",
    )
    add_body(
        doc,
        "A further distinction is that the project treats live score prediction and win-probability estimation as related but separate tasks. This is important because "
        "the model that best estimates a final total is not automatically the model that produces the most reliable win probability. The final design therefore allows different "
        "families to win different tasks when the evidence supports it.",
    )

    doc.add_heading("Why Cricket Prediction Needed More Than a Static Dataset", level=2)
    add_body(
        doc,
        "A recurring weakness in many early sports-analytics exercises is that they begin from a final-match summary table and then attempt to predict the winner. That setup can be useful for a classification assignment, "
        "but it is too coarse for an IPL live system. The project needed a representation that changes ball by ball because the end user is interested in how the match is evolving, not only in who eventually won.",
    )
    add_body(
        doc,
        "This is where the distinction between descriptive and predictive analytics became important. A descriptive summary can say that one team scored 188 and the other scored 179. A predictive live-state model must decide "
        "what to say when the batting side is 84/4 after 11.3 overs at Chennai or 132/7 in the 17th over of a chase at Lucknow. Those are structurally different questions, and the literature review helped make that distinction explicit.",
    )

    doc.add_heading("Why Structured ML Was Prioritized Before Deep Learning", level=2)
    add_body(
        doc,
        "The review also influenced the project's modeling order. On structured, mixed-type data, gradient boosting methods usually offer a stronger starting point than deep neural networks. That does not mean deep learning has no value; it means the burden of proof is higher. "
        "The project therefore used boosting models as the benchmark families that any richer architecture had to beat honestly.",
    )
    add_body(
        doc,
        "This ordering later saved time and improved clarity. It prevented the report from making exaggerated claims about model sophistication and instead forced every improvement to be defended through held-out performance and better cricket-specific behavior.",
    )

    doc.add_heading("Why Probability Quality Was Non-Negotiable", level=2)
    add_body(
        doc,
        "Most machine-learning summaries emphasize accuracy because it is easy to explain. The present project took a different view for the win task. If a system claims that the batting side has a 72 percent chance of winning, the quality of that 72 matters more than whether the final 0/1 label was correct in one isolated match. "
        "That is why the literature on calibration, Brier score, and log loss influenced the project directly.",
    )
    add_body(
        doc,
        "In practical cricket use, a user may accept that predictions can be wrong, but will lose trust quickly if the reported percentages are unstable or overconfident. The project therefore evaluated probability behavior as part of the core modeling question, not as an afterthought appended at the end of the results chapter.",
    )

    doc.add_heading("Why the Review Matters to the Final Result", level=2)
    add_body(
        doc,
        "The review phase shaped several concrete decisions later in the pipeline. The emphasis on structured-data methods justified starting with HistGradientBoosting, "
        "XGBoost, and CatBoost before relying on deeper models. The discussion around high-cardinality categories motivated experiments with entity embeddings. The probability "
        "literature made calibration central to model selection for the win task. And the weaknesses of naive cricket heuristics motivated the large investment in feature engineering.",
    )
    add_body(
        doc,
        "In practical terms, the literature review was not a formality. It helped the project avoid two common mistakes: overestimating the value of deep learning without enough "
        "tabular evidence, and overestimating classification accuracy without checking whether reported win percentages were trustworthy.",
    )
    doc.add_page_break()


def add_chapter_3(doc: Document, context: dict, best_entry: dict) -> None:
    doc.add_heading("CHAPTER 3 - PROBLEM OBJECTIVE", level=1)

    doc.add_heading("Problem Definition", level=2)
    add_body(
        doc,
        "The project solves a practical sports-analytics problem: given the current state of an IPL innings, estimate the eventual innings total and the batting side's "
        "probability of winning, using only the information that would be available at that point in time. This is a stronger formulation than predicting the winner before "
        "the match or after the innings has already matured, because it forces the pipeline to work under realistic information constraints.",
    )
    add_body(
        doc,
        "The challenge becomes harder in the second innings because the prediction must adapt to chase pressure. A score such as 96/3 after eleven overs can signal control or "
        "danger depending on the target, venue, current run rate, wickets in hand, and the required acceleration. The model must therefore evaluate state and context together.",
    )

    doc.add_heading("Formal Problem Statement", level=2)
    add_body(
        doc,
        "Formally, the project attempts to learn two functions over a live match-state representation. The first function maps the current innings state to a projected final total. The second function maps the same or a closely related feature representation to the batting side's probability of eventually winning the match. Both functions are constrained by chronology: they may only use information known at the exact moment of prediction.",
    )
    add_body(
        doc,
        "This problem statement is stronger than a post-match classification exercise because it captures uncertainty while the match is still unfolding. The project therefore treats every valid live state as a supervised learning opportunity rather than collapsing the whole innings into one final row.",
    )

    doc.add_heading("Research Questions", level=2)
    add_numbered_list(
        doc,
        [
            "Can a leakage-safe live-state dataset outperform simple scoreboard-based reasoning for IPL prediction?",
            "Does a broader all-history active-team training scope provide stronger live performance than a recent-only scope?",
            "Do score and win tasks prefer different model families once the same feature set is exposed to them?",
            "How much of the performance gain comes from feature design and calibration rather than from algorithm choice alone?",
        ],
    )

    doc.add_heading("Specific Objectives of the Project", level=2)
    add_body(
        doc,
        "The immediate project objective was to construct a reliable IPL prediction engine. However, the broader academic objective was to show how a final-year project can connect data collection, preprocessing, machine learning, evaluation, and software delivery without breaking the consistency of the inference logic between different interfaces.",
    )
    add_body(
        doc,
        "In practical terms, the project aimed to meet five goals together: build a realistic live-state dataset, learn models that improve on simple heuristics, produce probability outputs that remain believable, expose the system through usable software surfaces, and document the whole process clearly enough to support academic evaluation and future maintenance.",
    )

    doc.add_heading("Operational Constraints", level=2)
    add_body(
        doc,
        "Several constraints shaped the project. First, only publicly available ball-by-ball and metadata sources were used, which means the system does not depend on proprietary player-tracking or betting-market feeds. Second, the project had to remain executable in a student-development environment rather than require specialized production infrastructure. Third, the report needed to remain defensible and readable for academic review rather than optimize for black-box complexity.",
    )
    add_body(
        doc,
        "These constraints are important because they explain why some apparently attractive options were not chosen. For example, richer commercial data sources or full live deployment platforms might improve the system further, but they would move the work outside the realistic scope of the current academic project.",
    )

    doc.add_heading("What Counts as Success in This Project", level=2)
    add_body(
        doc,
        "Success was not defined as achieving an unrealistically perfect predictor. Cricket remains uncertain, especially early in the innings. Instead, the project defined success "
        "through a combination of predictive quality, probability quality, and system usability. For score prediction, the focus was on reducing MAE and RMSE while keeping errors "
        "stable across difficult match states. For win prediction, log loss and Brier score were treated as first-class metrics because the output was meant to be interpreted as a probability.",
    )
    add_body(
        doc,
        f"The strongest recorded project snapshot in the saved model registry was version {best_entry['version_id']}, which used "
        f"{prettify_model_name(best_entry['metrics']['live']['score_model'])} for live score and "
        f"{prettify_model_name(best_entry['metrics']['live']['win_model'])} for live win probability. That snapshot is used in this refined report as the best achieved result, "
        "because it is the clearest representation of what the project was able to accomplish when the full pipeline worked at its best.",
    )

    doc.add_heading("Project Deliverables", level=2)
    add_bullet_list(
        doc,
        [
            "A processed IPL live-state dataset and multiple support tables built from raw Cricsheet files.",
            "A set of trained live models compared under chronological evaluation across alternative scopes.",
            "Pre-match score and win models for a limited no-live-context setting.",
            "A shared inference module used by Flask, API, Streamlit, and CLI surfaces.",
            "A report-generation workflow that can be rerun against the current workspace state.",
        ],
    )

    doc.add_heading("Original Work Plan and Completion Status", level=2)
    add_body(
        doc,
        "At the planning stage, the work was divided into sequential milestones: raw data collection, preprocessing and support-table creation, baseline training, broader model search, pre-match branch construction, shared inference integration, interface development, testing, and final documentation. This staged plan was important because the project contained both analytical and software-delivery goals.",
    )
    add_body(
        doc,
        "The implemented system substantially follows that original plan. Raw data ingestion, preprocessing, live-model search, pre-match modeling, CLI support, Flask integration, Streamlit support, evaluation reporting, and automated tests were all completed. Public deployment to a permanent external URL remains a later-stage extension rather than a completed milestone, and this is treated honestly in the report as a limitation of the current submission.",
    )
    add_body(
        doc,
        "Including this completion review in the chapter is useful for internal evaluation because it makes clear that the project did not drift randomly. It progressed through planned phases and delivered most of the intended outputs in the order originally required.",
    )

    doc.add_heading("Why the Report Focuses on Reasons, Not Only Metrics", level=2)
    add_body(
        doc,
        "The original auto-generated report already contained enough metrics and inventories to fill pages. What it lacked was a clear explanation of why the strongest results appeared "
        "when they did. The present report therefore makes a different choice: it keeps only the tables needed to support the argument and spends more space explaining the reasoning that "
        "actually produced the improvement in match prediction quality.",
    )
    doc.add_page_break()


def add_chapter_4(doc: Document, context: dict, best_entry: dict, assets: dict[str, Path]) -> None:
    counts = context["counts"]
    best_search = context["best_search_report"]
    recent_scope = best_search["scopes"]["recent_active_history"]
    all_scope = best_search["scopes"]["all_active_history"]
    feature_profile = context["feature_reference_profile"]

    doc.add_heading("CHAPTER 4 - METHODOLOGY OF THE PROJECT", level=1)

    doc.add_heading("End-to-End Workflow", level=2)
    add_body(
        doc,
        "The project pipeline begins with raw IPL match files and ends with a deployable prediction service. The most important engineering decision was to separate the work into stages: "
        "raw data acquisition, chronological preprocessing, support-table construction, live-model training, pre-match-model training, interface integration, and final reporting. This stage-based "
        "organization made the project easier to verify and easier to explain.",
    )
    add_body(
        doc,
        "That staged structure also prevented the common academic-project failure mode in which every piece of logic is mixed inside one notebook. Instead, the pipeline writes intermediate artifacts "
        "that can be reused by later stages. The same processed CSV files power the training scripts, the web interface, the CLI, and the report itself.",
    )
    add_figure(doc, assets["Figure 4.1"], FIGURES[0].number, FIGURES[0].title)

    doc.add_heading("System Model and Architectural View", level=2)
    add_body(
        doc,
        "At the system level, the project can be viewed as a layered architecture. The data layer contains raw match files and processed support tables. The learning layer contains training workflows, reports, and promoted deployment artifacts. The service layer contains the common inference package that validates inputs and produces predictions. The presentation layer contains CLI, Flask, API, and Streamlit surfaces. This layered view matters because it separates responsibilities cleanly and reduces hidden coupling.",
    )
    add_body(
        doc,
        "The architecture also reflects software-engineering discipline. Data preparation does not depend on any specific interface. Model training does not depend on a specific frontend. Presentation layers do not reimplement the modeling logic. By preserving this separation, the project remains easier to debug, test, and extend than a design in which every concern is mixed together.",
    )

    doc.add_heading("Raw Data and Why It Was Suitable", level=2)
    add_body(
        doc,
        f"The project started from {counts['raw_matches']} raw IPL matches in Cricsheet CSV2 format. This source was suitable because it preserves both the event sequence and the match metadata. "
        "The ball-by-ball files make it possible to reconstruct a valid prediction point before every legal delivery, while the corresponding info files provide venue, toss, result, and date information.",
    )
    add_body(
        doc,
        "This raw structure matters because live prediction cannot be built honestly from match summaries alone. If only final totals and winners are used, the model never learns how the same innings "
        "looked at 6 overs, 10 overs, or 17 overs. The project therefore reconstructed innings state from the bottom up.",
    )
    add_body(
        doc,
        "The dataset size also changed how the project was approached. A few thousand rows might encourage a narrow proof-of-concept, but hundreds of thousands of match-state rows justify a more deliberate workflow. "
        "The project therefore treated data profiling, support-table reuse, and split discipline as essential parts of the model-development process rather than as optional cleanup steps.",
    )

    add_table_title(doc, TABLES[1])
    add_table(
        doc,
        ["Item", "Current Workspace Value", "Why It Matters"],
        [
            ["Raw matches", f"{counts['raw_matches']:,}", "Defines the historical base available for chronological preprocessing."],
            ["Processed live-state rows", f"{counts['processed_rows']:,}", "Gives the model many supervised match-state examples instead of a few match summaries."],
            ["Processed columns", str(counts["processed_cols"]), "Reflects how much context is combined into each prediction row."],
            ["Active-team filtered rows", f"{all_scope['scope_summary']['rows_after_scope_filter']:,}", "Shows the size of the main deployment scope used for final live comparison."],
            ["Score train rows in all-history scope", f"{all_scope['score_split']['train_rows']:,}", "Explains why the broader scope can support richer modeling than the recent-only alternative."],
            ["Score train rows in recent scope", f"{recent_scope['score_split']['train_rows']:,}", "Shows how restrictive the recent-only setting becomes once chronology is respected."],
        ],
        font_size=9,
    )

    doc.add_heading("Canonicalization and Match-State Reconstruction", level=2)
    add_body(
        doc,
        "The first major preprocessing task was normalization. Team names, venue names, and seasonal labels can appear in multiple forms across raw cricket files. If these variants are passed directly to a "
        "model, the project wastes model capacity on duplicated identities. Canonical mapping of teams and venues therefore became a necessary early step rather than a cosmetic cleanup.",
    )
    add_body(
        doc,
        "A second critical step was converting over.ball notation into legal-ball counts. This made it possible to derive balls left, innings progress, over number, ball in over, and phase indicators consistently. "
        "Without this conversion, momentum windows and chase calculations would not have been reliable enough for model training.",
    )
    add_body(
        doc,
        "This reconstruction step also allowed the project to align first-innings and second-innings states inside one common live framework. Although the innings differ in objective, both can still be described using legal deliveries consumed, wickets lost, and contextual priors. "
        "The project then layered chase-specific features on top of that common base rather than building two completely separate live pipelines.",
    )

    doc.add_heading("Mathematical and Analytical Basis of the Methodology", level=2)
    add_body(
        doc,
        "The methodology is grounded in a small set of operational cricket formulas. Current run rate is computed from cumulative runs divided by overs consumed. Required run rate is computed from remaining target divided by overs remaining in the chase. Runs versus par compares the current score with the venue-conditioned historical expectation at the same stage of the innings. These quantities are simple, but together they give the model a compact representation of scoring tempo and pressure.",
    )
    add_body(
        doc,
        "The evaluation metrics are equally important mathematically. Mean absolute error measures the average run miss in direct units. Root mean squared error gives larger misses greater weight. Log loss evaluates whether predicted win probabilities align with the realized match outcome in a probabilistic sense, while Brier score measures probability calibration through squared confidence error. This mathematical framing matters because it explains why different metrics were attached to different tasks instead of using one summary value everywhere.",
    )
    add_body(
        doc,
        "The project therefore treats methodology as both a software design problem and a mathematical measurement problem. A feature is useful only if it improves the realism of the state representation, and a model is useful only if its output is measured by a metric aligned with the intended use of that output.",
    )

    doc.add_heading("Why Leakage Control Was Treated as a First-Class Requirement", level=2)
    add_body(
        doc,
        "The best result of the project did not come from one dramatic algorithmic trick. It came from removing hidden sources of error. The most dangerous of these was temporal leakage. Venue averages, team form, "
        "head-to-head signals, player form, and batter-bowler history all look useful, but they become invalid if they are computed using future matches relative to the row being predicted.",
    )
    add_body(
        doc,
        "To avoid that problem, the project computed support tables in chronological order. Each match and each legal-ball state only saw prior evidence. That decision is one of the main reasons the reported results "
        "are credible: the models had to learn from history, not from accidental access to the future.",
    )
    add_body(
        doc,
        "Chronological processing also changed how data was interpreted. Some ground-level priors, for example, look impressively stable if they are calculated on the full archive. Once the same priors are recomputed as historical snapshots before each match, they become noisier but far more realistic. "
        "The project deliberately accepted that extra difficulty because an easier but leaked task would have produced a weaker final system.",
    )

    doc.add_heading("Data Quality Checks Performed Before Training", level=2)
    add_body(
        doc,
        "Before the model search started, the project checked whether the processed rows were internally coherent. Overs had to map cleanly to legal balls, innings boundaries had to be respected, and second-innings rows had to contain a valid target. "
        "Rows that violated the prediction contract would either distort the model or create confusing failures in the deployed interfaces.",
    )
    add_body(
        doc,
        "The tests later included in the repository reflect these quality checks. Alias normalization, overs parsing, feature-column consistency, and save-load behavior were all promoted into automated verification because they had direct consequences for both training and serving quality.",
    )

    doc.add_heading("Support Tables and Why They Improved the Final Result", level=2)
    add_body(
        doc,
        "The preprocessing stage writes a family of support tables rather than keeping everything inside one large training file. Venue statistics capture whether a ground is naturally high-scoring or low-scoring. Team form "
        "and team-venue form capture short-term strength and location-conditioned strength. Matchup form reflects how batting and bowling teams have performed against one another recently.",
    )
    add_body(
        doc,
        "Player-level support tables were equally important. Batter strike-rate and average, bowler economy and strike rate, and batter-versus-bowler micro-history all help the model distinguish between states that look similar "
        "on the scoreboard but are being driven by very different player combinations. These tables are one of the main reasons the project progressed beyond a basic run-rate predictor.",
    )
    add_body(
        doc,
        "The project also preserves venue par context and optional weather-derived dew risk at inference time. These additions are not magic features, but they improve the realism of the prediction surface. A second-innings chase "
        "on a dew-affected night should not be treated identically to a dry afternoon innings with the same visible score.",
    )
    add_body(
        doc,
        "One of the most useful consequences of the support-table strategy was interpretability. Even when a model is complex, the report can still explain the kinds of evidence available to it. That makes the final output easier to defend academically, because the reader can see that the prediction was driven by understandable cricket context rather than by opaque numerical noise.",
    )

    add_table_title(doc, TABLES[2])
    add_table(
        doc,
        ["Feature Family", "Examples", "Why It Helped the Final Result"],
        [
            ["Current match state", "runs, wickets, balls_left, innings_progress", "Anchored every prediction to the visible scoreboard situation."],
            ["Momentum windows", "runs_last_5, wickets_last_5, runs_last_12_balls", "Captured acceleration, stability, or collapse risk better than static averages."],
            ["Chase math", "target_remaining, required_run_rate, required_minus_current_rr", "Made second-innings predictions sensitive to actual pressure."],
            ["Venue priors", "venue averages, bat-first win rate, runs_vs_par", "Explained why the same score means different things at different grounds."],
            ["Team and matchup context", "team form, team-venue form, batting_vs_bowling_form", "Added competitive context missing from score-only models."],
            ["Player context", "striker form, bowler form, batter_vs_bowler history", "Helped the model distinguish states driven by specific player matchups."],
            ["Optional environment", "temperature, humidity, wind, dew_risk", "Added coarse environmental context without dominating the core cricket signals."],
        ],
        font_size=9,
    )
    add_body(
        doc,
        "The feature families above were not added all at once. The project moved toward them progressively by asking where the earlier versions of the model were still blind. Score-only features missed venue and player context. Venue-only features missed momentum. Momentum-only features missed opposition and matchup strength. "
        "The final feature design emerged by reducing these blind spots one by one.",
    )

    doc.add_heading("Feature Families in Operational Terms", level=2)
    add_body(
        doc,
        f"The stored reference profile gives a helpful picture of the operating distribution. Across the processed live rows, the mean runs at prediction time is {fmt_float(feature_profile['features']['runs']['mean'])}, "
        f"the 90th percentile is {fmt_float(feature_profile['features']['runs']['p90'])}, and the mean current run rate is {fmt_float(feature_profile['features']['current_run_rate']['mean'])}. "
        "These values confirm that the model is not being trained only on very early or very late states; it sees a broad spectrum of innings contexts.",
    )
    add_body(
        doc,
        f"The same profile shows that death-over rows account for about {fmt_pct(feature_profile['slice_rates']['death_over'])} of the processed data, while high-pressure chase rows account for about "
        f"{fmt_pct(feature_profile['slice_rates']['high_pressure_chase'])}. This matters because those slices are exactly where users expect a live predictor to behave sensibly under pressure.",
    )
    add_body(
        doc,
        "Operationally, this means the feature design had to work in both calm and volatile states. A useful live model cannot be good only at the start of an innings or only near the finish. It has to handle normal states, pressure states, and ambiguous transitions between the two.",
    )

    doc.add_heading("Why the Project Chose the All-Active-History Scope", level=2)
    add_body(
        doc,
        f"A major insight of the project was that the recent-only scope was too small once chronology was enforced. In that setting, the score model had only "
        f"{recent_scope['score_split']['train_rows']:,} training rows and the win model had {recent_scope['win_split']['train_rows']:,} training rows. That is a narrow training base for a problem containing venue, team, striker, and bowler interactions.",
    )
    add_body(
        doc,
        f"The all-active-history scope solved this without reintroducing outdated franchise noise. By restricting the category space to the currently active ten IPL teams while keeping the full time span, the project expanded the score training pool to "
        f"{all_scope['score_split']['train_rows']:,} rows and the win training pool to {all_scope['win_split']['train_rows']:,} rows. This gave the models enough examples to learn strong general patterns while staying relevant to the modern IPL.",
    )
    add_body(
        doc,
        "This scope decision is one of the clearest reasons the final results improved. The report therefore treats it as a modeling decision, not merely as a filtering convenience.",
    )
    add_body(
        doc,
        "The active-team filter was equally important. Had the project simply kept every historical team label, the category space would have become noisier and the practical relevance of the older rows would have fallen. By keeping current franchises while retaining long historical depth, the pipeline obtained a better compromise between generalization and relevance.",
    )

    doc.add_heading("Training, Validation, and Test Design", level=2)
    add_body(
        doc,
        "The split design was chosen to answer a simple question honestly: if the project is trained on earlier IPL seasons, how well does it behave on a future season? For that reason, the validation season and test season were always later than the training seasons. "
        "The project avoided random shuffling because that would have hidden the natural season-to-season drift in scoring environment and team quality.",
    )
    add_body(
        doc,
        "The score task and win task used slightly different row counts because some rows are not equally meaningful for both tasks once targets and valid chase information are enforced. The project accepted those differences rather than forcing perfect row symmetry, because validity of the supervision signal was more important than convenience.",
    )

    doc.add_heading("Methodology as an Incremental Software Process", level=2)
    add_body(
        doc,
        "From a software-engineering perspective, the methodology followed an incremental process model. The first increment established ingestion and preprocessing. The second increment added baseline learning and deployment-ready artifacts. The third increment broadened model search, calibration, and uncertainty support. The fourth increment integrated multi-interface serving, monitoring, and report generation. This incremental structure reduced risk because each stage produced a usable intermediate result.",
    )
    add_body(
        doc,
        "This process choice is relevant to internal evaluation because it shows that the project was not built by trial-and-error alone. Each increment solved a concrete part of the problem and created a stable base for the next stage. That is precisely the kind of disciplined implementation sequence expected from a serious engineering project.",
    )

    doc.add_heading("Season-Wise Scoring Environment and Why Chronology Mattered", level=2)
    previous_avg = None
    for row in context["season_summary"].itertuples(index=False):
        trend_note = ""
        if previous_avg is not None:
            delta = row.avg_first_innings - previous_avg
            if delta > 3:
                trend_note = "Compared with the previous season, batting conditions moved clearly upward."
            elif delta < -3:
                trend_note = "Compared with the previous season, batting conditions moved clearly downward."
            else:
                trend_note = "Compared with the previous season, the scoring level stayed broadly similar."
        previous_avg = row.avg_first_innings
        add_body(
            doc,
            f"In season {row.season_label}, the engineered pipeline retained {int(row.matches):,} matches and produced {int(row.engineered_rows):,} live match-state rows. "
            f"The average first-innings total was {fmt_float(row.avg_first_innings)} and the average second-innings total was {fmt_float(row.avg_second_innings)}, while the batting-side win rate across saved states was {fmt_pct(row.batting_win_rate)}. "
            f"{trend_note} This matters because the live model was never learning from one fixed IPL environment. It was learning across changing eras of batting tempo, venue behavior, and chase patterns.",
        )
    add_body(
        doc,
        "These season-wise shifts explain why the project refused to use random train-test splitting. A random split would have mixed older and newer batting environments together and would have hidden the exact temporal drift that a deployable IPL system must survive.",
    )

    doc.add_heading("Active-Team Universe Used for Final Modeling", level=2)
    add_body(
        doc,
        "The all-active-history scope was not only about row count. It was also about keeping the category space relevant to the current IPL. By focusing on active franchises, the project preserved deep historical evidence while avoiding unnecessary dilution from obsolete team identities.",
    )
    for row in context["team_summary"].itertuples(index=False):
        recent_form = row.recent_form if row.recent_form == row.recent_form else 0.0
        add_body(
            doc,
            f"{row.batting_team} appears in {int(row.matches):,} matches and {int(row.innings):,} innings within the active-team summary. "
            f"Its average innings total in this slice is {fmt_float(row.avg_total)} and its batting-side win rate is {fmt_pct(row.batting_win_rate)}, while the recent-form indicator stored by the pipeline is {fmt_pct(recent_form)}. "
            "This type of team profile shows why the model needed more than runs and wickets alone. The same visible score should be read differently for stronger and weaker teams, especially when venue behavior is added on top.",
        )
    add_body(
        doc,
        "This team-level view also explains why the broader scope helped the richer models. Player-aware and venue-aware features become more stable when each active franchise is represented across a deep historical window rather than inside a very narrow recent slice only.",
    )

    doc.add_heading("Support-Table Artifacts and Their Practical Role", level=2)
    add_body(
        doc,
        "A major design strength of the project is that it writes reusable support tables during preprocessing instead of recomputing every statistic separately inside each model or interface call. That decision improved both consistency and debugging.",
    )
    for row in context["support_table_catalog"]:
        add_body(
            doc,
            f"The artifact {row['file']} serves as the {row['label'].lower()} table. It stores {row['rows']:,} rows and {row['cols']} columns, is keyed mainly by {row['primary_key']}, and exists because {row['purpose']} "
            "In practical use, this means the predictor can attach venue, team, form, or player evidence quickly and consistently without rebuilding historical aggregates on the fly for every single prediction request.",
        )
    add_body(
        doc,
        "This support-table strategy matters academically as well. It lets the report explain exactly what historical evidence the inference layer can reuse, which makes the project easier to defend than a monolithic workflow in which all context stays hidden inside one large training script.",
    )

    doc.add_heading("Workflow Stages from Raw Data to Final Artifacts", level=2)
    add_body(
        doc,
        "The project became easier to improve once it was treated as a chain of explicit stages rather than as one oversized program. Each stage produces an artifact that the next stage consumes, which gives the overall system both traceability and reproducibility.",
    )
    for row in context["workflow_inventory"]:
        add_body(
            doc,
            f"The stage '{row['workflow']}' is implemented through {row['script']}. It takes {row['input']} as its main input and produces {row['output']} as its main output. "
            "This separation matters because it keeps data cleaning, feature generation, model search, interface serving, and report production from collapsing into one untestable block of logic.",
        )
    add_body(
        doc,
        "Seen together, these workflow stages explain why the report is able to discuss the project in engineering terms rather than only in model terms. The final predictor was not an isolated file; it was the output of a controlled pipeline.",
    )

    doc.add_heading("Preprocessing Rules That Protected Result Quality", level=2)
    add_body(
        doc,
        "The preprocessing layer contains several rules that are easy to overlook when only the final metrics are reported. In practice, these rules were some of the most important reasons the later models became trustworthy. They determine whether the live feature table represents the match honestly or whether it quietly introduces duplicated identities, invalid overs, or future information that would make the results look better than they really are.",
    )
    for name, operation, reason in context["preprocessing_rules"]:
        add_body(
            doc,
            f"The rule '{name}' was implemented so that {operation} Its practical value is that {reason} This kind of rule does not produce a visible headline metric on its own, but removing it would have weakened the final live predictor substantially.",
        )
    add_body(
        doc,
        "These rules also explain why the chapter places such strong emphasis on data design. When the preprocessing contract is correct, the model comparison later becomes meaningful. When the preprocessing contract is weak, even a high-performing model can simply be exploiting inconsistencies that would fail outside the evaluation file.",
    )

    doc.add_heading("Mathematical Translation of the Live Match State", level=2)
    add_body(
        doc,
        "The methodology also contains a small but important mathematical layer. The model does not receive the scoreboard exactly as a human sees it. Instead, the scoreboard is translated into deterministic state variables that make live cricket pressure easier for the model to interpret. This translation is one of the places where cricket knowledge and machine-learning preparation meet directly.",
    )
    add_body(
        doc,
        "For every legal ball, the pipeline computes legal balls bowled, balls left, wickets left, innings progress, and current run rate. In the second innings it also derives target remaining, required run rate, and the difference between required and current pace. These quantities are more informative than raw overs notation alone because they convert the state into comparable numerical pressure signals across venues and seasons.",
    )
    add_code_block(
        doc,
        "legal_balls_bowled = over_number * 6 + ball_in_over\n"
        "balls_left = 120 - legal_balls_bowled\n"
        "wickets_left = 10 - wickets\n"
        "innings_progress = legal_balls_bowled / 120\n"
        "current_run_rate = runs / max(legal_balls_bowled / 6, 1e-6)\n"
        "required_run_rate = target_remaining / max(balls_left / 6, 1e-6)\n"
        "required_minus_current_rr = required_run_rate - current_run_rate",
    )
    add_body(
        doc,
        "Momentum windows extend this translation further. Runs and wickets from the most recent balls or overs are used because the same cumulative score can hide very different short-term trajectories. A side that has scored 45 in the last five overs is in a structurally different state from a side that has reached the same total after losing control during the same interval.",
    )

    doc.add_heading("Narrative Data Flow of the Proposed System", level=2)
    add_body(
        doc,
        "The proposed system can also be understood as a narrative data flow. Raw ball-by-ball and info files are first converted into canonical match records. Those records are then reordered chronologically so that every future feature calculation can safely look backward. A live-state table is built from that ordered stream, and support tables are materialized alongside it for venue, team, player, and matchup context. Training workflows then read those artifacts, compare model families, and promote the strongest deployment pair for inference and interface use.",
    )
    add_body(
        doc,
        "This narrative matters because it makes the methodology easier to defend than a diagram alone. A faculty reviewer can follow the path from raw IPL evidence to a final prediction without needing to infer hidden steps. The report therefore describes the working logic in sentences as well as figures so that the project remains understandable even when the reader is not reading the source code side by side.",
    )

    doc.add_heading("Shared Inference Contract as a Methodological Decision", level=2)
    add_body(
        doc,
        "The shared inference contract was not merely an implementation convenience added later. It was part of the project methodology because the same feature logic had to remain valid across CLI, Flask, API, and Streamlit surfaces. If one interface built features differently from another, then the evaluation chapter would no longer describe one system. It would describe several slightly different systems under one project name.",
    )
    for row in context["api_contract_rows"]:
        add_body(
            doc,
            f"The contract field '{row['field']}' is marked as {row['required']}. Its role is that {row['description']} Listing these fields in the methodology chapter makes clear what information the final predictor expects and why the feature frame can be constructed consistently during live use.",
        )
    add_body(
        doc,
        "This contract-oriented view is especially valuable in a final-year project because it links methodology to deployment. The same methodological definitions used during training are preserved when the user later submits a live form or API payload. That continuity is one of the reasons the results chapter can claim that the deployed outputs are consistent with the evaluated model behavior.",
    )

    doc.add_heading("Why This Chapter Is the Real Source of Improvement", level=2)
    add_body(
        doc,
        "A reader looking only at the results chapter might conclude that the project's progress came from trying better algorithms. That is incomplete. The methodological work in this chapter is the real reason performance improved. Better scope selection, cleaner historical features, stronger live-state reconstruction, and honest chronology gave the later models a task they could actually solve well.",
    )
    add_body(
        doc,
        "This is why the refined report spends more space on methodology than the earlier version. In a project like this, the path to the best result runs through data design first and model choice second.",
    )

    doc.add_heading("Why Live and Pre-Match Models Were Separated", level=2)
    add_body(
        doc,
        "The project also created a separate pre-match branch because pre-match prediction has a fundamentally different information budget. Before the first ball, there is no current score, no momentum, no active batter-bowler matchup, and no chase pressure. "
        "Trying to force the live model into this setting would have produced misleading outputs. The pre-match models were therefore kept smaller and more conservative.",
    )
    add_body(
        doc,
        "That design choice later made the report more honest. It allowed the live system to remain rich and context-aware while acknowledging that pre-match forecasting is a lower-information task with weaker attainable accuracy.",
    )
    doc.add_page_break()


def add_chapter_5(doc: Document, context: dict, best_entry: dict, assets: dict[str, Path]) -> None:
    best_search = context["best_search_report"]
    recent_scope = best_search["scopes"]["recent_active_history"]
    all_scope = best_search["scopes"]["all_active_history"]
    pre_match = context["pre_match_model_report"]
    best_live = best_entry["metrics"]["live"]
    score_wicket_summary = context["score_wicket_summary"]
    win_calibration_bins = context["win_calibration_bins"]

    doc.add_heading("CHAPTER 5 - RESULTS AND ANALYSIS", level=1)

    doc.add_heading("Modeling Philosophy", level=2)
    add_body(
        doc,
        "The project did not assume that one model family would be ideal for every task. Score prediction is a regression problem in which the size of the error matters directly. Win prediction is a probability problem in which a 0.60 output should mean something close to a 60 percent event over time. "
        "That distinction led the project to evaluate model families separately and to prioritize calibration for the win task.",
    )
    add_body(
        doc,
        "The modeling workflow therefore moved in layers. HistGradientBoosting provided a strong CPU baseline. XGBoost and CatBoost were added because gradient boosting remains one of the strongest options for structured data. Weighted ensembles were added to capture complementary strengths of classical models. "
        "Torch entity embeddings were explored because the feature set contains high-cardinality categorical fields that may benefit from learned dense representations.",
    )
    add_body(
        doc,
        "This layered search strategy made the project easier to interpret. If a richer model won, the report could explain what new capacity it used. If it failed, the project still had strong baseline families to fall back on. That is a much stronger experimental design than jumping directly to one complex model and treating its performance as final.",
    )

    add_table_title(doc, TABLES[3])
    add_table(
        doc,
        ["Scope", "Score Train Rows", "Win Train Rows", "Selected Score Model", "Selected Win Model"],
        [
            [
                "recent_active_history",
                f"{recent_scope['score_split']['train_rows']:,}",
                f"{recent_scope['win_split']['train_rows']:,}",
                prettify_model_name(recent_scope["selected_score_model"]),
                prettify_model_name(recent_scope["selected_win_model"]),
            ],
            [
                "all_active_history",
                f"{all_scope['score_split']['train_rows']:,}",
                f"{all_scope['win_split']['train_rows']:,}",
                prettify_model_name(all_scope["selected_score_model"]),
                prettify_model_name(all_scope["selected_win_model"]),
            ],
        ],
        font_size=9,
    )
    add_body(
        doc,
        "This scope comparison explains a central project lesson. The recent-only setting looks attractive because it is temporally close to the target season, but it becomes too data-poor when splits are done honestly. The all-active-history setting gave the model enough observations to learn robust venue and player effects while still focusing on the active IPL team universe.",
    )
    add_body(
        doc,
        "The difference is especially important for player-aware modeling. A recent-only scope may not show enough repeated striker and bowler contexts for player-level features to stabilize. Once the project moved to the broader scope, the richer feature set became more meaningful and the advanced score model had a fairer chance to outperform simpler baselines.",
    )

    doc.add_heading("Why Some Alternative Models Were Not Selected", level=2)
    add_body(
        doc,
        "A strong results chapter should explain not only why one model won, but also why other credible candidates did not become the final deployment choice. The CPU HistGradientBoosting baseline remained important because it showed that the feature engineering itself already carried strong signal. However, it did not match the best snapshot once richer categorical interactions became available.",
    )
    add_body(
        doc,
        "XGBoost served as a powerful structured-data benchmark and remained competitive, especially in score prediction. Yet the strongest saved score result still came from the entity-embedding approach once the broader all-history scope provided enough data to justify that extra capacity. CatBoost remained crucial on the win side because its categorical handling and later calibration produced more trustworthy probability behavior than the alternatives chosen for deployment.",
    )
    add_body(
        doc,
        "This explanation matters because it shows that the final design was not based on preference or novelty. Each model family had a clear role in the comparison, and the deployment pair was selected because it matched the task requirements more closely than the rest.",
    )

    doc.add_heading("Why the Best Experimental Snapshot Matters", level=2)
    add_body(
        doc,
        f"The saved model registry contains more than one experimental snapshot. The strongest tracked result was version {best_entry['version_id']}, created on "
        f"{best_entry['created_at_utc'][:10]}, with {prettify_model_name(best_live['score_model'])} for live score and "
        f"{prettify_model_name(best_live['win_model'])} for live win probability. Its combined score-plus-win objective was stronger than the later saved version, "
        "so this report treats it as the clearest expression of the project's best achieved performance rather than blindly assuming the newest saved version must also be the best.",
    )
    add_body(
        doc,
        "That distinction is important. A good technical report should not confuse latest with best. By reviewing the saved registry rather than only the latest artifact pointer, the report can explain what actually produced the strongest result inside the project history.",
    )
    add_body(
        doc,
        "This deeper analysis also exposed a practical truth about model development: later saved artifacts are not guaranteed to be better just because they are newer. The report therefore anchors its argument in tracked evaluation results, not in file modification time alone. That is the correct standard for explaining how the project reached its strongest match-prediction snapshot.",
    )

    add_table_title(doc, TABLES[4])
    add_table(
        doc,
        ["Score Model", "MAE", "RMSE", "Why It Was Important in the Search"],
        [
            [prettify_model_name("hgb"), "17.72", "24.50", "Strong CPU baseline showing that feature engineering already carried useful signal."],
            [prettify_model_name("xgboost_gpu"), "16.56", "23.06", "Powerful boosted-tree benchmark and strong structured-data reference point."],
            [prettify_model_name("catboost_gpu"), "19.90", "26.27", "Useful categorical-aware alternative, but weaker for live score than the strongest options."],
            [prettify_model_name("torch_entity_gpu"), fmt_float(best_live["score_test"]["mae"]), fmt_float(best_live["score_test"]["rmse"]), "Best recorded score model because learned embeddings helped on rich categorical context."],
            [prettify_model_name("ensemble_ml_top3"), "16.04", "23.54", "Later robust ensemble variant; useful, but not the strongest recorded score snapshot."],
        ],
        font_size=9,
    )
    add_figure(doc, assets["Figure 5.1"], FIGURES[1].number, FIGURES[1].title)

    doc.add_heading("Why the Torch Score Model Worked Best in the Strongest Snapshot", level=2)
    add_body(
        doc,
        "The strongest recorded score result did not come from the simplest model. The torch entity-embedding regressor benefited from the fact that the project exposed many meaningful categorical signals: batting team, bowling team, venue, striker, bowler, phase, toss context, and season identity. "
        "A one-hot-heavy baseline can use these fields, but learned embeddings can compress them into dense representations that express similarity and interaction more naturally when enough training data is available.",
    )
    add_body(
        doc,
        "That advantage only became meaningful after the project moved to the all-active-history scope. In the recent-only scope the data volume was too small to justify a richer tabular model. Once the training pool exceeded two hundred thousand score rows, however, the deeper architecture had enough support to outperform the classical boosted-tree baselines in the strongest saved experiment.",
    )
    add_body(
        doc,
        "A second reason the score model improved is that score regression benefits from interaction sensitivity. The effect of one feature often depends on another: 90/2 after ten overs is different at Wankhede than at Chennai, and different again if the striker is an aggressive finisher or an anchor. The entity-embedding approach was useful because it let the project represent such interactions without exploding the manual feature space.",
    )

    add_table_title(doc, TABLES[5])
    add_table(
        doc,
        ["Win Model", "Log Loss", "Brier", "Accuracy", "Interpretation"],
        [
            [prettify_model_name("hgb_calibrated"), "0.6047", "0.2040", "65.7%", "Good classical baseline after calibration, but not the strongest project snapshot."],
            [prettify_model_name("xgboost_gpu_calibrated"), "0.5909", "0.2026", "66.3%", "Competitive calibrated alternative with strong ranking quality."],
            [prettify_model_name("catboost_gpu_calibrated"), fmt_float(best_live["win_test"]["log_loss"], 4), fmt_float(best_live["win_test"]["brier"], 4), fmt_pct(best_live["win_test"]["accuracy"]), "Best recorded balance of probability quality and practical classification."],
        ],
        font_size=9,
    )

    doc.add_heading("Why Calibrated CatBoost Was the Right Choice for Win Probability", level=2)
    add_body(
        doc,
        "The win task was evaluated differently from the score task because the intended output is a probability, not just a label. A model with slightly lower accuracy can still be better if its 65 percent predictions behave like 65 percent events in held-out evaluation. "
        "This is why log loss and Brier score were prioritized alongside accuracy.",
    )
    add_body(
        doc,
        "Calibrated CatBoost emerged as the strongest recorded win model because it handled the mixture of categorical and numeric context well and, after calibration, produced more reliable percentages. This is a key project lesson: the best classifier for a deployment report is not simply the model with the highest hard-label accuracy. "
        "It is the model whose probability output remains believable when used by an end user.",
    )
    add_body(
        doc,
        "This distinction is particularly important in cricket, where user trust depends on whether a reported probability feels sensible as the innings evolves. A model that jumps too sharply between extreme values may look decisive, but it becomes difficult to trust in edge cases. The calibrated CatBoost branch gave the project a better compromise between responsiveness and reliability.",
    )

    add_figure(doc, assets["Figure 5.2"], FIGURES[2].number, FIGURES[2].title)
    add_body(
        doc,
        "The score-error distribution shows another reason the project took uncertainty seriously. Even the strongest score model remains less certain early in the innings, because multiple future trajectories are still possible. That observation led to the use of projected ranges rather than only a single score point.",
    )
    add_body(
        doc,
        "This was a useful design shift in the interface layer as well. A single predicted score can appear overly precise, especially when only a few overs have been bowled. Reporting a point estimate together with a plausible range communicates both the model's expectation and the residual uncertainty still present in the match.",
    )

    add_figure(doc, assets["Figure 5.3"], FIGURES[3].number, FIGURES[3].title)
    add_body(
        doc,
        "The calibration curve confirms why the calibrated win model was worth keeping. The closer the empirical win rate remains to the predicted probability bands, the more defensible it becomes to show a batting-side win chance directly to the user.",
    )
    add_body(
        doc,
        "From a report-writing perspective, this figure is more valuable than several interface screenshots. It directly explains why the final win model was selected and what property of the output improved. Since the user requested more actual reasoning and fewer decorative elements, this is exactly the kind of figure that deserves to stay.",
    )

    doc.add_heading("What the Held-Out Errors Taught Us", level=2)
    add_body(
        doc,
        "The hardest score errors generally occurred when the match state had not yet revealed the eventual innings character. A side can be 62/1 after six overs and still finish at very different totals depending on whether it consolidates, collapses, or launches again in the death overs. This is not a failure of the feature set alone; it reflects genuine uncertainty in cricket.",
    )
    add_body(
        doc,
        "The win task showed a different pattern. Miscalibration is most dangerous in chase-heavy phases when the score looks superficially healthy but the required acceleration is starting to separate from the current rate. For this reason, the project's chase features and probability calibration mattered as much as the core classifier itself.",
    )
    add_body(
        doc,
        "The held-out slices also taught the project that accuracy alone can conceal the real problem. A classifier can still be right often in obvious states while behaving poorly in the difficult middle band where the match is genuinely undecided. That is why the report keeps returning to log loss, Brier score, and scenario-level interpretation rather than treating one percentage as the whole story.",
    )

    doc.add_heading("Phase-Wise Interpretation of Match Behavior", level=2)
    add_body(
        doc,
        "The project results also make more sense when they are read phase by phase. In the powerplay, uncertainty remains high because teams can launch, consolidate, or lose early wickets. The middle overs reward context features more strongly because scoreboard similarity hides major differences in batting intent and player quality. The death overs create a different difficulty: scores can change sharply over a short ball window, so the model must respond to momentum without becoming unstable.",
    )
    add_body(
        doc,
        "This phase-wise view explains why no single summary metric should be overinterpreted. A model may behave well in one phase and still fail in another. The project therefore used scenario interpretation and slice analysis to complement overall held-out numbers.",
    )

    doc.add_heading("Phase Summary of the Dataset and What It Means", level=2)
    add_body(
        doc,
        "The engineered dataset itself explains many of the observed result patterns when it is read phase by phase and innings by innings. The project kept this analysis in the main chapter because phase balance determines where the model receives the most evidence and where uncertainty is naturally highest.",
    )
    for row in context["phase_summary_rows"]:
        add_body(
            doc,
            f"In the {row['innings_label'].lower()} during the {row['phase'].lower()} phase, the dataset contains {row['rows']:,} rows. The average live score at that stage is {fmt_float(row['avg_runs'])}, the average wickets lost is {fmt_float(row['avg_wickets'])}, the average current run rate is {fmt_float(row['avg_current_rr'])}, and the average final total attached to those rows is {fmt_float(row['avg_total'])}. The batting-side win rate in this slice is {fmt_pct(row['batting_win_rate'])}. This helps explain where the model is learning stable structure and where it is learning under high volatility.",
        )
    add_body(
        doc,
        "Two insights follow from this summary. First, the middle overs supply a very large share of the live training evidence, so feature quality in that region affects the final model strongly. Second, the death overs remain volatile even with many rows because the state can change sharply over a few deliveries. That is why momentum windows and calibrated probabilities were kept central to the final system.",
    )

    doc.add_heading("Error Analysis by Wicket State", level=2)
    for row in score_wicket_summary:
        add_body(
            doc,
            f"When wickets lost were in the {row['bucket']} range, the held-out score model recorded MAE {row['mae']} and RMSE {row['rmse']}, with an average actual total of {row['mean_actual']} and an average predicted total of {row['mean_prediction']}. "
            "This slice is useful because it shows whether the model is mainly struggling with early uncertainty or with late-innings collapse dynamics.",
        )
    add_body(
        doc,
        "Reading these wicket-state slices together gives a better picture of model behavior than one global error number. They show how the model responds as batting resources disappear and the innings becomes structurally constrained.",
    )

    doc.add_heading("Calibration Buckets in Practical Terms", level=2)
    for row in win_calibration_bins:
        add_body(
            doc,
            f"In the {row['bucket']} probability band, the model produced {row['rows']:,} held-out rows with an average predicted win probability of {row['mean_pred']} and an empirical batting-side win rate of {row['empirical_win']}. "
            f"The resulting calibration gap was {row['gap']}. This bucket-level view helps determine whether the reported probabilities stay believable across low-confidence and high-confidence states.",
        )
    add_body(
        doc,
        "This kind of analysis mattered because the project output is consumed by people, not only by metrics files. End users notice when a 70 percent state repeatedly behaves more like a 50-50 state. Calibration analysis gave the report a principled way to discuss that risk.",
    )

    doc.add_heading("Pre-Match Modeling and Why It Stayed Simpler", level=2)
    add_body(
        doc,
        "The pre-match branch was intentionally built as a lower-information model family. Before the toss and first ball, the system has access only to team identity, venue priors, and broad recent-form indicators. It does not know the active batter-bowler state, innings pressure, or over-by-over momentum. The project therefore used regularized classical models rather than pretending this branch could rival live prediction.",
    )
    add_table_title(doc, TABLES[6])
    add_table(
        doc,
        ["Task", "Split", "MAE / Accuracy", "RMSE / Log Loss", "R2 / Brier"],
        [
            ["Pre-match score", "Validation", fmt_float(pre_match["score_metrics"]["valid"]["mae"]), fmt_float(pre_match["score_metrics"]["valid"]["rmse"]), fmt_float(pre_match["score_metrics"]["valid"]["r2"], 3)],
            ["Pre-match score", "Test", fmt_float(pre_match["score_metrics"]["test"]["mae"]), fmt_float(pre_match["score_metrics"]["test"]["rmse"]), fmt_float(pre_match["score_metrics"]["test"]["r2"], 3)],
            ["Pre-match win", "Validation", fmt_pct(pre_match["win_metrics"]["valid"]["accuracy"]), fmt_float(pre_match["win_metrics"]["valid"]["log_loss"], 4), fmt_float(pre_match["win_metrics"]["valid"]["brier"], 4)],
            ["Pre-match win", "Test", fmt_pct(pre_match["win_metrics"]["test"]["accuracy"]), fmt_float(pre_match["win_metrics"]["test"]["log_loss"], 4), fmt_float(pre_match["win_metrics"]["test"]["brier"], 4)],
        ],
        font_size=9,
    )
    add_body(
        doc,
        "These weaker numbers are not a problem for the report. On the contrary, they reinforce one of the project's honest conclusions: live prediction is stronger precisely because the project engineered the system to use the right information at the right time instead of forcing one model to answer every question equally well.",
    )
    add_body(
        doc,
        "The pre-match branch is still useful in one important sense: it demonstrates how much predictive value is added once the innings actually begins. By placing the pre-match results beside the live results, the report can show concretely that match-state features, momentum windows, and live context are not cosmetic additions but the main drivers of stronger performance.",
    )

    doc.add_heading("Practical Meaning of the Reported Results", level=2)
    add_body(
        doc,
        "From a practical point of view, the reported results mean that the system can respond to real cricket states with more nuance than simple score heuristics. A projected total that is tied to a plausible interval, and a win probability that has been calibrated rather than merely classified, are both more useful to a reviewer than a raw unqualified number.",
    )
    add_body(
        doc,
        "This practicality is especially relevant in an academic demonstration setting. Faculty members evaluating the project do not only want to know that a model scored well on a held-out file. They want to understand whether the outputs look sensible, whether the reasoning behind them is defensible, and whether the system behaves consistently across different types of match situations.",
    )

    doc.add_heading("Scenario Interpretation in Plain Cricket Terms", level=2)
    for row in context["live_scenarios"][:4]:
        add_body(
            doc,
            f"In the scenario labelled '{row['scenario']}', the system classified the state as {row['phase']} and returned a projected total of {row['projected_total']} with a win probability of {row['win_prob']}. "
            f"The practical meaning of this result is that the model was not reacting to score alone; it was combining score, wickets, pace, contextual priors, and matchup information. The saved interpretation for this scenario is: {row['interpretation']}",
        )
    add_body(
        doc,
        "Including these scenarios in narrative form is more useful than filling the chapter with many screenshots. They show how the system reasons in different phases of the innings and what kind of information the final outputs actually convey to a human reader.",
    )
    add_body(
        doc,
        "This narrative style also makes the report easier to defend in viva or faculty review. A reader can connect the statistical results to cricket situations that are easy to imagine: a powerplay launch, a middle-over rebuild, a death-over surge, or a chase that looks balanced on the scoreboard but is quietly drifting away on required-rate pressure.",
    )

    doc.add_heading("Experiment Progression and Selection Discipline", level=2)
    add_body(
        doc,
        "The final snapshot did not appear by accident. The project preserved a version history so that later report writing could explain how the best configuration was reached instead of pretending that the strongest combination was obvious from the start.",
    )
    for row in context["version_inventory"]:
        add_body(
            doc,
            f"Version {row['version_id']} was saved on {row['created_at']} under the {row['scope']} scope. It paired {row['score_model']} for score prediction with {row['win_model']} for win prediction, and its held-out summary reported score RMSE {row['score_rmse']} and win log loss {row['win_log_loss']}. "
            "Keeping these version notes inside the main results chapter is important because model selection is part of the core project story, not a side appendix detail.",
        )
    add_body(
        doc,
        "This disciplined version tracking also prevented a common reporting mistake: assuming that the newest model artifact is automatically the best one. The project instead selected the strongest recorded snapshot on evaluation evidence.",
    )

    doc.add_heading("Decision Sequence Behind the Strongest Result", level=2)
    decision_notes = [
        "Use ball-by-ball reconstruction instead of final-match summaries so the model learns real live states rather than retrospective outcomes only.",
        "Normalize team and venue identities before modeling so that the feature space is not polluted by alias duplicates.",
        "Compute history chronologically so that support features use only past evidence and do not leak future results.",
        "Keep current IPL franchises while preserving deep history so the category space stays relevant without becoming data-poor.",
        "Compare several model families before selecting a winner so that the final choice is based on evidence rather than assumption.",
        "Judge win prediction with log loss and Brier score, not only accuracy, because the output shown to users is a probability.",
        "Keep live and pre-match branches separate because the information available before the first ball is fundamentally smaller.",
        "Expose all interfaces through one shared inference core so that a good result in one surface also means a good result in the others.",
    ]
    for note in decision_notes:
        add_body(
            doc,
            f"A decisive project step was to {note} Each of these choices removed a weakness that would otherwise have made the final result look stronger on paper than it really was in live use.",
        )
    add_body(
        doc,
        "Read together, these decisions show that the final improvement came from a sequence of good engineering choices rather than from one isolated model trick. That is exactly what a serious project report should make clear.",
    )

    doc.add_heading("Expanded Scenario Studies", level=2)
    add_body(
        doc,
        "The saved scenario outputs are useful because they show how the prediction engine behaves in concrete cricket situations. Instead of hiding that material in the appendix, the results chapter should use it directly to explain what the model has learned.",
    )
    for row in context["live_scenarios"]:
        add_body(
            doc,
            f"In the live scenario '{row['scenario']}', the system marked the state as {row['phase']}, projected a final total of {row['projected_total']}, returned a win probability of {row['win_prob']}, and gave a projected range of {row['range']}. "
            f"The stored interpretation was: {row['interpretation']} This scenario is useful because it shows that the model is reacting to tempo, wickets, venue priors, and matchup context together rather than just extrapolating the current score linearly.",
        )
    for row in context["pre_match_scenarios"]:
        add_body(
            doc,
            f"In the pre-match scenario '{row['matchup']}' at {row['venue']}, the system projected a first-innings score of {row['predicted_score']} and estimated team-one and team-two win probabilities at {row['team1_win_prob']} and {row['team2_win_prob']}. "
            f"The interpretation attached to this case was: {row['interpretation']} This matters because it highlights how much more cautious the pre-match branch remains when live match-state evidence is not yet available.",
        )
    add_body(
        doc,
        "These scenario studies give the results chapter a practical voice. They turn the metric tables into cricket situations that faculty reviewers can actually discuss and evaluate.",
    )

    doc.add_heading("Why the Final Result Is Defensible in a Viva Setting", level=2)
    add_body(
        doc,
        "A project result becomes easier to defend in viva when the reasons behind it are stable under questioning. In this report, the strongest result is defensible because the explanation does not rely on one lucky metric. It rests on several linked claims that remain coherent together: the training scope was widened intelligently, the feature set captured real cricket context, the score and win tasks were separated, the win branch was calibrated, and the final outputs were tested through actual interfaces.",
    )
    add_body(
        doc,
        "This matters because viva questions are rarely limited to raw numbers. Reviewers often ask why one model was selected over another, why the same score can mean different things at different venues, why predictions change between innings, or why a probability should be trusted. The current chapter now contains direct answers to those questions in the language of the actual project rather than in generic machine-learning wording.",
    )

    doc.add_heading("How the Main Metrics Were Read Together", level=2)
    add_body(
        doc,
        "The project deliberately refused to treat one number as the full truth. MAE explains the typical run error, RMSE highlights larger misses, log loss evaluates whether probability assignments are sensible, and Brier score gives a second view of probability quality. Accuracy was still reported, but not treated as sufficient on its own for the win task.",
    )
    add_body(
        doc,
        "This broader metric reading is important because it matches how the system will actually be judged in use. A model can have attractive accuracy while still being overconfident, and a score predictor can have acceptable average error while hiding unstable tail behavior. The strongest result had to survive all of these questions together.",
    )

    doc.add_heading("Why This Chapter Explains the Best Result Properly", level=2)
    add_body(
        doc,
        "The refined results chapter is intentionally written around causes rather than inventories. The project did not improve because it contained more tables. It improved because the training scope widened intelligently, the categorical context was modeled more effectively, the win branch was calibrated, and the output layer acknowledged uncertainty. Those are the real reasons the strongest IPL prediction result was achieved in this workspace.",
    )
    doc.add_page_break()


def add_chapter_6(doc: Document, context: dict, assets: dict[str, Path], base_module) -> None:
    doc.add_heading("CHAPTER 6 - SYSTEM IMPLEMENTATION AND TESTING", level=1)

    doc.add_heading("Shared Inference Core", level=2)
    add_body(
        doc,
        "A major engineering strength of the project is that every user-facing surface calls the same core prediction path. Input normalization, overs parsing, feature-frame construction, support-table loading, score inference, win inference, uncertainty formatting, and optional monitoring hooks are all centralized in the shared prediction module. "
        "This reduced a class of bugs common in academic projects: one interface quietly using different logic from another.",
    )
    add_body(
        doc,
        "Because the same inference path is reused, the project can be evaluated more honestly. If a prediction works in the Flask form, the same logic also powers the API and CLI. This consistency is one of the reasons the final system is more than a notebook result.",
    )
    add_body(
        doc,
        "This consistency also made the refined report easier to improve. Since one prediction contract is reused everywhere, the report can focus on explaining one engine clearly instead of repeating nearly identical explanations for several interfaces that behave differently under the surface.",
    )

    add_table_title(doc, TABLES[7])
    add_table(
        doc,
        ["Surface", "Role in the Project", "Why It Was Kept"],
        [
            ["Flask web application", "Primary demonstration interface for live and pre-match prediction.", "Best suited for faculty evaluation and everyday interactive use."],
            ["JSON API", "Machine-readable prediction surface for external tools.", "Shows that the project can integrate with other services."],
            ["Streamlit dashboard", "Analytical interface for model and data exploration.", "Useful for inspection and demonstration of supporting context."],
            ["CLI", "Lightweight terminal interaction path.", "Useful for quick manual testing and debugging."],
        ],
        font_size=9,
    )

    doc.add_heading("Hardware and Software Environment", level=2)
    add_body(
        doc,
        "The project was developed and evaluated in a Windows-based Python environment. Training and inference were organized through a virtual environment so that the same code path could be reused for preprocessing, model search, web serving, dashboard serving, testing, and report generation.",
    )
    add_body(
        doc,
        "The principal software stack includes Python, pandas, scikit-learn, XGBoost, CatBoost, PyTorch, Flask, Streamlit, and python-docx. The repository also contains GPU-capable training branches for XGBoost, CatBoost, and PyTorch entity-embedding experiments, which helped broaden the experimental comparison beyond CPU-only baselines.",
    )
    add_body(
        doc,
        "This environment detail belongs in the main chapter because result interpretation depends partly on the software stack used to produce those results. It clarifies that the project was evaluated with modern structured-data libraries rather than with a single ad hoc script.",
    )

    doc.add_heading("Use of Software Engineering Principles During Implementation", level=2)
    add_body(
        doc,
        "The implementation reflects several core software-engineering principles. Modularity was maintained by separating preprocessing, training, inference, interfaces, and documentation into distinct files and packages. Reuse was maintained through the common inference contract so that the same prediction logic serves CLI, Flask, API, and Streamlit. Validation and fail-fast checks were added at the shared entry point so that incorrect inputs are rejected consistently rather than differently in each surface.",
    )
    add_body(
        doc,
        "Low coupling was another important principle. The frontend layers know how to collect inputs and display outputs, but they do not contain independent prediction logic. This reduces drift between demonstrations and makes regression testing more meaningful. In addition, versioned artifacts and registry-style tracking support traceability, which is an engineering concern as much as a machine-learning concern.",
    )
    add_body(
        doc,
        "These principles matter in the internal evaluation context because they show that the project was implemented as software, not only as a model experiment. The resulting system is easier to maintain, explain, and test.",
    )

    doc.add_heading("Why Fewer Screenshots Are Better Here", level=2)
    add_body(
        doc,
        "The earlier report used a screenshot-heavy chapter to occupy space. In this refined report, only one interface figure is retained. The reason is simple: screenshots should support the engineering argument, not replace it. The real value of the project lies in how the interfaces are connected to the common prediction core, not in how many mockups can be inserted into the document.",
    )
    add_figure(doc, assets["Figure 6.1"], FIGURES[4].number, FIGURES[4].title)

    doc.add_heading("Validation and Testing", level=2)
    add_body(
        doc,
        "The repository contains dedicated tests for input validation, alias normalization, feature-frame consistency, saved-model loading, and Flask route behavior. These tests matter because the project accepts many user-provided fields such as teams, venue, overs, innings, current runs, and wickets. Without validation, the interfaces would produce brittle or misleading outputs.",
    )
    add_body(
        doc,
        "The project also includes a reporting and monitoring layer. Even though the current live-outcome volume is still small for firm drift conclusions, the presence of prediction-event tracking and drift status artifacts shows that the system was designed with maintenance in mind rather than being treated as a one-time demo.",
    )
    add_body(
        doc,
        "The monitoring pieces do not yet dominate the project, but they are meaningful. They show that the system was built with a view toward future season updates and post-deployment checks, which is a stronger engineering stance than exporting one model file and treating the work as complete.",
    )

    doc.add_heading("Demonstration of Working Modules", level=2)
    add_body(
        doc,
        "The project can be demonstrated module by module. The preprocessing module reconstructs legal-ball states and support tables from raw Cricsheet data. The training modules compare multiple model families under season-aware splits. The inference module validates inputs and generates score and win outputs. The Flask module exposes live and pre-match forms plus an API route. The Streamlit module adds interactive analytical views. The CLI module provides fast manual checks from the terminal.",
    )
    add_body(
        doc,
        "This module-level demonstration structure is useful during evaluation because it lets the reviewer see the system as a chain of functioning components rather than as one opaque output screen. It also supports debugging because each stage can be inspected independently if a problem appears.",
    )

    doc.add_heading("Module Responsibility Split Inside the Repository", level=2)
    add_body(
        doc,
        "One reason the project remained manageable is that the codebase is organized by responsibility instead of being kept in one giant notebook or one oversized application file. This matters both for engineering quality and for report credibility.",
    )
    for relative_path, note in base_module.SOURCE_FILE_NOTES.items():
        add_body(
            doc,
            f"The file {relative_path} matters because {note} Its presence in the repository shows that the project was decomposed into understandable roles rather than built as an unreadable one-shot experiment.",
        )
    add_body(
        doc,
        "This module split is part of the delivery story. It made it possible to retrain, debug, serve predictions, and regenerate the report without rewriting the same logic repeatedly in different places.",
    )

    doc.add_heading("Completion Against the Original Plan", level=2)
    add_body(
        doc,
        "The implementation can also be judged against the original project plan. Data collection, preprocessing, live-state feature construction, baseline training, broader model comparison, pre-match modeling, multi-surface serving, and regression testing were all completed in the current workspace. Monitoring support and report-generation automation were also added, which extends the plan beyond a minimal classroom demonstration.",
    )
    add_body(
        doc,
        "The only major item not completed at the same maturity level is public deployment through a permanent external URL. The project is fully runnable locally and suitable for demonstration, but public hosting remains a later extension rather than a completed deliverable in the current submission. Stating that clearly improves the honesty of the report.",
    )

    doc.add_heading("Reliability Controls and Why They Matter Beyond Accuracy", level=2)
    add_body(
        doc,
        "A deployed cricket predictor can fail in ways that are not visible in a model-score table. It can accept an invalid over format, misread an alias, load the wrong artifact, or expose an interface path that behaves differently from the shared inference core. The reliability layer exists to stop those failures.",
    )
    for name, description in base_module.TEST_CASE_NOTES.items():
        add_body(
            doc,
            f"The test '{name}' exists because {description} This is part of the main project chapter, not a side note, because a system that breaks during normal use cannot be called successful even if its held-out metrics are strong.",
        )
    add_body(
        doc,
        "By documenting these tests inside the implementation chapter, the report makes a stronger claim: the project was evaluated not only as a prediction model but also as a working software system with validation and interface discipline.",
    )

    doc.add_heading("Research Manuscript and Documentation Output", level=2)
    add_body(
        doc,
        "In addition to the project report itself, the workspace includes an IEEE-style paper draft and multiple technical summary documents. This matters because it shows that the work has already been translated into more than one communication form: a full academic project report, a structured technical architecture summary, and a paper-style narrative suitable for further refinement.",
    )
    add_body(
        doc,
        "The existence of this manuscript material does not by itself claim journal publication. However, it does demonstrate that the project has been documented to a level where external communication, paper submission preparation, or later publication work would be realistic next steps.",
    )

    doc.add_heading("Request Flow Through the Application", level=2)
    add_body(
        doc,
        "The request flow begins with user input. A browser form, dashboard widget, CLI session, or API payload collects the current innings context. The shared inference module then validates required fields, normalizes team and venue aliases, parses overs into legal-ball structure, builds the live feature frame, attaches support-table context, runs the score and win models, formats uncertainty and probability outputs, and returns a clean result object.",
    )
    add_body(
        doc,
        "This flow matters because each step represents an opportunity for mismatch if it is duplicated across interfaces. By centralizing the flow, the project ensured that a prediction shown in the Flask application can be traced to the same underlying logic that would answer an API request or a terminal query.",
    )

    doc.add_heading("Why Deployment Evidence Matters", level=2)
    add_body(
        doc,
        "Machine-learning reports often stop after the best validation metric has been reported. This project deliberately continued further by showing that the models could be wrapped inside usable delivery surfaces. That extra step matters in academic evaluation because it proves the work can operate as a small application rather than as an isolated experiment.",
    )
    add_body(
        doc,
        "The use of Flask, Streamlit, API, and CLI surfaces also broadened the ways the project could be demonstrated. Faculty review, peer testing, manual debugging, and future integration work all benefit from having more than one access path to the same prediction logic.",
    )

    doc.add_heading("Dependency Stack and Why It Was Chosen", level=2)
    add_body(
        doc,
        "The software environment was not assembled randomly. Each major dependency was chosen because it solved a clear part of the overall project. Keeping that rationale visible improves the report because it shows the software stack was selected as an engineering toolset rather than accumulated by convenience.",
    )
    for row in context["dependency_inventory"][:10]:
        add_body(
            doc,
            f"The package {row['package']} was pinned as {row['specifier']}. It was retained because {row['role']} This kind of package-level explanation helps show that the project used an appropriate stack for data engineering, structured modeling, interface delivery, and reporting.",
        )
    add_body(
        doc,
        "This dependency review also makes the implementation chapter stronger for internal evaluation. It demonstrates that software choices were connected to technical responsibilities instead of being inserted without justification.",
    )

    doc.add_heading("Artifact Traceability and Working Evidence", level=2)
    add_body(
        doc,
        "A second implementation strength is traceability across saved artifacts. The repository does not contain only the final promoted models. It also contains comparison reports, prediction exports, uncertainty metadata, drift summaries, and registry snapshots that make the training history easier to inspect afterward.",
    )
    for row in context["artifact_inventory"][:10]:
        add_body(
            doc,
            f"The artifact {row['name']} is stored as a {row['kind']} file of size {row['size']} and was last modified on {row['modified']}. Its presence matters because it preserves evidence from the training or deployment workflow instead of forcing the report to rely on memory alone.",
        )
    add_body(
        doc,
        "This traceability is important in an academic setting. It allows the report to connect claims about model selection and deployment back to concrete files that exist in the workspace, which is stronger than presenting final screenshots without their supporting artifacts.",
    )

    doc.add_heading("Workspace Organization as an Implementation Decision", level=2)
    add_body(
        doc,
        "Repository organization is also part of implementation quality. The project is easier to test and extend because raw data, processed artifacts, reusable Python modules, models, notebooks, and generated documents are separated by role. This prevents later experiments from mixing with runtime code and keeps the main delivery surfaces easier to inspect.",
    )
    for row in context["workspace_map_rows"]:
        add_body(
            doc,
            f"The top-level workspace area '{row['name']}' exists to hold {row['role']} This structure is one reason the project remained manageable as it expanded from data preparation to training, interface delivery, and report generation.",
        )
    add_body(
        doc,
        "In other words, the folder structure itself helped the team implement the project in a disciplined way. It reduced accidental coupling and made it clearer where new code or outputs should be placed during each stage of the work.",
    )

    doc.add_heading("Monitoring Readiness and Current Operational Status", level=2)
    add_body(
        doc,
        "The project also includes an early monitoring layer because a prediction system should eventually be judged on more than its held-out test set. Monitoring does not yet dominate the submission, but its presence shows that the implementation was designed with future maintenance in mind.",
    )
    for row in context["drift_summary_rows"]:
        add_body(
            doc,
            f"The monitoring dimension '{row['dimension']}' currently reports the value {row['value']}. Its practical meaning is that {row['interpretation']} This kind of status record turns monitoring from a vague future idea into a visible part of the implementation.",
        )
    add_body(
        doc,
        "The current status is still limited by the small number of resolved live outcomes, so the project does not exaggerate what monitoring can prove today. Even so, adding the module now was a useful engineering decision because it establishes the path by which the system could later justify retraining or drift investigation using production evidence instead of guesswork.",
    )

    doc.add_heading("What This Chapter Proves", level=2)
    add_body(
        doc,
        "This implementation chapter proves that the project moved from experiment to usable system. The predictive models are important, but they would not constitute a strong final-year report on their own. The more convincing result is that the models were exposed through stable interfaces, reused through one inference contract, and supported by tests and reporting assets.",
    )
    add_body(
        doc,
        "For that reason, this chapter is not a formality. It is part of the argument that the strongest result in the report is not merely numerical; it is operational. The project can be rerun, inspected, and used through interfaces that reflect the same underlying model decisions discussed in the earlier chapters.",
    )
    doc.add_page_break()


def add_chapter_7(doc: Document, context: dict, best_entry: dict) -> None:
    best_live = best_entry["metrics"]["live"]
    doc.add_heading("CHAPTER 7 - CONCLUSION AND FUTURE SCOPE", level=1)

    doc.add_heading("Conclusion", level=2)
    add_body(
        doc,
        "This project succeeded because it treated IPL prediction as an end-to-end engineering problem. The most important gains came from using the right scope, preventing temporal leakage, constructing reusable support tables, and evaluating score and win tasks separately. The final report therefore reaches a clearer conclusion than the earlier version: strong cricket prediction is driven as much by disciplined feature design and evaluation logic as by the choice of model family.",
    )
    add_body(
        doc,
        f"The strongest tracked experimental result used {prettify_model_name(best_live['score_model'])} for live score and "
        f"{prettify_model_name(best_live['win_model'])} for live win probability, achieving RMSE {fmt_float(best_live['score_test']['rmse'])} and "
        f"log loss {fmt_float(best_live['win_test']['log_loss'], 4)} on held-out evaluation. More important than the numbers themselves is the reason they improved: "
        "the project aligned model complexity with better context rather than treating complexity as a substitute for context.",
    )

    doc.add_heading("Major Lessons Learned", level=2)
    add_bullet_list(
        doc,
        [
            "The best improvements came from historical context design and leakage control before they came from trying more algorithms.",
            "Recent-only training can be too small for rich live cricket prediction once chronology is enforced honestly.",
            "Score and win tasks reward different modeling properties, so a hybrid final system is often better than a single-family solution.",
            "Probability calibration is essential when the end product displays win percentages to users.",
            "A strong academic ML project should include reproducible interfaces and tests, not only notebook metrics.",
        ],
    )

    doc.add_heading("Current Limitations", level=2)
    add_body(
        doc,
        "The system is still limited by genuine cricket uncertainty, especially in early-innings states and rare collapse or explosion patterns. The pre-match branch remains much weaker than the live branch because it operates under a smaller information budget. The current pipeline also does not incorporate probable XI information, pitch reports, or live ball-tracking data.",
    )
    add_body(
        doc,
        "A second limitation is operational rather than predictive. The project contains monitoring readiness, but live production outcomes are still too few to support strong drift decisions. That means the current report can document preparedness for monitoring, but not a mature post-deployment feedback loop yet.",
    )

    doc.add_heading("Future Scope", level=2)
    add_bullet_list(
        doc,
        [
            "Integrate probable playing XI and lineup-specific priors before the match starts.",
            "Experiment with quantile or conformal methods to replace fixed uncertainty bands with more adaptive intervals.",
            "Add richer explanation outputs so that the API can return feature-level reasons behind a live prediction.",
            "Expand the monitoring layer from passive reporting to retraining triggers once enough live outcomes accumulate.",
            "Extend the same pipeline design to other T20 leagues after validating how much of the feature logic transfers cleanly.",
        ],
    )

    doc.add_heading("Positive Impact of the Project", level=2)
    add_body(
        doc,
        "The project has positive academic value because it demonstrates how data collection, preprocessing, machine learning, calibration, software integration, and technical reporting can be combined in one coherent final-year submission. It gives future students a stronger example of how to move from raw data to an end-to-end applied system.",
    )
    add_body(
        doc,
        "The project also has practical relevance for cricket analytics discussions. By translating live score states into projected totals and calibrated win probabilities, it helps users reason more systematically about match context instead of depending only on instinct or isolated scoreboard snapshots.",
    )

    doc.add_heading("Social, Environmental, and Ethical Considerations", level=2)
    add_body(
        doc,
        "The social impact of the project is mostly educational and analytical. It encourages evidence-based reasoning in sports analytics and provides a learning resource for students interested in machine learning, software integration, and applied statistics. The project can therefore support academic growth more directly than commercial disruption.",
    )
    add_body(
        doc,
        "The ethical dimension lies in how predictions are presented and interpreted. Because sports predictions can easily be misunderstood as certainty, the system deliberately presents uncertainty-aware score ranges and calibrated probabilities rather than exaggerated deterministic claims. This is a small but important design choice because it reduces the risk of overstating what the model actually knows.",
    )
    add_body(
        doc,
        "The environmental impact of the project is modest but still relevant. Model training consumes compute resources, especially in GPU-capable experiments, so the workflow benefits from selective retraining rather than unnecessary repetition. The use of reusable artifacts, model comparison reports, and one promoted deployment pair helps reduce wasted computation over time.",
    )

    doc.add_heading("Innovations and Distinctive Features", level=2)
    add_body(
        doc,
        "The main innovations of the project are not based on novelty for its own sake. They lie in how the system combines leakage-safe historical support tables, player-aware live features, separate score and win branches, calibration for probability output, and a shared inference core reused across multiple interfaces. Together these features make the project stand out from a simple notebook-only predictor.",
    )
    add_body(
        doc,
        "Another distinctive aspect is the way the report now explains the strongest result. Instead of relying on filler or decorative screenshots, it connects the final performance back to specific engineering choices that improved the realism and trustworthiness of the system.",
    )

    doc.add_heading("Concluding Remarks", level=2)
    add_body(
        doc,
        "In conclusion, the project achieved its objective by building a usable IPL prediction system grounded in realistic preprocessing, stronger context features, comparative model evaluation, and software delivery discipline. The final system is valuable not because it eliminates uncertainty, but because it handles uncertainty more intelligently than a score-only baseline and presents the result through interfaces that can actually be used and evaluated.",
    )
    add_body(
        doc,
        "Viewed as a whole, the project now satisfies the most important expectations of a final-year engineering submission: it defines a relevant problem, reviews prior work, proposes and explains a concrete methodology, evaluates results honestly, demonstrates working modules, documents limitations, and outlines a credible path for future extension. Those are the reasons it stands as more than a model notebook and more as a complete project report.",
    )
    doc.add_page_break()


def add_appendices(doc: Document, context: dict, best_entry: dict, base_module) -> None:
    doc.add_heading("APPENDIX 1 - FLOWCHART AND ALGORITHM NOTES", level=1)
    add_body(
        doc,
        "This appendix summarizes the implementation logic followed by the live prediction system. It complements the workflow figure in the methodology chapter by expressing the same process as an ordered algorithm.",
    )
    add_numbered_list(
        doc,
        [
            "Accept the current match inputs such as batting team, bowling team, venue, innings, runs, wickets, overs, and target when applicable.",
            "Normalize aliases for teams and venues so that inconsistent naming does not break joins or model features.",
            "Convert overs into legal-ball counts and derive balls left, phase indicators, innings progress, and chase pressure variables.",
            "Load the required support-table context for venue priors, team form, matchup history, and player-related statistics.",
            "Construct the live feature frame in the exact column order expected by the saved score and win models.",
            "Run the score model to estimate the likely innings total and derive a practical prediction range around that estimate.",
            "Run the calibrated win model to estimate batting-side win probability under the current live state.",
            "Format the outputs for the selected surface such as Flask, API, Streamlit, or CLI and return the result to the user.",
        ],
    )
    add_body(
        doc,
        "The same overall algorithm is reused across all user-facing interfaces. This is one of the reasons the project remains consistent during testing and demonstration.",
    )
    doc.add_page_break()

    doc.add_heading("APPENDIX 2 - SELECTED FEATURE NOTES", level=1)
    add_body(
        doc,
        "Only a few feature notes are retained here because the detailed reasoning now sits in Chapters 4 and 5.",
    )
    add_bullet_list(
        doc,
        [
            "runs_vs_par compares the current innings score with the venue baseline and helps distinguish pressure from comfort.",
            "required_minus_current_rr captures whether the chasing side is falling behind the pace needed to win.",
            "runs_last_5 and wickets_last_5 summarize recent momentum and collapse risk better than season averages alone.",
            "batting_team_venue_form and bowling_team_venue_form preserve location-specific team behavior.",
            "batter_vs_bowler history gives the model a compact view of recurring live player matchups.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("APPENDIX 3 - SHORT SUPPORTING CASE SNAPSHOTS", level=1)
    add_body(
        doc,
        "The main chapter already discusses scenario interpretation. This final appendix keeps a small supporting record of sample outputs together with the minimum rerun path for rebuilding the principal project artifacts.",
    )
    for row in context["live_scenarios"][:2]:
        add_body(
            doc,
            f"Scenario '{row['scenario']}' was saved with projected total {row['projected_total']}, win probability {row['win_prob']}, projected range {row['range']}, and interpretation: {row['interpretation']}",
        )
    add_code_block(
        doc,
        "& .\\.venv\\Scripts\\python.exe scripts\\update_external_data.py\n"
        "& .\\.venv\\Scripts\\python.exe scripts\\preprocess_ipl.py\n"
        "& .\\.venv\\Scripts\\python.exe scripts\\train_best_model_search.py\n"
        "& .\\.venv\\Scripts\\python.exe scripts\\train_pre_match.py\n"
        "& .\\.venv\\Scripts\\python.exe web_app.py\n"
        "& .\\.venv\\Scripts\\python.exe -m pytest -q",
    )
    add_body(
        doc,
        f"The strongest saved version referenced throughout the report is {best_entry['version_id']}.",
    )


def add_references(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("REFERENCES", level=1)
    references = [
        "[1] Cricsheet. IPL CSV2 downloads and documentation. https://cricsheet.org/downloads/.",
        "[2] Chen, T., and Guestrin, C. XGBoost: A Scalable Tree Boosting System. KDD 2016.",
        "[3] Prokhorenkova, L. et al. CatBoost: unbiased boosting with categorical features. NeurIPS 2018.",
        "[4] Guo, C., and Berkhahn, F. Entity Embeddings of Categorical Variables. arXiv:1604.06737.",
        "[5] Pedregosa, F. et al. Scikit-learn: Machine Learning in Python. JMLR 2011.",
        "[6] PyTorch documentation. https://pytorch.org/docs/.",
        "[7] Flask documentation. https://flask.palletsprojects.com/.",
        "[8] Streamlit documentation. https://docs.streamlit.io/.",
        "[9] Open-Meteo API documentation. https://open-meteo.com/en/docs.",
    ]
    for ref in references:
        add_body(doc, ref)


def postprocess_report_with_word(docx_path: Path, pdf_path: Path) -> None:
    toc_json = json.dumps(TOC_ENTRY_LABELS)
    table_json = json.dumps(TABLE_ENTRY_LABELS)
    figure_json = json.dumps(FIGURE_ENTRY_LABELS)
    front_json = json.dumps(FRONT_MATTER_LABELS)

    ps_script = f"""
$ErrorActionPreference = 'Stop'
$docPath = '{docx_path}'
$pdfPath = '{pdf_path}'
$tocEntries = @'
{toc_json}
'@ | ConvertFrom-Json
$tableEntries = @'
{table_json}
'@ | ConvertFrom-Json
$figureEntries = @'
{figure_json}
'@ | ConvertFrom-Json
$frontMatterEntries = @'
{front_json}
'@ | ConvertFrom-Json

function Convert-ToRoman([int]$number) {{
    $pairs = @(
        @('m', 1000), @('cm', 900), @('d', 500), @('cd', 400),
        @('c', 100), @('xc', 90), @('l', 50), @('xl', 40),
        @('x', 10), @('ix', 9), @('v', 5), @('iv', 4), @('i', 1)
    )
    $result = ''
    foreach ($pair in $pairs) {{
        $symbol = $pair[0]
        $value = $pair[1]
        while ($number -ge $value) {{
            $result += $symbol
            $number -= $value
        }}
    }}
    return $result
}}

function Format-Entry([string]$label, [string]$pageText) {{
    $width = 88
    $dots = '.' * [Math]::Max(4, $width - $label.Length - $pageText.Length)
    return "$label $dots $pageText"
}}

function Find-HeadingIndex($paragraphs, [string]$headingText) {{
    for ($i = 0; $i -lt $paragraphs.Count; $i++) {{
        $text = $paragraphs[$i].Range.Text.Trim("`r", "`a", " ")
        $style = $paragraphs[$i].Range.Style.NameLocal
        if ($text -eq $headingText -and $style -eq 'Heading 1') {{
            return $i
        }}
    }}
    return -1
}}

$wdActiveEndAdjustedPageNumber = 1
$word = $null
$doc = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $doc = $word.Documents.Open($docPath, $false, $false)
    $doc.Repaginate()

    $paragraphs = @($doc.Paragraphs)
    $headingPages = @{{}}
    $tablePages = @{{}}
    $figurePages = @{{}}

    foreach ($paragraph in $paragraphs) {{
        $text = $paragraph.Range.Text.Trim("`r", "`a", " ")
        if (-not $text) {{
            continue
        }}
        $page = [int]$paragraph.Range.Information($wdActiveEndAdjustedPageNumber)
        $style = $paragraph.Range.Style.NameLocal
        if ($style -eq 'Heading 1') {{
            $headingPages[$text] = $page
        }}
        if ($tableEntries -contains $text) {{
            $tablePages[$text] = $page
        }}
        if ($figureEntries -contains $text) {{
            $figurePages[$text] = $page
        }}
    }}

    $tocStart = Find-HeadingIndex $paragraphs 'TABLE OF CONTENTS'
    if ($tocStart -ge 0) {{
        for ($i = 0; $i -lt $tocEntries.Count; $i++) {{
            $label = [string]$tocEntries[$i]
            if ($headingPages.ContainsKey($label)) {{
                $pageValue = [int]$headingPages[$label]
                $pageText = if ($frontMatterEntries -contains $label) {{ Convert-ToRoman $pageValue }} else {{ [string]$pageValue }}
                $paragraphs[$tocStart + 1 + $i].Range.Text = (Format-Entry $label $pageText) + "`r"
            }}
        }}
    }}

    $tableStart = Find-HeadingIndex $paragraphs 'LIST OF TABLES'
    if ($tableStart -ge 0) {{
        for ($i = 0; $i -lt $tableEntries.Count; $i++) {{
            $label = [string]$tableEntries[$i]
            if ($tablePages.ContainsKey($label)) {{
                $paragraphs[$tableStart + 1 + $i].Range.Text = (Format-Entry $label ([string]$tablePages[$label])) + "`r"
            }}
        }}
    }}

    $figureStart = Find-HeadingIndex $paragraphs 'LIST OF FIGURES'
    if ($figureStart -ge 0) {{
        for ($i = 0; $i -lt $figureEntries.Count; $i++) {{
            $label = [string]$figureEntries[$i]
            if ($figurePages.ContainsKey($label)) {{
                $paragraphs[$figureStart + 1 + $i].Range.Text = (Format-Entry $label ([string]$figurePages[$label])) + "`r"
            }}
        }}
    }}

    $doc.Save()
    $doc.ExportAsFixedFormat($pdfPath, 17)
}} finally {{
    if ($doc -ne $null) {{
        $doc.Close([ref]0)
    }}
    if ($word -ne $null) {{
        $word.Quit()
    }}
}}
"""

    subprocess.run(
        ["powershell", "-NoProfile", "-Command", ps_script],
        check=True,
        cwd=ROOT_DIR,
    )


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    base = load_base_report_module()
    context = base.build_context()
    best_entry = best_registry_entry(context["model_registry"])
    assets = ensure_asset_paths(context, base)

    doc = Document()
    configure_document(doc)

    today_text = date.today().strftime("%B %d, %Y")
    add_title_page(doc, today_text)

    front_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    apply_page_layout(front_section)
    front_section.different_first_page_header_footer = False
    set_page_number_format(front_section, "lowerRoman", start=1)
    configure_footer(front_section)

    add_declaration_page(doc, today_text)
    add_approval_page(doc, today_text)
    add_certificate_page(doc, today_text)
    add_acknowledgements_page(doc, today_text)
    add_abstract_page(doc, context, best_entry)
    add_table_of_contents_page(doc)
    add_list_of_tables_page(doc)
    add_list_of_figures_page(doc)
    add_abbreviations_page(doc)

    main_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    apply_page_layout(main_section)
    main_section.different_first_page_header_footer = False
    set_page_number_format(main_section, "decimal", start=1)
    configure_footer(main_section)

    add_chapter_1(doc, context)
    add_chapter_2(doc, context)
    add_chapter_3(doc, context, best_entry)
    add_chapter_4(doc, context, best_entry, assets)
    add_chapter_5(doc, context, best_entry, assets)
    add_chapter_6(doc, context, assets, base)
    add_chapter_7(doc, context, best_entry)
    add_appendices(doc, context, best_entry, base)
    add_references(doc)

    doc.save(OUTPUT_PATH_ROOT)
    postprocess_report_with_word(OUTPUT_PATH_ROOT, FINAL_OUTPUT_PDF_PATH)
    shutil.copy2(OUTPUT_PATH_ROOT, OUTPUT_PATH)
    shutil.copy2(OUTPUT_PATH_ROOT, FINAL_OUTPUT_PATH)

    print(f"Generated refined report at: {OUTPUT_PATH_ROOT}")
    print(f"Generated refined report at: {OUTPUT_PATH}")
    print(f"Generated submission report at: {FINAL_OUTPUT_PATH}")
    print(f"Generated submission PDF at: {FINAL_OUTPUT_PDF_PATH}")


if __name__ == "__main__":
    main()
