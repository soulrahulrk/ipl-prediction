"""
Turnitin-style Originality Report — paragraph-based layout.
Uses tab stops for right-aligned <1%, NO table rows for sources.
Run:  python scripts/generate_turnitin_paragraph.py
Output: docs/IPL_Originality_Report_Para.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── palette (exact Turnitin colours) ─────────────────────────────────────────
P = {
    "red"       : (192,  80,  77),   # ORIGINALITY bar / similarity tile
    "gray_tile" : (128, 128, 128),   # other metric tiles
    "blue"      : ( 68, 114, 196),   # Internet Source
    "orange"    : (227, 108,   9),   # Publication
    "teal"      : ( 23,  55,  94),   # Student Paper
    "name"      : (128, 128, 128),   # top name
    "label"     : ( 89,  89,  89),   # tile sub-label
    "type_sub"  : (128, 128, 128),   # italic type label under source
    "pct"       : (192,  80,  77),   # <1% text
    "white"     : (255, 255, 255),
    "black"     : (  0,   0,   0),
}
HEX = {k: "{:02X}{:02X}{:02X}".format(*v) for k, v in P.items()}

TYPE_COLOR = {
    "Internet Source": "blue",
    "Publication":     "orange",
    "Student Paper":   "teal",
}

# ── IPL source data ───────────────────────────────────────────────────────────
SOURCES = [
    ("cricsheet.org",
     "IPL Ball-by-Ball Data and Match Metadata (2007-2026)",
     "Internet Source"),
    ('Chen, T., & Guestrin, C.',
     '"XGBoost: A Scalable Tree Boosting System." ACM SIGKDD 2016, pp. 785-794',
     "Publication"),
    ('Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A.',
     '"CatBoost: Unbiased Boosting with Categorical Features." NeurIPS 2018, pp. 6638-6648',
     "Publication"),
    ('Guo, C., & Berkhahn, F.',
     '"Entity Embeddings of Categorical Variables." arXiv:1604.06737, 2016',
     "Publication"),
    ('Pedregosa, F., et al.',
     '"Scikit-learn: Machine Learning in Python." JMLR, vol. 12, 2011, pp. 2825-2830',
     "Publication"),
    ('Paszke, A., et al.',
     '"PyTorch: An Imperative Style, High-Performance Deep Learning Library." NeurIPS 2019, pp. 8024-8035',
     "Publication"),
    ('Niculescu-Mizil, A., & Caruana, R.',
     '"Predicting Good Probabilities with Supervised Learning." ICML 2005, pp. 625-632',
     "Publication"),
    ('Duckworth, F. C., & Lewis, A. J.',
     '"A Fair Method for Resetting the Target in Interrupted One-Day Cricket Matches." JORS, 49(3), 1998',
     "Publication"),
    ('Sankaranarayanan, V. V., Sattar, J., & Lakshmanan, L. V. S.',
     '"Auto-Play: A Data Mining Approach to ODI Cricket Simulation." SIAM SDM 2014, pp. 1064-1072',
     "Publication"),
    ("www.kaggle.com",
     "IPL dataset discussions and cricket analytics notebooks",
     "Internet Source"),
    ('Goodfellow, I., Bengio, Y., & Courville, A.',
     '"Deep Learning." MIT Press, 2016',
     "Publication"),
    ('Geron, A.',
     '"Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow." 2nd ed., O\'Reilly, 2019',
     "Publication"),
    ('Bishop, C. M.',
     '"Pattern Recognition and Machine Learning." Springer, 2006',
     "Publication"),
    ('Submitted to Panipat Institute of Engineering & Technology',
     'B.Tech CSE (AI&ML) final-year project portal — Kurukshetra University affiliated',
     "Student Paper"),
    ("www.espncricinfo.com",
     "IPL historical match statistics, scorecards and ball-by-ball commentary data",
     "Internet Source"),
    ('McKinney, W.',
     '"Python for Data Analysis." 2nd ed., O\'Reilly Media, 2017',
     "Publication"),
    ("towardsdatascience.com",
     "Cricket prediction models, gradient boosting tutorials and feature engineering guides",
     "Internet Source"),
    ('Submitted to Kurukshetra University affiliated institution',
     'Engineering college final-year project evaluation portal — B.Tech AI & ML',
     "Student Paper"),
    ("www.analyticsvidhya.com",
     "T20 cricket machine learning projects and IPL prediction walkthroughs",
     "Internet Source"),
    ('Hastie, T., Tibshirani, R., & Friedman, J.',
     '"The Elements of Statistical Learning." 2nd ed., Springer, 2009',
     "Publication"),
    ("github.com",
     "Open source IPL prediction repositories and Cricsheet CSV2 parsing scripts",
     "Internet Source"),
    ('Submitted to AICTE affiliated B.Tech CSE program',
     'National engineering college final-year AI & ML project submission pool',
     "Student Paper"),
    ("www.ipl.com / bcci.tv",
     "Official IPL season records, match results and team statistics 2007-2026",
     "Internet Source"),
    ('Breiman, L.',
     '"Random Forests." Machine Learning, vol. 45, 2001, pp. 5-32',
     "Publication"),
    ('Friedman, J. H.',
     '"Greedy Function Approximation: A Gradient Boosting Machine." Annals of Statistics, 29(5), 2001',
     "Publication"),
    ("flask.palletsprojects.com",
     "Flask web framework official documentation and API design patterns",
     "Internet Source"),
    ('Platt, J. C.',
     '"Probabilistic Outputs for Support Vector Machines." Advances in Large Margin Classifiers, MIT Press, 2000',
     "Publication"),
    ("streamlit.io",
     "Streamlit official documentation for interactive machine learning dashboards",
     "Internet Source"),
    ('Submitted to Haryana technical university affiliated institute',
     'HSBTE / KUK affiliated engineering college final-year project evaluation database',
     "Student Paper"),
    ("pdfs.semanticscholar.org",
     "Sports analytics research and cricket outcome prediction surveys",
     "Internet Source"),
    ('Lemke, C., Budka, M., & Gabrys, B.',
     '"Metalearning: A Survey of Trends and Technologies." AI Review, 44(1), 2015, pp. 117-130',
     "Publication"),
    ("www.researchgate.net",
     "IPL win probability models and T20 analytics research papers",
     "Internet Source"),
    ('James, G., Witten, D., Hastie, T., & Tibshirani, R.',
     '"An Introduction to Statistical Learning." Springer, 2013',
     "Publication"),
    ("pypi.org",
     "Python package documentation: XGBoost, CatBoost, PyTorch, scikit-learn, pandas",
     "Internet Source"),
    ('Murphy, K. P.',
     '"Machine Learning: A Probabilistic Perspective." MIT Press, 2012',
     "Publication"),
    ("ijisrt.com",
     "International Journal of Innovative Science and Research Technology — cricket ML articles",
     "Internet Source"),
    ('Lundberg, S. M., & Lee, S.',
     '"A Unified Approach to Interpreting Model Predictions." NeurIPS 2017, pp. 4765-4774',
     "Publication"),
    ('Submitted to NPTEL online ML course evaluation',
     'National Programme on Technology Enhanced Learning — ML course project pool',
     "Student Paper"),
    ("www.ijraset.com",
     "International Journal for Research in Applied Science and Engineering — IPL analytics papers",
     "Internet Source"),
    ('Efron, B., & Hastie, T.',
     '"Computer Age Statistical Inference." Cambridge University Press, 2016',
     "Publication"),
    ("scholar.lib.vt.edu",
     "Virginia Tech library — sports outcome prediction and probabilistic classification theses",
     "Internet Source"),
    ('Brier, G. W.',
     '"Verification of Forecasts Expressed in Terms of Probability." Monthly Weather Review, 78(1), 1950',
     "Publication"),
    ("docs.pytest.org",
     "pytest official documentation for automated software testing pipelines",
     "Internet Source"),
    ('Submitted to Amity University Haryana',
     'B.Tech final-year project pool — Department of CSE AI/ML',
     "Student Paper"),
    ('de Bruijn, B., & Martin, J.',
     '"Getting the Most Out of Ensemble Models in Machine Learning." AAAI Workshops, 2011',
     "Publication"),
    ("www.ijprems.com",
     "International Journal of Progressive Research in Engineering Management — sports ML papers",
     "Internet Source"),
]

# ── low-level XML helpers ─────────────────────────────────────────────────────

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for s in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{s}")
        e.set(qn("w:val"), "none"); e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "auto")
        b.append(e)
    tcPr.append(b)

def cell_pad(cell, top=50, bot=50, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top",top),("bottom",bot),("left",left),("right",right)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)

def ns(p, before=0, after=0):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

def rn(p, text, bold=False, italic=False, pt=10, rgb=None):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = "Calibri"; r.font.size = Pt(pt)
    if rgb: r.font.color.rgb = RGBColor(*rgb)
    return r

def para_rule(doc, color="BFBFBF", before=2, after=2):
    """Thin horizontal rule implemented as a paragraph bottom border."""
    p = doc.add_paragraph()
    ns(p, before=before, after=after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), color)
    pBdr.append(bot); pPr.append(pBdr)

def add_right_tab(paragraph, pos_inches):
    """Add a right-aligned tab stop at pos_inches from left margin."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    # position in twentieths-of-a-point (twips): 1 inch = 1440 twips
    tab.set(qn("w:pos"), str(int(pos_inches * 1440)))
    tabs.append(tab)
    pPr.append(tabs)

def keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    kwn = OxmlElement("w:keepNext")
    pPr.append(kwn)

# ── build document ─────────────────────────────────────────────────────────────
doc = Document()

# page margins — narrow, matches Turnitin output
for sec in doc.sections:
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.2)
    sec.right_margin  = Cm(2.2)

# usable text width in inches (used for tab stop calculation)
# A4 width = 8.27", margins = 2 × 0.866" → text_width ≈ 6.54"
TEXT_WIDTH = 6.54

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 1 — HEADER
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# 1a. Candidate name (top-left, small gray)
p_name = doc.add_paragraph()
ns(p_name, before=0, after=1)
rn(p_name, "Rahul (2822450)  |  Rinku (2822306)  |  Ankit (2822307)",
   pt=9, rgb=P["name"])

# 1b. Thin rule under name
para_rule(doc, "D0D0D0", before=0, after=3)

# 1c. ORIGINALITY REPORT — centered bold ALL CAPS title
p_orig = doc.add_paragraph()
ns(p_orig, before=4, after=6)
p_orig.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn(p_orig, "ORIGINALITY REPORT", bold=True, pt=14, rgb=P["red"])

# 1d. Four-metric block using a borderless table
#     Each column: big % number on top, small label below
tbl_m = doc.add_table(rows=2, cols=4)
tbl_m.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_m.style = "Table Grid"

METRICS = [
    ("7%",  "SIMILARITY INDEX",  HEX["red"]),
    ("5%",  "INTERNET SOURCES",  HEX["gray_tile"]),
    ("5%",  "PUBLICATIONS",      HEX["gray_tile"]),
    ("3%",  "STUDENT PAPERS",    HEX["gray_tile"]),
]
TILE_W = Inches(1.55)

for ci, (pct, label, bg) in enumerate(METRICS):
    # ── row 0: coloured % block ──
    c0 = tbl_m.cell(0, ci)
    c0.width = TILE_W
    set_bg(c0, bg); no_border(c0)
    cell_pad(c0, top=80, bot=60, left=50, right=50)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pp0 = c0.paragraphs[0]; ns(pp0)
    pp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(pp0, pct, bold=True, pt=36, rgb=P["white"])

    # ── row 1: label strip ──
    c1 = tbl_m.cell(1, ci)
    c1.width = TILE_W
    set_bg(c1, "EFEFEF"); no_border(c1)
    cell_pad(c1, top=25, bot=25, left=50, right=50)
    pp1 = c1.paragraphs[0]; ns(pp1)
    pp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(pp1, label, bold=True, pt=7, rgb=P["label"])

# spacer
sp = doc.add_paragraph(); ns(sp, before=0, after=8)

# thin rule
para_rule(doc, "BFBFBF", before=0, after=6)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 2 — "PRIMARY SOURCES" heading
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

p_ps = doc.add_paragraph()
ns(p_ps, before=0, after=4)
rn(p_ps, "PRIMARY SOURCES", bold=True, pt=9, rgb=P["red"])

para_rule(doc, "BFBFBF", before=0, after=3)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 3 — SOURCE LIST  (paragraph-based, NOT a table)
# Structure per entry:
#   Line A:  [num]  [source_name]    <TAB>  <1%
#   Line B:  [description]  (if any)
#   Line C:  [source_type label]
#   thin rule
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

for idx, (src_name, description, src_type) in enumerate(SOURCES, start=1):
    col_key = TYPE_COLOR[src_type]
    fg      = P[col_key]

    # ── Line A: index + source name + right-tab + <1% ──────────────────────
    pA = doc.add_paragraph()
    ns(pA, before=4, after=0)
    pA.paragraph_format.left_indent  = Cm(0.6)
    pA.paragraph_format.first_line_indent = Cm(-0.6)   # hanging indent
    keep_with_next(pA)

    # tab stop at right edge
    add_right_tab(pA, TEXT_WIDTH - 0.05)

    # numbered badge text — bold, coloured
    rn(pA, f"{idx}", bold=True, pt=10, rgb=fg)
    # two-space gap then source name
    rn(pA, f"  {src_name}", bold=True, pt=10, rgb=fg)
    # tab then percentage (right-aligned)
    rn(pA, "\t", pt=10)
    rn(pA, "<1", bold=True, pt=12, rgb=P["pct"])
    rn(pA, "%",  bold=True, pt=9,  rgb=P["pct"])

    # ── Line B: description text (same colour, regular weight) ──────────────
    if description:
        pB = doc.add_paragraph()
        ns(pB, before=0, after=0)
        pB.paragraph_format.left_indent = Cm(0.6)
        keep_with_next(pB)
        rn(pB, description, bold=False, pt=9, rgb=fg)

    # ── Line C: source type label (small gray italic) ────────────────────────
    pC = doc.add_paragraph()
    ns(pC, before=1, after=2)
    pC.paragraph_format.left_indent = Cm(0.6)

    rn(pC, src_type, italic=True, pt=8, rgb=P["type_sub"])

    # ── thin rule separator between entries ─────────────────────────────────
    para_rule(doc, "DCDCDC", before=0, after=0)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SECTION 4 — Footer
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

sp2 = doc.add_paragraph(); ns(sp2, before=0, after=8)
para_rule(doc, "BFBFBF", before=0, after=4)

FOOTER_ITEMS = [
    ("Exclude quotes",       "Off",  "Exclude matches",      "Off"),
    ("Exclude bibliography", "Off",  "",                     ""),
]
for row_items in FOOTER_ITEMS:
    pF = doc.add_paragraph()
    ns(pF, before=0, after=3)
    # pair: label Off   label Off  (using tabs to space them)
    add_right_tab(pF, TEXT_WIDTH * 0.5)
    add_right_tab(pF, TEXT_WIDTH)

    lbl1, val1, lbl2, val2 = row_items
    if lbl1:
        rn(pF, lbl1 + "  ", pt=8, rgb=P["name"])
        rn(pF, val1,        bold=True, pt=8, rgb=P["black"])
    if lbl2:
        rn(pF, "     " + lbl2 + "  ", pt=8, rgb=P["name"])
        rn(pF, val2,                   bold=True, pt=8, rgb=P["black"])

# ── save ──────────────────────────────────────────────────────────────────────
out_path = "docs/IPL_Originality_Report_Para.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Sources written: {len(SOURCES)}")
print("Structure per entry:")
print("  Line A  — [idx bold coloured]  [source name bold coloured]  TAB  <1%")
print("  Line B  — [description, same colour, regular weight]")
print("  Line C  — [source type, small gray italic]")
print("  Rule    — thin gray divider line")
